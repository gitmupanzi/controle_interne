from __future__ import annotations

from dataclasses import dataclass, field, replace
from html import escape
from io import BytesIO
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import unicodedata
from typing import Any, Iterable

import numpy as np
import pandas as pd
from openpyxl.styles import Font, PatternFill

from credit_app.data_schema import (
    CURRENT_SAVINGS_SCHEMA,
    CUSTOMERS_SCHEMA,
    FIXED_SAVINGS_SCHEMA,
    G2_TRANSACTIONS_SCHEMA,
    LOANS_SCHEMA,
    MPESA_TRANSACTIONS_SCHEMA,
    PERFECT_CLIENTS_SCHEMA,
    DataSchema,
    validate_dataframe_schema,
)


TRANSACTION_REQUIRED_COLUMNS = set(MPESA_TRANSACTIONS_SCHEMA.required)
CURRENT_SAVINGS_REQUIRED_COLUMNS = set(CURRENT_SAVINGS_SCHEMA.required)
FIXED_SAVINGS_REQUIRED_COLUMNS = set(FIXED_SAVINGS_SCHEMA.required)
G2_TRANSACTION_REQUIRED_COLUMNS = set(G2_TRANSACTIONS_SCHEMA.required)
CUSTOMERS_REQUIRED_COLUMNS = set(CUSTOMERS_SCHEMA.required)
PERFECT_CLIENTS_REQUIRED_COLUMNS = set(PERFECT_CLIENTS_SCHEMA.required)

DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT = 11.0
DEFAULT_DAT_REPAYMENT_PREPARATION_HORIZON_DAYS = 30

CUSTOMER_STATEMENT_COLUMNS = [
    "date",
    "compte",
    "receipt_no",
    "devise",
    "description",
    "entree",
    "sortie",
    "solde",
]

CUSTOMER_DAT_INTEREST_COLUMNS = [
    "maturity_date",
    "savings_id",
    "customer_id",
    "msisdn",
    "currency_code",
    "product_name",
    "capital_place",
    "taux_interet_annuel_pct",
    "interet_client_constate",
    "voda_interest",
    "montant_echeance_client",
    "status",
    "reference_transaction_turbo",
    "date_ecriture_turbo",
    "statut_tracabilite",
    "source_interet",
    "impact_solde_mpesa",
]

CUSTOMER_ACTIVE_DAT_COLUMNS = [
    "date_situation",
    "savings_id",
    "customer_id",
    "msisdn",
    "currency_code",
    "product_name",
    "date_approved",
    "maturity_date",
    "duree_contractuelle_mois_estimee",
    "jours_avant_echeance",
    "balance",
    "taux_interet_annuel_pct",
    "interet_estime_echeance",
    "capital_plus_interet_estime",
    "situation_dat_client",
    "status",
]

CUSTOMER_STATEMENT_FOCUS_OPERATION_TYPES = frozenset(
    {
        "Sortie M-PESA_Turbo vers epargne",
        "Sortie M-PESA_Turbo vers DAT",
        "Entree M-PESA_Turbo depuis epargne",
        "Decaissement de credit",
        "Remboursement de credit",
        "Remboursement avec penalite",
    }
)

CUSTOMER_STATEMENT_LOGO_PATH = (
    Path(__file__).resolve().parents[2] / "skills" / "logo Bisou Bisou.PNG"
)


def _timestamp_plus(
    value: object,
    *,
    days: int = 0,
    hours: int = 0,
    minutes: int = 0,
    seconds: int = 0,
    microseconds: int = 0,
) -> pd.Timestamp:
    """Ajoute une duree avec des unites NumPy explicites."""
    total_microseconds = (
        (((days * 24 + hours) * 60 + minutes) * 60 + seconds) * 1_000_000
        + microseconds
    )
    base = np.datetime64(pd.Timestamp(value), "us")
    return pd.Timestamp(base + np.timedelta64(total_microseconds, "us"))


def _customer_statement_filename_token(value: object, fallback: str) -> str:
    """Nettoie un segment de nom de fichier sans supprimer les espaces du nom."""
    text = "" if _is_empty_text(value) else str(value).strip()
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", text)
    text = re.sub(r"\s+", " ", text).strip(" ._")
    return text or fallback


def build_customer_statement_filename(
    *,
    customer_id: object,
    customer_name: object = "",
    telephone: object = "",
    currency: object,
    period_start: object | None = None,
    period_end: object | None = None,
    g2_available: bool = False,
) -> str:
    """Construit le nom du Word client selon la disponibilite de G2.

    Turbo fournit toujours l'identifiant et le telephone. Quand G2 est charge,
    son nom client complete le nom du fichier sans modifier les mouvements.
    """
    customer_token = _customer_statement_filename_token(customer_id, "client")
    phone_token = _customer_statement_filename_token(telephone, "telephone_non_disponible")
    currency_token = _customer_statement_filename_token(str(currency).upper(), "devise")
    start = pd.to_datetime(period_start, errors="coerce")
    end = pd.to_datetime(period_end, errors="coerce")
    start_token = f"{start:%Y%m%d}" if pd.notna(start) else "debut"
    end_token = f"{end:%Y%m%d}" if pd.notna(end) else "fin"

    tokens = ["extrait_compte", customer_token]
    if g2_available:
        tokens.append(
            _customer_statement_filename_token(
                str(customer_name).upper(),
                "NOM NON DISPONIBLE",
            )
        )
    tokens.extend([phone_token, currency_token, start_token, end_token])
    return "_".join(tokens) + ".docx"

LOAN_USEFUL_COLUMNS = {
    "loan_id",
    "customer_id",
    "savings_account_id",
    "Nom_client",
    "customer",
    "msisdn1",
    "currency_code",
    "loan_amount",
    "loan_balance",
    "amount_paid",
    "outstanding_principle",
    "outstanding_setup_fees",
    "outstanding_interest",
    "outstanding_penalty_fees",
    "status_name",
    "due_date",
    "last_repayment_date",
    "created_at",
    "updated_at",
}

TEXT_COLUMNS = [
    "customer_id",
    "msisdn",
    "msisdn1",
    "account_type",
    "reference_id",
    "ref_no",
    "currency_code",
    "description",
    "loan_id",
    "customer",
    "product_name",
    "status_name",
]

DATE_COLUMNS = [
    "created_at",
    "updated_at",
    "date_approved",
    "maturity_date",
    "due_date",
    "last_repayment_date",
]

NUMERIC_COLUMNS = [
    "dr",
    "cr",
    "bal_before",
    "bal_after",
    "balance",
    "loan_amount",
    "loan_balance",
    "amount_paid",
    "outstanding_principle",
    "outstanding_setup_fees",
    "outstanding_interest",
    "outstanding_penalty_fees",
    "interest_earned",
    "voda_interest",
]

KNOWN_ACCOUNT_TYPES = {
    "MPESA ACCOUNT",
    "NORMAL SAVINGS",
    "FIXED SAVINGS",
    "LOAN ACCOUNT",
    "PRINCIPLE",
    "LOAN PORTFOLIO",
    "BISOU COLLECTION",
    "VODA COLLECTION A/C",
    "INTEREST EARNED",
    "LOAN AMOUNT A/C",
    "LOAN PENALTY FEES",
    "CUSTOMER USD WALLET PENALTY",
}

G2_OPERATION_CATEGORIES = [
    "DAT",
    "Depot normal",
    "Remboursement prets",
    "Paiement client B2C",
    "Demande de credit",
    "Operation interne Bisou",
    "Autre entree",
    "Autre sortie",
    "Flux a verifier",
]

G2_COMPLETED_STATUS_LABELS = frozenset({"completed", "complete", "successful", "success"})
# Conserver une fenetre assez large pour retrouver une sortie B2C Turbo, puis
# appliquer un seuil de controle plus strict pour le rapport d'anomalies.
G2_TURBO_OUTPUT_MATCH_TOLERANCE_MINUTES = 120.0
G2_TURBO_DATE_ANOMALY_TOLERANCE_MINUTES = 60.0

G2_CLASSIFIED_TRANSACTION_COLUMNS = [
    "date",
    "receipt_no",
    "currency_code",
    "details_rapport",
    "opposite_party",
    "duree",
    "compte_cree",
    "montant",
    "montant_entree",
    "montant_sortie",
    "balance_numeric",
]


@dataclass(frozen=True)
class MpesaPreparedData:
    transactions: pd.DataFrame
    current_savings: pd.DataFrame
    fixed_savings: pd.DataFrame
    loans: pd.DataFrame
    load_report: pd.DataFrame
    g2_transactions: pd.DataFrame = field(default_factory=pd.DataFrame)
    customers: pd.DataFrame = field(default_factory=pd.DataFrame)
    perfect_clients: pd.DataFrame = field(default_factory=pd.DataFrame)
    cache_fingerprint: str = ""
    fixed_savings_control: pd.DataFrame = field(default_factory=pd.DataFrame)


def _is_empty_text(value: Any) -> bool:
    if pd.isna(value):
        return True
    return str(value).strip().lower() in {"", "nan", "<na>", "none", "null"}


def normalize_label(value: Any) -> str:
    text = "" if pd.isna(value) else str(value)
    text = unicodedata.normalize("NFKD", text)
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = re.sub(r"\s+", " ", text).strip().lower()
    return text


def normalize_g2_transaction_status(value: Any) -> str:
    """Regroupe les statuts du portail G2 sans perdre leur valeur source."""
    status = normalize_label(value)
    if not status:
        return "Non renseigne"
    if status in G2_COMPLETED_STATUS_LABELS or "complete" in status or "success" in status:
        return "Completed"
    if any(token in status for token in ["declin", "reject", "refus", "failed", "failure", "echec"]):
        return "Declined"
    if any(token in status for token in ["cancel", "annul"]):
        return "Cancelled"
    if any(token in status for token in ["expir", "echeance depassee"]):
        return "Expired"
    if any(token in status for token in ["pending", "en attente", "processing", "en cours"]):
        return "Pending"
    return "Autre"


def g2_completed_transaction_mask(dataframe: pd.DataFrame | None) -> pd.Series:
    """Selectionne les transactions G2 terminees, avec compatibilite des anciens exports.

    Les anciens fichiers qui ne possedent aucun statut restent entierement
    exploitables. Des qu'au moins un statut est renseigne dans un fichier, seules
    les lignes explicitement terminees sont eligibles aux analyses metier.
    """
    if not isinstance(dataframe, pd.DataFrame):
        return pd.Series(dtype="bool")
    if dataframe.empty:
        return pd.Series(False, index=dataframe.index, dtype="bool")
    if "transaction_status" not in dataframe.columns:
        return pd.Series(True, index=dataframe.index, dtype="bool")
    normalized = dataframe["transaction_status"].apply(normalize_label)
    if not normalized.ne("").any():
        return pd.Series(True, index=dataframe.index, dtype="bool")
    return normalized.isin(G2_COMPLETED_STATUS_LABELS).astype(bool)


def classify_g2_business_operation(row: pd.Series, *, dat_matched: bool = False) -> str:
    """Classe une transaction G2 apres determination independante de son sens."""
    direction = normalize_label(row.get("sens_flux", ""))
    details = normalize_label(row.get("details", ""))
    reason_type = normalize_label(row.get("reason_type", ""))
    text = f"{reason_type} {details}"

    if "super transaction" in text or "supertransaction" in text:
        return "Operation interne Bisou"
    if direction not in {"entree", "sortie"}:
        return "Flux a verifier"
    if direction == "sortie":
        if any(token in text for token in ["bisoubisouloanrequest", "loan request", "loan payment", "loan payement"]):
            return "Demande de credit"
        if "bisoubisoub2c" in text or "b2c payment" in text:
            return "Paiement client B2C"
        return "Autre sortie"
    if "bisoubisouc2brepayment" in text or "bisoubisourepayment" in text or "repayment" in text:
        return "Remboursement prets"
    if dat_matched:
        return "DAT"
    if "bisoubisouc2b" in text or direction == "entree":
        return "Depot normal"
    return "Autre entree"


def clean_identifier(series: pd.Series) -> pd.Series:
    return (
        series.astype("string")
        .fillna("")
        .str.strip()
        .str.replace(r"\.0$", "", regex=True)
        .replace({"nan": "", "<NA>": "", "None": ""})
    )


def normalize_phone(series: pd.Series) -> pd.Series:
    digits = series.astype("string").fillna("").str.replace(r"\D+", "", regex=True)
    digits = digits.str.replace(r"^2430", "243", regex=True)
    digits = digits.str.replace(r"^0", "243", regex=True)
    needs_prefix = digits.ne("") & ~digits.str.startswith("243", na=False)
    digits = digits.where(~needs_prefix, "243" + digits)
    return digits.replace({"": pd.NA})


def clean_text(series: pd.Series) -> pd.Series:
    return series.astype("string").fillna("").str.strip()


def _normalize_column_name(column: object) -> str:
    text = str(column).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _normalize_business_column_key(column: object) -> str:
    text = normalize_label(_normalize_column_name(column))
    text = re.sub(r"[^0-9a-z]+", "_", text).strip("_")
    return text


def concat_unique(values: Iterable[object]) -> str:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        if _is_empty_text(value):
            continue
        text = str(value).strip()
        if text not in seen:
            result.append(text)
            seen.add(text)
    return " | ".join(result)


def first_non_empty(values: Iterable[object]) -> object:
    for value in values:
        if not _is_empty_text(value):
            return str(value).strip()
    return pd.NA


def extract_prefixed_reference(value: object, prefix: str) -> str:
    if _is_empty_text(value):
        return ""
    parts = [part.strip() for part in str(value).split("|")]
    matches = [part for part in parts if part.upper().startswith(prefix.upper())]
    return " | ".join(matches)


def concat_frames_stable(frames: Iterable[pd.DataFrame]) -> pd.DataFrame:
    prepared_frames: list[pd.DataFrame] = []
    expected_columns: list[object] = []

    for frame in frames:
        if frame is None or not isinstance(frame, pd.DataFrame) or frame.empty:
            continue
        for column in frame.columns:
            if column not in expected_columns:
                expected_columns.append(column)
        # Pandas warns when all-NA columns participate in dtype inference.
        # Drop them for concatenation, then restore the full column set below.
        useful_frame = frame.dropna(axis=1, how="all")
        if useful_frame.empty:
            continue
        prepared_frames.append(useful_frame)

    if not prepared_frames:
        return pd.DataFrame(columns=expected_columns)

    concatenated = pd.concat(prepared_frames, ignore_index=True)
    return concatenated.reindex(columns=expected_columns)


def validate_required_columns(
    dataframe: pd.DataFrame,
    required_columns: set[str],
    file_label: str,
) -> list[str]:
    known_schemas: tuple[DataSchema, ...] = (
        MPESA_TRANSACTIONS_SCHEMA,
        CURRENT_SAVINGS_SCHEMA,
        FIXED_SAVINGS_SCHEMA,
        G2_TRANSACTIONS_SCHEMA,
        CUSTOMERS_SCHEMA,
        LOANS_SCHEMA,
        PERFECT_CLIENTS_SCHEMA,
    )
    matching_schema = next(
        (candidate for candidate in known_schemas if set(candidate.required) == set(required_columns)),
        DataSchema(file_label, frozenset(required_columns)),
    )
    result = validate_dataframe_schema(dataframe, matching_schema, file_label)
    return list(result.missing)


def remove_export_index_columns(dataframe: pd.DataFrame) -> pd.DataFrame:
    if dataframe is None or dataframe.empty:
        return pd.DataFrame()
    frame = dataframe.copy()
    mask = frame.columns.astype(str).str.match(r"^Unnamed(:|$)", na=False)
    return frame.loc[:, ~mask].copy()


def _normalize_common_columns(dataframe: pd.DataFrame) -> pd.DataFrame:
    frame = remove_export_index_columns(dataframe)
    if frame.empty:
        return frame

    frame.columns = [_normalize_column_name(column) for column in frame.columns]
    for column in TEXT_COLUMNS:
        if column in frame.columns:
            frame[column] = clean_text(frame[column])
    for column in ["customer_id", "id", "reference_id", "ref_no", "loan_id"]:
        if column in frame.columns:
            frame[column] = clean_identifier(frame[column])
    for column in ["msisdn", "msisdn1"]:
        if column in frame.columns:
            frame[column] = normalize_phone(frame[column])
    if "currency_code" in frame.columns:
        frame["currency_code"] = clean_text(frame["currency_code"]).str.upper()
    if "account_type" in frame.columns:
        frame["account_type"] = clean_text(frame["account_type"]).str.upper()
    for column in DATE_COLUMNS:
        if column in frame.columns:
            frame[column] = pd.to_datetime(frame[column], errors="coerce")
    for column in NUMERIC_COLUMNS:
        if column in frame.columns:
            frame[column] = pd.to_numeric(frame[column], errors="coerce")
    return frame


def _deduplicate_multi_file_snapshot(
    dataframe: pd.DataFrame,
    *,
    source_column: str,
    source_trace_column: str,
    key_candidates: list[list[str]],
    recency_columns: list[str],
) -> pd.DataFrame:
    """Evite de recompter les chevauchements entre plusieurs exports d'une source."""
    if dataframe.empty:
        return dataframe.copy()
    frame = dataframe.copy().reset_index(drop=True)
    if source_column not in frame.columns or clean_text(frame[source_column]).nunique() <= 1:
        return frame.drop(columns=["ordre_fichier_import"], errors="ignore")

    frame["__import_row_order"] = np.arange(len(frame))
    selected_keys: list[str] = []
    valid_key = pd.Series(False, index=frame.index)
    for candidate in key_candidates:
        if not candidate or not set(candidate).issubset(frame.columns):
            continue
        candidate_valid = pd.Series(True, index=frame.index)
        for column in candidate:
            values = frame[column]
            if pd.api.types.is_datetime64_any_dtype(values) or pd.api.types.is_numeric_dtype(values):
                candidate_valid &= values.notna()
            else:
                candidate_valid &= clean_text(values).ne("")
        if candidate_valid.any():
            selected_keys = candidate
            valid_key = candidate_valid
            break

    provenance_exclusions = {
        source_column,
        source_trace_column,
        "ordre_fichier_import",
        "__import_row_order",
    }
    if not selected_keys:
        comparison_columns = [column for column in frame.columns if column not in provenance_exclusions]
        result = frame.drop_duplicates(subset=comparison_columns, keep="last")
        return result.drop(columns=["ordre_fichier_import", "__import_row_order"], errors="ignore").reset_index(drop=True)

    keyed = frame.loc[valid_key].copy()
    unkeyed = frame.loc[~valid_key].copy()
    keyed[source_trace_column] = keyed.groupby(selected_keys, dropna=False)[source_column].transform(concat_unique)
    sort_columns = [column for column in recency_columns if column in keyed.columns]
    if "ordre_fichier_import" in keyed.columns:
        sort_columns.append("ordre_fichier_import")
    sort_columns.append("__import_row_order")
    keyed = (
        keyed.sort_values(sort_columns, na_position="first")
        .drop_duplicates(selected_keys, keep="last")
    )

    if not unkeyed.empty:
        comparison_columns = [column for column in unkeyed.columns if column not in provenance_exclusions]
        unkeyed = unkeyed.drop_duplicates(subset=comparison_columns, keep="last")
        unkeyed[source_trace_column] = clean_text(unkeyed[source_column])

    result = concat_frames_stable([keyed, unkeyed]).sort_values("__import_row_order")
    return result.drop(columns=["ordre_fichier_import", "__import_row_order"], errors="ignore").reset_index(drop=True)


def prepare_transactions(dataframe: pd.DataFrame) -> pd.DataFrame:
    frame = _normalize_common_columns(dataframe)
    for column in ["dr", "cr", "bal_before", "bal_after"]:
        if column in frame.columns:
            frame[column] = pd.to_numeric(frame[column], errors="coerce").fillna(0.0)
    frame = _deduplicate_multi_file_snapshot(
        frame,
        source_column="fichier_source_transactions_turbo",
        source_trace_column="fichiers_sources_transactions_turbo",
        key_candidates=[
            ["id"],
            ["ref_no", "account_type", "customer_id", "currency_code", "dr", "cr", "created_at"],
        ],
        recency_columns=["created_at"],
    )
    if "created_at" in frame.columns:
        frame = frame.sort_values(["created_at", "id"], na_position="last").reset_index(drop=True)
    return frame


def prepare_savings_accounts(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    """Prépare Savings Account ou, à défaut, les deux synthèses historiques.

    La source complète est toujours prioritaire lorsqu'elle est chargée avec les
    exports ``Customers with Current/Fixed Savings Account``. Cela évite de
    recompter les mêmes comptes positifs et conserve les comptes à solde nul.
    """
    frame = _normalize_common_columns(dataframe if dataframe is not None else pd.DataFrame())
    if "msisdn" not in frame.columns and "msisdn1" in frame.columns:
        frame["msisdn"] = normalize_phone(frame["msisdn1"])

    # ``savings_id`` appartient à la source maître. Les deux exports résumés
    # historiques n'en disposent pas. Si les deux familles sont téléversées
    # ensemble, garder tous les enregistrements des fichiers maîtres et ignorer
    # les synthèses qui reprendraient les mêmes soldes positifs.
    savings_id_available = clean_text(
        frame.get("savings_id", pd.Series("", index=frame.index))
    ).ne("")
    source_column = "fichier_source_epargne_turbo"
    if source_column in frame.columns:
        source_names = clean_text(frame[source_column])
        master_source_names = set(source_names.loc[savings_id_available & source_names.ne("")])
        master_rows = source_names.isin(master_source_names)
        master_rows |= source_names.eq("") & savings_id_available
    else:
        master_rows = savings_id_available
    source_complete_available = bool(master_rows.any())
    if source_complete_available:
        frame = frame.loc[master_rows].copy()
    frame["source_savings_account_complete"] = source_complete_available

    account_type = clean_text(
        frame.get("account_type", pd.Series("", index=frame.index))
    ).str.upper()
    normalized_account_type = account_type.apply(normalize_label)
    explicit_current = normalized_account_type.str.contains(
        r"\bcurrent account\b|\bnormal savings\b",
        regex=True,
        na=False,
    )
    explicit_fixed = normalized_account_type.str.contains(
        r"\bfixed account\b|\bfixed savings\b",
        regex=True,
        na=False,
    )
    product_text = (
        clean_text(frame.get("product_name", pd.Series("", index=frame.index)))
        + " "
        + clean_text(
            frame.get("product_description", pd.Series("", index=frame.index))
        )
    ).apply(normalize_label)
    inferred_current = product_text.str.contains(
        r"\bopen savings\b|\bcurrent account\b|\bnormal savings\b",
        regex=True,
        na=False,
    )
    inferred_fixed = product_text.str.contains(
        r"\bfixed account\b|\bfixed savings\b",
        regex=True,
        na=False,
    )
    frame["account_type"] = np.select(
        [explicit_current, explicit_fixed, account_type.ne(""), inferred_current, inferred_fixed],
        ["NORMAL SAVINGS", "FIXED SAVINGS", account_type, "NORMAL SAVINGS", "FIXED SAVINGS"],
        default="",
    )
    if "created_at" not in frame.columns:
        frame["created_at"] = pd.NaT
    if "date_approved" in frame.columns:
        fixed_without_creation_date = (
            frame["account_type"].eq("FIXED SAVINGS")
            & frame["created_at"].isna()
            & frame["date_approved"].notna()
        )
        frame.loc[fixed_without_creation_date, "created_at"] = frame.loc[
            fixed_without_creation_date, "date_approved"
        ]
    if "balance" in frame.columns:
        frame["balance"] = pd.to_numeric(frame["balance"], errors="coerce").fillna(0.0)

    account_types = frame.get("account_type", pd.Series("", index=frame.index))
    current = frame.loc[account_types.eq("NORMAL SAVINGS")].copy()
    fixed = frame.loc[account_types.eq("FIXED SAVINGS")].copy()
    other = frame.loc[~account_types.isin(["NORMAL SAVINGS", "FIXED SAVINGS"])].copy()

    prepared_parts = [
        _deduplicate_multi_file_snapshot(
            current,
            source_column=source_column,
            source_trace_column="fichiers_sources_epargne_turbo",
            key_candidates=[
                ["savings_id"],
                ["customer_id", "currency_code", "account_type", "product_name", "created_at"],
            ],
            recency_columns=["updated_at", "created_at"],
        ),
        _deduplicate_multi_file_snapshot(
            fixed,
            source_column=source_column,
            source_trace_column="fichiers_sources_epargne_turbo",
            key_candidates=[
                ["savings_id"],
                [
                    "customer_id", "currency_code", "account_type", "product_name",
                    "date_approved", "maturity_date",
                ],
            ],
            recency_columns=["updated_at", "date_approved", "created_at"],
        ),
        _deduplicate_multi_file_snapshot(
            other,
            source_column=source_column,
            source_trace_column="fichiers_sources_epargne_turbo",
            key_candidates=[
                ["savings_id"],
                ["customer_id", "currency_code", "account_type", "product_name", "created_at"],
            ],
            recency_columns=["updated_at", "created_at"],
        ),
    ]
    return concat_frames_stable(prepared_parts).reset_index(drop=True)


def prepare_current_savings(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    frame = prepare_savings_accounts(dataframe)
    # Une vue résumée ancienne peut ne pas porter de type reconnu. Dans ce cas,
    # conserver son comportement historique; l'export complet, lui, est scindé.
    if frame.get("account_type", pd.Series("", index=frame.index)).eq("NORMAL SAVINGS").any():
        frame = frame.loc[frame["account_type"].eq("NORMAL SAVINGS")].copy()
    return frame.reset_index(drop=True)


def prepare_fixed_savings_from_accounts(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    """Extrait les DAT actifs ou historiques de la source Savings Account."""
    frame = prepare_savings_accounts(dataframe)
    if frame.empty or "account_type" not in frame.columns:
        return pd.DataFrame()
    return frame.loc[frame["account_type"].eq("FIXED SAVINGS")].reset_index(drop=True)


def prepare_fixed_savings(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    frame = _normalize_common_columns(dataframe if dataframe is not None else pd.DataFrame())
    if "balance" in frame.columns:
        frame["balance"] = pd.to_numeric(frame["balance"], errors="coerce").fillna(0.0)
    if "created_at" not in frame.columns and "date_approved" in frame.columns:
        frame["created_at"] = frame["date_approved"]
    return _deduplicate_multi_file_snapshot(
        frame,
        source_column="fichier_source_dat_turbo",
        source_trace_column="fichiers_sources_dat_turbo",
        key_candidates=[
            ["customer_id", "currency_code", "account_type", "date_approved", "maturity_date"],
        ],
        recency_columns=["date_approved"],
    )


def build_savings_accounts_reconciliation(
    prepared: MpesaPreparedData,
) -> dict[str, pd.DataFrame]:
    """Décrit Savings Account et rapproche un ancien export DAT s'il est fourni."""
    master = prepared.fixed_savings.copy()
    control = prepared.fixed_savings_control.copy()
    empty = {"synthese": pd.DataFrame(), "ecarts": pd.DataFrame()}
    if master.empty:
        return empty

    master["balance"] = numeric_column(master, "balance")
    positive = master.loc[master["balance"].gt(0)].copy()
    zero = master.loc[master["balance"].eq(0)].copy()
    if "source_savings_account_complete" in master.columns:
        source_complete_available = bool(
            master["source_savings_account_complete"].fillna(False).astype(bool).any()
        )
    else:
        # Compatibilité avec les objets préparés avant l'ajout du marqueur.
        source_complete_available = "fichier_source_epargne_turbo" in master.columns

    key_columns = [
        "customer_id",
        "currency_code",
        "date_approved",
        "maturity_date",
        "balance",
    ]

    def comparison_keys(frame: pd.DataFrame) -> pd.Series:
        if frame.empty or not set(key_columns).issubset(frame.columns):
            return pd.Series(dtype="string")
        work = frame[key_columns].copy()
        work["customer_id"] = clean_identifier(work["customer_id"])
        work["currency_code"] = clean_text(work["currency_code"]).str.upper()
        for column in ["date_approved", "maturity_date"]:
            work[column] = pd.to_datetime(work[column], errors="coerce").dt.strftime(
                "%Y-%m-%d %H:%M:%S"
            ).fillna("")
        work["balance"] = pd.to_numeric(work["balance"], errors="coerce").round(2)
        return work.astype("string").fillna("").agg("|".join, axis=1)

    if control.empty:
        summary = pd.DataFrame(
            [
                {
                    "source_savings_account_complete_disponible": source_complete_available,
                    "comptes_courants": int(len(prepared.current_savings)),
                    "dat_total_source_complete": int(len(master)),
                    "dat_solde_positif": int(len(positive)),
                    "dat_solde_nul": int(len(zero)),
                    "dat_export_resume": 0,
                    "dat_export_retrouves": 0,
                    "dat_export_absents_source_complete": 0,
                    "dat_positifs_absents_export_resume": 0,
                    "statut_rapprochement": (
                        "Source autonome"
                        if source_complete_available
                        else "Syntheses positives de compatibilite"
                    ),
                }
            ]
        )
        return {"synthese": summary, "ecarts": pd.DataFrame()}

    positive_keys = comparison_keys(positive)
    control_keys = comparison_keys(control)
    control_key_set = set(control_keys.astype(str))
    positive_key_set = set(positive_keys.astype(str))
    positive_missing_mask = ~positive_keys.astype(str).isin(control_key_set)
    control_missing_mask = ~control_keys.astype(str).isin(positive_key_set)
    positive_missing = positive.loc[positive_missing_mask].copy()
    control_missing = control.loc[control_missing_mask].copy()
    ecarts = concat_frames_stable(
        [
            positive_missing.assign(
                type_ecart="DAT positif de Savings Account absent de l'export DAT"
            ),
            control_missing.assign(
                type_ecart="DAT de l'export résumé absent de Savings Account"
            ),
        ]
    )
    summary = pd.DataFrame(
        [
            {
                "source_savings_account_complete_disponible": source_complete_available,
                "comptes_courants": int(len(prepared.current_savings)),
                "dat_total_source_complete": int(len(master)),
                "dat_solde_positif": int(len(positive)),
                "dat_solde_nul": int(len(zero)),
                "dat_export_resume": int(len(control)),
                "dat_export_retrouves": int((~control_missing_mask).sum()),
                "dat_export_absents_source_complete": int(control_missing_mask.sum()),
                "dat_positifs_absents_export_resume": int(positive_missing_mask.sum()),
                "statut_rapprochement": (
                    "Concordance exacte"
                    if not control.empty
                    and not positive_missing_mask.any()
                    and not control_missing_mask.any()
                    else "Ecarts a verifier"
                ),
            }
        ]
    )
    return {"synthese": summary, "ecarts": ecarts.reset_index(drop=True)}


def build_large_dat_summary(
    fixed_savings: pd.DataFrame | None,
    *,
    percentile: float = 0.90,
    as_of_date: Any | None = None,
) -> dict[str, pd.DataFrame]:
    """Classe les clients DAT par devise et identifie les encours les plus eleves."""
    empty_result = {"clients": pd.DataFrame(), "portefeuille": pd.DataFrame()}
    if not isinstance(fixed_savings, pd.DataFrame) or fixed_savings.empty:
        return empty_result

    required = {"customer_id", "currency_code", "balance"}
    if not required.issubset(fixed_savings.columns):
        return empty_result

    percentile = min(max(float(percentile), 0.0), 1.0)
    reporting_date = pd.Timestamp(as_of_date if as_of_date is not None else pd.Timestamp.now()).normalize()
    dat = fixed_savings.copy()
    dat["customer_id"] = clean_identifier(dat["customer_id"])
    dat["currency_code"] = clean_text(dat["currency_code"]).str.upper()
    dat["balance"] = pd.to_numeric(dat["balance"], errors="coerce").fillna(0.0)
    dat = dat.loc[
        dat["customer_id"].ne("")
        & dat["currency_code"].ne("")
        & dat["balance"].gt(0)
    ].copy()
    if dat.empty:
        return empty_result

    for column in ["date_approved", "maturity_date"]:
        dat[column] = pd.to_datetime(dat.get(column, pd.Series(pd.NaT, index=dat.index)), errors="coerce")
    for column in ["Nom_client", "msisdn", "product_name", "account_type"]:
        if column not in dat.columns:
            dat[column] = pd.NA

    dat["dat_echu"] = dat["maturity_date"].notna() & dat["maturity_date"].lt(reporting_date)
    dat["echeance_30j"] = (
        dat["maturity_date"].notna()
        & dat["maturity_date"].ge(reporting_date)
        & dat["maturity_date"].le(_timestamp_plus(reporting_date, days=30))
    )
    dat["solde_dat_echu"] = dat["balance"].where(dat["dat_echu"], 0.0)
    dat["solde_echeance_30j"] = dat["balance"].where(dat["echeance_30j"], 0.0)

    clients = (
        dat.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
        .agg(
            Nom_client=("Nom_client", concat_unique),
            telephone=("msisdn", concat_unique),
            produits_dat=("product_name", concat_unique),
            types_comptes_dat=("account_type", concat_unique),
            nb_comptes_dat=("balance", "size"),
            solde_dat_total=("balance", "sum"),
            plus_fort_dat=("balance", "max"),
            solde_dat_echu=("solde_dat_echu", "sum"),
            solde_echeance_30j=("solde_echeance_30j", "sum"),
            nb_dat_echus=("dat_echu", "sum"),
            nb_echeances_30j=("echeance_30j", "sum"),
            date_premier_dat=("date_approved", "min"),
            date_dernier_dat=("date_approved", "max"),
            prochaine_echeance=("maturity_date", "min"),
        )
    )
    clients["rang_devise"] = (
        clients.groupby("currency_code")["solde_dat_total"]
        .rank(method="dense", ascending=False)
        .astype("int64")
    )
    clients["total_portefeuille_devise"] = clients.groupby("currency_code")["solde_dat_total"].transform("sum")
    clients["part_portefeuille_pct"] = (
        clients["solde_dat_total"]
        .div(clients["total_portefeuille_devise"].replace(0, pd.NA))
        .mul(100)
        .fillna(0.0)
    )
    clients["seuil_fort_dat"] = clients.groupby("currency_code")["solde_dat_total"].transform(
        lambda values: values.quantile(percentile)
    )
    clients["est_fort_dat"] = clients["solde_dat_total"].ge(clients["seuil_fort_dat"])
    clients = clients.sort_values(
        ["currency_code", "solde_dat_total", "customer_id"],
        ascending=[True, False, True],
    ).reset_index(drop=True)
    clients["part_cumulee_pct"] = clients.groupby("currency_code")["part_portefeuille_pct"].cumsum()

    portefeuille = (
        clients.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            total_dat=("solde_dat_total", "sum"),
            nb_clients_dat=("customer_id", "nunique"),
            nb_comptes_dat=("nb_comptes_dat", "sum"),
            solde_median_client=("solde_dat_total", "median"),
            seuil_fort_dat=("seuil_fort_dat", "first"),
            nb_clients_forts=("est_fort_dat", "sum"),
            solde_dat_echu=("solde_dat_echu", "sum"),
            solde_echeance_30j=("solde_echeance_30j", "sum"),
        )
    )
    concentration = (
        clients.loc[clients["est_fort_dat"]]
        .groupby("currency_code", as_index=False)["solde_dat_total"]
        .sum()
        .rename(columns={"solde_dat_total": "total_clients_forts"})
    )
    portefeuille = portefeuille.merge(concentration, on="currency_code", how="left")
    portefeuille["total_clients_forts"] = portefeuille["total_clients_forts"].fillna(0.0)
    portefeuille["concentration_clients_forts_pct"] = (
        portefeuille["total_clients_forts"]
        .div(portefeuille["total_dat"].replace(0, pd.NA))
        .mul(100)
        .fillna(0.0)
    )
    portefeuille["percentile_fort_dat"] = percentile * 100
    portefeuille["date_analyse"] = reporting_date
    return {"clients": clients, "portefeuille": portefeuille}


def prepare_loans(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    frame = _normalize_common_columns(dataframe if dataframe is not None else pd.DataFrame())
    if not frame.empty and "outstanding_principal" in frame.columns and "outstanding_principle" not in frame.columns:
        frame["outstanding_principle"] = frame["outstanding_principal"]
    for column in NUMERIC_COLUMNS:
        if column in frame.columns:
            frame[column] = pd.to_numeric(frame[column], errors="coerce").fillna(0.0)
    return _deduplicate_multi_file_snapshot(
        frame,
        source_column="fichier_source_credits_turbo",
        source_trace_column="fichiers_sources_credits_turbo",
        key_candidates=[["loan_id"], ["id"]],
        recency_columns=["updated_at", "created_at"],
    )


def prepare_customers(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    frame = remove_export_index_columns(dataframe if dataframe is not None else pd.DataFrame())
    frame = _normalize_common_columns(frame)
    return _deduplicate_multi_file_snapshot(
        frame,
        source_column="fichier_source_clients_turbo",
        source_trace_column="fichiers_sources_clients_turbo",
        key_candidates=[["customer_id"], ["msisdn1", "created_at"]],
        recency_columns=["updated_at", "created_at"],
    )


def prepare_perfect_clients(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    """Prepare l'export 122 Perfect sans transformer les numeros invalides en cles de jointure."""
    frame = remove_export_index_columns(dataframe if dataframe is not None else pd.DataFrame())
    if frame.empty:
        return frame

    frame = frame.rename(columns={column: _normalize_business_column_key(column) for column in frame.columns}).copy()
    for column in ["id_client", "code_client", "num_manuel"]:
        if column in frame.columns:
            frame[column] = clean_identifier(frame[column])
    for column in [
        "nom_complet",
        "phone_brut",
        "phone_prefixe",
        "statut_phone",
        "commentaire_phone",
        "type_client",
        "categorie_client",
        "gestionnaire",
        "collecteur",
    ]:
        if column in frame.columns:
            frame[column] = clean_text(frame[column])

    source_phone = frame.get("phone_prefixe", pd.Series(pd.NA, index=frame.index, dtype="string"))
    frame["phone_prefixe_source"] = source_phone
    normalized_phone = normalize_phone(source_phone)
    valid_phone = normalized_phone.astype("string").str.fullmatch(r"243[89]\d{8}", na=False)
    frame["phone_prefixe"] = normalized_phone.where(valid_phone, pd.NA)
    frame["source_perfect"] = "Perfect"
    return _deduplicate_multi_file_snapshot(
        frame,
        source_column="fichier_source_clients_perfect",
        source_trace_column="fichiers_sources_clients_perfect",
        key_candidates=[["id_client"], ["code_client"], ["num_manuel", "nom_complet"]],
        recency_columns=[],
    )


def _mpesa_identity_source(
    dataframe: pd.DataFrame | None,
    *,
    source: str,
    system: str,
    phone_column: str,
    customer_column: str | None = "customer_id",
    name_column: str = "Nom_client",
) -> pd.DataFrame:
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        return pd.DataFrame()
    frame = dataframe.copy()
    phone_values = frame.get(phone_column, pd.Series(pd.NA, index=frame.index, dtype="string"))
    result = pd.DataFrame(index=frame.index)
    result["phone_prefixe"] = normalize_phone(phone_values)
    valid_phone = result["phone_prefixe"].astype("string").str.fullmatch(r"243[89]\d{8}", na=False)
    result["phone_prefixe"] = result["phone_prefixe"].where(valid_phone, pd.NA)
    result["customer_id_turbo"] = (
        clean_identifier(frame[customer_column])
        if customer_column and customer_column in frame.columns
        else ""
    )
    result["nom_client_mpesa"] = (
        clean_text(frame[name_column]) if name_column in frame.columns else ""
    )
    result["source_mpesa"] = source
    result["systeme_mpesa"] = system
    fallback = (
        result["customer_id_turbo"].where(result["customer_id_turbo"].ne(""), frame.index.astype("string"))
    )
    result["cle_rapprochement"] = result["phone_prefixe"].astype("string")
    no_phone = result["phone_prefixe"].isna()
    result.loc[no_phone, "cle_rapprochement"] = (
        "SANS-TELEPHONE::" + system + "::" + fallback.loc[no_phone].astype("string")
    )
    return result.reset_index(drop=True)


def _build_mpesa_identity_population(prepared: MpesaPreparedData) -> pd.DataFrame:
    completed_g2 = prepared.g2_transactions
    if isinstance(completed_g2, pd.DataFrame) and not completed_g2.empty:
        completed_g2 = completed_g2.loc[g2_completed_transaction_mask(completed_g2)].copy()
    frames = [
        _mpesa_identity_source(prepared.transactions, source="Turbo - Transactions", system="Turbo", phone_column="msisdn1"),
        _mpesa_identity_source(prepared.customers, source="Turbo - Clients", system="Turbo", phone_column="msisdn1"),
        _mpesa_identity_source(prepared.current_savings, source="Turbo - Epargne courante", system="Turbo", phone_column="msisdn"),
        _mpesa_identity_source(prepared.fixed_savings, source="Turbo - DAT", system="Turbo", phone_column="msisdn"),
        _mpesa_identity_source(prepared.loans, source="Turbo - Credits", system="Turbo", phone_column="msisdn1"),
        _mpesa_identity_source(
            completed_g2,
            source="G2 - Transactions",
            system="G2",
            phone_column="phone_prefixe",
            customer_column=None,
        ),
    ]
    population = concat_frames_stable(frames)
    if population.empty:
        return pd.DataFrame()
    population["present_dans_turbo"] = population["systeme_mpesa"].astype("string").eq("Turbo")
    population["present_dans_g2"] = population["systeme_mpesa"].astype("string").eq("G2")
    output_columns = [
        "cle_rapprochement", "phone_prefixe", "customer_ids_turbo", "noms_clients_mpesa",
        "systemes_mpesa", "sources_mpesa", "present_dans_turbo", "present_dans_g2",
        "nb_lignes_sources_mpesa",
    ]
    key_counts = population.groupby("cle_rapprochement", dropna=False)["cle_rapprochement"].transform("size")
    unique_rows = population.loc[
        key_counts.eq(1),
        [
            "cle_rapprochement", "phone_prefixe", "customer_id_turbo", "nom_client_mpesa",
            "systeme_mpesa", "source_mpesa", "present_dans_turbo", "present_dans_g2",
        ],
    ].rename(
        columns={
            "customer_id_turbo": "customer_ids_turbo",
            "nom_client_mpesa": "noms_clients_mpesa",
            "systeme_mpesa": "systemes_mpesa",
            "source_mpesa": "sources_mpesa",
        }
    )
    unique_rows["nb_lignes_sources_mpesa"] = 1

    duplicates = population.loc[key_counts.gt(1)].copy()
    if duplicates.empty:
        return unique_rows[output_columns].reset_index(drop=True)
    duplicate_summary = (
        duplicates.groupby("cle_rapprochement", as_index=False, dropna=False)
        .agg(
            phone_prefixe=("phone_prefixe", first_non_empty),
            customer_ids_turbo=("customer_id_turbo", concat_unique),
            noms_clients_mpesa=("nom_client_mpesa", concat_unique),
            systemes_mpesa=("systeme_mpesa", concat_unique),
            sources_mpesa=("source_mpesa", concat_unique),
            present_dans_turbo=("present_dans_turbo", "max"),
            present_dans_g2=("present_dans_g2", "max"),
            nb_lignes_sources_mpesa=("source_mpesa", "size"),
        )
    )
    return concat_frames_stable([unique_rows[output_columns], duplicate_summary[output_columns]]).reset_index(drop=True)


def _build_mpesa_operation_detail(
    prepared: MpesaPreparedData,
    daily_detail: pd.DataFrame | None = None,
) -> pd.DataFrame:
    blocks: list[pd.DataFrame] = []
    transactions = prepared.transactions
    if isinstance(transactions, pd.DataFrame) and not transactions.empty:
        turbo = transactions.copy()
        if "account_type" in turbo.columns and turbo["account_type"].astype("string").eq("MPESA ACCOUNT").any():
            turbo = turbo.loc[turbo["account_type"].astype("string").eq("MPESA ACCOUNT")].copy()
        turbo["phone_prefixe"] = normalize_phone(turbo.get("msisdn1", pd.Series(pd.NA, index=turbo.index)))
        valid_phone = turbo["phone_prefixe"].astype("string").str.fullmatch(r"243[89]\d{8}", na=False)
        turbo["phone_prefixe"] = turbo["phone_prefixe"].where(valid_phone, pd.NA)
        turbo["customer_id_turbo"] = clean_identifier(
            turbo.get("customer_id", pd.Series("", index=turbo.index))
        )
        turbo["nom_client_mpesa"] = clean_text(
            turbo.get("Nom_client", pd.Series("", index=turbo.index))
        )
        turbo["date_operation"] = pd.to_datetime(turbo.get("created_at"), errors="coerce")
        turbo["operation_reference"] = clean_identifier(
            turbo.get("ref_no", pd.Series("", index=turbo.index))
        )
        fallback_reference = clean_identifier(
            turbo.get("reference_id", pd.Series("", index=turbo.index))
        )
        turbo["operation_reference"] = turbo["operation_reference"].where(
            turbo["operation_reference"].ne(""), fallback_reference
        )
        missing_reference = turbo["operation_reference"].eq("")
        turbo.loc[missing_reference, "operation_reference"] = (
            "TURBO-LIGNE-" + turbo.index[missing_reference].astype("string")
        )
        turbo["currency_code"] = clean_text(
            turbo.get("currency_code", pd.Series("", index=turbo.index))
        ).str.upper()
        turbo["entree_mpesa"] = numeric_column(turbo, "cr")
        turbo["sortie_mpesa"] = numeric_column(turbo, "dr")
        turbo["mouvement_net_mpesa"] = turbo["entree_mpesa"] - turbo["sortie_mpesa"]
        turbo["description_operation"] = clean_text(
            turbo.get("description", pd.Series("", index=turbo.index))
        )
        turbo["type_compte"] = clean_text(
            turbo.get("account_type", pd.Series("", index=turbo.index))
        )
        turbo["source_operation"] = "Turbo"
        turbo["statut_operation"] = ""
        group_keys = [
            "phone_prefixe", "customer_id_turbo", "nom_client_mpesa", "date_operation",
            "operation_reference", "currency_code", "source_operation", "statut_operation",
        ]
        turbo_detail = (
            turbo.groupby(group_keys, as_index=False, dropna=False)
            .agg(
                description_operation=("description_operation", concat_unique),
                type_compte=("type_compte", concat_unique),
                entree_mpesa=("entree_mpesa", "sum"),
                sortie_mpesa=("sortie_mpesa", "sum"),
                mouvement_net_mpesa=("mouvement_net_mpesa", "sum"),
                nb_lignes_source=("operation_reference", "size"),
            )
        )
        turbo_detail["montant_operation"] = turbo_detail["mouvement_net_mpesa"].abs()
        turbo_detail["sens_operation"] = np.select(
            [turbo_detail["mouvement_net_mpesa"].gt(0), turbo_detail["mouvement_net_mpesa"].lt(0)],
            ["Entree M-PESA_Turbo", "Sortie M-PESA_Turbo"],
            default="Mouvement nul",
        )
        turbo_detail["type_operation"] = turbo_detail.apply(
            lambda row: classify_mpesa_operation(
                row["description_operation"], row["type_compte"], row["mouvement_net_mpesa"]
            ),
            axis=1,
        )
        blocks.append(turbo_detail)

    if isinstance(prepared.g2_transactions, pd.DataFrame) and not prepared.g2_transactions.empty:
        g2_report = (
            daily_detail.copy()
            if isinstance(daily_detail, pd.DataFrame)
            else build_g2_daily_savings_report(prepared).get("detail", pd.DataFrame())
        )
        g2 = pd.DataFrame()
        if not g2_report.empty:
            g2 = g2_report.copy()
            if "incluse_synthese" in g2.columns:
                eligible = g2["incluse_synthese"].astype("boolean").fillna(False).astype(bool)
                g2 = g2.loc[eligible].copy()
            else:
                g2 = g2.loc[g2_completed_transaction_mask(g2)].copy()
        if not g2.empty:
            g2["phone_prefixe"] = normalize_phone(g2.get("phone_prefixe", pd.Series(pd.NA, index=g2.index)))
            valid_phone = g2["phone_prefixe"].astype("string").str.fullmatch(r"243[89]\d{8}", na=False)
            g2["phone_prefixe"] = g2["phone_prefixe"].where(valid_phone, pd.NA)
            g2["customer_id_turbo"] = ""
            g2["nom_client_mpesa"] = clean_text(g2.get("Nom_client", pd.Series("", index=g2.index)))
            g2["date_operation"] = pd.to_datetime(g2.get("date"), errors="coerce")
            g2["operation_reference"] = clean_identifier(g2.get("receipt_no", pd.Series("", index=g2.index)))
            g2["currency_code"] = clean_text(g2.get("currency_code", pd.Series("", index=g2.index))).str.upper()
            g2["type_operation"] = clean_text(g2.get("details_rapport", pd.Series("", index=g2.index)))
            g2["description_operation"] = g2["type_operation"]
            g2["type_compte"] = "G2"
            g2["montant_operation"] = numeric_column(g2, "montant").abs()
            g2["entree_mpesa"] = np.nan
            g2["sortie_mpesa"] = np.nan
            g2["mouvement_net_mpesa"] = np.nan
            g2["sens_operation"] = clean_text(
                g2.get("sens_flux", pd.Series("Indetermine", index=g2.index))
            ).map({"Entree": "Entree G2", "Sortie": "Sortie G2"}).fillna("Flux G2 a verifier")
            g2["source_operation"] = "G2"
            g2["statut_operation"] = clean_text(
                g2.get("transaction_status", pd.Series("", index=g2.index))
            )
            g2["nb_lignes_source"] = 1
            blocks.append(
                g2[[
                    "phone_prefixe", "customer_id_turbo", "nom_client_mpesa", "date_operation",
                    "operation_reference", "currency_code", "type_operation", "description_operation",
                    "type_compte", "montant_operation", "entree_mpesa", "sortie_mpesa",
                    "mouvement_net_mpesa", "sens_operation", "source_operation",
                    "statut_operation", "nb_lignes_source",
                ]]
            )

    detail = concat_frames_stable(blocks)
    if detail.empty:
        return detail
    detail["cle_rapprochement"] = detail["phone_prefixe"].astype("string")
    no_phone = detail["phone_prefixe"].isna()
    fallback = detail["customer_id_turbo"].astype("string").where(
        detail["customer_id_turbo"].astype("string").ne(""), detail.index.astype("string")
    )
    detail.loc[no_phone, "cle_rapprochement"] = (
        "SANS-TELEPHONE::" + detail.loc[no_phone, "source_operation"].astype("string") + "::" + fallback.loc[no_phone]
    )
    return detail.sort_values(["date_operation", "source_operation"], na_position="last").reset_index(drop=True)


def _aggregate_perfect_clients(perfect_clients: pd.DataFrame) -> pd.DataFrame:
    if not isinstance(perfect_clients, pd.DataFrame) or perfect_clients.empty or "phone_prefixe" not in perfect_clients.columns:
        return pd.DataFrame()
    perfect = perfect_clients.dropna(subset=["phone_prefixe"]).copy()
    if perfect.empty:
        return pd.DataFrame()
    for column in [
        "id_client", "code_client", "num_manuel", "nom_complet", "statut_phone",
        "commentaire_phone", "type_client", "categorie_client", "gestionnaire", "collecteur",
    ]:
        if column not in perfect.columns:
            perfect[column] = ""
    output_map = {
        "id_client": "ids_clients_perfect",
        "code_client": "codes_clients_perfect",
        "num_manuel": "numeros_manuels_perfect",
        "nom_complet": "noms_clients_perfect",
        "statut_phone": "statuts_phone_perfect",
        "commentaire_phone": "commentaires_phone_perfect",
        "type_client": "types_clients_perfect",
        "categorie_client": "categories_clients_perfect",
        "gestionnaire": "gestionnaires_perfect",
        "collecteur": "collecteurs_perfect",
    }
    output_columns = ["phone_prefixe", "nb_clients_perfect", *output_map.values()]
    phone_counts = perfect.groupby("phone_prefixe", dropna=False)["phone_prefixe"].transform("size")

    # La plupart des numeros sont uniques : les copier directement evite des milliers
    # d'appels Python groupby inutiles. Seuls les numeros partages sont agreges.
    unique_rows = perfect.loc[phone_counts.eq(1), ["phone_prefixe", *output_map]].copy()
    unique_rows = unique_rows.rename(columns=output_map)
    unique_rows["nb_clients_perfect"] = 1
    unique_rows = unique_rows[output_columns]

    duplicate_rows = perfect.loc[phone_counts.gt(1)].copy()
    if duplicate_rows.empty:
        return unique_rows.reset_index(drop=True)
    duplicate_summary = (
        duplicate_rows.groupby("phone_prefixe", as_index=False, dropna=False)
        .agg(
            nb_clients_perfect=("id_client", lambda values: max(1, clean_identifier(values).replace("", pd.NA).nunique())),
            ids_clients_perfect=("id_client", concat_unique),
            codes_clients_perfect=("code_client", concat_unique),
            numeros_manuels_perfect=("num_manuel", concat_unique),
            noms_clients_perfect=("nom_complet", concat_unique),
            statuts_phone_perfect=("statut_phone", concat_unique),
            commentaires_phone_perfect=("commentaire_phone", concat_unique),
            types_clients_perfect=("type_client", concat_unique),
            categories_clients_perfect=("categorie_client", concat_unique),
            gestionnaires_perfect=("gestionnaire", concat_unique),
            collecteurs_perfect=("collecteur", concat_unique),
        )
    )
    return concat_frames_stable([unique_rows, duplicate_summary[output_columns]]).reset_index(drop=True)


def _perfect_match_status(frame: pd.DataFrame, perfect_available: bool) -> pd.Series:
    if not perfect_available:
        status = pd.Series("Fichier Perfect absent", index=frame.index, dtype="string")
    else:
        matches = numeric_column(frame, "nb_clients_perfect").astype(int)
        status = pd.Series(
            np.select(
                [matches.eq(1), matches.gt(1)],
                ["Trouve dans Perfect - correspondance unique", "Trouve dans Perfect - plusieurs clients"],
                default="Non trouve dans Perfect",
            ),
            index=frame.index,
            dtype="string",
        )
    return status.where(frame["phone_prefixe"].notna(), "Telephone Turbo/G2 inexploitable")


def _add_system_presence_columns(frame: pd.DataFrame) -> pd.DataFrame:
    """Ajoute les indicateurs de présence Turbo, G2 et Perfect sans modifier le grain."""
    output = frame.copy()
    turbo = (
        output["present_dans_turbo"].astype("boolean").fillna(False).astype(bool)
        if "present_dans_turbo" in output.columns
        else pd.Series(False, index=output.index)
    )
    g2 = (
        output["present_dans_g2"].astype("boolean").fillna(False).astype(bool)
        if "present_dans_g2" in output.columns
        else pd.Series(False, index=output.index)
    )
    perfect = numeric_column(output, "nb_clients_perfect").gt(0)
    output["present_dans_turbo"] = turbo
    output["present_dans_g2"] = g2
    output["present_dans_perfect"] = perfect
    output["present_dans_les_3_systemes"] = turbo & g2 & perfect
    output["statut_presence_systemes"] = np.select(
        [
            turbo & g2 & perfect,
            turbo & g2,
            turbo & perfect,
            g2 & perfect,
            turbo,
            g2,
            perfect,
        ],
        [
            "Present dans G2, Turbo et Perfect",
            "Present dans G2 et Turbo - absent Perfect",
            "Present dans Turbo et Perfect - absent G2",
            "Present dans G2 et Perfect - absent Turbo",
            "Present dans Turbo uniquement",
            "Present dans G2 uniquement",
            "Present dans Perfect uniquement",
        ],
        default="Presence systeme indeterminee",
    )
    return output


def build_perfect_client_crosscheck(prepared: MpesaPreparedData) -> dict[str, pd.DataFrame]:
    """Croise la population unifiee Turbo + G2 avec Perfect, une ligne par telephone."""
    population = _build_mpesa_identity_population(prepared)
    operations = _build_mpesa_operation_detail(prepared)
    perfect_by_phone = _aggregate_perfect_clients(prepared.perfect_clients)
    perfect_available = isinstance(prepared.perfect_clients, pd.DataFrame) and not prepared.perfect_clients.empty
    if population.empty:
        return {
            "synthese": pd.DataFrame(),
            "operations": operations,
            "perfect_agrege": perfect_by_phone,
            "clients_perfect_dans_mpesa": pd.DataFrame(),
            "clients_perfect_dans_turbo": pd.DataFrame(),
            "clients_perfect_dans_turbo_et_mpesa": pd.DataFrame(),
            "clients_trois_systemes": pd.DataFrame(),
        }

    if not operations.empty:
        operation_summary = (
            operations.groupby("cle_rapprochement", as_index=False, dropna=False)
            .agg(
                types_operations_mpesa=("type_operation", concat_unique),
                sources_operations=("source_operation", concat_unique),
                devises_operations=("currency_code", concat_unique),
                premiere_operation=("date_operation", "min"),
                derniere_operation=("date_operation", "max"),
                nombre_operations_observees=("operation_reference", "size"),
                references_operations=("operation_reference", concat_unique),
            )
        )
        turbo_counts = operations["source_operation"].eq("Turbo").groupby(operations["cle_rapprochement"]).sum()
        g2_counts = operations["source_operation"].eq("G2").groupby(operations["cle_rapprochement"]).sum()
        operation_summary["nombre_operations_turbo"] = operation_summary["cle_rapprochement"].map(turbo_counts).fillna(0).astype(int)
        operation_summary["nombre_operations_g2"] = operation_summary["cle_rapprochement"].map(g2_counts).fillna(0).astype(int)
        summary = population.merge(operation_summary, on="cle_rapprochement", how="left")
    else:
        summary = population.copy()
        summary["nombre_operations_observees"] = 0
        summary["nombre_operations_turbo"] = 0
        summary["nombre_operations_g2"] = 0

    if not perfect_by_phone.empty:
        summary = summary.merge(perfect_by_phone, on="phone_prefixe", how="left")
    summary["statut_rapprochement_perfect"] = _perfect_match_status(summary, perfect_available)
    summary["nb_clients_perfect"] = numeric_column(summary, "nb_clients_perfect").astype(int)
    summary["trouve_dans_perfect"] = summary["nb_clients_perfect"].gt(0)
    summary = _add_system_presence_columns(summary)

    if not operations.empty:
        identity_columns = [
            "cle_rapprochement", "customer_ids_turbo", "noms_clients_mpesa",
            "systemes_mpesa", "sources_mpesa", "present_dans_turbo", "present_dans_g2",
        ]
        operations = operations.merge(summary[identity_columns], on="cle_rapprochement", how="left")
        if not perfect_by_phone.empty:
            operations = operations.merge(perfect_by_phone, on="phone_prefixe", how="left")
        operations["statut_rapprochement_perfect"] = _perfect_match_status(operations, perfect_available)
        operations["nb_clients_perfect"] = numeric_column(operations, "nb_clients_perfect").astype(int)
        operations = _add_system_presence_columns(operations)

    summary = summary.sort_values(
        ["statut_rapprochement_perfect", "phone_prefixe"], na_position="last"
    ).reset_index(drop=True)
    clients_perfect_dans_mpesa = summary.loc[
        summary["present_dans_g2"] & summary["present_dans_perfect"]
    ].reset_index(drop=True)
    clients_perfect_dans_turbo = summary.loc[
        summary["present_dans_turbo"] & summary["present_dans_perfect"]
    ].reset_index(drop=True)
    clients_perfect_dans_turbo_et_mpesa = summary.loc[
        summary["present_dans_les_3_systemes"]
    ].reset_index(drop=True)
    return {
        "synthese": summary,
        "operations": operations,
        "perfect_agrege": perfect_by_phone,
        "clients_perfect_dans_mpesa": clients_perfect_dans_mpesa,
        "clients_perfect_dans_turbo": clients_perfect_dans_turbo,
        "clients_perfect_dans_turbo_et_mpesa": clients_perfect_dans_turbo_et_mpesa,
        "clients_trois_systemes": clients_perfect_dans_turbo_et_mpesa,
    }


def _parse_money_series(series: pd.Series) -> pd.Series:
    cleaned = (
        series.astype("string")
        .fillna("")
        .str.replace(r"[^\d,\.\-]", "", regex=True)
        .str.replace(",", "", regex=False)
        .replace("", pd.NA)
    )
    return pd.to_numeric(cleaned, errors="coerce")


def numeric_column(frame: pd.DataFrame, column: str, default: float = 0.0) -> pd.Series:
    if isinstance(frame, pd.DataFrame) and column in frame.columns:
        return pd.to_numeric(frame[column], errors="coerce").fillna(default)
    index = frame.index if isinstance(frame, pd.DataFrame) else None
    return pd.Series(default, index=index, dtype="float64")


def _extract_phone_from_opposite_party(series: pd.Series) -> pd.Series:
    raw_phone = (
        series.astype("string")
        .fillna("")
        .str.split("-", n=1)
        .str[0]
        .str.replace(r"[\r\n\t]+", " ", regex=True)
        .str.replace(r"\s+", " ", regex=True)
        .str.strip()
    )
    return raw_phone.replace({"": pd.NA})


def _extract_customer_name_from_opposite_party(series: pd.Series) -> pd.Series:
    customer_name = (
        series.astype("string")
        .fillna("")
        .str.split("-", n=1)
        .str[1]
        .astype("string")
        .str.replace(r"[\r\n\t]+", " ", regex=True)
        .str.replace(r"\s+", " ", regex=True)
        .str.strip()
    )
    return customer_name.replace({"": pd.NA})


def promote_g2_statement_header(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    """Promouvoir l'en-tête métier des relevés organisation G2.

    Les exports 1441/15558 commencent par cinq lignes d'identification du
    compte. ``pandas.read_excel`` utilise alors la première ligne comme noms
    de colonnes. Cette fonction retrouve la ligne contenant Receipt No.,
    Currency et Opposite Party, y compris après concaténation de plusieurs
    fichiers dont la deuxième colonne porte un nom d'organisation différent.
    """
    if dataframe is None or not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        return pd.DataFrame()

    required_header_keys = {"receipt_no", "currency", "opposite_party"}
    provenance_columns = {
        "fichier_source_g2",
        "ordre_fichier_import",
    }

    def promote_block(block: pd.DataFrame) -> pd.DataFrame:
        candidate = block.dropna(axis=1, how="all").copy()
        available_keys = {
            _normalize_business_column_key(column)
            for column in candidate.columns
            if column not in provenance_columns
        }
        if required_header_keys.issubset(available_keys):
            return candidate.reset_index(drop=True)

        business_columns = [
            column for column in candidate.columns if column not in provenance_columns
        ]
        for position in range(min(20, len(candidate))):
            header_values = candidate.iloc[position]
            row_keys = {
                _normalize_business_column_key(header_values[column])
                for column in business_columns
                if not _is_empty_text(header_values[column])
            }
            if not required_header_keys.issubset(row_keys):
                continue
            rename_map = {
                column: _normalize_column_name(header_values[column])
                for column in business_columns
                if not _is_empty_text(header_values[column])
            }
            promoted = candidate.iloc[position + 1 :].copy().rename(columns=rename_map)
            return promoted.reset_index(drop=True)
        return candidate.reset_index(drop=True)

    if "fichier_source_g2" not in dataframe.columns:
        return promote_block(dataframe)

    blocks = [
        promote_block(group)
        for _, group in dataframe.groupby(
            "fichier_source_g2", dropna=False, sort=False
        )
    ]
    return concat_frames_stable(blocks).reset_index(drop=True)


def prepare_g2_transactions(dataframe: pd.DataFrame | None) -> pd.DataFrame:
    frame = promote_g2_statement_header(
        dataframe if dataframe is not None else pd.DataFrame()
    )
    frame = remove_export_index_columns(frame)
    if frame.empty:
        return frame

    rename_map: dict[object, str] = {}
    for column in frame.columns:
        key = _normalize_business_column_key(column)
        standard = {
            "receipt_no": "receipt_no",
            "receipt_no_": "receipt_no",
            "completion_time": "completion_time",
            "initiation_time": "initiation_time",
            "details": "details",
            "opposite_party": "opposite_party",
            "transaction_status": "transaction_status",
            "currency": "currency_code",
            "transaction_amount": "transaction_amount",
            "paid_in": "paid_in",
            "withdrawn": "withdrawn",
            "balance": "balance",
            "operation": "operation",
            "reason_type": "reason_type",
            "linked_transaction_id": "linked_transaction_id",
        }.get(key)
        rename_map[column] = standard or key
    frame = frame.rename(columns=rename_map).copy()

    for column in ["receipt_no", "details", "opposite_party", "transaction_status", "currency_code", "operation", "reason_type"]:
        if column in frame.columns:
            frame[column] = clean_text(frame[column])
    if "receipt_no" in frame.columns:
        frame["receipt_no"] = clean_identifier(frame["receipt_no"])
    if "currency_code" in frame.columns:
        frame["currency_code"] = clean_text(frame["currency_code"]).str.upper()
    for column in ["completion_time", "initiation_time"]:
        if column in frame.columns:
            frame[column] = pd.to_datetime(frame[column], errors="coerce", format="mixed", dayfirst=True)
    if "opposite_party" in frame.columns:
        frame["phone"] = _extract_phone_from_opposite_party(frame["opposite_party"])
        frame["phone_prefixe"] = normalize_phone(frame["phone"])
        frame["Nom_client"] = _extract_customer_name_from_opposite_party(frame["opposite_party"])
    else:
        frame["phone"] = pd.NA
        frame["phone_prefixe"] = pd.NA
        frame["Nom_client"] = pd.NA
    for column in ["transaction_amount", "paid_in", "withdrawn", "balance"]:
        if column in frame.columns:
            frame[f"{column}_numeric"] = _parse_money_series(frame[column])

    original_amount = frame.get("transaction_amount_numeric", pd.Series(np.nan, index=frame.index, dtype="float64"))
    paid_in = frame.get("paid_in_numeric", pd.Series(np.nan, index=frame.index, dtype="float64"))
    withdrawn = frame.get("withdrawn_numeric", pd.Series(np.nan, index=frame.index, dtype="float64"))
    paid_in_available = paid_in.notna() & (paid_in.ne(0) | withdrawn.isna())
    withdrawn_available = withdrawn.notna() & ~paid_in_available
    derived_amount = paid_in.where(paid_in_available, withdrawn)
    frame["transaction_amount_numeric"] = original_amount.where(original_amount.notna(), derived_amount)
    frame["transaction_amount_source"] = np.select(
        [original_amount.notna(), original_amount.isna() & paid_in_available, original_amount.isna() & withdrawn_available],
        ["Transaction Amount", "Paid In", "Withdrawn"],
        default="Montant absent",
    )
    paid_in_non_zero = paid_in.notna() & paid_in.ne(0)
    withdrawn_non_zero = withdrawn.notna() & withdrawn.ne(0)
    both_directions = paid_in_non_zero & withdrawn_non_zero
    fallback_amount = frame["transaction_amount_numeric"]
    frame["sens_flux"] = np.select(
        [
            both_directions,
            withdrawn_non_zero,
            paid_in_non_zero,
            fallback_amount.lt(0).fillna(False),
            fallback_amount.gt(0).fillna(False),
        ],
        ["A verifier", "Sortie", "Entree", "Sortie", "Entree"],
        default="Indetermine",
    )
    absolute_amount = fallback_amount.abs()
    frame["montant_entree"] = absolute_amount.where(frame["sens_flux"].eq("Entree"), 0.0)
    frame["montant_sortie"] = absolute_amount.where(frame["sens_flux"].eq("Sortie"), 0.0)
    frame["type_operation_g2"] = frame.apply(classify_g2_business_operation, axis=1)
    if "transaction_amount" not in frame.columns:
        frame["transaction_amount"] = derived_amount
    else:
        frame["transaction_amount"] = frame["transaction_amount"].where(original_amount.notna(), derived_amount)
    frame["source_g2"] = "G2"
    frame = frame.drop(columns=["ordre_fichier_import"], errors="ignore")
    sort_columns = [column for column in ["completion_time", "receipt_no"] if column in frame.columns]
    if sort_columns:
        frame = frame.sort_values(sort_columns, na_position="last").reset_index(drop=True)
    return frame


def filter_g2_transactions_by_completion_time(
    dataframe: pd.DataFrame | None,
    start_date: Any | None = None,
    end_date: Any | None = None,
    start_time: Any | None = None,
    end_time: Any | None = None,
) -> pd.DataFrame:
    """Filtre G2 sur Completion Time avec des bornes de date et d'heure inclusives."""
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        return pd.DataFrame() if dataframe is None else dataframe.copy()
    if "completion_time" not in dataframe.columns:
        return dataframe.copy()

    frame = dataframe.copy()
    completion_time = pd.to_datetime(frame["completion_time"], errors="coerce")
    mask = completion_time.notna()
    if start_date is not None:
        start_bound = pd.Timestamp(start_date).normalize()
        if start_time is not None:
            start_bound = _timestamp_plus(
                start_bound,
                hours=start_time.hour,
                minutes=start_time.minute,
                seconds=start_time.second,
                microseconds=start_time.microsecond,
            )
        mask &= completion_time.ge(start_bound)
    if end_date is not None:
        end_bound = pd.Timestamp(end_date).normalize()
        if end_time is None:
            mask &= completion_time.lt(_timestamp_plus(end_bound, days=1))
        else:
            end_bound = _timestamp_plus(
                end_bound,
                hours=end_time.hour,
                minutes=end_time.minute,
                seconds=end_time.second,
                microseconds=end_time.microsecond,
            )
            mask &= completion_time.le(end_bound)
    return frame.loc[mask].reset_index(drop=True)


def filter_g2_transactions_by_direction(
    dataframe: pd.DataFrame | None,
    directions: str | Iterable[str] | None = None,
) -> pd.DataFrame:
    """Filtre G2 sur Entree/Sortie; None conserve tous les sens et les anomalies."""
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        return pd.DataFrame() if dataframe is None else dataframe.copy()
    if directions is None or "sens_flux" not in dataframe.columns:
        return dataframe.copy()

    requested = [directions] if isinstance(directions, str) else list(directions)
    canonical = {
        "entree": "Entree",
        "entrees": "Entree",
        "sortie": "Sortie",
        "sorties": "Sortie",
    }
    selected = {canonical.get(normalize_label(value), str(value).strip()) for value in requested}
    selected.discard("")
    if not selected:
        return dataframe.copy()
    return dataframe.loc[dataframe["sens_flux"].astype("string").isin(selected)].reset_index(drop=True)


def enrich_turbo_with_g2_customer_names(
    dataframe: pd.DataFrame | None,
    g2_transactions: pd.DataFrame | None,
    *,
    phone_column: str,
    reference_column: str | None = None,
) -> pd.DataFrame:
    """Ajoute le nom G2 a une source Turbo, par telephone puis par reference."""
    result = dataframe.copy() if isinstance(dataframe, pd.DataFrame) else pd.DataFrame()
    if result.empty:
        return result

    existing_name = result.get("Nom_client", pd.Series(pd.NA, index=result.index)).copy()
    has_existing_name = existing_name.astype("string").fillna("").str.strip().ne("")
    result["Nom_client"] = existing_name
    result["mode_rapprochement_nom_client"] = np.where(has_existing_name, "Nom existant Turbo", "Fichier G2 absent")
    if not isinstance(g2_transactions, pd.DataFrame) or g2_transactions.empty or "Nom_client" not in g2_transactions.columns:
        return result

    g2 = g2_transactions.copy()
    g2["__phone_key"] = normalize_phone(
        g2.get("phone_prefixe", g2.get("phone", pd.Series("", index=g2.index)))
    )
    g2["__receipt_key"] = clean_identifier(g2.get("receipt_no", pd.Series("", index=g2.index)))
    g2["Nom_client"] = clean_text(g2["Nom_client"]).replace("", pd.NA)

    phone_rows = g2.dropna(subset=["__phone_key", "Nom_client"])
    if not phone_rows.empty:
        phone_lookup = (
            phone_rows.groupby("__phone_key", as_index=False, dropna=False)
            .agg(__nom_par_telephone=("Nom_client", concat_unique))
        )
    else:
        phone_lookup = pd.DataFrame(columns=["__phone_key", "__nom_par_telephone"])

    reference_rows = g2.loc[g2["__receipt_key"].ne("")].dropna(subset=["Nom_client"])
    if not reference_rows.empty:
        reference_lookup = (
            reference_rows.groupby("__receipt_key", as_index=False, dropna=False)
            .agg(__nom_par_reference=("Nom_client", concat_unique))
        )
    else:
        reference_lookup = pd.DataFrame(columns=["__receipt_key", "__nom_par_reference"])

    result["__row_order"] = np.arange(len(result))
    result["__nom_existant"] = existing_name
    result["__phone_key"] = normalize_phone(result.get(phone_column, pd.Series("", index=result.index)))
    result["__receipt_key"] = clean_identifier(
        result.get(reference_column, pd.Series("", index=result.index))
        if reference_column
        else pd.Series("", index=result.index)
    )
    result = result.merge(phone_lookup, on="__phone_key", how="left")
    result = result.merge(reference_lookup, on="__receipt_key", how="left")

    has_phone_name = result["__nom_par_telephone"].astype("string").fillna("").str.strip().ne("")
    has_reference_name = result["__nom_par_reference"].astype("string").fillna("").str.strip().ne("")
    matched_name = result["__nom_par_telephone"].where(has_phone_name, result["__nom_par_reference"])
    has_existing_name = result["__nom_existant"].astype("string").fillna("").str.strip().ne("")
    result["Nom_client"] = matched_name.where(has_phone_name | has_reference_name, result["__nom_existant"])
    result["mode_rapprochement_nom_client"] = np.select(
        [has_phone_name, ~has_phone_name & has_reference_name, ~has_phone_name & ~has_reference_name & has_existing_name],
        [f"Telephone G2 = {phone_column} Turbo", f"Receipt No G2 = {reference_column} Turbo", "Nom existant Turbo"],
        default="Nom G2 non rapproche",
    )
    return (
        result.sort_values("__row_order")
        .drop(columns=["__row_order", "__nom_existant", "__phone_key", "__receipt_key", "__nom_par_telephone", "__nom_par_reference"])
        .reset_index(drop=True)
    )


def enrich_transactions_with_g2_customer_names(
    transactions: pd.DataFrame | None,
    g2_transactions: pd.DataFrame | None,
) -> pd.DataFrame:
    return enrich_turbo_with_g2_customer_names(
        transactions,
        g2_transactions,
        phone_column="msisdn1",
        reference_column="ref_no",
    )


def build_load_report(files: dict[str, pd.DataFrame], missing: dict[str, list[str]]) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for label, frame in files.items():
        source_names: list[str] = []
        if isinstance(frame, pd.DataFrame) and not frame.empty:
            source_columns = [
                column
                for column in frame.columns
                if str(column).startswith(("fichier_source_", "fichiers_sources_"))
            ]
            for column in source_columns:
                for value in clean_text(frame[column]):
                    for part in str(value).split("|"):
                        name = part.strip()
                        if name and name not in source_names:
                            source_names.append(name)
        rows.append(
            {
                "fichier": label,
                "nombre_fichiers": len(source_names),
                "fichiers_sources": " | ".join(source_names),
                "lignes": int(len(frame)) if isinstance(frame, pd.DataFrame) else 0,
                "colonnes": int(frame.shape[1]) if isinstance(frame, pd.DataFrame) else 0,
                "colonnes_manquantes": ", ".join(missing.get(label, [])),
                "statut": "OK" if not missing.get(label) else "Colonnes manquantes",
            }
        )
    return pd.DataFrame(rows)


def classify_mpesa_operation(descriptions: object, account_types: object = "", movement_net: float = 0.0) -> str:
    text = normalize_label(f"{descriptions} {account_types}")
    if "penalite" in text and "remboursement" in text:
        return "Remboursement avec penalite"
    if "montant pret" in text or "decaissement" in text or "loan disbur" in text:
        return "Decaissement de credit"
    if "remboursement" in text or "repayment" in text:
        return "Remboursement de credit"
    if "retrait vers m-pesa" in text or "retrait vers mpesa" in text:
        return "Entree M-PESA_Turbo depuis epargne"
    if "m-pesa depot" in text or "mpesa depot" in text:
        return "Sortie M-PESA_Turbo vers epargne"
    if "depot bloque" in text or "depot bloque" in text or "fixed savings" in text or "m-pesa compte" in text:
        return "Sortie M-PESA_Turbo vers DAT" if movement_net < 0 else "Entree M-PESA_Turbo depuis DAT"
    if movement_net > 0:
        return "Autre entree M-PESA_Turbo"
    if movement_net < 0:
        return "Autre sortie M-PESA_Turbo"
    return "Autre mouvement M-PESA_Turbo"


def build_account_events(transactions_client: pd.DataFrame, account_type: str) -> pd.DataFrame:
    if transactions_client.empty or "account_type" not in transactions_client.columns:
        return pd.DataFrame(columns=["currency_code", "created_at", "variation", "references", "descriptions"])
    lines = transactions_client.loc[transactions_client["account_type"].eq(account_type)].copy()
    if lines.empty:
        return pd.DataFrame(columns=["currency_code", "created_at", "variation", "references", "descriptions"])
    lines["variation"] = pd.to_numeric(lines["bal_after"], errors="coerce").fillna(0) - pd.to_numeric(
        lines["bal_before"], errors="coerce"
    ).fillna(0)
    event_aggregations: dict[str, tuple[str, object]] = {
        "variation": ("variation", "sum"),
        "references": ("reference_id", concat_unique),
        "descriptions": ("description", concat_unique),
    }
    if "Nom_client" in lines.columns:
        event_aggregations["Nom_client"] = ("Nom_client", concat_unique)
    if "mode_rapprochement_nom_client" in lines.columns:
        event_aggregations["mode_rapprochement_nom_client"] = ("mode_rapprochement_nom_client", concat_unique)
    return (
        lines.groupby(["currency_code", "created_at"], as_index=False, dropna=False)
        .agg(**event_aggregations)
        .sort_values(["currency_code", "created_at"])
    )


def add_reconstructed_balance(
    operations: pd.DataFrame,
    events: pd.DataFrame,
    final_balances: dict[str, float],
    column_name: str,
) -> pd.DataFrame:
    result = operations.copy()
    result[column_name] = np.nan
    result[f"{column_name}_ouverture_estimee"] = np.nan
    if result.empty:
        return result

    blocks: list[pd.DataFrame] = []
    currencies = set(result["currency_code"].dropna().astype(str))
    currencies |= set(str(currency) for currency in final_balances)
    if not events.empty and "currency_code" in events.columns:
        currencies |= set(events["currency_code"].dropna().astype(str))

    for currency in sorted(currencies):
        ops = result.loc[result["currency_code"].astype(str).eq(currency)].sort_values("created_at").copy()
        if ops.empty:
            continue
        ev = events.loc[events["currency_code"].astype(str).eq(currency)].sort_values("created_at").copy() if not events.empty else pd.DataFrame()
        final_balance = float(final_balances.get(currency, 0.0))
        total_variation = float(ev["variation"].sum()) if not ev.empty else 0.0
        opening_estimate = final_balance - total_variation
        ops[f"{column_name}_ouverture_estimee"] = opening_estimate
        if ev.empty:
            ops[column_name] = opening_estimate
        else:
            ev["variation_cumulee"] = pd.to_numeric(ev["variation"], errors="coerce").fillna(0).cumsum()
            ops = pd.merge_asof(
                ops.sort_values("created_at"),
                ev[["created_at", "variation_cumulee"]],
                on="created_at",
                direction="backward",
                allow_exact_matches=True,
            )
            ops[column_name] = opening_estimate + ops["variation_cumulee"].fillna(0)
            ops = ops.drop(columns=["variation_cumulee"])
        blocks.append(ops)

    if not blocks:
        return result
    return concat_frames_stable(blocks).sort_values(["created_at", "operation_reference"]).reset_index(drop=True)


def build_savings_final(current_savings: pd.DataFrame, customer_id: str) -> dict[str, float]:
    required_columns = {"customer_id", "currency_code", "balance"}
    if current_savings.empty or not required_columns.issubset(current_savings.columns):
        return {}
    frame = current_savings.loc[current_savings["customer_id"].eq(customer_id)].copy()
    if frame.empty:
        return {}
    frame["balance"] = numeric_column(frame, "balance")
    return frame.groupby("currency_code", dropna=False)["balance"].sum().to_dict()


def build_dat_final(fixed_savings: pd.DataFrame, customer_id: str) -> dict[str, float]:
    required_columns = {"customer_id", "currency_code", "balance"}
    if fixed_savings.empty or not required_columns.issubset(fixed_savings.columns):
        return {}
    frame = fixed_savings.loc[fixed_savings["customer_id"].eq(customer_id)].copy()
    if frame.empty:
        return {}
    frame["balance"] = numeric_column(frame, "balance")
    return frame.groupby("currency_code", dropna=False)["balance"].sum().to_dict()


def filter_customer_frame(frame: pd.DataFrame, customer_id: str) -> pd.DataFrame:
    if frame is None or not isinstance(frame, pd.DataFrame) or frame.empty or "customer_id" not in frame.columns:
        return pd.DataFrame()
    return frame.loc[frame["customer_id"].astype("string").eq(str(customer_id))].copy()


def _deduplicate_g2_transactions(g2: pd.DataFrame) -> pd.DataFrame:
    """Conserve une ligne canonique par Receipt No. et trace les doublons sources."""
    if g2.empty:
        return g2.copy()
    frame = g2.copy().reset_index(drop=True)
    frame["receipt_no"] = clean_identifier(frame.get("receipt_no", pd.Series("", index=frame.index)))
    frame["__row_order"] = np.arange(len(frame))
    frame["__receipt_key"] = frame["receipt_no"].where(
        frame["receipt_no"].ne(""),
        "__sans_receipt_" + frame["__row_order"].astype(str),
    )
    status = frame.get("transaction_status", pd.Series("", index=frame.index)).apply(normalize_label)
    frame["__completed_priority"] = status.isin({"completed", "complete", "successful", "success"}).astype(int)
    frame["__completion_sort"] = pd.to_datetime(
        frame.get("completion_time", pd.Series(pd.NaT, index=frame.index)), errors="coerce"
    )

    grouped = frame.groupby("__receipt_key", dropna=False, sort=False)
    frame["nombre_lignes_g2_reference"] = grouped["__receipt_key"].transform("size").astype(int)
    frame["devises_g2_reference"] = grouped["currency_code"].transform(concat_unique) if "currency_code" in frame.columns else ""
    frame["statuts_g2_reference"] = grouped["transaction_status"].transform(concat_unique) if "transaction_status" in frame.columns else ""
    amount_text = numeric_column(frame, "transaction_amount_numeric").map(
        lambda value: "" if pd.isna(value) else f"{float(value):.2f}"
    )
    frame["__amount_text"] = amount_text
    frame["montants_g2_reference"] = grouped["__amount_text"].transform(concat_unique)
    frame["doublon_receipt_no"] = frame["receipt_no"].ne("") & frame["nombre_lignes_g2_reference"].gt(1)

    canonical = (
        frame.sort_values(
            ["__receipt_key", "__completed_priority", "__completion_sort", "__row_order"],
            ascending=[True, False, False, True],
            na_position="last",
        )
        .drop_duplicates("__receipt_key", keep="first")
        .sort_values("__row_order")
        .drop(columns=["__row_order", "__receipt_key", "__completed_priority", "__completion_sort", "__amount_text"])
        .reset_index(drop=True)
    )
    return canonical


def _portal_operation_flags(group: pd.DataFrame) -> tuple[bool, bool, bool]:
    account_types = " ".join(group.get("account_type", pd.Series(dtype="string")).fillna("").astype(str))
    descriptions = " ".join(group.get("description", pd.Series(dtype="string")).fillna("").astype(str))
    text = normalize_label(f"{account_types} {descriptions}")
    has_loan = any(
        token in text
        for token in [
            "loan account",
            "loan portfolio",
            "principle",
            "remboursement",
            "repayment",
        ]
    )
    has_fixed = "fixed savings" in text or "depot bloque" in text
    has_normal = "normal savings" in text or "epargne depot" in text
    return has_loan, has_fixed, has_normal


def _build_portal_reference_controls(transactions: pd.DataFrame) -> pd.DataFrame:
    """Agrège les écritures Portal au grain ref_no sans additionner les miroirs comptables."""
    if transactions.empty or "ref_no" not in transactions.columns:
        return pd.DataFrame()
    tx = transactions.copy()
    tx["ref_no"] = clean_identifier(tx["ref_no"])
    tx = tx.loc[tx["ref_no"].ne("")].copy()
    if tx.empty:
        return pd.DataFrame()
    tx["currency_code"] = clean_text(tx.get("currency_code", pd.Series("", index=tx.index))).str.upper()
    tx["phone_portal_normalise"] = normalize_phone(tx.get("msisdn1", pd.Series("", index=tx.index)))
    tx["created_at"] = pd.to_datetime(tx.get("created_at", pd.Series(pd.NaT, index=tx.index)), errors="coerce")
    line_amounts = pd.concat(
        [
            numeric_column(tx, "dr").abs(),
            numeric_column(tx, "cr").abs(),
            (numeric_column(tx, "bal_after") - numeric_column(tx, "bal_before")).abs(),
        ],
        axis=1,
    )
    tx["montant_ligne_portal"] = line_amounts.max(axis=1)
    tx["est_compte_mpesa"] = clean_text(
        tx.get("account_type", pd.Series("", index=tx.index))
    ).apply(normalize_label).eq("mpesa account")

    rows: list[dict[str, object]] = []
    for ref_no, group in tx.groupby("ref_no", dropna=False, sort=False):
        has_loan, has_fixed, has_normal = _portal_operation_flags(group)
        source_files = concat_unique(
            group.get(
                "fichiers_sources_transactions_turbo",
                group.get(
                    "fichier_source_transactions_turbo",
                    pd.Series(dtype="string"),
                ),
            )
        )
        mpesa_amounts = group.loc[group["est_compte_mpesa"], "montant_ligne_portal"]
        control_amounts = mpesa_amounts.loc[mpesa_amounts.gt(0)]
        if control_amounts.empty:
            control_amounts = group.loc[group["montant_ligne_portal"].gt(0), "montant_ligne_portal"]
        portal_amount = float(control_amounts.max()) if not control_amounts.empty else np.nan
        target_account = (
            "LOAN ACCOUNT / PRINCIPLE / LOAN PORTFOLIO"
            if has_loan
            else "FIXED SAVINGS"
            if has_fixed
            else "NORMAL SAVINGS"
            if has_normal
            else "MPESA ACCOUNT"
        )
        rows.append(
            {
                "ref_no_portal": ref_no,
                "nombre_ecritures_portal": int(len(group)),
                "customer_id_portal": first_non_empty(group.get("customer_id", pd.Series(dtype="string"))),
                "customer_ids_portal": concat_unique(group.get("customer_id", pd.Series(dtype="string"))),
                "telephones_portal": concat_unique(group["phone_portal_normalise"]),
                "devises_portal": concat_unique(group["currency_code"]),
                "account_types_portal": concat_unique(group.get("account_type", pd.Series(dtype="string"))),
                "descriptions_portal": concat_unique(group.get("description", pd.Series(dtype="string"))),
                "references_internes_portal": concat_unique(group.get("reference_id", pd.Series(dtype="string"))),
                "date_portal_min": group["created_at"].min(),
                "date_portal_max": group["created_at"].max(),
                "montant_portal_controle": portal_amount,
                "portal_has_loan": bool(has_loan),
                "portal_has_fixed": bool(has_fixed),
                "portal_has_normal": bool(has_normal),
                "account_type_cible": target_account,
                "fichiers_sources_turbo": source_files,
            }
        )
    return pd.DataFrame(rows)


def _build_turbo_output_controls(transactions: pd.DataFrame) -> pd.DataFrame:
    """Agrège les retraits Turbo sans ``ref_no`` par compte et horodatage.

    Les sorties G2 ``BisouBisouB2C`` observées dans les exports réels ne
    transmettent pas leur ``Receipt No`` dans Turbo. Elles apparaissent comme
    deux écritures miroir ``Retrait Vers M-Pesa`` partageant un ``reference_id``
    de type SA. Comme ce compte peut être réutilisé, ``created_at`` distingue
    les opérations. Cette table fournit une clé secondaire réservée aux sorties.
    """
    if transactions.empty or "reference_id" not in transactions.columns:
        return pd.DataFrame()

    tx = transactions.copy()
    tx["reference_id"] = clean_identifier(tx["reference_id"])
    tx["description_normalisee"] = clean_text(
        tx.get("description", pd.Series("", index=tx.index))
    ).apply(normalize_label)
    tx = tx.loc[
        tx["reference_id"].ne("")
        & tx["description_normalisee"].eq("retrait vers m-pesa")
    ].copy()
    if tx.empty:
        return pd.DataFrame()

    tx["currency_code"] = clean_text(
        tx.get("currency_code", pd.Series("", index=tx.index))
    ).str.upper()
    tx["phone_portal_normalise"] = normalize_phone(
        tx.get("msisdn1", pd.Series("", index=tx.index))
    )
    tx["created_at"] = pd.to_datetime(
        tx.get("created_at", pd.Series(pd.NaT, index=tx.index)), errors="coerce"
    )
    line_amounts = pd.concat(
        [
            numeric_column(tx, "dr").abs(),
            numeric_column(tx, "cr").abs(),
            (numeric_column(tx, "bal_after") - numeric_column(tx, "bal_before")).abs(),
        ],
        axis=1,
    )
    tx["montant_ligne_portal"] = line_amounts.max(axis=1)

    rows: list[dict[str, object]] = []
    group_columns = [
        "reference_id",
        "created_at",
        "phone_portal_normalise",
        "currency_code",
    ]
    for group_key, group in tx.groupby(group_columns, dropna=False, sort=False):
        reference_id, created_at, _, _ = group_key
        has_loan, has_fixed, has_normal = _portal_operation_flags(group)
        source_files = concat_unique(
            group.get(
                "fichiers_sources_transactions_turbo",
                group.get(
                    "fichier_source_transactions_turbo",
                    pd.Series(dtype="string"),
                ),
            )
        )
        positive_amounts = group.loc[
            group["montant_ligne_portal"].gt(0), "montant_ligne_portal"
        ]
        portal_amount = (
            float(positive_amounts.max()) if not positive_amounts.empty else np.nan
        )
        target_account = (
            "LOAN ACCOUNT / PRINCIPLE / LOAN PORTFOLIO"
            if has_loan
            else "FIXED SAVINGS"
            if has_fixed
            else "NORMAL SAVINGS"
            if has_normal
            else "MPESA ACCOUNT"
        )
        rows.append(
            {
                "reference_sortie_turbo": reference_id,
                "cle_sortie_turbo": (
                    f"{reference_id} @ {created_at:%Y-%m-%d %H:%M:%S}"
                    if pd.notna(created_at)
                    else str(reference_id)
                ),
                "nombre_ecritures_portal": int(len(group)),
                "customer_id_portal": first_non_empty(
                    group.get("customer_id", pd.Series(dtype="string"))
                ),
                "customer_ids_portal": concat_unique(
                    group.get("customer_id", pd.Series(dtype="string"))
                ),
                "telephones_portal": concat_unique(group["phone_portal_normalise"]),
                "devises_portal": concat_unique(group["currency_code"]),
                "account_types_portal": concat_unique(
                    group.get("account_type", pd.Series(dtype="string"))
                ),
                "descriptions_portal": concat_unique(
                    group.get("description", pd.Series(dtype="string"))
                ),
                "references_internes_portal": reference_id,
                "date_portal_min": group["created_at"].min(),
                "date_portal_max": group["created_at"].max(),
                "montant_portal_controle": portal_amount,
                "portal_has_loan": bool(has_loan),
                "portal_has_fixed": bool(has_fixed),
                "portal_has_normal": bool(has_normal),
                "account_type_cible": target_account,
                "operation_turbo_confirmee": "Retrait epargne vers M-PESA",
                "fichiers_sources_turbo": source_files,
            }
        )
    return pd.DataFrame(rows)


def build_turbo_only_g2_transactions(transactions: pd.DataFrame | None) -> pd.DataFrame:
    """Construit le périmètre G2/DAT depuis Turbo lorsque le relevé G2 manque.

    Le dataset synthétique reste limité aux opérations M-PESA démontrables :
    entrées portant un ``ref_no`` et classables en épargne/DAT/remboursement,
    puis sorties ``Retrait Vers M-Pesa``. Il ne simule ni statut, ni solde, ni
    horodatage G2 et conserve ``source_mode_analyse = Turbo seul``.
    """
    if not isinstance(transactions, pd.DataFrame) or transactions.empty:
        return pd.DataFrame()

    direct_controls = _build_portal_reference_controls(transactions)
    output_controls = _build_turbo_output_controls(transactions)
    rows: list[dict[str, object]] = []

    def first_pipe_value(value: object) -> str:
        if _is_empty_text(value):
            return ""
        return next(
            (part.strip() for part in str(value).split("|") if part.strip()),
            "",
        )

    if not direct_controls.empty:
        for _, control in direct_controls.iterrows():
            has_loan = bool(control.get("portal_has_loan", False))
            has_fixed = bool(control.get("portal_has_fixed", False))
            has_normal = bool(control.get("portal_has_normal", False))
            if not (has_loan or has_fixed or has_normal):
                continue
            reason_type = (
                "BisouBisouC2BRepayment" if has_loan else "BisouBisouC2B"
            )
            details = (
                "Remboursement prets Turbo"
                if has_loan
                else "Depot Bloque Turbo"
                if has_fixed
                else "Epargne depot Turbo"
            )
            source_files = control.get("fichiers_sources_turbo")
            if _is_empty_text(source_files):
                source_files = "Transactions M-PESA_Turbo"
            rows.append(
                {
                    "Receipt No.": control.get("ref_no_portal"),
                    "Initiation Time": control.get("date_portal_min"),
                    "Completion Time": control.get("date_portal_max"),
                    "Details": details,
                    "Reason Type": reason_type,
                    "Transaction Status": "Comptabilisee Turbo",
                    "Currency": first_pipe_value(control.get("devises_portal")),
                    "Paid In": control.get("montant_portal_controle"),
                    "Withdrawn": 0.0,
                    "Opposite Party": first_pipe_value(
                        control.get("telephones_portal")
                    ),
                    "Linked Transaction ID": control.get("ref_no_portal"),
                    "fichier_source_g2": source_files,
                    "fichier_source_analyse": source_files,
                    "source_mode_analyse": "Turbo seul",
                }
            )

    if not output_controls.empty:
        for _, control in output_controls.iterrows():
            reference_id = str(control.get("reference_sortie_turbo", "")).strip()
            created_at = pd.to_datetime(control.get("date_portal_min"), errors="coerce")
            timestamp_key = (
                f"{created_at:%Y%m%d%H%M%S}" if pd.notna(created_at) else "sansdate"
            )
            source_files = control.get("fichiers_sources_turbo")
            if _is_empty_text(source_files):
                source_files = "Transactions M-PESA_Turbo"
            rows.append(
                {
                    "Receipt No.": f"TURBO-{reference_id}-{timestamp_key}",
                    "Initiation Time": created_at,
                    "Completion Time": pd.to_datetime(
                        control.get("date_portal_max"), errors="coerce"
                    ),
                    "Details": "Retrait Vers M-Pesa",
                    "Reason Type": "BisouBisouB2C",
                    "Transaction Status": "Comptabilisee Turbo",
                    "Currency": first_pipe_value(control.get("devises_portal")),
                    "Paid In": 0.0,
                    "Withdrawn": -abs(
                        float(control.get("montant_portal_controle", 0) or 0)
                    ),
                    "Opposite Party": first_pipe_value(
                        control.get("telephones_portal")
                    ),
                    "Linked Transaction ID": control.get("cle_sortie_turbo"),
                    "fichier_source_g2": source_files,
                    "fichier_source_analyse": source_files,
                    "source_mode_analyse": "Turbo seul",
                }
            )

    if not rows:
        return pd.DataFrame()
    proxy = prepare_g2_transactions(pd.DataFrame(rows))
    proxy["source_mode_analyse"] = "Turbo seul"
    proxy["source_g2"] = "Turbo"
    proxy["fichier_source_analyse"] = proxy.get(
        "fichier_source_g2",
        pd.Series("Transactions M-PESA_Turbo", index=proxy.index),
    )
    return proxy


def _contains_pipe_value(values: object, expected: object) -> bool:
    if _is_empty_text(values) or _is_empty_text(expected):
        return False
    expected_text = str(expected).strip()
    return any(part.strip() == expected_text for part in str(values).split("|") if part.strip())


def _scope_portal_transactions_for_g2(
    g2: pd.DataFrame,
    transactions: pd.DataFrame,
) -> pd.DataFrame:
    """Limite le rapprochement aux references et sorties candidates de G2.

    Les controles directs n'ont besoin que des ``ref_no`` figurant dans G2.
    Pour une sortie B2C sans ``ref_no``, conserver les lignes Turbo portant le
    libelle de retrait et compatibles avec les telephones, devises et la plage
    horaire G2. Cette reduction est sans incidence sur la tolerance appliquee
    ensuite au candidat retenu.
    """
    if transactions.empty or g2.empty:
        return transactions.copy()

    tx = transactions.copy()
    ref_no = clean_identifier(
        tx.get("ref_no", pd.Series("", index=tx.index))
    )
    receipt_numbers = set(
        clean_identifier(g2.get("receipt_no", pd.Series(dtype="string")))
        .loc[lambda values: values.ne("")]
        .astype(str)
    )
    direct_mask = ref_no.isin(receipt_numbers) if receipt_numbers else pd.Series(False, index=tx.index)

    descriptions = clean_text(
        tx.get("description", pd.Series("", index=tx.index))
    ).str.strip().str.casefold()
    output_candidate_mask = descriptions.eq("retrait vers m-pesa")
    g2_directions = clean_text(
        g2.get("sens_flux", pd.Series("", index=g2.index))
    )
    g2_outputs = g2.loc[g2_directions.eq("Sortie")].copy()
    if g2_outputs.empty:
        return tx.loc[direct_mask].copy()

    output_candidates = tx.loc[output_candidate_mask].copy()
    output_mask = pd.Series(True, index=output_candidates.index)

    output_phones = set(
        normalize_phone(
            g2_outputs.get(
                "phone_prefixe",
                g2_outputs.get("phone", pd.Series("", index=g2_outputs.index)),
            )
        )
        .dropna()
        .astype(str)
    )
    if output_phones:
        output_mask &= normalize_phone(
            output_candidates.get(
                "msisdn1", pd.Series("", index=output_candidates.index)
            )
        ).astype("string").isin(output_phones)

    output_currencies = set(
        clean_text(
            g2_outputs.get("currency_code", pd.Series("", index=g2_outputs.index))
        )
        .str.upper()
        .loc[lambda values: values.ne("")]
        .astype(str)
    )
    if output_currencies:
        output_mask &= clean_text(
            output_candidates.get(
                "currency_code", pd.Series("", index=output_candidates.index)
            )
        ).str.upper().isin(output_currencies)

    initiation = pd.to_datetime(
        g2_outputs.get("initiation_time", pd.Series(pd.NaT, index=g2_outputs.index)),
        errors="coerce",
    )
    completion = pd.to_datetime(
        g2_outputs.get("completion_time", pd.Series(pd.NaT, index=g2_outputs.index)),
        errors="coerce",
    )
    output_dates = initiation.combine_first(completion).dropna()
    if not output_dates.empty:
        tolerance = pd.Timedelta(minutes=G2_TURBO_OUTPUT_MATCH_TOLERANCE_MINUTES)
        transaction_dates = pd.to_datetime(
            output_candidates.get(
                "created_at", pd.Series(pd.NaT, index=output_candidates.index)
            ),
            errors="coerce",
        )
        output_mask &= transaction_dates.between(
            output_dates.min() - tolerance,
            output_dates.max() + tolerance,
            inclusive="both",
        )

    selected_output = pd.Series(False, index=tx.index)
    selected_output.loc[output_mask.index] = output_mask
    return tx.loc[direct_mask | selected_output].copy()


def _enrich_g2_with_portal_controls(g2: pd.DataFrame, transactions: pd.DataFrame) -> pd.DataFrame:
    output = _deduplicate_g2_transactions(g2)
    if output.empty:
        return output
    scoped_transactions = _scope_portal_transactions_for_g2(output, transactions)
    portal = _build_portal_reference_controls(scoped_transactions)
    if not portal.empty:
        output = output.merge(portal, left_on="receipt_no", right_on="ref_no_portal", how="left")
    else:
        for column, default in {
            "ref_no_portal": pd.NA,
            "nombre_ecritures_portal": 0,
            "customer_id_portal": pd.NA,
            "customer_ids_portal": "",
            "telephones_portal": "",
            "devises_portal": "",
            "account_types_portal": "",
            "descriptions_portal": "",
            "references_internes_portal": "",
            "date_portal_min": pd.NaT,
            "date_portal_max": pd.NaT,
            "montant_portal_controle": np.nan,
            "portal_has_loan": False,
            "portal_has_fixed": False,
            "portal_has_normal": False,
            "account_type_cible": "",
        }.items():
            output[column] = default

    output["reference_sortie_turbo"] = pd.NA
    output["cle_sortie_turbo"] = pd.NA
    output["nombre_candidats_sortie_turbo"] = 0
    output["operation_turbo_confirmee"] = ""
    direct_reference = output["ref_no_portal"].astype("string").fillna("").ne("")
    output["methode_rapprochement_turbo"] = np.where(
        direct_reference,
        "Receipt No = ref_no",
        "Non rapproche",
    )

    # Les sorties B2C ne portent pas le Receipt No dans ref_no côté Turbo.
    # Utiliser une clé secondaire stricte, uniquement quand la clé principale
    # est absente : téléphone + devise + montant + proximité horaire, sur une
    # opération Turbo explicitement libellée Retrait Vers M-Pesa.
    output_candidates = _build_turbo_output_controls(scoped_transactions)
    if not output_candidates.empty:
        g2_direction = clean_text(
            output.get("sens_flux", pd.Series("", index=output.index))
        )
        g2_reason = clean_text(
            output.get("reason_type", pd.Series("", index=output.index))
        ).apply(normalize_label)
        g2_details = clean_text(
            output.get("details", pd.Series("", index=output.index))
        ).apply(normalize_label)
        g2_phone_for_match = normalize_phone(
            output.get(
                "phone_prefixe",
                output.get("phone", pd.Series("", index=output.index)),
            )
        )
        g2_currency_for_match = clean_text(
            output.get("currency_code", pd.Series("", index=output.index))
        ).str.upper()
        g2_amount_for_match = numeric_column(
            output, "transaction_amount_numeric"
        ).abs()
        initiation_for_match = pd.to_datetime(
            output.get("initiation_time", pd.Series(pd.NaT, index=output.index)),
            errors="coerce",
        )
        completion_for_match = pd.to_datetime(
            output.get("completion_time", pd.Series(pd.NaT, index=output.index)),
            errors="coerce",
        )
        g2_date_for_match = initiation_for_match.combine_first(completion_for_match)
        source_mode_for_match = clean_text(
            output.get("source_mode_analyse", pd.Series("", index=output.index))
        ).apply(normalize_label)
        linked_key_for_match = clean_text(
            output.get("linked_transaction_id", pd.Series("", index=output.index))
        )

        portal_columns = [
            "cle_sortie_turbo",
            "nombre_ecritures_portal",
            "customer_id_portal",
            "customer_ids_portal",
            "telephones_portal",
            "devises_portal",
            "account_types_portal",
            "descriptions_portal",
            "references_internes_portal",
            "date_portal_min",
            "date_portal_max",
            "montant_portal_controle",
            "portal_has_loan",
            "portal_has_fixed",
            "portal_has_normal",
            "account_type_cible",
            "operation_turbo_confirmee",
        ]
        candidate_lookup: dict[tuple[str, str, int], list[object]] = {}
        for candidate_index, candidate in output_candidates.iterrows():
            amount = pd.to_numeric(
                pd.Series([candidate.get("montant_portal_controle")]),
                errors="coerce",
            ).iloc[0]
            if pd.isna(amount):
                continue
            phones = [
                value.strip()
                for value in str(candidate.get("telephones_portal", "")).split("|")
                if value.strip()
            ]
            currencies = [
                value.strip()
                for value in str(candidate.get("devises_portal", "")).split("|")
                if value.strip()
            ]
            amount_cents = int(round(float(amount) * 100))
            for phone in phones:
                for currency in currencies:
                    candidate_lookup.setdefault(
                        (phone, currency, amount_cents), []
                    ).append(candidate_index)
        for index in output.index:
            if direct_reference.loc[index] or g2_direction.loc[index] != "Sortie":
                continue
            business_text = f"{g2_reason.loc[index]} {g2_details.loc[index]}"
            if "bisoubisoub2c" not in business_text and "b2c payment" not in business_text:
                continue
            g2_date = g2_date_for_match.loc[index]
            g2_amount = g2_amount_for_match.loc[index]
            if pd.isna(g2_date) or pd.isna(g2_amount):
                continue
            is_turbo_proxy = source_mode_for_match.loc[index] == "turbo seul"
            if is_turbo_proxy and linked_key_for_match.loc[index]:
                candidates = output_candidates.loc[
                    output_candidates["cle_sortie_turbo"]
                    .astype("string")
                    .eq(linked_key_for_match.loc[index])
                ].copy()
            else:
                amount_cents = int(round(float(g2_amount) * 100))
                candidate_indexes: list[object] = []
                for cents in range(amount_cents - 1, amount_cents + 2):
                    candidate_indexes.extend(
                        candidate_lookup.get(
                            (
                                str(g2_phone_for_match.loc[index]),
                                str(g2_currency_for_match.loc[index]),
                                cents,
                            ),
                            [],
                        )
                    )
                candidate_indexes = list(dict.fromkeys(candidate_indexes))
                candidates = output_candidates.loc[candidate_indexes].copy()
                if not candidates.empty:
                    candidates = candidates.loc[
                        pd.to_numeric(
                            candidates["montant_portal_controle"], errors="coerce"
                        )
                        .sub(float(g2_amount))
                        .abs()
                        .le(0.01)
                    ].copy()
            candidates["__ecart_minutes"] = (
                pd.to_datetime(candidates["date_portal_min"], errors="coerce")
                - g2_date
            ).dt.total_seconds() / 60
            if not is_turbo_proxy:
                candidates = candidates.loc[
                    candidates["__ecart_minutes"]
                    .abs()
                    .le(G2_TURBO_OUTPUT_MATCH_TOLERANCE_MINUTES)
                ].copy()
            if candidates.empty:
                continue
            candidates = candidates.sort_values(
                "__ecart_minutes", key=lambda values: values.abs()
            )
            best = candidates.iloc[0]
            output.at[index, "reference_sortie_turbo"] = best[
                "reference_sortie_turbo"
            ]
            output.at[index, "nombre_candidats_sortie_turbo"] = int(
                len(candidates)
            )
            output.at[index, "methode_rapprochement_turbo"] = (
                "Telephone + devise + montant + heure (sortie)"
            )
            for column in portal_columns:
                output.at[index, column] = best[column]

    output["nombre_ecritures_portal"] = pd.to_numeric(
        output["nombre_ecritures_portal"], errors="coerce"
    ).fillna(0).astype(int)
    has_reference = (
        output["ref_no_portal"].astype("string").fillna("").ne("")
        | output["reference_sortie_turbo"].astype("string").fillna("").ne("")
    )
    output["cle_rapprochement_turbo"] = output["ref_no_portal"].where(
        output["ref_no_portal"].astype("string").fillna("").ne(""),
        output["cle_sortie_turbo"],
    )
    g2_currency = clean_text(output.get("currency_code", pd.Series("", index=output.index))).str.upper()
    g2_phone = normalize_phone(output.get("phone_prefixe", output.get("phone", pd.Series("", index=output.index))))
    output["controle_devise"] = np.select(
        [
            ~has_reference,
            [
                _contains_pipe_value(values, expected)
                for values, expected in zip(output["devises_portal"], g2_currency)
            ],
        ],
        ["Non controlable", "Conforme"],
        default="Ecart",
    )
    output["controle_telephone"] = np.select(
        [
            ~has_reference,
            [
                _contains_pipe_value(values, expected)
                for values, expected in zip(output["telephones_portal"], g2_phone)
            ],
        ],
        ["Non controlable", "Conforme"],
        default="Ecart",
    )
    g2_amount = numeric_column(output, "transaction_amount_numeric").abs()
    portal_amount = pd.to_numeric(output["montant_portal_controle"], errors="coerce").abs()
    output["ecart_montant"] = g2_amount - portal_amount
    amount_comparable = has_reference & g2_amount.notna() & portal_amount.notna()
    output["controle_montant"] = np.select(
        [
            (~has_reference | ~amount_comparable).fillna(True).astype(bool),
            output["ecart_montant"].abs().le(0.01).fillna(False).astype(bool),
        ],
        ["Non controlable", "Conforme"],
        default="Ecart",
    )
    initiation_date = pd.to_datetime(
        output.get("initiation_time", pd.Series(pd.NaT, index=output.index)),
        errors="coerce",
    )
    completion_date = pd.to_datetime(
        output.get("completion_time", pd.Series(pd.NaT, index=output.index)),
        errors="coerce",
    )
    portal_date = pd.to_datetime(output["date_portal_min"], errors="coerce")
    output["date_creation_g2"] = initiation_date.combine_first(completion_date)
    output["source_date_creation_g2"] = np.select(
        [initiation_date.notna(), completion_date.notna()],
        ["Initiation Time", "Completion Time (repli)"],
        default="Date G2 absente",
    )
    output["date_creation_turbo"] = portal_date
    output["date_finalisation_g2"] = completion_date
    output["ecart_creation_minutes"] = (
        output["date_creation_g2"] - output["date_creation_turbo"]
    ).dt.total_seconds() / 60
    output["ecart_finalisation_minutes"] = (
        output["date_finalisation_g2"] - output["date_creation_turbo"]
    ).dt.total_seconds() / 60
    output["delai_traitement_g2_minutes"] = (
        output["date_finalisation_g2"] - initiation_date
    ).dt.total_seconds() / 60
    creation_comparable = (
        has_reference
        & output["date_creation_g2"].notna()
        & output["date_creation_turbo"].notna()
    )
    creation_same_date = (
        output["date_creation_g2"]
        .dt.date.eq(output["date_creation_turbo"].dt.date)
        .fillna(False)
        .astype(bool)
    )
    creation_within_tolerance = (
        creation_comparable
        & output["ecart_creation_minutes"]
        .abs().le(G2_TURBO_DATE_ANOMALY_TOLERANCE_MINUTES)
        .fillna(False)
    )
    creation_cross_date_tolerated = (
        creation_within_tolerance
        & ~creation_same_date
    )
    output["controle_date_creation"] = np.select(
        [
            (~creation_comparable).fillna(True).astype(bool),
            creation_within_tolerance & creation_same_date,
            creation_cross_date_tolerated,
        ],
        ["Non controlable", "Conforme", "Conforme - passage de date"],
        default="Ecart de date",
    )
    finalisation_comparable = (
        has_reference
        & output["date_finalisation_g2"].notna()
        & output["date_creation_turbo"].notna()
    )
    finalisation_same_date = (
        output["date_finalisation_g2"]
        .dt.date.eq(output["date_creation_turbo"].dt.date)
        .fillna(False)
        .astype(bool)
    )
    finalisation_within_tolerance = (
        finalisation_comparable
        & output["ecart_finalisation_minutes"]
        .abs().le(G2_TURBO_DATE_ANOMALY_TOLERANCE_MINUTES)
        .fillna(False)
    )
    finalisation_cross_date_tolerated = (
        finalisation_within_tolerance
        & ~finalisation_same_date
    )
    output["controle_date_finalisation"] = np.select(
        [
            (~finalisation_comparable).fillna(True).astype(bool),
            finalisation_within_tolerance & finalisation_same_date,
            finalisation_cross_date_tolerated,
        ],
        ["Non controlable", "Conforme", "Conforme - passage de date"],
        default="Ecart de date",
    )
    # Colonnes historiques conservees pour les filtres et exports existants.
    output["ecart_date_minutes"] = output["ecart_creation_minutes"]
    output["controle_date"] = output["controle_date_creation"]

    def date_gap_observation(row: pd.Series) -> str:
        creation_g2 = pd.to_datetime(row.get("date_creation_g2"), errors="coerce")
        creation_turbo = pd.to_datetime(row.get("date_creation_turbo"), errors="coerce")
        finalisation_g2 = pd.to_datetime(row.get("date_finalisation_g2"), errors="coerce")
        portal_max = pd.to_datetime(row.get("date_portal_max"), errors="coerce")
        creation_g2_text = (
            creation_g2.strftime("%d/%m/%Y %H:%M:%S")
            if pd.notna(creation_g2)
            else "date indisponible"
        )
        creation_turbo_text = (
            creation_turbo.strftime("%d/%m/%Y %H:%M:%S")
            if pd.notna(creation_turbo)
            else "date indisponible"
        )
        if pd.notna(portal_max) and portal_max != creation_turbo:
            creation_turbo_text += f" au {portal_max:%d/%m/%Y %H:%M:%S}"
        creation_gap = pd.to_numeric(
            pd.Series([row.get("ecart_creation_minutes")]), errors="coerce"
        ).iloc[0]
        creation_gap_text = (
            f" | Decalage creation : {abs(float(creation_gap)):.0f} minute(s)"
            if pd.notna(creation_gap)
            else ""
        )
        if row.get("controle_date_creation") == "Ecart de date":
            return (
                f"Creation G2 : {creation_g2_text} | "
                f"Creation Turbo : {creation_turbo_text}{creation_gap_text}"
            )

        delay = pd.to_numeric(
            pd.Series([row.get("delai_traitement_g2_minutes")]), errors="coerce"
        ).iloc[0]
        delay_text = (
            f" | Delai traitement G2 : {float(delay):.0f} minute(s)"
            if pd.notna(delay)
            else ""
        )
        if row.get("controle_date_creation") == "Conforme - passage de date":
            finalisation_text = (
                finalisation_g2.strftime("%d/%m/%Y %H:%M:%S")
                if pd.notna(finalisation_g2)
                else "date indisponible"
            )
            return (
                f"Creation G2 : {creation_g2_text} | "
                f"Creation Turbo : {creation_turbo_text}{creation_gap_text} | "
                f"Finalisation G2 : {finalisation_text}{delay_text} | "
                f"Passage de date tolere (seuil {G2_TURBO_DATE_ANOMALY_TOLERANCE_MINUTES:.0f} minutes)"
            )
        if row.get("controle_date_finalisation") in {
            "Ecart de date",
            "Conforme - passage de date",
        }:
            finalisation_text = (
                finalisation_g2.strftime("%d/%m/%Y %H:%M:%S")
                if pd.notna(finalisation_g2)
                else "date indisponible"
            )
            return (
                f"Creation G2 : {creation_g2_text} | "
                f"Creation Turbo : {creation_turbo_text} | "
                f"Finalisation G2 : {finalisation_text}{delay_text} | "
                "Creation conforme; finalisation sur une autre date"
            )
        if pd.notna(delay) and float(delay) < 0:
            return (
                f"Creation G2 : {creation_g2_text} | "
                f"Creation Turbo : {creation_turbo_text}{delay_text} | "
                "Delai de traitement G2 negatif"
            )
        return ""

    output["Observation"] = output.apply(date_gap_observation, axis=1)

    fallback_category = output.apply(
        lambda row: classify_g2_business_operation(row, dat_matched=False), axis=1
    )
    direction = clean_text(output.get("sens_flux", pd.Series("", index=output.index)))
    internal = fallback_category.eq("Operation interne Bisou")
    output["categorie_operation"] = fallback_category
    entry = direction.eq("Entree") & has_reference & ~internal
    portal_has_loan = output["portal_has_loan"].astype("boolean").fillna(False).astype(bool)
    portal_has_fixed = output["portal_has_fixed"].astype("boolean").fillna(False).astype(bool)
    portal_has_normal = output["portal_has_normal"].astype("boolean").fillna(False).astype(bool)
    output.loc[entry & portal_has_loan, "categorie_operation"] = "Remboursement prets"
    output.loc[entry & ~portal_has_loan & portal_has_fixed, "categorie_operation"] = "DAT"
    output.loc[
        entry
        & ~portal_has_loan
        & ~portal_has_fixed
        & portal_has_normal,
        "categorie_operation",
    ] = "Depot normal"
    output["description_metier"] = output["categorie_operation"]

    turbo_only = clean_text(
        output.get("source_mode_analyse", pd.Series("", index=output.index))
    ).apply(normalize_label).eq("turbo seul")

    raw_status = clean_text(
        output.get("transaction_status", pd.Series("", index=output.index))
    )
    output["transaction_status"] = raw_status
    output["statut_transaction_g2"] = raw_status.apply(normalize_g2_transaction_status)
    output["est_transaction_terminee"] = g2_completed_transaction_mask(output)
    output.loc[turbo_only, "statut_transaction_g2"] = "Comptabilisee Turbo"
    output.loc[turbo_only, "est_transaction_terminee"] = True
    output["incluse_synthese"] = output["est_transaction_terminee"]
    output["traitement_statut_g2"] = np.where(
        output["incluse_synthese"],
        "Incluse dans les analyses",
        "Controle uniquement",
    )
    output.loc[turbo_only, "traitement_statut_g2"] = (
        "Incluse - operation comptabilisee dans Turbo"
    )
    output.loc[turbo_only, "date_creation_g2"] = pd.NaT
    output.loc[turbo_only, "source_date_creation_g2"] = "G2 absent"
    output.loc[turbo_only, "date_finalisation_g2"] = pd.NaT
    output.loc[turbo_only, "ecart_creation_minutes"] = np.nan
    output.loc[turbo_only, "ecart_finalisation_minutes"] = np.nan
    output.loc[turbo_only, "delai_traitement_g2_minutes"] = np.nan
    output.loc[turbo_only, "ecart_date_minutes"] = np.nan
    for column in [
        "controle_telephone",
        "controle_devise",
        "controle_montant",
        "controle_date_creation",
        "controle_date_finalisation",
        "controle_date",
    ]:
        output.loc[turbo_only, column] = "Non applicable - Turbo seul"
    output.loc[turbo_only, "ecart_montant"] = np.nan
    output.loc[turbo_only, "Observation"] = (
        "Analyse Turbo seule : controles independants G2/Turbo non applicables."
    )
    turbo_direct = turbo_only & output["ref_no_portal"].astype("string").fillna("").ne("")
    turbo_output = turbo_only & output["reference_sortie_turbo"].astype("string").fillna("").ne("")
    output.loc[turbo_direct, "methode_rapprochement_turbo"] = (
        "Agregation Turbo par ref_no"
    )
    output.loc[turbo_output, "methode_rapprochement_turbo"] = (
        "Agregation Turbo par reference_id + created_at"
    )
    has_control_gap = (
        output[["controle_devise", "controle_telephone", "controle_montant", "controle_date"]]
        .eq("Ecart")
        .any(axis=1)
        | output["controle_date"].eq("Ecart de date")
        | pd.to_numeric(
            output["nombre_candidats_sortie_turbo"], errors="coerce"
        ).fillna(0).gt(1)
    )
    output["statut_rapprochement"] = np.select(
        [turbo_only, internal & ~has_reference, ~has_reference, has_control_gap],
        [
            "Non applicable - Turbo seul",
            "Non applicable - operation interne",
            "Non rapproche",
            "Rapproche avec ecart",
        ],
        default="Rapproche exact",
    )

    def anomaly_reason(row: pd.Series) -> str:
        reasons: list[str] = []
        if _is_empty_text(row.get("receipt_no")):
            reasons.append("Receipt No manquant")
        if bool(row.get("doublon_receipt_no", False)):
            reasons.append(f"Receipt No duplique ({int(row.get('nombre_lignes_g2_reference', 0))} lignes G2)")
        if not bool(row.get("est_transaction_terminee", True)):
            status_label = row.get("transaction_status", "") or row.get("statut_transaction_g2", "")
            reasons.append(f"Statut G2 non termine : {status_label}")
        if row.get("statut_rapprochement") == "Non rapproche":
            reasons.append("Receipt No non trouve dans ref_no")
        for column, label in [
            ("controle_devise", "Ecart de devise"),
            ("controle_telephone", "Ecart de telephone"),
            ("controle_montant", "Ecart de montant"),
        ]:
            if row.get(column) == "Ecart":
                reasons.append(label)
        if row.get("controle_date") == "Ecart de date":
            reasons.append("Ecart de date de creation")
        candidates = pd.to_numeric(
            pd.Series([row.get("nombre_candidats_sortie_turbo")]), errors="coerce"
        ).iloc[0]
        if pd.notna(candidates) and int(candidates) > 1:
            reasons.append(
                f"Plusieurs candidats Turbo pour la sortie ({int(candidates)})"
            )
        delay = pd.to_numeric(
            pd.Series([row.get("delai_traitement_g2_minutes")]), errors="coerce"
        ).iloc[0]
        if pd.notna(delay) and float(delay) < 0:
            reasons.append("Delai de traitement G2 negatif")
        if row.get("categorie_operation") in {"Flux a verifier", "Autre entree", "Autre sortie"}:
            reasons.append("Operation non classee")
        return " | ".join(reasons)

    output["motif_anomalie"] = output.apply(anomaly_reason, axis=1)
    output["est_anomalie"] = output["motif_anomalie"].ne("")
    start = pd.to_datetime(output.get("completion_time", pd.Series(pd.NaT, index=output.index)), errors="coerce").min()
    end = pd.to_datetime(output.get("completion_time", pd.Series(pd.NaT, index=output.index)), errors="coerce").max()
    output["source_analytique"] = np.select(
        [turbo_only, has_reference],
        ["Turbo seul", "G2 + Turbo"],
        default="G2",
    )
    output["identifiant_lot"] = (
        f"G2_{start:%Y%m%d}_{end:%Y%m%d}" if pd.notna(start) and pd.notna(end) else "G2_sans_periode"
    )
    return output


def _filter_g2_crosscheck_by_customer(
    frame: pd.DataFrame,
    customer_id: str | None,
) -> pd.DataFrame:
    """Limite un controle G2 aux identifiants Turbo effectivement rapproches."""
    if customer_id is None or frame.empty:
        return frame
    customer_text = str(customer_id).strip()
    matches = pd.Series(False, index=frame.index)
    for column in ["customer_id_dat", "customer_id_ref_no", "customer_id_portal"]:
        if column not in frame.columns:
            continue
        matches = matches | frame[column].apply(
            lambda value: (
                False
                if _is_empty_text(value)
                else any(part.strip() == customer_text for part in str(value).split("|"))
            )
        )
    return frame.loc[matches].copy().reset_index(drop=True)


def build_g2_dat_crosscheck(prepared: MpesaPreparedData, customer_id: str | None = None) -> pd.DataFrame:
    g2 = prepared.g2_transactions
    fixed = prepared.fixed_savings
    transactions = prepared.transactions
    if g2.empty:
        return pd.DataFrame()

    output = _enrich_g2_with_portal_controls(g2, transactions)
    if not transactions.empty and "ref_no" in transactions.columns:
        tx = transactions.copy()
        tx["ref_no"] = clean_identifier(tx["ref_no"])
        receipt_numbers = set(
            clean_identifier(output.get("receipt_no", pd.Series(dtype="string")))
            .loc[lambda values: values.ne("")]
            .astype(str)
        )
        tx = tx.loc[tx["ref_no"].isin(receipt_numbers)].copy()
        tx["currency_code"] = clean_text(tx["currency_code"]).str.upper() if "currency_code" in tx.columns else ""
        tx["reference_dat_ligne"] = tx["reference_id"].apply(lambda value: extract_prefixed_reference(value, "FA")) if "reference_id" in tx.columns else ""
        tx_summary = (
            tx.groupby(["ref_no", "currency_code"], as_index=False, dropna=False)
            .agg(
                customer_id_ref_no=("customer_id", first_non_empty),
                customer_ids_ref_no=("customer_id", concat_unique),
                telephone_ref_no=("msisdn1", first_non_empty),
                telephones_ref_no=("msisdn1", concat_unique),
                references_transactions=("reference_id", concat_unique),
                references_dat_transactions=("reference_dat_ligne", concat_unique),
                account_types_transactions=("account_type", concat_unique),
                descriptions_transactions=("description", concat_unique),
                date_transaction_ref_no=("created_at", "min"),
                nb_lignes_transactions=("ref_no", "size"),
            )
        )
        fixed_tx = tx.loc[tx["account_type"].eq("FIXED SAVINGS")].copy() if "account_type" in tx.columns else pd.DataFrame()
        if not fixed_tx.empty:
            fixed_tx["variation_dat_operation"] = (
                numeric_column(fixed_tx, "bal_after")
                - numeric_column(fixed_tx, "bal_before")
            )
            fixed_tx["reference_dat_operation_ligne"] = fixed_tx["reference_id"].apply(
                lambda value: extract_prefixed_reference(value, "FA")
            ) if "reference_id" in fixed_tx.columns else ""
            fixed_tx_summary = (
                fixed_tx.groupby(["ref_no", "currency_code"], as_index=False, dropna=False)
                .agg(
                    reference_dat_operation=("reference_dat_operation_ligne", concat_unique),
                    solde_dat_operation_avant=("bal_before", "sum"),
                    solde_dat_operation_apres=("bal_after", "sum"),
                    variation_dat_operation=("variation_dat_operation", "sum"),
                    nb_lignes_fixed_savings=("ref_no", "size"),
                    descriptions_dat_operation=("description", concat_unique),
                )
            )
            tx_summary = tx_summary.merge(
                fixed_tx_summary,
                on=["ref_no", "currency_code"],
                how="left",
            )
        else:
            tx_summary["reference_dat_operation"] = ""
            tx_summary["solde_dat_operation_avant"] = np.nan
            tx_summary["solde_dat_operation_apres"] = np.nan
            tx_summary["variation_dat_operation"] = np.nan
            tx_summary["nb_lignes_fixed_savings"] = 0
            tx_summary["descriptions_dat_operation"] = ""
        output = output.merge(
            tx_summary,
            left_on=["receipt_no", "currency_code"],
            right_on=["ref_no", "currency_code"],
            how="left",
        )
    else:
        output["customer_id_ref_no"] = pd.NA
        output["telephone_ref_no"] = pd.NA
        output["references_transactions"] = ""
        output["references_dat_transactions"] = ""
        output["account_types_transactions"] = ""
        output["descriptions_transactions"] = ""
        output["date_transaction_ref_no"] = pd.NaT
        output["nb_lignes_transactions"] = 0
        output["reference_dat_operation"] = ""
        output["solde_dat_operation_avant"] = np.nan
        output["solde_dat_operation_apres"] = np.nan
        output["variation_dat_operation"] = np.nan
        output["nb_lignes_fixed_savings"] = 0
        output["descriptions_dat_operation"] = ""

    # Propager le rapprochement secondaire des sorties dans les colonnes déjà
    # utilisées par les rapports clients. La provenance reste explicite grâce
    # à methode_rapprochement_turbo et reference_sortie_turbo.
    output_match = output.get(
        "reference_sortie_turbo", pd.Series(pd.NA, index=output.index)
    ).astype("string").fillna("").ne("")
    fallback_values = {
        "customer_id_ref_no": "customer_id_portal",
        "customer_ids_ref_no": "customer_ids_portal",
        "telephone_ref_no": "telephones_portal",
        "telephones_ref_no": "telephones_portal",
        "references_transactions": "reference_sortie_turbo",
        "account_types_transactions": "account_types_portal",
        "descriptions_transactions": "descriptions_portal",
        "date_transaction_ref_no": "date_creation_turbo",
        "nb_lignes_transactions": "nombre_ecritures_portal",
    }
    for target, source in fallback_values.items():
        if target not in output.columns:
            output[target] = pd.NA
        if source not in output.columns:
            continue
        current = output[target]
        if pd.api.types.is_datetime64_any_dtype(current):
            current_missing = pd.to_datetime(current, errors="coerce").isna()
        elif pd.api.types.is_numeric_dtype(current):
            current_missing = pd.to_numeric(current, errors="coerce").isna()
        else:
            current_missing = current.astype("string").fillna("").eq("")
        replace_mask = output_match & current_missing
        output.loc[replace_mask, target] = output.loc[replace_mask, source]

    if fixed.empty or "msisdn" not in fixed.columns:
        output["customer_id_dat"] = pd.NA
        output["dat_final_client_devise"] = np.nan
        output["nombre_dat_client_devise"] = 0
        output["produits_dat"] = ""
        output["maturites_dat"] = ""
        output["mode_rapprochement"] = np.select(
            [
                output_match,
                output["customer_id_ref_no"].astype("string").fillna("").ne(""),
            ],
            [
                "Sortie Turbo: telephone + devise + montant + heure",
                "Receipt No = ref_no",
            ],
            default="Non rapproche",
        )
        output["statut_rapprochement_dat"] = "Fichier DAT absent"
        return _filter_g2_crosscheck_by_customer(output, customer_id)

    dat = fixed.copy()
    dat["phone_prefixe"] = normalize_phone(dat["msisdn"])
    dat["balance"] = numeric_column(dat, "balance")
    dat["currency_code"] = clean_text(dat["currency_code"]).str.upper() if "currency_code" in dat.columns else ""
    dat_by_customer = (
        dat.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
        .agg(
            dat_final_par_ref_no=("balance", "sum"),
            nombre_dat_par_ref_no=("balance", "size"),
            produits_dat_par_ref_no=("product_name", concat_unique),
            maturites_dat_par_ref_no=("maturity_date", concat_unique),
        )
    )
    dat_by_phone = (
        dat.groupby(["phone_prefixe", "currency_code"], as_index=False, dropna=False)
        .agg(
            customer_id_dat_phone=("customer_id", concat_unique),
            dat_final_par_phone=("balance", "sum"),
            nombre_dat_par_phone=("balance", "size"),
            produits_dat_par_phone=("product_name", concat_unique),
            maturites_dat_par_phone=("maturity_date", concat_unique),
        )
    )
    output = output.merge(
        dat_by_customer,
        left_on=["customer_id_ref_no", "currency_code"],
        right_on=["customer_id", "currency_code"],
        how="left",
    )
    if "customer_id" in output.columns:
        output = output.drop(columns=["customer_id"])
    output = output.merge(dat_by_phone, on=["phone_prefixe", "currency_code"], how="left")

    output["nb_lignes_fixed_savings"] = numeric_column(output, "nb_lignes_fixed_savings").astype(int)
    for column in [
        "reference_dat_operation",
        "descriptions_dat_operation",
        "references_dat_transactions",
        "references_transactions",
        "account_types_transactions",
        "descriptions_transactions",
    ]:
        if column in output.columns:
            output[column] = output[column].fillna("")
    has_ref_match = output["customer_id_ref_no"].astype("string").fillna("").ne("")
    has_dat_operation = output["reference_dat_operation"].astype("string").fillna("").ne("") | output["nb_lignes_fixed_savings"].gt(0)
    has_dat_by_ref = pd.to_numeric(output["nombre_dat_par_ref_no"], errors="coerce").fillna(0).gt(0)
    has_phone_match = output["customer_id_dat_phone"].astype("string").fillna("").ne("")
    output["mode_rapprochement"] = np.select(
        [
            output_match,
            has_ref_match & has_dat_operation,
            has_ref_match & has_dat_by_ref,
            has_ref_match,
            has_phone_match,
        ],
        [
            "Sortie Turbo: telephone + devise + montant + heure",
            "Receipt No = ref_no + DAT operation",
            "Receipt No = ref_no + DAT final client",
            "Receipt No = ref_no sans DAT",
            "Telephone G2 = telephone DAT",
        ],
        default="Non rapproche",
    )
    output["customer_id_dat"] = output["customer_id_ref_no"].where(has_ref_match, output["customer_id_dat_phone"])
    output["dat_final_client_devise"] = output["dat_final_par_ref_no"].where(has_dat_by_ref, output["dat_final_par_phone"])
    output["nombre_dat_client_devise"] = output["nombre_dat_par_ref_no"].where(has_dat_by_ref, output["nombre_dat_par_phone"])
    output["produits_dat"] = output["produits_dat_par_ref_no"].where(has_dat_by_ref, output["produits_dat_par_phone"])
    output["maturites_dat"] = output["maturites_dat_par_ref_no"].where(has_dat_by_ref, output["maturites_dat_par_phone"])
    output["dat_operation"] = output["reference_dat_operation"].astype("string").fillna("").replace("", pd.NA)
    output["solde_dat_operation"] = pd.to_numeric(output["solde_dat_operation_apres"], errors="coerce")
    output["dat_final"] = pd.to_numeric(output["dat_final_client_devise"], errors="coerce")

    output = _filter_g2_crosscheck_by_customer(output, customer_id)
    output["nombre_dat_client_devise"] = pd.to_numeric(output["nombre_dat_client_devise"], errors="coerce").fillna(0).astype(int)
    has_ref_match = output["customer_id_ref_no"].astype("string").fillna("").ne("")
    has_dat_operation = output["reference_dat_operation"].astype("string").fillna("").ne("") | output["nb_lignes_fixed_savings"].gt(0)
    output["statut_rapprochement_dat"] = np.select(
        [
            output["customer_id_ref_no"].astype("string").fillna("").eq("") & output["customer_id_dat"].astype("string").fillna("").eq(""),
            has_ref_match & has_dat_operation,
            output["mode_rapprochement"].eq("Receipt No = ref_no sans DAT"),
            pd.to_numeric(output["dat_final_client_devise"], errors="coerce").fillna(0).gt(0),
            output["nombre_dat_client_devise"].gt(0),
        ],
        [
            "Receipt No non trouve dans ref_no et telephone non trouve dans DAT",
            "DAT de l'operation retrouve via ref_no",
            "Receipt No trouve dans ref_no mais aucune ligne FIXED SAVINGS",
            "DAT trouve avec solde",
            "DAT trouve sans solde positif",
        ],
        default="A verifier",
    )
    preferred_columns = [
        "receipt_no",
        "completion_time",
        "currency_code",
        "transaction_amount",
        "transaction_amount_numeric",
        "categorie_operation",
        "description_metier",
        "ref_no_portal",
        "reference_sortie_turbo",
        "cle_sortie_turbo",
        "cle_rapprochement_turbo",
        "methode_rapprochement_turbo",
        "nombre_candidats_sortie_turbo",
        "operation_turbo_confirmee",
        "nombre_ecritures_portal",
        "account_types_portal",
        "descriptions_portal",
        "statut_rapprochement",
        "controle_telephone",
        "controle_devise",
        "montant_portal_controle",
        "ecart_montant",
        "controle_montant",
        "date_creation_g2",
        "source_date_creation_g2",
        "date_creation_turbo",
        "ecart_creation_minutes",
        "controle_date_creation",
        "date_finalisation_g2",
        "delai_traitement_g2_minutes",
        "ecart_finalisation_minutes",
        "controle_date_finalisation",
        "ecart_date_minutes",
        "controle_date",
        "Observation",
        "nombre_lignes_g2_reference",
        "devises_g2_reference",
        "statuts_g2_reference",
        "montants_g2_reference",
        "doublon_receipt_no",
        "incluse_synthese",
        "motif_anomalie",
        "est_anomalie",
        "balance",
        "balance_numeric",
        "opposite_party",
        "phone",
        "phone_prefixe",
        "ref_no",
        "customer_id_ref_no",
        "customer_ids_ref_no",
        "telephone_ref_no",
        "telephones_ref_no",
        "references_transactions",
        "references_dat_transactions",
        "dat_operation",
        "solde_dat_operation",
        "dat_final",
        "reference_dat_operation",
        "solde_dat_operation_avant",
        "solde_dat_operation_apres",
        "variation_dat_operation",
        "nb_lignes_fixed_savings",
        "descriptions_dat_operation",
        "account_types_transactions",
        "date_transaction_ref_no",
        "customer_id_dat",
        "dat_final_client_devise",
        "nombre_dat_client_devise",
        "produits_dat",
        "maturites_dat",
        "mode_rapprochement",
        "transaction_status",
        "details",
        "operation",
        "statut_rapprochement_dat",
    ]
    present = [column for column in preferred_columns if column in output.columns]
    rest = [column for column in output.columns if column not in present]
    return output[present + rest].reset_index(drop=True)


def classify_g2_entry_report(row: pd.Series) -> str:
    reference_value = row.get("reference_dat_operation", "")
    reference_dat = "" if _is_empty_text(reference_value) else str(reference_value).strip()
    references = normalize_label(row.get("references_transactions", ""))
    account_types = normalize_label(row.get("account_types_transactions", ""))
    descriptions = normalize_label(row.get("descriptions_transactions", ""))
    text = f"{references} {account_types} {descriptions}"
    classified = classify_g2_business_operation(row, dat_matched=bool(reference_dat))
    if classified not in {"Depot normal", "Autre entree"}:
        return classified
    if "ln" in references or "loan" in account_types or "principle" in account_types or "remboursement" in text:
        return "Remboursement prets"
    return classified


def build_entry_pivot(detail: pd.DataFrame) -> pd.DataFrame:
    if detail is None or not isinstance(detail, pd.DataFrame) or detail.empty:
        return pd.DataFrame()
    required_columns = {"currency_code", "details_rapport", "montant"}
    if not required_columns.issubset(detail.columns):
        return pd.DataFrame()

    work = detail.copy()
    work["currency_code"] = clean_text(work["currency_code"]).str.upper()
    work["details_rapport"] = clean_text(work["details_rapport"])
    work["montant"] = numeric_column(work, "montant")
    amount_pivot = (
        work.pivot_table(
            index="currency_code",
            columns="details_rapport",
            values="montant",
            aggfunc="sum",
            fill_value=0,
        )
        .reset_index()
        .rename_axis(None, axis=1)
    )
    count_pivot = (
        work.pivot_table(
            index="currency_code",
            columns="details_rapport",
            values="montant",
            aggfunc="count",
            fill_value=0,
        )
        .reset_index()
        .rename_axis(None, axis=1)
    )
    amount_categories = [column for column in amount_pivot.columns if column != "currency_code"]
    count_categories = [column for column in count_pivot.columns if column != "currency_code"]
    amount_pivot = amount_pivot.rename(columns={column: f"montant_{column}" for column in amount_categories})
    count_pivot = count_pivot.rename(columns={column: f"nombre_{column}" for column in count_categories})
    pivot = amount_pivot.merge(count_pivot, on="currency_code", how="outer").fillna(0)

    preferred_details = G2_OPERATION_CATEGORIES
    for detail_name in preferred_details:
        amount_column = f"montant_{detail_name}"
        count_column = f"nombre_{detail_name}"
        if amount_column not in pivot.columns:
            pivot[amount_column] = 0.0
        if count_column not in pivot.columns:
            pivot[count_column] = 0

    amount_columns = [column for column in pivot.columns if column.startswith("montant_")]
    count_columns = [column for column in pivot.columns if column.startswith("nombre_")]
    pivot["montant_total"] = pivot[amount_columns].sum(axis=1)
    pivot["nombre_total"] = pivot[count_columns].sum(axis=1)

    if "sens_flux" in work.columns:
        flow_summary = (
            work.groupby(["currency_code", "sens_flux"], dropna=False)
            .agg(montant=("montant", "sum"), nombre=("montant", "size"))
            .reset_index()
        )
        flow_amounts = flow_summary.pivot_table(
            index="currency_code", columns="sens_flux", values="montant", aggfunc="sum", fill_value=0
        ).reset_index().rename_axis(None, axis=1)
        flow_counts = flow_summary.pivot_table(
            index="currency_code", columns="sens_flux", values="nombre", aggfunc="sum", fill_value=0
        ).reset_index().rename_axis(None, axis=1)
        flow_amounts = flow_amounts.rename(
            columns={"Entree": "montant_total_entrees", "Sortie": "montant_total_sorties", "A verifier": "montant_flux_a_verifier", "Indetermine": "montant_flux_indetermine"}
        )
        flow_counts = flow_counts.rename(
            columns={"Entree": "nombre_entrees", "Sortie": "nombre_sorties", "A verifier": "nombre_flux_a_verifier", "Indetermine": "nombre_flux_indetermine"}
        )
        pivot = pivot.merge(flow_amounts, on="currency_code", how="left").merge(flow_counts, on="currency_code", how="left")
    for column in ["montant_total_entrees", "montant_total_sorties", "nombre_entrees", "nombre_sorties"]:
        if column not in pivot.columns:
            pivot[column] = 0
    pivot["solde_net_flux"] = numeric_column(pivot, "montant_total_entrees") - numeric_column(pivot, "montant_total_sorties")

    ordered = ["currency_code"]
    for detail_name in preferred_details:
        ordered.extend([f"montant_{detail_name}", f"nombre_{detail_name}"])
    ordered.extend(
        [
            "nombre_entrees", "montant_total_entrees", "nombre_sorties", "montant_total_sorties",
            "solde_net_flux", "montant_total", "nombre_total",
        ]
    )
    rest = [column for column in pivot.columns if column not in ordered]
    return pivot[ordered + rest].sort_values("currency_code").reset_index(drop=True)


def build_entry_count_summary(detail: pd.DataFrame) -> pd.DataFrame:
    """Compte chaque categorie G2, avec entrees et sorties separees par devise."""
    rename_map = {category: f"Nombre de {category.lower()}" for category in G2_OPERATION_CATEGORIES}
    rename_map.update(
        {
            "DAT": "Nombre de DAT",
            "Depot normal": "Nombre de depot normal",
            "Remboursement prets": "Nombre de remboursement de pret",
            "Paiement client B2C": "Nombre de paiement client B2C",
            "Demande de credit": "Nombre de demande de credit",
            "Operation interne Bisou": "Nombre d'operation interne Bisou",
        }
    )
    category_columns = [rename_map[category] for category in G2_OPERATION_CATEGORIES]
    columns = ["currency_code", *category_columns, "Nombre d'entrees", "Nombre de sorties", "Nombre total"]
    if detail is None or not isinstance(detail, pd.DataFrame) or detail.empty:
        return pd.DataFrame(columns=columns)
    if not {"currency_code", "details_rapport"}.issubset(detail.columns):
        return pd.DataFrame(columns=columns)

    work = detail.copy()
    work["currency_code"] = clean_text(work["currency_code"]).str.upper()
    work["details_rapport"] = clean_text(work["details_rapport"])
    grouped = (
        work.groupby(["currency_code", "details_rapport"], dropna=False)
        .size()
        .unstack(fill_value=0)
    )
    for category in G2_OPERATION_CATEGORIES:
        if category not in grouped.columns:
            grouped[category] = 0
    summary = grouped[G2_OPERATION_CATEGORIES].reset_index().rename(columns=rename_map)
    if "sens_flux" in work.columns:
        direction_counts = work.groupby(["currency_code", "sens_flux"]).size().unstack(fill_value=0)
        summary["Nombre d'entrees"] = summary["currency_code"].map(direction_counts.get("Entree", pd.Series(dtype="int64"))).fillna(0).astype(int)
        summary["Nombre de sorties"] = summary["currency_code"].map(direction_counts.get("Sortie", pd.Series(dtype="int64"))).fillna(0).astype(int)
    else:
        summary["Nombre d'entrees"] = summary[category_columns].sum(axis=1)
        summary["Nombre de sorties"] = 0
    summary["Nombre total"] = summary[category_columns].sum(axis=1)
    return summary[columns].sort_values("currency_code").reset_index(drop=True)


def build_entry_vertical_summary(detail: pd.DataFrame) -> pd.DataFrame:
    if detail is None or not isinstance(detail, pd.DataFrame) or detail.empty:
        return pd.DataFrame(columns=["currency_code", "details_rapport", "montant"])
    required_columns = {"currency_code", "details_rapport", "montant"}
    if not required_columns.issubset(detail.columns):
        return pd.DataFrame(columns=["currency_code", "details_rapport", "montant"])

    work = detail.copy()
    work["currency_code"] = clean_text(work["currency_code"]).str.upper()
    work["details_rapport"] = clean_text(work["details_rapport"])
    work["montant"] = numeric_column(work, "montant")
    grouped = (
        work.groupby(["currency_code", "details_rapport"], as_index=False, dropna=False)
        .agg(montant=("montant", "sum"))
    )
    currencies = sorted(value for value in grouped["currency_code"].dropna().astype(str).unique() if value)
    observed_categories = [
        value for value in grouped["details_rapport"].dropna().astype(str).unique() if value not in G2_OPERATION_CATEGORIES
    ]
    categories = [*G2_OPERATION_CATEGORIES, *sorted(observed_categories)]
    rows: list[dict[str, object]] = []
    for currency in currencies:
        currency_group = grouped.loc[grouped["currency_code"].eq(currency)]
        values = {
            str(row["details_rapport"]): float(row["montant"])
            for _, row in currency_group.iterrows()
        }
        for category in categories:
            rows.append(
                {
                    "currency_code": currency,
                    "details_rapport": category,
                    "montant": values.get(category, 0.0),
                }
            )
        rows.append(
            {
                "currency_code": currency,
                "details_rapport": f"Total {currency}",
                "montant": sum(values.values()),
            }
        )
    return pd.DataFrame(rows)


def build_g2_entry_report(prepared: MpesaPreparedData) -> dict[str, pd.DataFrame]:
    detail = build_g2_dat_crosscheck(prepared)
    if detail.empty:
        return {
            "detail": pd.DataFrame(),
            "synthese": pd.DataFrame(),
            "pivot": pd.DataFrame(),
            "vertical_summary": pd.DataFrame(),
            "comptages": build_entry_count_summary(pd.DataFrame()),
        }

    report = detail.copy()
    report["details_rapport"] = report.apply(classify_g2_entry_report, axis=1)
    report["montant"] = numeric_column(report, "transaction_amount_numeric").abs()
    report["sens_flux"] = clean_text(report.get("sens_flux", pd.Series("Indetermine", index=report.index)))
    report["montant_entree"] = report["montant"].where(report["sens_flux"].eq("Entree"), 0.0)
    report["montant_sortie"] = report["montant"].where(report["sens_flux"].eq("Sortie"), 0.0)
    report["date"] = pd.to_datetime(report.get("completion_time"), errors="coerce")
    report["duree"] = np.where(
        report["details_rapport"].eq("DAT"),
        report.get("produits_dat", pd.Series("", index=report.index)).astype("string").fillna(""),
        "-",
    )
    report["compte_cree"] = pd.to_datetime(report.get("date_transaction_ref_no"), errors="coerce")
    report["receipt_no"] = report.get("receipt_no", pd.Series("", index=report.index))
    report["opposite_party"] = report.get("opposite_party", pd.Series("", index=report.index))
    report["currency_code"] = report.get("currency_code", pd.Series("", index=report.index))

    detail_columns = [
        "date",
        "receipt_no",
        "sens_flux",
        "details_rapport",
        "reason_type",
        "opposite_party",
        "Nom_client",
        "duree",
        "compte_cree",
        "montant",
        "montant_entree",
        "montant_sortie",
        "currency_code",
        "customer_id_ref_no",
        "reference_dat_operation",
        "solde_dat_operation_avant",
        "solde_dat_operation_apres",
        "variation_dat_operation",
        "mode_rapprochement",
        "methode_rapprochement_turbo",
        "reference_sortie_turbo",
        "cle_sortie_turbo",
        "operation_turbo_confirmee",
        "statut_rapprochement_dat",
    ]
    detail_columns = [column for column in detail_columns if column in report.columns]
    report_detail = report[detail_columns].sort_values(["currency_code", "date"], ascending=[True, False]).reset_index(drop=True)

    synthese = (
        report.groupby(["currency_code", "sens_flux", "details_rapport"], as_index=False, dropna=False)
        .agg(nombre=("receipt_no", "count"), montant=("montant", "sum"))
        .sort_values(["currency_code", "sens_flux", "details_rapport"])
    )
    totals = (
        report.groupby("currency_code", as_index=False, dropna=False)
        .agg(nombre=("receipt_no", "count"), montant=("montant", "sum"))
    )
    totals["sens_flux"] = "Tous flux"
    totals["details_rapport"] = "Total " + totals["currency_code"].astype("string").fillna("")
    synthese = concat_frames_stable([synthese, totals[synthese.columns]]).reset_index(drop=True)

    return {
        "detail": report_detail,
        "synthese": synthese,
        "pivot": build_entry_pivot(report_detail),
        "vertical_summary": build_entry_vertical_summary(report_detail),
        "comptages": build_entry_count_summary(report_detail),
    }


def _is_round_savings_amount(amount: object, currency: object) -> bool:
    value = pd.to_numeric(pd.Series([amount]), errors="coerce").iloc[0]
    if pd.isna(value) or value <= 0:
        return False
    currency_text = str(currency or "").upper()
    if currency_text == "CDF":
        return abs(value % 500) < 0.01 or abs(value % 500 - 500) < 0.01
    if currency_text == "USD":
        return abs(value % 1) < 0.01
    return True


def _best_subset_for_target(items: list[tuple[int, float]], target: float, tolerance: float = 0.01) -> list[int]:
    if not items or pd.isna(target) or target <= 0:
        return []
    best: list[int] = []
    best_diff = float("inf")

    def walk(position: int, selected: list[int], total: float) -> bool:
        nonlocal best, best_diff
        diff = abs(total - target)
        if diff < best_diff:
            best = selected.copy()
            best_diff = diff
        if diff <= tolerance:
            return True
        if position >= len(items) or total > target + tolerance:
            return False
        for index in range(position, len(items)):
            row_index, amount = items[index]
            if walk(index + 1, selected + [row_index], total + amount):
                return True
        return False

    walk(0, [], 0.0)
    return best if best_diff <= tolerance else []


def _build_daily_dat_assignments(g2: pd.DataFrame, fixed: pd.DataFrame) -> pd.DataFrame:
    columns = [
        "receipt_no",
        "dat_customer_id",
        "duree",
        "compte_cree",
        "maturity_date",
        "dat_balance",
        "dat_match_rule",
    ]
    if g2.empty:
        return pd.DataFrame(columns=columns)

    assignments = pd.DataFrame({"receipt_no": g2.get("receipt_no", pd.Series(dtype="string"))})
    for column in columns[1:]:
        assignments[column] = pd.NA
    if fixed.empty or "msisdn" not in fixed.columns:
        return assignments

    dat = fixed.copy()
    dat["phone_prefixe"] = normalize_phone(dat["msisdn"])
    dat["currency_code"] = clean_text(dat.get("currency_code", pd.Series("", index=dat.index))).str.upper()
    dat["balance"] = numeric_column(dat, "balance")
    dat["compte_cree"] = pd.to_datetime(
        dat.get("created_at", dat.get("date_approved", pd.Series(pd.NaT, index=dat.index))),
        errors="coerce",
    )
    if "product_name" not in dat.columns:
        dat["product_name"] = dat.get("account_type", pd.Series("", index=dat.index))
    dat["jour_creation"] = dat["compte_cree"].dt.date

    tx = g2.copy()
    tx["receipt_no"] = clean_identifier(tx.get("receipt_no", pd.Series("", index=tx.index)))
    tx["phone_prefixe"] = normalize_phone(tx.get("phone_prefixe", tx.get("phone", pd.Series("", index=tx.index))))
    tx["currency_code"] = clean_text(tx.get("currency_code", pd.Series("", index=tx.index))).str.upper()
    tx["montant"] = numeric_column(tx, "transaction_amount_numeric").abs()
    tx["date"] = pd.to_datetime(tx.get("completion_time", pd.Series(pd.NaT, index=tx.index)), errors="coerce")
    tx["jour"] = tx["date"].dt.date
    if "sens_flux" in tx.columns:
        tx = tx.loc[tx["sens_flux"].astype("string").eq("Entree")].copy()

    assignment_by_receipt: dict[str, dict[str, object]] = {}
    group_columns = ["phone_prefixe", "currency_code", "jour"]
    relevant_dat_keys = pd.MultiIndex.from_frame(
        tx[group_columns].drop_duplicates()
    )
    dat_keys = pd.MultiIndex.from_frame(
        dat[["phone_prefixe", "currency_code", "jour_creation"]]
    )
    dat = dat.loc[dat_keys.isin(relevant_dat_keys)].copy()
    dat_groups = {
        key: group.copy()
        for key, group in dat.groupby(
            ["phone_prefixe", "currency_code", "jour_creation"],
            dropna=False,
            sort=False,
        )
    }
    for group_key, group in tx.groupby(group_columns, dropna=False):
        phone, currency, day = group_key
        if _is_empty_text(phone) or _is_empty_text(currency) or pd.isna(day):
            continue
        dat_group = dat_groups.get((phone, currency, day), pd.DataFrame()).copy()
        if dat_group.empty:
            continue

        unassigned = {
            int(index): float(value)
            for index, value in group["montant"].items()
            if not pd.isna(value) and float(value) > 0
        }
        dat_group = dat_group.sort_values(["compte_cree", "balance"], ascending=[False, False])
        for dat_index, dat_row in dat_group.iterrows():
            items = sorted(unassigned.items(), key=lambda item: item[1], reverse=True)
            selected_indexes = _best_subset_for_target(items, float(dat_row.get("balance", 0)))
            if not selected_indexes:
                continue
            for tx_index in selected_indexes:
                receipt_no = str(tx.loc[tx_index, "receipt_no"])
                assignment_by_receipt[receipt_no] = {
                    "dat_customer_id": dat_row.get("customer_id", pd.NA),
                    "duree": dat_row.get("product_name", pd.NA),
                    "compte_cree": dat_row.get("compte_cree", pd.NaT),
                    "maturity_date": dat_row.get("maturity_date", pd.NaT),
                    "dat_balance": dat_row.get("balance", np.nan),
                    "dat_match_rule": "telephone + devise + montant DAT du jour",
                }
                unassigned.pop(tx_index, None)

        has_dat_assignment = any(
            str(receipt_no) in assignment_by_receipt
            for receipt_no in group["receipt_no"].astype("string").fillna("")
        )
        if has_dat_assignment:
            latest_dat = dat_group.iloc[0]
            for tx_index, amount in unassigned.items():
                if _is_round_savings_amount(amount, currency):
                    receipt_no = str(tx.loc[tx_index, "receipt_no"])
                    assignment_by_receipt[receipt_no] = {
                        "dat_customer_id": latest_dat.get("customer_id", pd.NA),
                        "duree": latest_dat.get("product_name", pd.NA),
                        "compte_cree": latest_dat.get("compte_cree", pd.NaT),
                        "maturity_date": latest_dat.get("maturity_date", pd.NaT),
                        "dat_balance": latest_dat.get("balance", np.nan),
                        "dat_match_rule": "telephone + devise + DAT du jour + montant arrondi",
                    }

    for index, receipt_no in assignments["receipt_no"].astype("string").fillna("").items():
        match = assignment_by_receipt.get(str(receipt_no))
        if not match:
            continue
        for column, value in match.items():
            assignments.loc[index, column] = value
    return assignments[columns]


def build_g2_transaction_status_summary(detail: pd.DataFrame | None) -> pd.DataFrame:
    """Compte tous les statuts G2 et indique lesquels alimentent les analyses."""
    columns = [
        "currency_code",
        "fichier_source_g2",
        "statut_transaction_g2",
        "transaction_status_source",
        "nombre_transactions",
        "part_transactions_pct",
        "prise_en_compte_analyse",
    ]
    if not isinstance(detail, pd.DataFrame) or detail.empty:
        return pd.DataFrame(columns=columns)
    frame = detail.copy()
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["fichier_source_g2"] = clean_text(
        frame.get("fichier_source_g2", pd.Series("", index=frame.index))
    ).replace("", "Source non renseignee")
    source_status = clean_text(
        frame.get("transaction_status", pd.Series("", index=frame.index))
    ).replace("", "Non renseigne")
    frame["transaction_status_source"] = source_status
    frame["statut_transaction_g2"] = frame.get(
        "statut_transaction_g2", source_status.apply(normalize_g2_transaction_status)
    )
    frame["incluse_synthese"] = frame.get(
        "incluse_synthese", g2_completed_transaction_mask(frame)
    ).astype("boolean").fillna(False).astype(bool)
    summary = (
        frame.groupby(
            [
                "currency_code",
                "fichier_source_g2",
                "statut_transaction_g2",
                "transaction_status_source",
                "incluse_synthese",
            ],
            as_index=False,
            dropna=False,
        )
        .size()
        .rename(columns={"size": "nombre_transactions"})
    )
    currency_totals = summary.groupby("currency_code")["nombre_transactions"].transform("sum")
    summary["part_transactions_pct"] = (
        summary["nombre_transactions"].div(currency_totals.replace(0, pd.NA)).mul(100)
    )
    summary["prise_en_compte_analyse"] = np.where(
        summary["incluse_synthese"], "Oui", "Non - controle uniquement"
    )
    return (
        summary[columns]
        .sort_values(
            ["currency_code", "fichier_source_g2", "prise_en_compte_analyse", "nombre_transactions"],
            ascending=[True, True, False, False],
        )
        .reset_index(drop=True)
    )


def build_g2_daily_savings_report(prepared: MpesaPreparedData) -> dict[str, pd.DataFrame]:
    g2 = prepared.g2_transactions
    if g2.empty:
        g2 = build_turbo_only_g2_transactions(prepared.transactions)
        if g2.empty:
            return {
                "detail": pd.DataFrame(),
                "anomalies": pd.DataFrame(),
                "synthese": pd.DataFrame(),
                "pivot": pd.DataFrame(),
                "vertical_summary": pd.DataFrame(),
                "comptages": build_entry_count_summary(pd.DataFrame()),
                "statuts": build_g2_transaction_status_summary(pd.DataFrame()),
            }
        prepared = replace(prepared, g2_transactions=g2)

    report = build_g2_dat_crosscheck(prepared)
    if report.empty:
        report = _enrich_g2_with_portal_controls(g2, prepared.transactions)
    report["receipt_no"] = clean_identifier(report.get("receipt_no", pd.Series("", index=report.index)))
    report["date"] = pd.to_datetime(report.get("completion_time"), errors="coerce")
    report["montant"] = numeric_column(report, "transaction_amount_numeric").abs()
    report["sens_flux"] = clean_text(
        report.get("sens_flux", pd.Series("Indetermine", index=report.index))
    )
    report["montant_entree"] = report["montant"].where(report["sens_flux"].eq("Entree"), 0.0)
    report["montant_sortie"] = report["montant"].where(report["sens_flux"].eq("Sortie"), 0.0)
    report["currency_code"] = clean_text(report.get("currency_code", pd.Series("", index=report.index))).str.upper()
    report["phone_prefixe"] = normalize_phone(report.get("phone_prefixe", report.get("phone", pd.Series("", index=report.index))))

    assignments = _build_daily_dat_assignments(report, prepared.fixed_savings)
    if not assignments.empty:
        report = report.merge(assignments, on="receipt_no", how="left")
    else:
        report["dat_customer_id"] = pd.NA
        report["duree"] = pd.NA
        report["compte_cree"] = pd.NaT
        report["maturity_date"] = pd.NaT
        report["dat_balance"] = np.nan
        report["dat_match_rule"] = pd.NA

    has_portal_reference = report.get(
        "cle_rapprochement_turbo",
        report.get("ref_no_portal", pd.Series(pd.NA, index=report.index)),
    ).astype("string").fillna("").ne("")
    heuristic_dat = (
        ~has_portal_reference
        & report["dat_customer_id"].astype("string").fillna("").ne("")
        & report["sens_flux"].eq("Entree")
    )
    report["details_rapport"] = clean_text(
        report.get("categorie_operation", pd.Series("", index=report.index))
    )
    missing_category = report["details_rapport"].eq("")
    if missing_category.any():
        report.loc[missing_category, "details_rapport"] = report.loc[missing_category].apply(
            lambda row: classify_g2_business_operation(row, dat_matched=False), axis=1
        )
    report.loc[heuristic_dat, "details_rapport"] = "DAT"
    is_dat = report["details_rapport"].eq("DAT") & report["sens_flux"].eq("Entree")
    report["duree"] = report["duree"].where(is_dat, "-")
    report["compte_cree_dat"] = pd.to_datetime(report["compte_cree"], errors="coerce")

    customers = prepared.customers
    if not customers.empty and {"msisdn1", "created_at"}.issubset(customers.columns):
        customer_accounts = customers.copy()
        customer_accounts["phone_prefixe"] = normalize_phone(customer_accounts["msisdn1"])
        customer_accounts["compte_cree_client"] = pd.to_datetime(customer_accounts["created_at"], errors="coerce")
        customer_summary = (
            customer_accounts.dropna(subset=["phone_prefixe", "compte_cree_client"])
            .sort_values("compte_cree_client")
            .groupby("phone_prefixe", as_index=False, dropna=False)
            .agg(compte_cree_client=("compte_cree_client", "min"))
        )
        report = report.merge(customer_summary, on="phone_prefixe", how="left")
    else:
        report["compte_cree_client"] = pd.NaT

    current = prepared.current_savings
    if not current.empty and {"msisdn", "currency_code", "created_at"}.issubset(current.columns):
        current_accounts = current.copy()
        current_accounts["phone_prefixe"] = normalize_phone(current_accounts["msisdn"])
        current_accounts["currency_code"] = clean_text(current_accounts["currency_code"]).str.upper()
        current_accounts["compte_cree_epargne_courante"] = pd.to_datetime(current_accounts["created_at"], errors="coerce")
        current_summary = (
            current_accounts.dropna(subset=["phone_prefixe", "currency_code", "compte_cree_epargne_courante"])
            .sort_values("compte_cree_epargne_courante")
            .groupby(["phone_prefixe", "currency_code"], as_index=False, dropna=False)
            .agg(
                current_customer_id=("customer_id", first_non_empty),
                compte_cree_epargne_courante=("compte_cree_epargne_courante", "min"),
            )
        )
        report = report.merge(current_summary, on=["phone_prefixe", "currency_code"], how="left")
    else:
        report["current_customer_id"] = pd.NA
        report["compte_cree_epargne_courante"] = pd.NaT
    report["compte_cree"] = (
        pd.to_datetime(report["compte_cree_client"], errors="coerce")
        .combine_first(pd.to_datetime(report["compte_cree_epargne_courante"], errors="coerce"))
        .combine_first(report["compte_cree_dat"])
    )

    has_same_day_dat_phone = pd.Series(False, index=report.index)
    fixed = prepared.fixed_savings
    if not fixed.empty and "msisdn" in fixed.columns:
        dat = fixed.copy()
        dat["phone_prefixe"] = normalize_phone(dat["msisdn"])
        dat["currency_code"] = clean_text(dat.get("currency_code", pd.Series("", index=dat.index))).str.upper()
        dat["compte_cree"] = pd.to_datetime(
            dat.get("created_at", dat.get("date_approved", pd.Series(pd.NaT, index=dat.index))),
            errors="coerce",
        )
        dat_keys = set(
            dat.dropna(subset=["phone_prefixe", "currency_code", "compte_cree"])
            .assign(jour=lambda frame: frame["compte_cree"].dt.date)
            [["phone_prefixe", "currency_code", "jour"]]
            .astype("string")
            .itertuples(index=False, name=None)
        )
        tx_keys = (
            report.assign(jour=report["date"].dt.date)[["phone_prefixe", "currency_code", "jour"]]
            .astype("string")
            .itertuples(index=False, name=None)
        )
        has_same_day_dat_phone = pd.Series([key in dat_keys for key in tx_keys], index=report.index)

    unusual_same_day_dat_amount = (
        ~has_portal_reference
        & report["sens_flux"].eq("Entree")
        & report["details_rapport"].eq("Depot normal")
        & ~is_dat
        & has_same_day_dat_phone
        & ~report.apply(lambda row: _is_round_savings_amount(row.get("montant"), row.get("currency_code")), axis=1)
    )
    report.loc[unusual_same_day_dat_amount, "details_rapport"] = "Remboursement prets"

    detail_columns = [
        "date",
        "receipt_no",
        "fichier_source_analyse",
        "fichier_source_g2",
        "sens_flux",
        "details_rapport",
        "reason_type",
        "details",
        "opposite_party",
        "Nom_client",
        "phone_prefixe",
        "duree",
        "compte_cree",
        "montant",
        "montant_entree",
        "montant_sortie",
        "currency_code",
        "current_customer_id",
        "dat_customer_id",
        "compte_cree_client",
        "compte_cree_epargne_courante",
        "compte_cree_dat",
        "maturity_date",
        "dat_balance",
        "dat_match_rule",
        "transaction_status",
        "statut_transaction_g2",
        "traitement_statut_g2",
        "transaction_amount_source",
        "paid_in_numeric",
        "withdrawn_numeric",
        "balance_numeric",
        "ref_no_portal",
        "reference_sortie_turbo",
        "cle_sortie_turbo",
        "cle_rapprochement_turbo",
        "methode_rapprochement_turbo",
        "nombre_candidats_sortie_turbo",
        "operation_turbo_confirmee",
        "customer_id_portal",
        "nombre_ecritures_portal",
        "account_types_portal",
        "descriptions_portal",
        "account_type_cible",
        "statut_rapprochement",
        "controle_telephone",
        "controle_devise",
        "montant_portal_controle",
        "ecart_montant",
        "controle_montant",
        "date_creation_g2",
        "source_date_creation_g2",
        "date_creation_turbo",
        "ecart_creation_minutes",
        "controle_date_creation",
        "date_finalisation_g2",
        "delai_traitement_g2_minutes",
        "ecart_finalisation_minutes",
        "controle_date_finalisation",
        "ecart_date_minutes",
        "controle_date",
        "Observation",
        "nombre_lignes_g2_reference",
        "devises_g2_reference",
        "statuts_g2_reference",
        "montants_g2_reference",
        "doublon_receipt_no",
        "incluse_synthese",
        "motif_anomalie",
        "est_anomalie",
        "source_analytique",
        "identifiant_lot",
    ]
    detail_columns = [column for column in detail_columns if column in report.columns]
    detail = report[detail_columns].sort_values(["currency_code", "date"], ascending=[True, False]).reset_index(drop=True)

    summary_detail = detail.loc[
        detail.get("incluse_synthese", pd.Series(True, index=detail.index)).fillna(False).astype(bool)
    ].copy()
    synthese = (
        summary_detail.groupby(["currency_code", "sens_flux", "details_rapport"], as_index=False, dropna=False)
        .agg(nombre=("receipt_no", "count"), montant=("montant", "sum"))
        .sort_values(["currency_code", "sens_flux", "details_rapport"])
    )
    direction_totals = (
        summary_detail.groupby(["currency_code", "sens_flux"], as_index=False, dropna=False)
        .agg(nombre=("receipt_no", "count"), montant=("montant", "sum"))
    )
    direction_totals["details_rapport"] = "Total " + direction_totals["sens_flux"].astype("string").fillna("")
    totals = (
        summary_detail.groupby("currency_code", as_index=False, dropna=False)
        .agg(nombre=("receipt_no", "count"), montant=("montant", "sum"))
    )
    totals["sens_flux"] = "Tous flux"
    totals["details_rapport"] = "Total " + totals["currency_code"].astype("string").fillna("")
    synthese = concat_frames_stable(
        [synthese, direction_totals[synthese.columns], totals[synthese.columns]]
    ).reset_index(drop=True)
    return {
        "detail": detail,
        "anomalies": detail.loc[
            detail.get("est_anomalie", pd.Series(False, index=detail.index)).fillna(False).astype(bool)
        ].reset_index(drop=True),
        "synthese": synthese,
        "pivot": build_entry_pivot(summary_detail),
        "vertical_summary": build_entry_vertical_summary(summary_detail),
        "comptages": build_entry_count_summary(summary_detail),
        "statuts": build_g2_transaction_status_summary(detail),
    }


def build_g2_transaction_time_analysis(daily_detail: pd.DataFrame | None) -> dict[str, pd.DataFrame]:
    """Compte les transactions G2 terminees par date et par heure.

    ``daily_detail`` est le detail canonique produit par
    :func:`build_g2_daily_savings_report`. Les lignes non terminees sont
    exclues lorsque ``incluse_synthese`` est disponible, afin que les volumes
    affiches restent coherents avec la synthese G2/DAT.
    """
    daily_columns = ["date_transaction", "currency_code", "sens_flux", "nombre_transactions"]
    weekday_columns = [
        "jour_semaine_num",
        "jour_semaine",
        "currency_code",
        "sens_flux",
        "nombre_transactions",
    ]
    hourly_columns = ["heure_num", "heure", "currency_code", "sens_flux", "nombre_transactions"]
    day_hour_columns = [
        "date_transaction",
        "heure_num",
        "heure",
        "currency_code",
        "sens_flux",
        "nombre_transactions",
    ]
    empty_result = {
        "par_jour": pd.DataFrame(columns=daily_columns),
        "par_jour_semaine": pd.DataFrame(columns=weekday_columns),
        "par_heure": pd.DataFrame(columns=hourly_columns),
        "jour_heure": pd.DataFrame(columns=day_hour_columns),
    }
    if not isinstance(daily_detail, pd.DataFrame) or daily_detail.empty or "date" not in daily_detail.columns:
        return empty_result

    work = daily_detail.copy()
    if "incluse_synthese" in work.columns:
        eligible = work["incluse_synthese"].astype("boolean").fillna(False).astype(bool)
        work = work.loc[eligible].copy()
    work["date_transaction_complete"] = pd.to_datetime(work["date"], errors="coerce")
    work = work.dropna(subset=["date_transaction_complete"])
    if work.empty:
        return empty_result

    # Le detail journalier est deja canonique. Cette protection evite toutefois
    # de recompter un meme Receipt No. si la fonction est reutilisee ailleurs.
    if "receipt_no" in work.columns:
        receipt = clean_identifier(work["receipt_no"])
        with_receipt = work.loc[receipt.ne("")].copy()
        without_receipt = work.loc[receipt.eq("")].copy()
        work = concat_frames_stable(
            [with_receipt.drop_duplicates(subset=["receipt_no"], keep="first"), without_receipt]
        ).reset_index(drop=True)

    work["date_transaction"] = work["date_transaction_complete"].dt.normalize()
    work["jour_semaine_num"] = work["date_transaction_complete"].dt.dayofweek.astype(int)
    work["heure_num"] = work["date_transaction_complete"].dt.hour.astype(int)
    work["heure"] = work["heure_num"].map(lambda value: f"{int(value):02d}h")
    work["currency_code"] = clean_text(
        work.get("currency_code", pd.Series("", index=work.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    work["sens_flux"] = clean_text(
        work.get("sens_flux", pd.Series("", index=work.index))
    ).replace("", "Indetermine")

    observed_combinations = work[["currency_code", "sens_flux"]].drop_duplicates().reset_index(drop=True)

    daily_counts = (
        work.groupby(["date_transaction", "currency_code", "sens_flux"], as_index=False, dropna=False)
        .size()
        .rename(columns={"size": "nombre_transactions"})
    )
    dates = pd.DataFrame(
        {"date_transaction": pd.date_range(work["date_transaction"].min(), work["date_transaction"].max(), freq="D")}
    )
    daily_grid = observed_combinations.merge(dates, how="cross")
    par_jour = (
        daily_grid.merge(
            daily_counts,
            on=["date_transaction", "currency_code", "sens_flux"],
            how="left",
        )
        .fillna({"nombre_transactions": 0})
        .sort_values(["date_transaction", "currency_code", "sens_flux"])
        .reset_index(drop=True)
    )
    par_jour["nombre_transactions"] = par_jour["nombre_transactions"].astype(int)
    par_jour = par_jour[daily_columns]

    weekday_names = {
        0: "Lundi",
        1: "Mardi",
        2: "Mercredi",
        3: "Jeudi",
        4: "Vendredi",
        5: "Samedi",
        6: "Dimanche",
    }
    weekday_counts = (
        work.groupby(["jour_semaine_num", "currency_code", "sens_flux"], as_index=False, dropna=False)
        .size()
        .rename(columns={"size": "nombre_transactions"})
    )
    weekdays = pd.DataFrame({"jour_semaine_num": range(7)})
    weekday_grid = observed_combinations.merge(weekdays, how="cross")
    par_jour_semaine = (
        weekday_grid.merge(
            weekday_counts,
            on=["jour_semaine_num", "currency_code", "sens_flux"],
            how="left",
        )
        .fillna({"nombre_transactions": 0})
        .sort_values(["jour_semaine_num", "currency_code", "sens_flux"])
        .reset_index(drop=True)
    )
    par_jour_semaine["nombre_transactions"] = par_jour_semaine["nombre_transactions"].astype(int)
    par_jour_semaine["jour_semaine"] = par_jour_semaine["jour_semaine_num"].map(weekday_names)
    par_jour_semaine = par_jour_semaine[weekday_columns]

    hourly_counts = (
        work.groupby(["heure_num", "currency_code", "sens_flux"], as_index=False, dropna=False)
        .size()
        .rename(columns={"size": "nombre_transactions"})
    )
    hours = pd.DataFrame({"heure_num": range(24)})
    hourly_grid = observed_combinations.merge(hours, how="cross")
    par_heure = (
        hourly_grid.merge(
            hourly_counts,
            on=["heure_num", "currency_code", "sens_flux"],
            how="left",
        )
        .fillna({"nombre_transactions": 0})
        .sort_values(["heure_num", "currency_code", "sens_flux"])
        .reset_index(drop=True)
    )
    par_heure["nombre_transactions"] = par_heure["nombre_transactions"].astype(int)
    par_heure["heure"] = par_heure["heure_num"].map(lambda value: f"{int(value):02d}h")
    par_heure = par_heure[hourly_columns]

    jour_heure = (
        work.groupby(
            ["date_transaction", "heure_num", "heure", "currency_code", "sens_flux"],
            as_index=False,
            dropna=False,
        )
        .size()
        .rename(columns={"size": "nombre_transactions"})
        .sort_values(["date_transaction", "heure_num", "currency_code", "sens_flux"])
        .reset_index(drop=True)
    )
    jour_heure["nombre_transactions"] = jour_heure["nombre_transactions"].astype(int)
    return {
        "par_jour": par_jour,
        "par_jour_semaine": par_jour_semaine,
        "par_heure": par_heure,
        "jour_heure": jour_heure[day_hour_columns],
    }


def build_g2_retention_report(
    prepared: MpesaPreparedData,
    daily_detail: pd.DataFrame | None = None,
) -> dict[str, pd.DataFrame]:
    """Mesure le retour des clients G2 actifs par mois, devise et operation.

    La retention M+1 exige une activite pendant le mois civil suivant. La
    retention a 90 jours exige au moins une activite dans les 90 jours suivant
    la fin du mois de base. Les taux restent vides tant que la fenetre complete
    n'est pas observable dans le perimetre analyse.
    """
    empty_result = {
        "mensuelle": pd.DataFrame(),
        "operations": pd.DataFrame(),
        "detail_clients": pd.DataFrame(),
        "definitions": pd.DataFrame(
            [
                {
                    "indicateur": "Retention M+1",
                    "definition": "Clients actifs pendant le mois de base et actifs de nouveau pendant le mois civil suivant.",
                    "denominateur": "Telephones clients distincts actifs pendant le mois de base, par devise.",
                    "numerateur": "Clients du denominateur avec au moins une operation eligible le mois suivant.",
                },
                {
                    "indicateur": "Retention 90 jours",
                    "definition": "Clients actifs pendant le mois de base et actifs de nouveau dans les 90 jours suivant la fin de ce mois.",
                    "denominateur": "Telephones clients distincts actifs pendant le mois de base, par devise.",
                    "numerateur": "Clients du denominateur avec au moins une operation eligible avant l'echeance des 90 jours.",
                },
                {
                    "indicateur": "Population eligible",
                    "definition": "Operations G2 avec Completion Time et telephone valides, hors operations internes et statuts explicitement en echec/annules/inverses.",
                    "denominateur": "Toutes les operations du perimetre filtre.",
                    "numerateur": "Operations respectant les criteres d'eligibilite.",
                },
            ]
        ),
    }
    if prepared.g2_transactions.empty and not (
        isinstance(daily_detail, pd.DataFrame) and not daily_detail.empty
    ):
        return empty_result

    activity = (
        daily_detail.copy()
        if isinstance(daily_detail, pd.DataFrame)
        else build_g2_daily_savings_report(prepared).get("detail", pd.DataFrame()).copy()
    )
    if activity.empty:
        return empty_result
    if "incluse_synthese" in activity.columns:
        eligible_activity = (
            activity["incluse_synthese"]
            .astype("boolean")
            .fillna(False)
            .astype(bool)
        )
        activity = activity.loc[eligible_activity].copy()
        if activity.empty:
            return empty_result

    activity["date"] = pd.to_datetime(activity.get("date"), errors="coerce")
    activity["phone_prefixe"] = normalize_phone(
        activity.get("phone_prefixe", pd.Series("", index=activity.index))
    )
    activity["currency_code"] = clean_text(
        activity.get("currency_code", pd.Series("", index=activity.index))
    ).str.upper()
    activity["details_rapport"] = clean_text(
        activity.get("details_rapport", pd.Series("Non classe", index=activity.index))
    ).replace("", "Non classe")
    activity["transaction_status"] = clean_text(
        activity.get("transaction_status", pd.Series("", index=activity.index))
    )
    rejected_status = activity["transaction_status"].apply(normalize_label).str.contains(
        r"fail|echec|cancel|annul|reverse|revers|reject|rejet",
        regex=True,
        na=False,
    )
    internal_operation = activity["details_rapport"].apply(normalize_label).str.contains(
        "operation interne",
        regex=False,
        na=False,
    )
    activity = activity.loc[
        activity["date"].notna()
        & activity["phone_prefixe"].notna()
        & activity["currency_code"].ne("")
        & ~rejected_status
        & ~internal_operation
    ].copy()
    if activity.empty:
        return empty_result

    activity["receipt_no"] = clean_identifier(
        activity.get("receipt_no", pd.Series("", index=activity.index))
    )
    activity["__transaction_key"] = activity["receipt_no"].where(
        activity["receipt_no"].ne(""),
        "ligne_" + activity.index.astype(str),
    )
    activity = activity.drop_duplicates(
        subset=["currency_code", "__transaction_key", "phone_prefixe", "date"]
    )
    activity["periode"] = activity["date"].dt.to_period("M").dt.to_timestamp()
    observation_start = activity["date"].min()
    observation_end = activity["date"].max()

    rows: list[dict[str, object]] = []
    client_activity = {
        (currency, phone): group.sort_values("date")
        for (currency, phone), group in activity.groupby(
            ["currency_code", "phone_prefixe"], dropna=False, sort=False
        )
    }
    base_groups = activity.groupby(["periode", "currency_code", "phone_prefixe"], dropna=False)
    for (period, currency, phone), base in base_groups:
        # Les cles issues d'un Period peuvent conserver l'unite NumPy generique.
        # Forcer les nanosecondes evite l'addition de timedelta sans unite explicite.
        period = pd.Timestamp(period).as_unit("ns")
        month_end = period + pd.offsets.MonthEnd(1)
        next_month_start = period + pd.offsets.MonthBegin(1)
        next_month_end = next_month_start + pd.offsets.MonthEnd(1)
        deadline_90 = _timestamp_plus(month_end, days=90)
        eligible_m1 = observation_end >= next_month_end
        eligible_90 = observation_end >= deadline_90
        client_history = client_activity[(currency, phone)]
        future = client_history.loc[client_history["date"].gt(month_end)]
        first_return = future["date"].min() if not future.empty else pd.NaT
        returned_m1 = bool(
            not future.empty
            and future["date"].between(next_month_start, next_month_end, inclusive="both").any()
        )
        returned_90 = bool(
            not future.empty
            and future["date"].le(deadline_90).any()
        )
        rows.append(
            {
                "periode": period,
                "annee": int(period.year),
                "mois": period.strftime("%Y-%m"),
                "currency_code": currency,
                "phone_prefixe": phone,
                "Nom_client": concat_unique(base.get("Nom_client", pd.Series(dtype="string"))),
                "types_operations": concat_unique(base["details_rapport"]),
                "nombre_operations_mois_base": int(len(base)),
                "montant_entrees_mois_base": float(numeric_column(base, "montant_entree").sum()),
                "montant_sorties_mois_base": float(numeric_column(base, "montant_sortie").sum()),
                "premier_retour": first_return,
                "delai_premier_retour_jours": (
                    int((first_return.normalize() - month_end.normalize()).days)
                    if pd.notna(first_return)
                    else pd.NA
                ),
                "eligible_retention_m1": bool(eligible_m1),
                "retenu_m1": returned_m1 if eligible_m1 else pd.NA,
                "eligible_retention_90j": bool(eligible_90),
                "retenu_90j": returned_90 if eligible_90 else pd.NA,
                "debut_observation": observation_start,
                "fin_observation": observation_end,
            }
        )

    detail = pd.DataFrame(rows).sort_values(
        ["periode", "currency_code", "phone_prefixe"], ascending=[False, True, True]
    ).reset_index(drop=True)

    def _summarize(frame: pd.DataFrame, group_columns: list[str]) -> pd.DataFrame:
        summary_rows: list[dict[str, object]] = []
        for keys, group in frame.groupby(group_columns, dropna=False, sort=True):
            key_values = keys if isinstance(keys, tuple) else (keys,)
            result = dict(zip(group_columns, key_values))
            clients_base = int(group["phone_prefixe"].nunique())
            eligible_m1 = bool(group["eligible_retention_m1"].all())
            eligible_90 = bool(group["eligible_retention_90j"].all())
            retained_m1 = int(group["retenu_m1"].eq(True).sum()) if eligible_m1 else pd.NA
            retained_90 = int(group["retenu_90j"].eq(True).sum()) if eligible_90 else pd.NA
            result.update(
                {
                    "clients_actifs_mois_base": clients_base,
                    "clients_retenus_m1": retained_m1,
                    "retention_m1_pct": round(100 * retained_m1 / clients_base, 2) if eligible_m1 and clients_base else np.nan,
                    "clients_retenus_90j": retained_90,
                    "retention_90j_pct": round(100 * retained_90 / clients_base, 2) if eligible_90 and clients_base else np.nan,
                    "eligible_retention_m1": eligible_m1,
                    "eligible_retention_90j": eligible_90,
                    "debut_observation": observation_start,
                    "fin_observation": observation_end,
                }
            )
            summary_rows.append(result)
        return pd.DataFrame(summary_rows)

    monthly = _summarize(detail, ["periode", "annee", "mois", "currency_code"])
    monthly = monthly.sort_values(["currency_code", "periode"]).reset_index(drop=True)

    operation_detail = detail.assign(
        details_rapport=detail["types_operations"].str.split(r"\s*\|\s*", regex=True)
    ).explode("details_rapport")
    operation_detail["details_rapport"] = clean_text(operation_detail["details_rapport"]).replace("", "Non classe")
    operations = _summarize(
        operation_detail,
        ["periode", "annee", "mois", "currency_code", "details_rapport"],
    )
    operations = operations.sort_values(
        ["currency_code", "details_rapport", "periode"]
    ).reset_index(drop=True)

    return {
        "mensuelle": monthly,
        "operations": operations,
        "detail_clients": detail,
        "definitions": empty_result["definitions"],
    }


def _mpesa_analysis_date(prepared: MpesaPreparedData, as_of_date: Any | None = None) -> pd.Timestamp:
    if as_of_date is not None:
        parsed = pd.to_datetime(as_of_date, errors="coerce")
        if pd.notna(parsed):
            return pd.Timestamp(parsed).normalize()

    candidates: list[pd.Series] = []
    source_columns = [
        (prepared.g2_transactions, "completion_time"),
        (prepared.transactions, "created_at"),
        (prepared.loans, "updated_at"),
        (prepared.loans, "created_at"),
        (prepared.fixed_savings, "updated_at"),
        (prepared.fixed_savings, "date_approved"),
        (prepared.current_savings, "updated_at"),
        (prepared.current_savings, "created_at"),
    ]
    for frame, column in source_columns:
        if isinstance(frame, pd.DataFrame) and not frame.empty and column in frame.columns:
            values = pd.to_datetime(frame[column], errors="coerce").dropna()
            if not values.empty:
                candidates.append(values)
    if candidates:
        return pd.Timestamp(pd.concat(candidates, ignore_index=True).max()).normalize()
    return pd.Timestamp.now().normalize()


def build_mpesa_credit_risk_analysis(
    loans: pd.DataFrame | None,
    *,
    as_of_date: Any | None = None,
) -> dict[str, pd.DataFrame]:
    """Construit les indicateurs de remboursement et de retard par devise."""
    empty = {"synthese": pd.DataFrame(), "detail": pd.DataFrame()}
    if not isinstance(loans, pd.DataFrame) or loans.empty or "loan_id" not in loans.columns:
        return empty

    analysis_date = pd.Timestamp(
        pd.to_datetime(as_of_date, errors="coerce") if as_of_date is not None else pd.Timestamp.now()
    ).normalize()
    frame = loans.copy()
    frame["loan_id"] = clean_identifier(frame["loan_id"])
    frame["customer_id"] = clean_identifier(frame.get("customer_id", pd.Series("", index=frame.index)))
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    for column in ["created_at", "updated_at", "due_date", "last_repayment_date"]:
        frame[column] = pd.to_datetime(frame.get(column, pd.Series(pd.NaT, index=frame.index)), errors="coerce")
    sort_columns = [column for column in ["updated_at", "created_at"] if column in frame.columns]
    if sort_columns:
        frame = frame.sort_values(sort_columns, ascending=False, na_position="last")
    frame = frame.loc[frame["loan_id"].ne("")].drop_duplicates("loan_id", keep="first").copy()
    if frame.empty:
        return empty

    numeric_sources = [
        "loan_amount",
        "loan_balance",
        "amount_paid",
        "outstanding_principle",
        "outstanding_setup_fees",
        "outstanding_interest",
        "outstanding_penalty_fees",
    ]
    source_presence = {column: column in loans.columns for column in numeric_sources}
    for column in numeric_sources:
        frame[column] = numeric_column(frame, column).clip(lower=0)
    computed_outstanding = (
        frame["outstanding_principle"]
        + frame["outstanding_setup_fees"]
        + frame["outstanding_interest"]
        + frame["outstanding_penalty_fees"]
    )
    frame["encours_total"] = frame["loan_balance"] if source_presence["loan_balance"] else computed_outstanding
    encours_reference = computed_outstanding
    frame["ecart_encours_sources"] = encours_reference - frame["loan_balance"]
    frame["encours_sources_incoherents"] = (
        source_presence["loan_balance"]
        and any(
            source_presence[column]
            for column in [
                "outstanding_principle",
                "outstanding_setup_fees",
                "outstanding_interest",
                "outstanding_penalty_fees",
            ]
        )
    ) & frame["ecart_encours_sources"].abs().gt(
        np.maximum(1.0, frame[["loan_balance"]].abs().max(axis=1) * 0.50)
    )
    has_outstanding_source = source_presence["loan_balance"] or any(
        source_presence[column]
        for column in ["outstanding_principle", "outstanding_interest", "outstanding_penalty_fees"]
    )
    frame["donnee_encours_disponible"] = has_outstanding_source
    frame["donnee_echeance_disponible"] = frame["due_date"].notna()
    overdue_days = (analysis_date - frame["due_date"].dt.normalize()).dt.days
    frame["jours_retard"] = overdue_days.where(
        frame["encours_total"].gt(0) & overdue_days.gt(0), 0
    ).fillna(0).astype(int)
    frame["encours_retard_1j"] = frame["encours_total"].where(frame["jours_retard"].ge(1), 0.0)
    frame["encours_retard_7j"] = frame["encours_total"].where(frame["jours_retard"].ge(7), 0.0)
    frame["encours_retard_30j"] = frame["encours_total"].where(frame["jours_retard"].ge(30), 0.0)
    frame["encours_sans_echeance"] = frame["encours_total"].where(frame["due_date"].isna(), 0.0)
    frame["statut_risque"] = np.select(
        [
            ~frame["donnee_encours_disponible"],
            frame["encours_total"].le(0),
            ~frame["donnee_echeance_disponible"],
            frame["jours_retard"].ge(30),
            frame["jours_retard"].ge(7),
            frame["jours_retard"].ge(1),
            frame["due_date"].dt.normalize().eq(analysis_date),
        ],
        [
            "Encours non renseigne",
            "Solde nul / rembourse",
            "Echeance non renseignee",
            "En retard 30 jours et plus",
            "En retard 7 a 29 jours",
            "En retard 1 a 6 jours",
            "Echu aujourd'hui",
        ],
        default="A jour",
    )
    frame["date_analyse"] = analysis_date
    for column in ["Nom_client", "customer", "msisdn1", "status_name"]:
        if column not in frame.columns:
            frame[column] = pd.NA

    summary = (
        frame.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            nombre_credits=("loan_id", "nunique"),
            nombre_clients=("customer_id", lambda values: clean_identifier(values).replace("", pd.NA).nunique()),
            montant_credits=("loan_amount", "sum"),
            montant_rembourse=("amount_paid", "sum"),
            encours_total=("encours_total", "sum"),
            encours_retard_1j=("encours_retard_1j", "sum"),
            encours_retard_7j=("encours_retard_7j", "sum"),
            encours_retard_30j=("encours_retard_30j", "sum"),
            encours_sans_echeance=("encours_sans_echeance", "sum"),
            credits_retard_1j=("jours_retard", lambda values: int(pd.Series(values).ge(1).sum())),
            credits_retard_30j=("jours_retard", lambda values: int(pd.Series(values).ge(30).sum())),
            echeances_renseignees=("donnee_echeance_disponible", "sum"),
            incoherences_encours=("encours_sources_incoherents", "sum"),
        )
    )
    denominator = summary["encours_total"].replace(0, pd.NA)
    summary["par_1j_pct"] = summary["encours_retard_1j"].div(denominator).mul(100)
    summary["par_7j_pct"] = summary["encours_retard_7j"].div(denominator).mul(100)
    summary["par_30j_pct"] = summary["encours_retard_30j"].div(denominator).mul(100)
    if source_presence["amount_paid"] and source_presence["loan_amount"]:
        summary["taux_remboursement_pct"] = summary["montant_rembourse"].div(
            summary["montant_credits"].replace(0, pd.NA)
        ).mul(100)
    else:
        summary["taux_remboursement_pct"] = np.nan
    if not has_outstanding_source:
        summary[["par_1j_pct", "par_7j_pct", "par_30j_pct"]] = np.nan
    summary.loc[
        summary["encours_sans_echeance"].gt(0),
        ["par_1j_pct", "par_7j_pct", "par_30j_pct"],
    ] = np.nan
    summary["date_analyse"] = analysis_date

    detail_columns = [
        "loan_id", "customer_id", "Nom_client", "customer", "msisdn1", "currency_code",
        "loan_amount", "amount_paid", "encours_total", "outstanding_principle", "outstanding_setup_fees",
        "outstanding_interest", "outstanding_penalty_fees", "status_name", "created_at",
        "due_date", "last_repayment_date", "jours_retard", "statut_risque", "date_analyse",
        "ecart_encours_sources", "encours_sources_incoherents",
    ]
    detail = frame[detail_columns].sort_values(
        ["currency_code", "jours_retard", "encours_total"], ascending=[True, False, False]
    ).reset_index(drop=True)
    return {"synthese": summary, "detail": detail}


def build_loan_savings_reconciliation(
    loans: pd.DataFrame | None,
    current_savings: pd.DataFrame | None,
    fixed_savings: pd.DataFrame | None = None,
) -> dict[str, pd.DataFrame]:
    """Rapproche les credits avec l'epargne au grain client x devise.

    ``savings_account_id`` fournit une liaison directe lorsqu'il correspond a
    ``id`` ou ``savings_id`` du compte courant. A defaut, une correspondance
    est deduite uniquement si ``customer_id + currency_code`` identifie un
    compte courant unique. La vue consolidee juxtapose les positions sans
    compenser comptablement l'epargne et le credit.
    """
    empty = {
        "synthese": pd.DataFrame(),
        "clients": pd.DataFrame(),
        "detail": pd.DataFrame(),
        "controles": pd.DataFrame(),
        "sources": pd.DataFrame(),
    }
    loans_available = isinstance(loans, pd.DataFrame) and not loans.empty
    current_available = isinstance(current_savings, pd.DataFrame) and not current_savings.empty
    fixed_available = isinstance(fixed_savings, pd.DataFrame) and not fixed_savings.empty
    if not loans_available:
        return empty

    credit = loans.copy()
    credit["loan_id"] = clean_identifier(
        credit.get("loan_id", pd.Series("", index=credit.index))
    )
    credit["customer_id"] = clean_identifier(
        credit.get("customer_id", pd.Series("", index=credit.index))
    )
    credit["currency_code"] = clean_text(
        credit.get("currency_code", pd.Series("", index=credit.index))
    ).str.upper()
    credit["telephone_credit"] = normalize_phone(
        credit.get("msisdn1", pd.Series("", index=credit.index))
    )
    credit["savings_account_id_source"] = clean_identifier(
        credit.get("savings_account_id", pd.Series("", index=credit.index))
    )
    for column in ["created_at", "updated_at"]:
        credit[column] = pd.to_datetime(
            credit.get(column, pd.Series(pd.NaT, index=credit.index)), errors="coerce"
        )
    sort_columns = [column for column in ["updated_at", "created_at"] if column in credit.columns]
    if sort_columns:
        credit = credit.sort_values(sort_columns, na_position="first")
    credit = credit.loc[credit["loan_id"].ne("")].drop_duplicates("loan_id", keep="last").copy()
    if credit.empty:
        return empty

    monetary_columns = [
        "loan_amount",
        "loan_balance",
        "amount_paid",
        "outstanding_principle",
        "outstanding_setup_fees",
        "outstanding_interest",
        "outstanding_penalty_fees",
    ]
    source_presence = {column: column in loans.columns for column in monetary_columns}
    for column in monetary_columns:
        credit[column] = numeric_column(credit, column).clip(lower=0)
    credit["encours_calcule"] = (
        credit["outstanding_principle"]
        + credit["outstanding_setup_fees"]
        + credit["outstanding_interest"]
        + credit["outstanding_penalty_fees"]
    )
    credit["encours_credit"] = (
        credit["loan_balance"]
        if source_presence["loan_balance"]
        else credit["encours_calcule"]
    )
    for column in ["Nom_client", "customer", "status_name", "due_date", "last_repayment_date"]:
        if column not in credit.columns:
            credit[column] = pd.NA
    credit["due_date"] = pd.to_datetime(credit["due_date"], errors="coerce")
    credit["last_repayment_date"] = pd.to_datetime(
        credit["last_repayment_date"], errors="coerce"
    )

    current = current_savings.copy() if current_available else pd.DataFrame()
    current_source_complete_available = False
    if current_available:
        current["__compte_row"] = np.arange(len(current))
        current["customer_id"] = clean_identifier(
            current.get("customer_id", pd.Series("", index=current.index))
        )
        current["currency_code"] = clean_text(
            current.get("currency_code", pd.Series("", index=current.index))
        ).str.upper()
        current["id_compte_turbo"] = clean_identifier(
            current.get("id", pd.Series("", index=current.index))
        )
        current["savings_id"] = clean_identifier(
            current.get("savings_id", pd.Series("", index=current.index))
        )
        current["telephone_epargne"] = normalize_phone(
            current.get(
                "msisdn",
                current.get("msisdn1", pd.Series("", index=current.index)),
            )
        )
        current["solde_epargne_courante"] = numeric_column(current, "balance")
        current["source_epargne_complete"] = (
            current.get(
                "source_savings_account_complete",
                pd.Series(False, index=current.index),
            )
            .fillna(False)
            .astype(bool)
        )
        current_source_complete_available = bool(current["source_epargne_complete"].any())

        # Le rapprochement porte uniquement sur les clients de Loans Account.
        # Restreindre ici les quelque 77 000 comptes courants aux paires client /
        # devise utiles (plus les identifiants directs cites par un pret) evite
        # des centaines de milliers d'agregations Python sans perdre les controles
        # de liaison directe, y compris lorsqu'un identifiant pointe vers un autre
        # client ou une autre devise.
        credit_pairs = pd.MultiIndex.from_frame(
            credit[["customer_id", "currency_code"]].drop_duplicates()
        )
        current_pairs = pd.MultiIndex.from_arrays(
            [current["customer_id"], current["currency_code"]]
        )
        pair_needed = current_pairs.isin(credit_pairs)
        direct_identifiers = set(
            credit.loc[
                credit["savings_account_id_source"].ne(""),
                "savings_account_id_source",
            ].astype(str)
        )
        direct_needed = (
            current["id_compte_turbo"].isin(direct_identifiers)
            | current["savings_id"].isin(direct_identifiers)
        )
        current = current.loc[pair_needed | direct_needed].copy()

    valid_current = (
        current.loc[
            current["customer_id"].ne("") & current["currency_code"].ne("")
        ].copy()
        if current_available
        else pd.DataFrame()
    )
    if valid_current.empty:
        fallback_accounts = pd.DataFrame(
            columns=[
                "customer_id",
                "currency_code",
                "nb_comptes_courants_candidats",
                "savings_id_deduit",
                "id_compte_turbo_deduit",
                "telephone_epargne_deduit",
                "solde_epargne_courante_deduit",
                "source_epargne_complete_deduite",
            ]
        )
    else:
        fallback_accounts = (
            valid_current.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
            .agg(
                nb_comptes_courants_candidats=("__compte_row", "nunique"),
                savings_id_deduit=("savings_id", concat_unique),
                id_compte_turbo_deduit=("id_compte_turbo", concat_unique),
                telephone_epargne_deduit=("telephone_epargne", concat_unique),
                solde_epargne_courante_deduit=("solde_epargne_courante", "sum"),
                source_epargne_complete_deduite=("source_epargne_complete", "max"),
            )
        )

    direct_candidates = pd.DataFrame()
    if current_available:
        candidate_frames: list[pd.DataFrame] = []
        direct_columns = [
            "__compte_row",
            "customer_id",
            "currency_code",
            "savings_id",
            "id_compte_turbo",
            "telephone_epargne",
            "solde_epargne_courante",
            "source_epargne_complete",
        ]
        for identifier_column in ["id_compte_turbo", "savings_id"]:
            candidates = current.loc[current[identifier_column].ne(""), direct_columns].copy()
            if candidates.empty:
                continue
            candidates["cle_compte_directe"] = candidates[identifier_column]
            candidate_frames.append(candidates)
        if candidate_frames:
            direct_rows = concat_frames_stable(candidate_frames).drop_duplicates(
                ["cle_compte_directe", "__compte_row"]
            )
            direct_candidates = (
                direct_rows.groupby("cle_compte_directe", as_index=False, dropna=False)
                .agg(
                    nb_correspondances_identifiant=("__compte_row", "nunique"),
                    customer_id_direct=("customer_id", concat_unique),
                    currency_code_direct=("currency_code", concat_unique),
                    savings_id_direct=("savings_id", concat_unique),
                    id_compte_turbo_direct=("id_compte_turbo", concat_unique),
                    telephone_epargne_direct=("telephone_epargne", concat_unique),
                    solde_epargne_courante_direct=("solde_epargne_courante", "sum"),
                    source_epargne_complete_direct=("source_epargne_complete", "max"),
                )
            )

    detail = credit.merge(
        fallback_accounts,
        on=["customer_id", "currency_code"],
        how="left",
    )
    if direct_candidates.empty:
        detail["nb_correspondances_identifiant"] = 0
        for column in [
            "customer_id_direct",
            "currency_code_direct",
            "savings_id_direct",
            "id_compte_turbo_direct",
            "telephone_epargne_direct",
        ]:
            detail[column] = ""
        detail["solde_epargne_courante_direct"] = 0.0
        detail["source_epargne_complete_direct"] = False
    else:
        detail = detail.merge(
            direct_candidates,
            left_on="savings_account_id_source",
            right_on="cle_compte_directe",
            how="left",
        )

    for column in ["nb_comptes_courants_candidats", "nb_correspondances_identifiant"]:
        detail[column] = pd.to_numeric(detail.get(column), errors="coerce").fillna(0).astype(int)
    for column in [
        "savings_id_deduit",
        "id_compte_turbo_deduit",
        "telephone_epargne_deduit",
        "customer_id_direct",
        "currency_code_direct",
        "savings_id_direct",
        "id_compte_turbo_direct",
        "telephone_epargne_direct",
    ]:
        detail[column] = clean_text(detail.get(column, pd.Series("", index=detail.index)))
    for column in ["solde_epargne_courante_deduit", "solde_epargne_courante_direct"]:
        detail[column] = pd.to_numeric(detail.get(column), errors="coerce").fillna(0.0)
    for column in ["source_epargne_complete_deduite", "source_epargne_complete_direct"]:
        detail[column] = (
            detail.get(column, pd.Series(False, index=detail.index))
            .astype("boolean")
            .fillna(False)
            .astype(bool)
        )

    source_id_present = detail["savings_account_id_source"].ne("")
    direct_unique = source_id_present & detail["nb_correspondances_identifiant"].eq(1)
    direct_ambiguous = source_id_present & detail["nb_correspondances_identifiant"].gt(1)
    fallback_unique = detail["nb_comptes_courants_candidats"].eq(1)
    fallback_ambiguous = detail["nb_comptes_courants_candidats"].gt(1)
    use_fallback = ~direct_unique & ~direct_ambiguous & fallback_unique

    detail["methode_rapprochement_epargne"] = np.select(
        [
            direct_unique,
            direct_ambiguous,
            use_fallback & source_id_present,
            use_fallback,
            fallback_ambiguous,
        ],
        [
            "Directe - savings_account_id",
            "Ambigue - savings_account_id",
            "Deduite - identifiant source non retrouve",
            "Deduite - customer_id + devise",
            "Ambigue - customer_id + devise",
        ],
        default="Non rapproche",
    )
    detail["savings_id_correspondant"] = detail["savings_id_direct"].where(
        direct_unique, detail["savings_id_deduit"].where(use_fallback, "")
    )
    detail["id_compte_turbo_correspondant"] = detail["id_compte_turbo_direct"].where(
        direct_unique, detail["id_compte_turbo_deduit"].where(use_fallback, "")
    )
    detail["telephone_epargne"] = detail["telephone_epargne_direct"].where(
        direct_unique, detail["telephone_epargne_deduit"].where(use_fallback, "")
    )
    detail["solde_epargne_courante"] = detail["solde_epargne_courante_direct"].where(
        direct_unique, detail["solde_epargne_courante_deduit"].where(use_fallback, 0.0)
    )
    detail["source_epargne_complete"] = detail["source_epargne_complete_direct"].where(
        direct_unique,
        detail["source_epargne_complete_deduite"].where(use_fallback, False),
    )
    detail["liaison_directe_source"] = direct_unique
    detail["correspondance_deduite"] = use_fallback

    loan_phone = detail["telephone_credit"].astype("string").fillna("")
    savings_phone = detail["telephone_epargne"].astype("string").fillna("")
    phone_comparable = loan_phone.ne("") & savings_phone.ne("") & (direct_unique | use_fallback)
    phone_mismatch = phone_comparable & loan_phone.ne(savings_phone)
    direct_customer_mismatch = direct_unique & detail["customer_id_direct"].ne(detail["customer_id"])
    direct_currency_mismatch = direct_unique & detail["currency_code_direct"].ne(
        detail["currency_code"]
    )
    source_id_not_found = source_id_present & detail["nb_correspondances_identifiant"].eq(0)

    reasons: list[list[str]] = [[] for _ in range(len(detail))]

    def add_reason(mask: pd.Series, label: str) -> None:
        for position in np.flatnonzero(mask.to_numpy(dtype=bool)):
            reasons[int(position)].append(label)

    if not current_available:
        add_reason(pd.Series(True, index=detail.index), "Savings Account non charge")
    else:
        add_reason(direct_ambiguous, "savings_account_id correspond a plusieurs comptes")
        add_reason(source_id_not_found, "savings_account_id non retrouve dans Savings Account")
        add_reason(
            ~source_id_present & detail["nb_comptes_courants_candidats"].eq(0),
            "Aucun compte courant pour le client et la devise",
        )
        add_reason(
            ~direct_unique & fallback_ambiguous,
            "Plusieurs comptes courants pour le client et la devise",
        )
        add_reason(direct_customer_mismatch, "Ecart de customer_id sur la liaison directe")
        add_reason(direct_currency_mismatch, "Ecart de devise sur la liaison directe")
        add_reason(phone_mismatch, "Ecart de telephone entre credit et epargne")
    detail["motif_controle"] = [" | ".join(items) for items in reasons]
    detail["statut_controle"] = np.select(
        [
            pd.Series(not current_available, index=detail.index),
            detail["motif_controle"].ne(""),
            direct_unique,
            use_fallback,
        ],
        [
            "Non calculable - Savings Account absent",
            "A revoir",
            "Conforme - correspondance directe",
            "Conforme - correspondance deduite",
        ],
        default="A revoir",
    )
    detail["correspondance_reussie"] = detail["statut_controle"].str.startswith("Conforme")

    fixed_positions = pd.DataFrame(
        columns=["customer_id", "currency_code", "nb_dat_positifs", "solde_dat_positif"]
    )
    if fixed_available:
        fixed = fixed_savings.copy()
        fixed["customer_id"] = clean_identifier(
            fixed.get("customer_id", pd.Series("", index=fixed.index))
        )
        fixed["currency_code"] = clean_text(
            fixed.get("currency_code", pd.Series("", index=fixed.index))
        ).str.upper()
        fixed["balance"] = numeric_column(fixed, "balance")
        fixed = fixed.loc[
            fixed["customer_id"].ne("")
            & fixed["currency_code"].ne("")
            & fixed["balance"].gt(0)
        ].copy()
        if not fixed.empty:
            fixed_positions = (
                fixed.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
                .agg(
                    nb_dat_positifs=("balance", "size"),
                    solde_dat_positif=("balance", "sum"),
                )
            )

    def client_control_status(values: pd.Series) -> str:
        statuses = clean_text(values)
        if statuses.eq("Non calculable - Savings Account absent").all():
            return "Non calculable - Savings Account absent"
        if statuses.eq("A revoir").any():
            return "A revoir"
        if statuses.eq("Conforme - correspondance directe").all():
            return "Conforme - correspondance directe"
        if statuses.eq("Conforme - correspondance deduite").all():
            return "Conforme - correspondance deduite"
        return "Conforme - correspondance mixte"

    fallback_client_positions = fallback_accounts[
        [
            "customer_id",
            "currency_code",
            "nb_comptes_courants_candidats",
            "savings_id_deduit",
            "telephone_epargne_deduit",
            "solde_epargne_courante_deduit",
            "source_epargne_complete_deduite",
        ]
    ].rename(
        columns={
            "savings_id_deduit": "savings_id_correspondant",
            "telephone_epargne_deduit": "telephone_epargne",
            "solde_epargne_courante_deduit": "solde_epargne_courante",
            "source_epargne_complete_deduite": "source_epargne_complete",
        }
    )
    clients = (
        detail.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
        .agg(
            Nom_client=("Nom_client", concat_unique),
            customer=("customer", concat_unique),
            telephone_credit=("telephone_credit", concat_unique),
            nombre_credits=("loan_id", "nunique"),
            loan_ids=("loan_id", concat_unique),
            statuts_credit=("status_name", concat_unique),
            montant_credits=("loan_amount", "sum"),
            montant_rembourse=("amount_paid", "sum"),
            encours_credit=("encours_credit", "sum"),
            principal_restant=("outstanding_principle", "sum"),
            frais_mise_en_place_restants=("outstanding_setup_fees", "sum"),
            interets_restants=("outstanding_interest", "sum"),
            penalites_restantes=("outstanding_penalty_fees", "sum"),
            statut_rapprochement=("statut_controle", client_control_status),
            motifs_controle=("motif_controle", concat_unique),
        )
        .merge(
            fallback_client_positions,
            on=["customer_id", "currency_code"],
            how="left",
        )
        .merge(fixed_positions, on=["customer_id", "currency_code"], how="left")
    )
    clients["savings_id_correspondant"] = clean_text(
        clients.get("savings_id_correspondant", pd.Series("", index=clients.index))
    )
    clients["telephone_epargne"] = clean_text(
        clients.get("telephone_epargne", pd.Series("", index=clients.index))
    )
    clients["nb_comptes_courants_candidats"] = pd.to_numeric(
        clients.get("nb_comptes_courants_candidats"), errors="coerce"
    ).fillna(0).astype(int)
    clients["solde_epargne_courante"] = pd.to_numeric(
        clients.get("solde_epargne_courante"), errors="coerce"
    ).fillna(0.0)
    clients["source_epargne_complete"] = (
        clients.get("source_epargne_complete", pd.Series(False, index=clients.index))
        .astype("boolean")
        .fillna(False)
        .astype(bool)
    )
    clients["nb_dat_positifs"] = pd.to_numeric(
        clients.get("nb_dat_positifs"), errors="coerce"
    ).fillna(0).astype(int)
    clients["solde_dat_positif"] = pd.to_numeric(
        clients.get("solde_dat_positif"), errors="coerce"
    ).fillna(0.0)
    clients["epargne_totale_observee"] = (
        clients["solde_epargne_courante"] + clients["solde_dat_positif"]
    )
    clients["interpretation_position"] = (
        "Montants juxtaposes - aucune compensation comptable credit/epargne"
    )
    clients = clients.sort_values(
        ["currency_code", "encours_credit", "customer_id"],
        ascending=[True, False, True],
    ).reset_index(drop=True)

    summary_credit = (
        detail.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            nombre_credits=("loan_id", "nunique"),
            nombre_clients=("customer_id", lambda values: clean_identifier(values).replace("", pd.NA).nunique()),
            montant_credits=("loan_amount", "sum"),
            montant_rembourse=("amount_paid", "sum"),
            encours_credit=("encours_credit", "sum"),
            credits_rapproches=("correspondance_reussie", "sum"),
            liaisons_directes=("liaison_directe_source", "sum"),
            liaisons_deduites=("correspondance_deduite", "sum"),
            credits_a_revoir=("statut_controle", lambda values: int(clean_text(values).eq("A revoir").sum())),
            credits_non_calculables=(
                "statut_controle",
                lambda values: int(clean_text(values).str.startswith("Non calculable").sum()),
            ),
            savings_account_id_renseignes=(
                "savings_account_id_source",
                lambda values: int(clean_identifier(values).ne("").sum()),
            ),
        )
    )
    summary_positions = (
        clients.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            solde_epargne_courante_clients_credit=("solde_epargne_courante", "sum"),
            solde_dat_clients_credit=("solde_dat_positif", "sum"),
            epargne_totale_clients_credit=("epargne_totale_observee", "sum"),
            clients_a_revoir=("statut_rapprochement", lambda values: int(clean_text(values).eq("A revoir").sum())),
        )
    )
    summary = summary_credit.merge(summary_positions, on="currency_code", how="left")
    summary["taux_rapprochement_pct"] = summary["credits_rapproches"].div(
        summary["nombre_credits"].replace(0, pd.NA)
    ).mul(100)
    if not current_available:
        summary["taux_rapprochement_pct"] = np.nan
    summary["interpretation"] = (
        "Vue de pilotage client; aucune compensation comptable entre epargne et credit"
    )

    detail_columns = [
        "loan_id",
        "customer_id",
        "Nom_client",
        "customer",
        "telephone_credit",
        "currency_code",
        "loan_amount",
        "amount_paid",
        "encours_credit",
        "outstanding_principle",
        "outstanding_setup_fees",
        "outstanding_interest",
        "outstanding_penalty_fees",
        "status_name",
        "due_date",
        "last_repayment_date",
        "savings_account_id_source",
        "savings_id_correspondant",
        "id_compte_turbo_correspondant",
        "telephone_epargne",
        "solde_epargne_courante",
        "nb_comptes_courants_candidats",
        "nb_correspondances_identifiant",
        "methode_rapprochement_epargne",
        "liaison_directe_source",
        "correspondance_deduite",
        "source_epargne_complete",
        "statut_controle",
        "motif_controle",
    ]
    detail = detail[detail_columns].sort_values(
        ["currency_code", "statut_controle", "encours_credit"],
        ascending=[True, False, False],
    ).reset_index(drop=True)
    controls = detail.loc[detail["statut_controle"].eq("A revoir")].reset_index(drop=True)
    source_complete_available = current_source_complete_available
    sources = pd.DataFrame(
        [
            {
                "loans_account_disponible": True,
                "savings_account_courant_disponible": current_available,
                "dat_disponible": fixed_available,
                "source_savings_account_complete": source_complete_available,
                "savings_account_id_renseignes": int(source_id_present.sum()),
                "grain_vue_client": "customer_id x currency_code",
                "regle_liaison_deduite": "customer_id + currency_code avec un compte courant unique",
            }
        ]
    )
    return {
        "synthese": summary.reset_index(drop=True),
        "clients": clients,
        "detail": detail,
        "controles": controls,
        "sources": sources,
    }


def build_mpesa_liquidity_analysis(
    daily_detail: pd.DataFrame | None,
    *,
    as_of_date: Any | None = None,
) -> dict[str, pd.DataFrame]:
    """Mesure la pression de liquidite et produit une projection mecanique a sept jours."""
    empty = {"synthese": pd.DataFrame(), "journalier": pd.DataFrame()}
    if not isinstance(daily_detail, pd.DataFrame) or daily_detail.empty or "date" not in daily_detail.columns:
        return empty
    frame = daily_detail.copy()
    if "incluse_synthese" in frame.columns:
        frame = frame.loc[frame["incluse_synthese"].astype("boolean").fillna(False)].copy()
    frame["date_complete"] = pd.to_datetime(frame["date"], errors="coerce")
    frame = frame.dropna(subset=["date_complete"])
    if frame.empty:
        return empty
    frame["date_transaction"] = frame["date_complete"].dt.normalize()
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["montant_entree"] = numeric_column(frame, "montant_entree").clip(lower=0)
    frame["montant_sortie"] = numeric_column(frame, "montant_sortie").clip(lower=0)
    frame["balance_numeric"] = pd.to_numeric(
        frame.get("balance_numeric", pd.Series(np.nan, index=frame.index)), errors="coerce"
    )

    observed_daily = (
        frame.groupby(["currency_code", "date_transaction"], as_index=False, dropna=False)
        .agg(
            nombre_transactions=("date_complete", "size"),
            montant_entrees=("montant_entree", "sum"),
            montant_sorties=("montant_sortie", "sum"),
        )
    )
    dates = pd.DataFrame(
        {"date_transaction": pd.date_range(frame["date_transaction"].min(), frame["date_transaction"].max(), freq="D")}
    )
    currencies = frame[["currency_code"]].drop_duplicates()
    daily = currencies.merge(dates, how="cross").merge(
        observed_daily, on=["currency_code", "date_transaction"], how="left"
    )
    daily[["nombre_transactions", "montant_entrees", "montant_sorties"]] = daily[
        ["nombre_transactions", "montant_entrees", "montant_sorties"]
    ].fillna(0)
    daily["nombre_transactions"] = daily["nombre_transactions"].astype(int)
    daily["flux_net"] = daily["montant_entrees"] - daily["montant_sorties"]

    balance_daily = (
        frame.dropna(subset=["balance_numeric"])
        .sort_values("date_complete")
        .groupby(["currency_code", "date_transaction"], as_index=False, dropna=False)
        .agg(solde_cloture_observe=("balance_numeric", "last"), solde_min_observe=("balance_numeric", "min"))
    )
    daily = daily.merge(balance_daily, on=["currency_code", "date_transaction"], how="left")

    hourly_outputs = (
        frame.assign(heure=frame["date_complete"].dt.floor("h"))
        .groupby(["currency_code", "heure"], as_index=False)["montant_sortie"]
        .sum()
    )
    peak_hourly = hourly_outputs.groupby("currency_code")["montant_sortie"].max()
    rows: list[dict[str, Any]] = []
    for currency, group in daily.groupby("currency_code", sort=True):
        group = group.sort_values("date_transaction")
        valid_balances = group.dropna(subset=["solde_cloture_observe"])
        latest_balance = valid_balances.iloc[-1]["solde_cloture_observe"] if not valid_balances.empty else np.nan
        min_balance = valid_balances["solde_min_observe"].min() if not valid_balances.empty else np.nan
        average_out = float(group["montant_sorties"].mean())
        average_net = float(group["flux_net"].mean())
        enough_history = len(group) >= 7
        rows.append(
            {
                "currency_code": currency,
                "jours_observes": int(len(group)),
                "nombre_transactions": int(group["nombre_transactions"].sum()),
                "montant_entrees": float(group["montant_entrees"].sum()),
                "montant_sorties": float(group["montant_sorties"].sum()),
                "flux_net": float(group["flux_net"].sum()),
                "sortie_journaliere_moyenne": average_out,
                "sortie_journaliere_max": float(group["montant_sorties"].max()),
                "sortie_horaire_max": float(peak_hourly.get(currency, 0.0)),
                "solde_plus_recent": latest_balance,
                "solde_min_observe": min_balance,
                "couverture_sorties_jours": latest_balance / average_out if pd.notna(latest_balance) and average_out > 0 else np.nan,
                "projection_solde_7j": latest_balance + (7 * average_net) if pd.notna(latest_balance) and enough_history else np.nan,
                "projection_7j_calculable": bool(pd.notna(latest_balance) and enough_history),
                "date_analyse": pd.Timestamp(as_of_date).normalize() if as_of_date is not None else frame["date_transaction"].max(),
            }
        )
    return {
        "synthese": pd.DataFrame(rows),
        "journalier": daily.sort_values(["currency_code", "date_transaction"]).reset_index(drop=True),
    }


def _build_unified_mpesa_operations(
    prepared: MpesaPreparedData,
    daily_detail: pd.DataFrame | None = None,
) -> pd.DataFrame:
    operations = _build_mpesa_operation_detail(prepared, daily_detail=daily_detail)
    if operations.empty:
        return operations
    operations = operations.copy()
    operations["date_operation"] = pd.to_datetime(operations["date_operation"], errors="coerce")
    operations["phone_prefixe"] = normalize_phone(operations["phone_prefixe"])
    operations["operation_reference"] = clean_identifier(operations["operation_reference"])
    operations["currency_code"] = clean_text(operations["currency_code"]).str.upper().replace("", "NON RENSEIGNEE")
    operations["montant_operation"] = numeric_column(operations, "montant_operation").abs()
    rejected = operations["statut_operation"].apply(normalize_label).str.contains(
        r"failed|failure|cancel|annul|reject|revers|inverse|expire", regex=True, na=False
    )
    operations = operations.loc[~(operations["source_operation"].eq("G2") & rejected)].copy()
    operations["priorite_source"] = operations["source_operation"].map({"G2": 0, "Turbo": 1}).fillna(2)
    reference_key = operations["operation_reference"].where(
        operations["operation_reference"].ne(""),
        operations["source_operation"].astype("string") + "-LIGNE-" + operations.index.astype("string"),
    )
    operations["cle_operation_unique"] = operations["currency_code"].astype("string") + "::" + reference_key
    operations = (
        operations.sort_values(["priorite_source", "date_operation"], na_position="last")
        .drop_duplicates("cle_operation_unique", keep="first")
        .drop(columns=["priorite_source"])
    )
    return operations.sort_values("date_operation", na_position="last").reset_index(drop=True)


def build_mpesa_client_activity_analysis(
    operations: pd.DataFrame | None,
    *,
    as_of_date: Any | None = None,
) -> dict[str, pd.DataFrame]:
    empty = {"synthese": pd.DataFrame(), "clients": pd.DataFrame()}
    if not isinstance(operations, pd.DataFrame) or operations.empty:
        return empty
    analysis_date = pd.Timestamp(
        pd.to_datetime(as_of_date, errors="coerce") if as_of_date is not None else pd.Timestamp.now()
    ).normalize()
    frame = operations.copy()
    frame["date_operation"] = pd.to_datetime(frame.get("date_operation"), errors="coerce")
    frame["phone_prefixe"] = normalize_phone(frame.get("phone_prefixe", pd.Series(pd.NA, index=frame.index)))
    frame = frame.dropna(subset=["date_operation", "phone_prefixe"])
    if frame.empty:
        return empty
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["montant_operation"] = numeric_column(frame, "montant_operation").abs()

    group_keys = ["phone_prefixe", "currency_code"]
    clients = (
        frame.groupby(group_keys, as_index=False, dropna=False)
        .agg(
            customer_ids_turbo=("customer_id_turbo", concat_unique),
            Nom_client=("nom_client_mpesa", concat_unique),
            premiere_operation=("date_operation", "min"),
            derniere_operation=("date_operation", "max"),
            nombre_operations=("cle_operation_unique", "nunique"),
            jours_actifs=("date_operation", lambda values: pd.to_datetime(values).dt.normalize().nunique()),
            montant_total=("montant_operation", "sum"),
            montant_median=("montant_operation", "median"),
            sources=("source_operation", concat_unique),
            types_operations=("type_operation", concat_unique),
        )
    )
    max_gaps = (
        frame.sort_values("date_operation")
        .groupby(group_keys, dropna=False)["date_operation"]
        .apply(lambda values: pd.to_datetime(values).sort_values().diff().dt.days.max())
        .rename("ecart_max_jours")
        .reset_index()
    )
    clients = clients.merge(max_gaps, on=group_keys, how="left")
    clients["jours_depuis_derniere_operation"] = (
        analysis_date - clients["derniere_operation"].dt.normalize()
    ).dt.days.clip(lower=0)
    clients["anciennete_jours"] = (
        analysis_date - clients["premiere_operation"].dt.normalize()
    ).dt.days.clip(lower=0)
    clients["est_nouveau_30j"] = clients["anciennete_jours"].le(30)
    clients["est_reactive_30j"] = (
        clients["jours_depuis_derniere_operation"].le(30)
        & clients["ecart_max_jours"].fillna(0).gt(90)
    )
    clients["statut_activite"] = np.select(
        [
            clients["jours_depuis_derniere_operation"].le(30),
            clients["jours_depuis_derniere_operation"].le(60),
            clients["jours_depuis_derniere_operation"].le(90),
        ],
        ["Actif 30 jours", "Dormant 31 a 60 jours", "Dormant 61 a 90 jours"],
        default="Inactif plus de 90 jours",
    )
    clients["date_analyse"] = analysis_date
    summary = (
        clients.groupby(["currency_code", "statut_activite"], as_index=False, dropna=False)
        .agg(
            nombre_clients=("phone_prefixe", "nunique"),
            nombre_operations=("nombre_operations", "sum"),
            montant_total=("montant_total", "sum"),
            nouveaux_30j=("est_nouveau_30j", "sum"),
            reactives_30j=("est_reactive_30j", "sum"),
        )
    )
    summary["date_analyse"] = analysis_date
    return {
        "synthese": summary,
        "clients": clients.sort_values(
            ["statut_activite", "jours_depuis_derniere_operation", "montant_total"],
            ascending=[True, False, False],
        ).reset_index(drop=True),
    }


def build_mpesa_savings_conversion_analysis(
    daily_detail: pd.DataFrame | None,
) -> dict[str, pd.DataFrame]:
    """Suit la conversion observee d'un depot normal vers un DAT."""
    empty = {"synthese": pd.DataFrame(), "clients": pd.DataFrame()}
    if not isinstance(daily_detail, pd.DataFrame) or daily_detail.empty:
        return empty
    frame = daily_detail.copy()
    if "incluse_synthese" in frame.columns:
        frame = frame.loc[frame["incluse_synthese"].astype("boolean").fillna(False)].copy()
    frame["date"] = pd.to_datetime(frame.get("date"), errors="coerce")
    frame["phone_prefixe"] = normalize_phone(frame.get("phone_prefixe", pd.Series(pd.NA, index=frame.index)))
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["details_rapport"] = clean_text(
        frame.get("details_rapport", pd.Series("", index=frame.index))
    )
    frame["montant"] = numeric_column(frame, "montant").abs()
    frame = frame.loc[
        frame["date"].notna()
        & frame["phone_prefixe"].notna()
        & frame["details_rapport"].isin(["Depot normal", "DAT"])
    ].copy()
    if frame.empty:
        return empty

    rows: list[dict[str, Any]] = []
    for (phone, currency), group in frame.groupby(["phone_prefixe", "currency_code"], dropna=False):
        deposits = group.loc[group["details_rapport"].eq("Depot normal")].sort_values("date")
        dat_rows = group.loc[group["details_rapport"].eq("DAT")].sort_values("date")
        first_deposit = deposits["date"].min() if not deposits.empty else pd.NaT
        eligible_dat = dat_rows.loc[dat_rows["date"].ge(first_deposit)] if pd.notna(first_deposit) else dat_rows.iloc[0:0]
        conversion_date = eligible_dat["date"].min() if not eligible_dat.empty else pd.NaT
        converted = pd.notna(first_deposit) and pd.notna(conversion_date)
        rows.append(
            {
                "phone_prefixe": phone,
                "currency_code": currency,
                "Nom_client": concat_unique(group.get("Nom_client", pd.Series("", index=group.index))),
                "nombre_depots_normaux": int(len(deposits)),
                "montant_depots_normaux": float(deposits["montant"].sum()),
                "nombre_dat": int(len(dat_rows)),
                "montant_dat": float(dat_rows["montant"].sum()),
                "premier_depot_normal": first_deposit,
                "premier_dat_apres_depot": conversion_date,
                "conversion_observee": bool(converted),
                "delai_conversion_jours": int((conversion_date - first_deposit).days) if converted else np.nan,
                "montant_dat_apres_depot": float(eligible_dat["montant"].sum()) if converted else 0.0,
            }
        )
    clients = pd.DataFrame(rows)
    summary_rows: list[dict[str, Any]] = []
    for currency, group in clients.groupby("currency_code", sort=True):
        deposit_clients = group.loc[group["nombre_depots_normaux"].gt(0)]
        converted_clients = deposit_clients.loc[deposit_clients["conversion_observee"]]
        summary_rows.append(
            {
                "currency_code": currency,
                "clients_avec_depot_normal": int(deposit_clients["phone_prefixe"].nunique()),
                "clients_avec_dat": int(group.loc[group["nombre_dat"].gt(0), "phone_prefixe"].nunique()),
                "clients_convertis_vers_dat": int(converted_clients["phone_prefixe"].nunique()),
                "taux_conversion_pct": (
                    100 * converted_clients["phone_prefixe"].nunique() / deposit_clients["phone_prefixe"].nunique()
                    if deposit_clients["phone_prefixe"].nunique()
                    else np.nan
                ),
                "delai_median_conversion_jours": converted_clients["delai_conversion_jours"].median(),
                "montant_depots_normaux": float(deposit_clients["montant_depots_normaux"].sum()),
                "montant_dat_apres_depot": float(converted_clients["montant_dat_apres_depot"].sum()),
            }
        )
    return {
        "synthese": pd.DataFrame(summary_rows),
        "clients": clients.sort_values(
            ["currency_code", "conversion_observee", "montant_depots_normaux"],
            ascending=[True, False, False],
        ).reset_index(drop=True),
    }


def build_mpesa_transaction_concentration_analysis(
    daily_detail: pd.DataFrame | None,
) -> dict[str, pd.DataFrame]:
    empty = {"synthese": pd.DataFrame(), "clients": pd.DataFrame()}
    if not isinstance(daily_detail, pd.DataFrame) or daily_detail.empty:
        return empty
    frame = daily_detail.copy()
    if "incluse_synthese" in frame.columns:
        frame = frame.loc[frame["incluse_synthese"].astype("boolean").fillna(False)].copy()
    frame["phone_prefixe"] = normalize_phone(frame.get("phone_prefixe", pd.Series(pd.NA, index=frame.index)))
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["montant_entree"] = numeric_column(frame, "montant_entree").clip(lower=0)
    frame["montant_sortie"] = numeric_column(frame, "montant_sortie").clip(lower=0)
    frame = frame.dropna(subset=["phone_prefixe"])
    if frame.empty:
        return empty
    clients = (
        frame.groupby(["phone_prefixe", "currency_code"], as_index=False, dropna=False)
        .agg(
            Nom_client=("Nom_client", concat_unique),
            nombre_transactions=("receipt_no", "size"),
            montant_entrees=("montant_entree", "sum"),
            montant_sorties=("montant_sortie", "sum"),
        )
    )
    clients["volume_total"] = clients["montant_entrees"] + clients["montant_sorties"]
    clients["rang_volume"] = clients.groupby("currency_code")["volume_total"].rank(
        method="first", ascending=False
    ).astype(int)
    total_by_currency = clients.groupby("currency_code")["volume_total"].transform("sum")
    clients["part_volume_pct"] = clients["volume_total"].div(total_by_currency.replace(0, pd.NA)).mul(100)
    clients = clients.sort_values(["currency_code", "volume_total"], ascending=[True, False]).reset_index(drop=True)

    summary_rows: list[dict[str, Any]] = []
    for currency, group in clients.groupby("currency_code", sort=True):
        total_volume = float(group["volume_total"].sum())
        total_entries = float(group["montant_entrees"].sum())
        total_outputs = float(group["montant_sorties"].sum())
        summary_rows.append(
            {
                "currency_code": currency,
                "nombre_clients": int(group["phone_prefixe"].nunique()),
                "volume_total": total_volume,
                "part_top_5_volume_pct": 100 * group.head(5)["volume_total"].sum() / total_volume if total_volume else np.nan,
                "part_top_10_volume_pct": 100 * group.head(10)["volume_total"].sum() / total_volume if total_volume else np.nan,
                "part_top_5_entrees_pct": 100 * group.nlargest(5, "montant_entrees")["montant_entrees"].sum() / total_entries if total_entries else np.nan,
                "part_top_5_sorties_pct": 100 * group.nlargest(5, "montant_sorties")["montant_sorties"].sum() / total_outputs if total_outputs else np.nan,
            }
        )
    return {"synthese": pd.DataFrame(summary_rows), "clients": clients}


def build_mpesa_transaction_quality_analysis(
    daily_detail: pd.DataFrame | None,
) -> dict[str, pd.DataFrame]:
    empty = {"synthese": pd.DataFrame(), "alertes": pd.DataFrame()}
    if not isinstance(daily_detail, pd.DataFrame) or daily_detail.empty:
        return empty
    frame = daily_detail.copy()
    frame["date"] = pd.to_datetime(frame.get("date"), errors="coerce")
    frame["receipt_no"] = clean_identifier(frame.get("receipt_no", pd.Series("", index=frame.index)))
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["phone_prefixe"] = normalize_phone(frame.get("phone_prefixe", pd.Series(pd.NA, index=frame.index)))
    frame["montant"] = numeric_column(frame, "montant").abs()
    frame["incluse_synthese"] = frame.get(
        "incluse_synthese", pd.Series(True, index=frame.index)
    ).astype("boolean").fillna(False)
    frame["est_anomalie"] = frame.get(
        "est_anomalie", pd.Series(False, index=frame.index)
    ).astype("boolean").fillna(False)
    frame["doublon_receipt_no"] = frame.get(
        "doublon_receipt_no", pd.Series(False, index=frame.index)
    ).astype("boolean").fillna(False)
    frame["non_rapproche"] = clean_text(
        frame.get("statut_rapprochement", pd.Series("", index=frame.index))
    ).eq("Non rapproche")
    frame["qualite_conforme"] = frame["incluse_synthese"] & ~frame["est_anomalie"]

    summary = (
        frame.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            transactions=("receipt_no", "size"),
            transactions_terminees=("incluse_synthese", "sum"),
            transactions_non_terminees=("incluse_synthese", lambda values: int((~pd.Series(values).astype(bool)).sum())),
            anomalies=("est_anomalie", "sum"),
            doublons=("doublon_receipt_no", "sum"),
            non_rapprochees=("non_rapproche", "sum"),
            conformes=("qualite_conforme", "sum"),
        )
    )
    summary["taux_succes_pct"] = summary["transactions_terminees"].div(
        summary["transactions"].replace(0, pd.NA)
    ).mul(100)
    summary["taux_anomalie_pct"] = summary["anomalies"].div(
        summary["transactions"].replace(0, pd.NA)
    ).mul(100)
    summary["taux_qualite_pct"] = summary["conformes"].div(
        summary["transactions"].replace(0, pd.NA)
    ).mul(100)

    eligible = frame.loc[frame["incluse_synthese"] & frame["date"].notna()].copy()
    if eligible.empty:
        return {"synthese": summary, "alertes": frame.loc[frame["est_anomalie"]].reset_index(drop=True)}
    thresholds = eligible.groupby("currency_code")["montant"].quantile(0.95)
    medians = eligible.groupby("currency_code")["montant"].median()
    eligible["alerte_montant_eleve"] = (
        eligible["montant"].ge(eligible["currency_code"].map(thresholds))
        & eligible["montant"].gt(eligible["currency_code"].map(medians))
        & eligible["montant"].gt(0)
    )
    eligible["heure_num"] = eligible["date"].dt.hour
    hour_frequency = eligible.groupby(["currency_code", "heure_num"])["receipt_no"].transform("size")
    currency_frequency = eligible.groupby("currency_code")["receipt_no"].transform("size").replace(0, pd.NA)
    eligible["alerte_horaire_atypique"] = (
        (eligible["heure_num"].lt(6) | eligible["heure_num"].ge(22))
        & hour_frequency.div(currency_frequency).le(0.10)
    )
    eligible["tranche_10_minutes"] = eligible["date"].dt.floor("10min")
    velocity = eligible.groupby(["phone_prefixe", "tranche_10_minutes"], dropna=False)["receipt_no"].transform("size")
    eligible["alerte_rafale_transactions"] = eligible["phone_prefixe"].notna() & velocity.ge(3)
    eligible["motif_alerte_comportement"] = ""
    eligible.loc[eligible["alerte_montant_eleve"], "motif_alerte_comportement"] += "Montant dans les 5% les plus eleves | "
    eligible.loc[eligible["alerte_horaire_atypique"], "motif_alerte_comportement"] += "Operation entre 22h et 06h | "
    eligible.loc[eligible["alerte_rafale_transactions"], "motif_alerte_comportement"] += "Au moins 3 operations en 10 minutes | "
    eligible["motif_alerte_comportement"] = eligible["motif_alerte_comportement"].str.rstrip(" |")
    behavioral = eligible.loc[eligible["motif_alerte_comportement"].ne("")].copy()
    operational = frame.loc[frame["est_anomalie"]].copy()
    operational["motif_alerte_comportement"] = "Anomalie de controle : " + clean_text(
        operational.get("motif_anomalie", pd.Series("", index=operational.index))
    )
    alerts = concat_frames_stable([behavioral, operational])
    if not alerts.empty:
        alerts["cle_alerte"] = clean_identifier(alerts["receipt_no"]).where(
            clean_identifier(alerts["receipt_no"]).ne(""),
            "ALERTE-LIGNE-" + alerts.index.astype("string"),
        )
        alerts = alerts.drop_duplicates("cle_alerte", keep="first").sort_values("date", ascending=False, na_position="last")
        alert_columns = [
            "date", "receipt_no", "currency_code", "sens_flux", "details_rapport", "phone_prefixe",
            "Nom_client", "montant", "transaction_status", "statut_rapprochement",
            "motif_alerte_comportement", "motif_anomalie", "Observation",
        ]
        alerts = alerts[[column for column in alert_columns if column in alerts.columns]].reset_index(drop=True)
    return {"synthese": summary, "alertes": alerts}


def build_mpesa_dat_maturity_analysis(
    fixed_savings: pd.DataFrame | None,
    *,
    as_of_date: Any | None = None,
    annual_interest_rate_pct: float | None = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
    preparation_horizon_days: int = DEFAULT_DAT_REPAYMENT_PREPARATION_HORIZON_DAYS,
) -> dict[str, pd.DataFrame]:
    empty = {"synthese": pd.DataFrame(), "detail": pd.DataFrame()}
    if not isinstance(fixed_savings, pd.DataFrame) or fixed_savings.empty:
        return empty
    analysis_date = pd.Timestamp(
        pd.to_datetime(as_of_date, errors="coerce") if as_of_date is not None else pd.Timestamp.now()
    ).normalize()
    frame = fixed_savings.copy()
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper().replace("", "NON RENSEIGNEE")
    frame["balance"] = numeric_column(frame, "balance").clip(lower=0)
    frame["maturity_date"] = pd.to_datetime(frame.get("maturity_date"), errors="coerce")
    frame["date_approved"] = pd.to_datetime(frame.get("date_approved"), errors="coerce")
    frame = frame.loc[frame["balance"].gt(0)].copy()
    if frame.empty:
        return empty
    invalid_date_order = (
        frame["date_approved"].notna()
        & frame["maturity_date"].notna()
        & frame["maturity_date"].lt(frame["date_approved"])
    )
    frame["controle_date_dat"] = np.select(
        [
            frame["date_approved"].isna() | frame["maturity_date"].isna(),
            invalid_date_order,
        ],
        ["Date manquante", "Echeance anterieure a l'approbation"],
        default="Dates coherentes",
    )
    frame["Observation"] = ""
    frame.loc[invalid_date_order, "Observation"] = frame.loc[invalid_date_order].apply(
        lambda row: (
            f"Date d'approbation : {row['date_approved']:%d/%m/%Y} | "
            f"Date d'echeance : {row['maturity_date']:%d/%m/%Y}"
        ),
        axis=1,
    )
    interest_rate = pd.to_numeric(
        pd.Series([annual_interest_rate_pct]), errors="coerce"
    ).iloc[0]
    interest_enabled = pd.notna(interest_rate) and float(interest_rate) > 0
    horizon_value = pd.to_numeric(
        pd.Series([preparation_horizon_days]), errors="coerce"
    ).iloc[0]
    preparation_horizon = (
        max(int(horizon_value), 0) if pd.notna(horizon_value) else 0
    )
    frame["duree_contractuelle_jours"] = (
        frame["maturity_date"].dt.normalize() - frame["date_approved"].dt.normalize()
    ).dt.days
    frame["duree_contractuelle_mois_estimee"] = (
        frame["duree_contractuelle_jours"] / (365.0 / 12.0)
    ).round(1)
    valid_interest_period = frame["duree_contractuelle_jours"].ge(0)
    frame["taux_interet_annuel_pct"] = float(interest_rate) if interest_enabled else np.nan
    frame["interet_estime_echeance"] = np.nan
    if interest_enabled:
        frame.loc[valid_interest_period, "interet_estime_echeance"] = (
            frame.loc[valid_interest_period, "balance"]
            * float(interest_rate)
            / 100
            * frame.loc[valid_interest_period, "duree_contractuelle_jours"]
            / 365
        )
    frame["capital_plus_interet_estime"] = frame["balance"] + frame["interet_estime_echeance"]
    frame["jours_avant_echeance"] = (frame["maturity_date"].dt.normalize() - analysis_date).dt.days
    frame["tranche_echeance"] = np.select(
        [
            frame["maturity_date"].isna(),
            frame["jours_avant_echeance"].lt(0),
            frame["jours_avant_echeance"].le(7),
            frame["jours_avant_echeance"].le(30),
            frame["jours_avant_echeance"].le(60),
            frame["jours_avant_echeance"].le(90),
        ],
        ["Date manquante", "Echu", "0 a 7 jours", "8 a 30 jours", "31 a 60 jours", "61 a 90 jours"],
        default="Plus de 90 jours",
    )
    frame["a_preparer_remboursement"] = (
        frame["maturity_date"].notna()
        & frame["jours_avant_echeance"].le(preparation_horizon)
    )
    frame["statut_preparation_remboursement"] = np.select(
        [
            frame["maturity_date"].isna(),
            frame["jours_avant_echeance"].lt(0),
            frame["jours_avant_echeance"].eq(0),
            frame["jours_avant_echeance"].le(preparation_horizon),
        ],
        [
            "Date d'echeance manquante",
            "Echu - remboursement a traiter",
            "Echeance aujourd'hui",
            f"A preparer sous {preparation_horizon} jours",
        ],
        default="Hors horizon de preparation",
    )
    frame["horizon_preparation_jours"] = preparation_horizon
    frame["montant_estime_a_rembourser"] = frame["capital_plus_interet_estime"]
    frame["date_analyse"] = analysis_date
    for column in [
        "savings_id",
        "customer_id",
        "Nom_client",
        "msisdn",
        "product_name",
        "product_description",
        "account_type",
        "status",
        "fichier_source_epargne_turbo",
    ]:
        if column not in frame.columns:
            frame[column] = pd.NA
    bucket_order = ["Echu", "0 a 7 jours", "8 a 30 jours", "31 a 60 jours", "61 a 90 jours", "Plus de 90 jours", "Date manquante"]
    frame["ordre_tranche"] = frame["tranche_echeance"].map({value: index for index, value in enumerate(bucket_order)})
    summary = (
        frame.groupby(["currency_code", "tranche_echeance", "ordre_tranche"], as_index=False, dropna=False)
        .agg(
            nombre_dat=("balance", "size"),
            nombre_clients=("customer_id", "nunique"),
            montant_dat=("balance", "sum"),
            interet_estime_echeance=("interet_estime_echeance", lambda values: values.sum(min_count=1)),
            capital_plus_interet_estime=("capital_plus_interet_estime", lambda values: values.sum(min_count=1)),
        )
        .sort_values(["currency_code", "ordre_tranche"])
        .drop(columns="ordre_tranche")
        .reset_index(drop=True)
    )
    detail_columns = [
        "savings_id", "customer_id", "Nom_client", "msisdn", "currency_code", "product_name",
        "product_description", "account_type", "status", "balance", "date_approved", "maturity_date",
        "controle_date_dat", "Observation", "duree_contractuelle_jours",
        "duree_contractuelle_mois_estimee", "taux_interet_annuel_pct", "interet_estime_echeance",
        "capital_plus_interet_estime", "montant_estime_a_rembourser", "jours_avant_echeance",
        "tranche_echeance", "a_preparer_remboursement", "statut_preparation_remboursement",
        "horizon_preparation_jours", "date_analyse", "fichier_source_epargne_turbo",
    ]
    detail = frame.sort_values(["currency_code", "ordre_tranche", "maturity_date"])[detail_columns].reset_index(drop=True)
    return {"synthese": summary, "detail": detail}


def build_mpesa_perfect_adoption_analysis(
    prepared: MpesaPreparedData,
    operations: pd.DataFrame | None,
    *,
    as_of_date: Any | None = None,
) -> dict[str, pd.DataFrame]:
    empty = {"synthese": pd.DataFrame(), "statuts": pd.DataFrame(), "detail": pd.DataFrame()}
    perfect = _aggregate_perfect_clients(prepared.perfect_clients)
    if perfect.empty:
        return empty
    analysis_date = pd.Timestamp(
        pd.to_datetime(as_of_date, errors="coerce") if as_of_date is not None else pd.Timestamp.now()
    ).normalize()
    crosscheck = _build_mpesa_identity_population(prepared)
    presence_columns = [
        column for column in ["phone_prefixe", "present_dans_turbo", "present_dans_g2", "present_dans_perfect"]
        if column in crosscheck.columns
    ]
    if presence_columns:
        presence = crosscheck[presence_columns].drop_duplicates("phone_prefixe")
        detail = perfect.merge(presence, on="phone_prefixe", how="left")
    else:
        detail = perfect.copy()
    for column in ["present_dans_turbo", "present_dans_g2"]:
        detail[column] = detail.get(column, pd.Series(False, index=detail.index)).astype("boolean").fillna(False)

    if isinstance(operations, pd.DataFrame) and not operations.empty:
        activity = operations.dropna(subset=["phone_prefixe"]).copy()
        activity["date_operation"] = pd.to_datetime(activity["date_operation"], errors="coerce")
        activity_summary = (
            activity.groupby("phone_prefixe", as_index=False, dropna=False)
            .agg(
                premiere_operation=("date_operation", "min"),
                derniere_operation=("date_operation", "max"),
                nombre_operations=("cle_operation_unique", "nunique"),
                devises_mpesa=("currency_code", concat_unique),
                types_operations_mpesa=("type_operation", concat_unique),
            )
        )
        detail = detail.merge(activity_summary, on="phone_prefixe", how="left")
    else:
        detail["premiere_operation"] = pd.NaT
        detail["derniere_operation"] = pd.NaT
        detail["nombre_operations"] = 0
    detail["nombre_operations"] = numeric_column(detail, "nombre_operations").astype(int)
    detail["present_dans_mpesa"] = detail["present_dans_turbo"] | detail["present_dans_g2"] | detail["nombre_operations"].gt(0)
    detail["jours_depuis_derniere_operation"] = (
        analysis_date - pd.to_datetime(detail["derniere_operation"], errors="coerce").dt.normalize()
    ).dt.days
    detail["statut_adoption"] = np.select(
        [
            ~detail["present_dans_mpesa"],
            detail["derniere_operation"].isna(),
            detail["jours_depuis_derniere_operation"].le(30),
            detail["jours_depuis_derniere_operation"].le(90),
        ],
        [
            "Jamais observe dans Turbo + G2",
            "Present Turbo/G2 sans operation datee",
            "Actif Turbo + G2 30 jours",
            "Actif Turbo + G2 31 a 90 jours",
        ],
        default="Inactif Turbo + G2 plus de 90 jours",
    )
    detail["date_analyse"] = analysis_date
    total = int(len(detail))
    present = int(detail["present_dans_mpesa"].sum())
    active30 = int(detail["statut_adoption"].eq("Actif Turbo + G2 30 jours").sum())
    summary = pd.DataFrame(
        [
            {
                "telephones_perfect_valides": total,
                "clients_perfect_dans_mpesa": present,
                "clients_perfect_actifs_30j": active30,
                "clients_perfect_jamais_observes": int(
                    detail["statut_adoption"].eq("Jamais observe dans Turbo + G2").sum()
                ),
                "taux_adoption_mpesa_pct": 100 * present / total if total else np.nan,
                "taux_activite_30j_pct": 100 * active30 / present if present else np.nan,
                "date_analyse": analysis_date,
            }
        ]
    )
    statuses = (
        detail.groupby("statut_adoption", as_index=False, dropna=False)
        .agg(nombre_clients=("phone_prefixe", "nunique"), nombre_fiches_perfect=("nb_clients_perfect", "sum"))
        .sort_values("nombre_clients", ascending=False)
        .reset_index(drop=True)
    )
    return {"synthese": summary, "statuts": statuses, "detail": detail.sort_values(["statut_adoption", "phone_prefixe"]).reset_index(drop=True)}


MPESA_ACCOUNTING_AUXILIARY_ACCOUNTS = {
    "NORMAL SAVINGS": ("Epargne courante", "Passif client observe"),
    "FIXED SAVINGS": ("DAT", "Passif client observe"),
    "PRINCIPLE": ("Principal credit", "Actif credit observe"),
}

MPESA_ACCOUNTING_FINANCIAL_PRODUCTS = {
    "INTEREST EARNED": "Interets comptabilises observes",
    "LOAN PENALTY FEES": "Penalites comptabilisees observees",
    "BISOU COLLECTION": "Part Bisou observee",
    "VODA COLLECTION A/C": "Part Voda observee",
}


def _empty_mpesa_accounting_analysis() -> dict[str, pd.DataFrame]:
    return {
        "periode": pd.DataFrame(),
        "synthese": pd.DataFrame(),
        "balance_clients": pd.DataFrame(),
        "balance_auxiliaire_clients": pd.DataFrame(),
        "balance_comptes": pd.DataFrame(),
        "journal_operations": pd.DataFrame(),
        "journal_ecritures": pd.DataFrame(),
        "controles_operations": pd.DataFrame(),
        "controles_soldes": pd.DataFrame(),
        "flux_mpesa": pd.DataFrame(),
        "produits_financiers": pd.DataFrame(),
        "positions_portefeuille": pd.DataFrame(),
        "controle_g2": pd.DataFrame(),
    }


def build_mpesa_accounting_analysis(
    prepared: MpesaPreparedData,
    *,
    date_start: object | None = None,
    date_end: object | None = None,
) -> dict[str, pd.DataFrame]:
    """Construit les analyses comptables observables dans Transactions Turbo.

    Turbo reste la source unique des mouvements et des montants. G2 enrichit le
    nom et mesure la couverture du rapprochement direct ``Receipt No = ref_no``.
    Les soldes ``bal_before``/``bal_after`` sont presentes comme des positions
    observees; ils ne constituent pas une balance generale certifiee sans plan
    comptable complet ni soldes d'ouverture officiels.
    """
    result = _empty_mpesa_accounting_analysis()
    transactions = prepared.transactions
    required = {
        "id", "customer_id", "msisdn1", "account_type", "reference_id",
        "currency_code", "dr", "cr", "bal_before", "bal_after", "ref_no",
        "description", "created_at",
    }
    if not isinstance(transactions, pd.DataFrame) or transactions.empty or not required.issubset(transactions.columns):
        return result

    full = transactions.copy()
    for column in ["id", "customer_id", "msisdn1", "account_type", "reference_id", "currency_code", "ref_no", "description"]:
        full[column] = clean_text(full[column])
    full["currency_code"] = full["currency_code"].str.upper().replace("", "NON RENSEIGNEE")
    full["created_at"] = pd.to_datetime(full["created_at"], errors="coerce")
    full = full.dropna(subset=["created_at"]).copy()
    if full.empty:
        return result
    for column in ["dr", "cr", "bal_before", "bal_after"]:
        full[column] = numeric_column(full, column)
    if "Nom_client" not in full.columns:
        full["Nom_client"] = ""
    full["Nom_client"] = clean_text(full["Nom_client"])
    full["__row_order"] = np.arange(len(full))

    available_dates = full["created_at"].dropna()
    start = pd.to_datetime(date_start, errors="coerce") if date_start is not None else available_dates.min()
    end = pd.to_datetime(date_end, errors="coerce") if date_end is not None else available_dates.max()
    if pd.isna(start):
        start = available_dates.min()
    if pd.isna(end):
        end = available_dates.max()
    start = pd.Timestamp(start)
    end = pd.Timestamp(end)
    if start == start.normalize():
        start = start.normalize()
    if end == end.normalize():
        end = _timestamp_plus(end.normalize(), days=1, microseconds=-1)
    if start > end:
        start, end = end.normalize(), _timestamp_plus(start.normalize(), days=1, microseconds=-1)

    scoped = full.loc[full["created_at"].between(start, end, inclusive="both")].copy()
    if scoped.empty:
        result["periode"] = pd.DataFrame(
            [{"date_debut": start, "date_fin": end, "source": "Transactions M-PESA_Turbo", "nombre_lignes": 0}]
        )
        return result

    reference = clean_identifier(scoped["ref_no"])
    timestamp_token = scoped["created_at"].dt.strftime("%Y-%m-%d %H:%M:%S.%f")
    scoped["cle_operation_turbo"] = np.where(
        reference.ne(""),
        "REF|" + reference,
        "HORODATAGE|" + scoped["customer_id"] + "|" + scoped["currency_code"] + "|" + timestamp_token,
    )
    scoped["reference_operation"] = reference.where(
        reference.ne(""),
        "TURBO-" + scoped["customer_id"] + "-" + timestamp_token,
    )
    scoped["date_operation"] = scoped["created_at"].dt.normalize()

    operation_group = ["customer_id", "currency_code", "cle_operation_turbo"]
    journal_operations = (
        scoped.groupby(operation_group, as_index=False, dropna=False)
        .agg(
            date_operation=("date_operation", "min"),
            created_at=("created_at", "min"),
            reference_operation=("reference_operation", first_non_empty),
            ref_no=("ref_no", concat_unique),
            Nom_client=("Nom_client", concat_unique),
            telephone=("msisdn1", concat_unique),
            comptes_turbo=("account_type", concat_unique),
            descriptions_turbo=("description", concat_unique),
            total_debit=("dr", "sum"),
            total_credit=("cr", "sum"),
            nombre_lignes=("id", "size"),
        )
        .sort_values(["currency_code", "created_at", "reference_operation"])
        .reset_index(drop=True)
    )
    journal_operations["ecart_debit_credit"] = journal_operations["total_debit"] - journal_operations["total_credit"]
    operation_tolerance = np.maximum(
        0.01,
        journal_operations[["total_debit", "total_credit"]].abs().max(axis=1) * 1e-6,
    )
    journal_operations["operation_symetrique"] = journal_operations["ecart_debit_credit"].abs().le(operation_tolerance)
    journal_operations["statut_controle_operation"] = np.where(
        journal_operations["operation_symetrique"],
        "Symetrique dans l'export Turbo",
        "A verifier - sous-registres ou ecritures incompletes",
    )
    journal_operations["solde_debiteur_mouvement"] = journal_operations["ecart_debit_credit"].clip(lower=0)
    journal_operations["solde_crediteur_mouvement"] = (-journal_operations["ecart_debit_credit"]).clip(lower=0)

    operation_controls = journal_operations.loc[
        ~journal_operations["operation_symetrique"]
    ].copy().reset_index(drop=True)

    client_control = (
        journal_operations.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
        .agg(
            nombre_operations=("cle_operation_turbo", "size"),
            operations_a_verifier=("operation_symetrique", lambda values: int((~values.astype(bool)).sum())),
        )
    )
    balance_clients = (
        scoped.groupby(["customer_id", "currency_code"], as_index=False, dropna=False)
        .agg(
            Nom_client=("Nom_client", concat_unique),
            telephone=("msisdn1", concat_unique),
            total_debit=("dr", "sum"),
            total_credit=("cr", "sum"),
            nombre_lignes=("id", "size"),
            nombre_types_comptes=("account_type", "nunique"),
            premiere_ecriture=("created_at", "min"),
            derniere_ecriture=("created_at", "max"),
        )
        .merge(client_control, on=["customer_id", "currency_code"], how="left")
    )
    balance_clients["ecart_debit_credit"] = balance_clients["total_debit"] - balance_clients["total_credit"]
    balance_clients["solde_debiteur_mouvement"] = balance_clients["ecart_debit_credit"].clip(lower=0)
    balance_clients["solde_crediteur_mouvement"] = (-balance_clients["ecart_debit_credit"]).clip(lower=0)

    balance_comptes = (
        scoped.groupby(["currency_code", "account_type"], as_index=False, dropna=False)
        .agg(
            total_debit=("dr", "sum"),
            total_credit=("cr", "sum"),
            nombre_lignes=("id", "size"),
            nombre_operations=("cle_operation_turbo", "nunique"),
            nombre_clients=("customer_id", "nunique"),
            premiere_ecriture=("created_at", "min"),
            derniere_ecriture=("created_at", "max"),
        )
    )
    balance_comptes["ecart_debit_credit"] = balance_comptes["total_debit"] - balance_comptes["total_credit"]
    balance_comptes["solde_debiteur_mouvement"] = balance_comptes["ecart_debit_credit"].clip(lower=0)
    balance_comptes["solde_crediteur_mouvement"] = (-balance_comptes["ecart_debit_credit"]).clip(lower=0)
    balance_comptes = balance_comptes.sort_values(["currency_code", "account_type"]).reset_index(drop=True)

    # Balance auxiliaire : uniquement les comptes produits qui representent une
    # position client exploitable. Les comptes techniques restent dans la
    # balance des mouvements et le journal des operations.
    auxiliary_full = full.loc[full["account_type"].isin(MPESA_ACCOUNTING_AUXILIARY_ACCOUNTS)].copy()
    auxiliary_balance = pd.DataFrame()
    if not auxiliary_full.empty:
        reference_keys = ["customer_id", "currency_code", "account_type"]
        known_references = (
            auxiliary_full.loc[auxiliary_full["reference_id"].ne("")]
            .groupby(reference_keys, as_index=False, dropna=False)
            .agg(
                reference_connue=("reference_id", first_non_empty),
                nombre_references_connues=("reference_id", "nunique"),
            )
        )
        auxiliary_scoped = scoped.loc[
            scoped["account_type"].isin(MPESA_ACCOUNTING_AUXILIARY_ACCOUNTS)
        ].copy()
        auxiliary_scoped = auxiliary_scoped.merge(known_references, on=reference_keys, how="left")
        auxiliary_scoped["nombre_references_connues"] = pd.to_numeric(
            auxiliary_scoped["nombre_references_connues"], errors="coerce"
        ).fillna(0).astype(int)
        exact_reference = auxiliary_scoped["reference_id"].ne("")
        resolvable_reference = ~exact_reference & auxiliary_scoped["nombre_references_connues"].eq(1)
        auxiliary_scoped["cle_compte_turbo"] = np.select(
            [exact_reference, resolvable_reference],
            [auxiliary_scoped["reference_id"], auxiliary_scoped["reference_connue"]],
            default="REFERENCE_A_VERIFIER",
        )
        auxiliary_scoped["statut_reference_compte"] = np.select(
            [exact_reference, resolvable_reference],
            ["Reference Turbo presente", "Reference resolue par le compte unique du client"],
            default="Reference compte ambigue ou absente",
        )
        auxiliary_scoped["famille_position"] = auxiliary_scoped["account_type"].map(
            lambda value: MPESA_ACCOUNTING_AUXILIARY_ACCOUNTS[value][0]
        )
        auxiliary_scoped["nature_comptable_indicative"] = auxiliary_scoped["account_type"].map(
            lambda value: MPESA_ACCOUNTING_AUXILIARY_ACCOUNTS[value][1]
        )
        auxiliary_scoped = auxiliary_scoped.sort_values(["created_at", "__row_order"])
        account_keys = [
            "customer_id", "currency_code", "account_type", "famille_position",
            "nature_comptable_indicative", "cle_compte_turbo",
        ]
        account_positions = (
            auxiliary_scoped.groupby(account_keys, as_index=False, dropna=False)
            .agg(
                Nom_client=("Nom_client", concat_unique),
                telephone=("msisdn1", concat_unique),
                statut_reference_compte=("statut_reference_compte", concat_unique),
                solde_ouverture_observe=("bal_before", "first"),
                total_debit=("dr", "sum"),
                total_credit=("cr", "sum"),
                solde_cloture_observe=("bal_after", "last"),
                nombre_lignes=("id", "size"),
                nombre_operations=("cle_operation_turbo", "nunique"),
                premiere_ecriture=("created_at", "min"),
                derniere_ecriture=("created_at", "max"),
            )
        )
        account_positions["variation_solde_observee"] = (
            account_positions["solde_cloture_observe"] - account_positions["solde_ouverture_observe"]
        )
        account_positions["reference_a_verifier"] = account_positions["cle_compte_turbo"].eq("REFERENCE_A_VERIFIER")
        auxiliary_balance = (
            account_positions.groupby(
                ["customer_id", "currency_code", "famille_position", "nature_comptable_indicative"],
                as_index=False,
                dropna=False,
            )
            .agg(
                Nom_client=("Nom_client", concat_unique),
                telephone=("telephone", concat_unique),
                nombre_comptes_observes=("cle_compte_turbo", "nunique"),
                references_a_verifier=("reference_a_verifier", "sum"),
                solde_ouverture_observe=("solde_ouverture_observe", "sum"),
                total_debit=("total_debit", "sum"),
                total_credit=("total_credit", "sum"),
                variation_solde_observee=("variation_solde_observee", "sum"),
                solde_cloture_observe=("solde_cloture_observe", "sum"),
                nombre_lignes=("nombre_lignes", "sum"),
                nombre_operations=("nombre_operations", "sum"),
                premiere_ecriture=("premiere_ecriture", "min"),
                derniere_ecriture=("derniere_ecriture", "max"),
            )
            .sort_values(["currency_code", "customer_id", "famille_position"])
            .reset_index(drop=True)
        )

        position_pivot = auxiliary_balance.pivot_table(
            index=["customer_id", "currency_code"],
            columns="famille_position",
            values="solde_cloture_observe",
            aggfunc="sum",
            fill_value=0.0,
        ).reset_index()
        position_pivot.columns.name = None
        position_pivot = position_pivot.rename(
            columns={
                "Epargne courante": "solde_epargne_courante_observe",
                "DAT": "solde_dat_observe",
                "Principal credit": "encours_principal_observe",
            }
        )
        for column in ["solde_epargne_courante_observe", "solde_dat_observe", "encours_principal_observe"]:
            if column not in position_pivot.columns:
                position_pivot[column] = 0.0
        position_pivot["avoirs_epargne_observes"] = (
            position_pivot["solde_epargne_courante_observe"] + position_pivot["solde_dat_observe"]
        )
        balance_clients = balance_clients.merge(
            position_pivot,
            on=["customer_id", "currency_code"],
            how="left",
        )

    for column in ["solde_epargne_courante_observe", "solde_dat_observe", "avoirs_epargne_observes", "encours_principal_observe"]:
        if column not in balance_clients.columns:
            balance_clients[column] = 0.0
        balance_clients[column] = pd.to_numeric(balance_clients[column], errors="coerce").fillna(0.0)
    balance_clients = balance_clients.sort_values(
        ["currency_code", "avoirs_epargne_observes", "encours_principal_observe", "customer_id"],
        ascending=[True, False, False, True],
    ).reset_index(drop=True)

    scoped["variation_solde_observee"] = scoped["bal_after"] - scoped["bal_before"]
    scoped["amplitude_variation_solde"] = scoped["variation_solde_observee"].abs()
    scoped["amplitude_mouvement_ligne"] = scoped["dr"].abs() + scoped["cr"].abs()
    scoped["ecart_amplitude_solde_mouvement"] = (
        scoped["amplitude_variation_solde"] - scoped["amplitude_mouvement_ligne"]
    )
    line_tolerance = np.maximum(0.01, scoped["amplitude_mouvement_ligne"] * 1e-6)
    scoped["variation_solde_coherente"] = scoped["ecart_amplitude_solde_mouvement"].abs().le(line_tolerance)
    scoped["statut_controle_solde"] = np.where(
        scoped["variation_solde_coherente"],
        "Amplitude de solde conforme au mouvement",
        "A verifier - variation de solde differente du mouvement",
    )
    balance_control_columns = [
        "created_at", "id", "customer_id", "Nom_client", "msisdn1", "currency_code",
        "account_type", "reference_id", "ref_no", "description", "dr", "cr",
        "bal_before", "bal_after", "variation_solde_observee", "amplitude_mouvement_ligne",
        "ecart_amplitude_solde_mouvement", "statut_controle_solde",
    ]
    balance_controls = scoped.loc[
        ~scoped["variation_solde_coherente"], balance_control_columns
    ].copy().reset_index(drop=True)

    cash = scoped.loc[scoped["account_type"].eq("MPESA ACCOUNT")].copy()
    if not cash.empty:
        cash["entree_bisou_observee"] = cash["dr"]
        cash["sortie_bisou_observee"] = cash["cr"]
        cash_flow = (
            cash.groupby("currency_code", as_index=False, dropna=False)
            .agg(
                entrees_bisou_observees=("entree_bisou_observee", "sum"),
                sorties_bisou_observees=("sortie_bisou_observee", "sum"),
                nombre_operations=("cle_operation_turbo", "nunique"),
                nombre_clients=("customer_id", "nunique"),
            )
        )
        cash_flow["flux_net_bisou_observe"] = (
            cash_flow["entrees_bisou_observees"] - cash_flow["sorties_bisou_observees"]
        )
    else:
        cash_flow = pd.DataFrame()

    product_rows = scoped.loc[
        scoped["account_type"].isin(MPESA_ACCOUNTING_FINANCIAL_PRODUCTS)
    ].copy()
    if not product_rows.empty:
        product_rows["produit_financier_observe"] = product_rows["account_type"].map(
            MPESA_ACCOUNTING_FINANCIAL_PRODUCTS
        )
        product_rows["montant_observe"] = product_rows[["dr", "cr"]].abs().max(axis=1)
        financial_products = (
            product_rows.groupby(
                ["currency_code", "account_type", "produit_financier_observe"],
                as_index=False,
                dropna=False,
            )
            .agg(
                montant_observe=("montant_observe", "sum"),
                nombre_lignes=("id", "size"),
                nombre_operations=("cle_operation_turbo", "nunique"),
            )
            .sort_values(["currency_code", "account_type"])
            .reset_index(drop=True)
        )
    else:
        financial_products = pd.DataFrame()

    portfolio_parts: list[pd.DataFrame] = []
    for frame, source_label, value_column, output_column, date_columns in [
        (prepared.current_savings, "Epargne courante_Turbo", "balance", "solde_epargne_courante_reference", ["updated_at", "created_at"]),
        (prepared.fixed_savings, "DAT_Turbo", "balance", "solde_dat_reference", ["date_approved"]),
        (prepared.loans, "Credits_Turbo", "loan_balance", "encours_credit_reference", ["updated_at", "created_at"]),
    ]:
        if not isinstance(frame, pd.DataFrame) or frame.empty or "currency_code" not in frame.columns or value_column not in frame.columns:
            continue
        snapshot = frame.copy()
        snapshot["currency_code"] = clean_text(snapshot["currency_code"]).str.upper().replace("", "NON RENSEIGNEE")
        snapshot[value_column] = numeric_column(snapshot, value_column)
        available_date_column = next((column for column in date_columns if column in snapshot.columns), None)
        if available_date_column:
            snapshot["__source_date"] = pd.to_datetime(snapshot[available_date_column], errors="coerce")
        else:
            snapshot["__source_date"] = pd.NaT
        aggregated = (
            snapshot.groupby("currency_code", as_index=False, dropna=False)
            .agg(**{
                output_column: (value_column, "sum"),
                f"nombre_lignes_{output_column}": (value_column, "size"),
                f"date_source_{output_column}": ("__source_date", "max"),
            })
        )
        aggregated[f"source_{output_column}"] = source_label
        portfolio_parts.append(aggregated)
    portfolio_positions = pd.DataFrame()
    for part in portfolio_parts:
        portfolio_positions = part if portfolio_positions.empty else portfolio_positions.merge(
            part, on="currency_code", how="outer"
        )
    if not portfolio_positions.empty:
        for column in ["solde_epargne_courante_reference", "solde_dat_reference", "encours_credit_reference"]:
            if column not in portfolio_positions.columns:
                portfolio_positions[column] = 0.0
            portfolio_positions[column] = pd.to_numeric(portfolio_positions[column], errors="coerce").fillna(0.0)
        portfolio_positions["depots_clients_reference"] = (
            portfolio_positions["solde_epargne_courante_reference"]
            + portfolio_positions["solde_dat_reference"]
        )
        portfolio_positions["ratio_credits_depots_pct"] = (
            portfolio_positions["encours_credit_reference"]
            .div(portfolio_positions["depots_clients_reference"].replace(0, pd.NA))
            .mul(100)
        )
        portfolio_positions = portfolio_positions.sort_values("currency_code").reset_index(drop=True)

    g2_control_rows: list[dict[str, object]] = []
    g2 = prepared.g2_transactions.copy() if isinstance(prepared.g2_transactions, pd.DataFrame) else pd.DataFrame()
    if not g2.empty:
        g2_time = pd.to_datetime(
            g2.get("completion_time", pd.Series(pd.NaT, index=g2.index)), errors="coerce"
        )
        g2 = g2.loc[g2_time.between(start, end, inclusive="both")].copy()
        g2["__completed"] = g2_completed_transaction_mask(g2)
        g2["__receipt_key"] = clean_identifier(g2.get("receipt_no", pd.Series("", index=g2.index)))
        g2["currency_code"] = clean_text(g2.get("currency_code", pd.Series("", index=g2.index))).str.upper()
        turbo_reference_keys = set(
            zip(
                scoped.loc[scoped["ref_no"].ne(""), "currency_code"],
                clean_identifier(scoped.loc[scoped["ref_no"].ne(""), "ref_no"]),
            )
        )
        g2["__direct_match"] = [
            (currency, receipt) in turbo_reference_keys and bool(receipt)
            for currency, receipt in zip(g2["currency_code"], g2["__receipt_key"])
        ]
    currencies = sorted(
        set(scoped["currency_code"].dropna().astype(str))
        | (set(g2["currency_code"].dropna().astype(str)) if not g2.empty else set())
    )
    for currency in currencies:
        turbo_currency = scoped.loc[scoped["currency_code"].eq(currency)]
        client_names = (
            turbo_currency.groupby("customer_id", as_index=False)
            .agg(Nom_client=("Nom_client", concat_unique))
        )
        named_clients = int(client_names["Nom_client"].ne("").sum())
        total_clients = int(len(client_names))
        g2_currency = g2.loc[g2["currency_code"].eq(currency)] if not g2.empty else pd.DataFrame()
        completed_g2 = g2_currency.loc[g2_currency["__completed"]] if not g2_currency.empty else pd.DataFrame()
        matched_g2 = int(completed_g2["__direct_match"].sum()) if not completed_g2.empty else 0
        g2_control_rows.append(
            {
                "currency_code": currency,
                "clients_turbo": total_clients,
                "clients_turbo_avec_nom_g2": named_clients,
                "taux_clients_nommes_g2_pct": 100 * named_clients / total_clients if total_clients else np.nan,
                "operations_turbo_avec_ref_no": int(
                    journal_operations.loc[
                        journal_operations["currency_code"].eq(currency)
                        & journal_operations["ref_no"].ne("")
                    ].shape[0]
                ),
                "transactions_g2_chargees": int(len(g2_currency)),
                "transactions_g2_terminees": int(len(completed_g2)),
                "references_g2_turbo_rapprochees": matched_g2,
                "taux_rapprochement_g2_pct": (
                    100 * matched_g2 / len(completed_g2) if len(completed_g2) else np.nan
                ),
            }
        )
    g2_control = pd.DataFrame(g2_control_rows)

    summary = (
        balance_clients.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            nombre_clients=("customer_id", "nunique"),
            nombre_lignes=("nombre_lignes", "sum"),
            nombre_operations=("nombre_operations", "sum"),
            operations_a_verifier=("operations_a_verifier", "sum"),
            total_debit=("total_debit", "sum"),
            total_credit=("total_credit", "sum"),
            solde_debiteur_mouvement=("solde_debiteur_mouvement", "sum"),
            solde_crediteur_mouvement=("solde_crediteur_mouvement", "sum"),
            avoirs_epargne_observes=("avoirs_epargne_observes", "sum"),
            encours_principal_observe=("encours_principal_observe", "sum"),
        )
    )
    summary["operations_symetriques"] = summary["nombre_operations"] - summary["operations_a_verifier"]
    summary["taux_operations_symetriques_pct"] = (
        summary["operations_symetriques"]
        .div(summary["nombre_operations"].replace(0, pd.NA))
        .mul(100)
    )
    line_quality = (
        scoped.groupby("currency_code", as_index=False, dropna=False)
        .agg(
            lignes_variation_solde_conforme=("variation_solde_coherente", "sum"),
            lignes_variation_solde_total=("variation_solde_coherente", "size"),
        )
    )
    line_quality["taux_variation_solde_conforme_pct"] = (
        line_quality["lignes_variation_solde_conforme"]
        .div(line_quality["lignes_variation_solde_total"].replace(0, pd.NA))
        .mul(100)
    )
    summary = summary.merge(line_quality, on="currency_code", how="left")
    if not g2_control.empty:
        summary = summary.merge(
            g2_control[
                ["currency_code", "taux_clients_nommes_g2_pct", "taux_rapprochement_g2_pct"]
            ],
            on="currency_code",
            how="left",
        )
    summary["date_debut"] = start
    summary["date_fin"] = end
    summary["source_mouvements"] = "Transactions M-PESA_Turbo"
    summary["role_g2"] = "Nom client et controle uniquement"

    journal_entry_columns = [
        "created_at", "date_operation", "id", "customer_id", "Nom_client", "msisdn1",
        "currency_code", "account_type", "reference_id", "ref_no", "reference_operation",
        "description", "dr", "cr", "bal_before", "bal_after", "variation_solde_observee",
        "statut_controle_solde", "cle_operation_turbo",
    ]
    period = pd.DataFrame(
        [
            {
                "date_debut": start,
                "date_fin": end,
                "source_mouvements": "Transactions M-PESA_Turbo",
                "source_identite_controle": "Transactions M-PESA_G2" if not prepared.g2_transactions.empty else "Non chargee",
                "nombre_lignes": int(len(scoped)),
                "nombre_clients": int(scoped["customer_id"].nunique()),
                "nombre_devises": int(scoped["currency_code"].nunique()),
                "limite": "Balance observee; pas de balance generale certifiee sans plan comptable et soldes d'ouverture.",
            }
        ]
    )

    result.update(
        {
            "periode": period,
            "synthese": summary.sort_values("currency_code").reset_index(drop=True),
            "balance_clients": balance_clients,
            "balance_auxiliaire_clients": auxiliary_balance,
            "balance_comptes": balance_comptes,
            "journal_operations": journal_operations,
            "journal_ecritures": scoped[journal_entry_columns].sort_values(
                ["currency_code", "created_at", "id"]
            ).reset_index(drop=True),
            "controles_operations": operation_controls,
            "controles_soldes": balance_controls,
            "flux_mpesa": cash_flow,
            "produits_financiers": financial_products,
            "positions_portefeuille": portfolio_positions,
            "controle_g2": g2_control,
        }
    )
    return result


def build_mpesa_management_dashboard(
    prepared: MpesaPreparedData,
    *,
    date_start: Any | None = None,
    as_of_date: Any | None = None,
    frequency: str = "Jour",
    dat_annual_interest_rate_pct: float | None = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
    fractionation_thresholds: dict[str, float] | None = None,
    large_transaction_thresholds: dict[str, float] | None = None,
    turbo_events: pd.DataFrame | None = None,
    turbo_transaction_lines: pd.DataFrame | None = None,
) -> dict[str, Any]:
    """Assemble le cockpit financier Turbo sans melanger les devises ni les grains."""
    report = build_mpesa_turbo_financial_analysis(
        prepared,
        date_start=date_start,
        date_end=as_of_date,
        frequency=frequency,
        dat_annual_interest_rate_pct=dat_annual_interest_rate_pct,
        fractionation_thresholds=fractionation_thresholds,
        large_transaction_thresholds=large_transaction_thresholds,
        turbo_events=turbo_events,
        turbo_transaction_lines=turbo_transaction_lines,
    )
    alerts = report.get("alertes_transactions", pd.DataFrame())
    quality_summary = (
        alerts.groupby(["currency_code", "alerte"], as_index=False, dropna=False)
        .agg(nombre_alertes=("alerte", "size"), montant_concerne=("montant", "sum"))
        if not alerts.empty
        else pd.DataFrame()
    )
    report.update(
        {
            "date_analyse": report["date_fin"],
            # Alias conserves pour les exports et integrations historiques.
            "liquidite_synthese": report["flux_synthese"],
            "liquidite_journaliere": report["flux_evolution"],
            "activite_synthese": pd.DataFrame(),
            "activite_clients": report["activite_epargne_clients"],
            "conversion_synthese": report["tranches_depots"],
            "conversion_clients": report["activite_epargne_clients"],
            "concentration_synthese": report["concentration_transactions_synthese"],
            "concentration_clients": report["concentration_transactions_clients"],
            "qualite_synthese": quality_summary,
            "perfect_adoption_synthese": pd.DataFrame(),
            "perfect_adoption_statuts": pd.DataFrame(),
            "perfect_adoption_detail": pd.DataFrame(),
        }
    )
    return report


def search_customers(query: object, prepared: MpesaPreparedData) -> pd.DataFrame:
    text = str(query).strip()
    if not text:
        return pd.DataFrame(columns=["customer_id", "Nom_client", "telephone", "source"])
    normalized_phone = normalize_phone(pd.Series([text])).iloc[0]
    frames: list[pd.DataFrame] = []
    source_map = [
        ("Transactions M-PESA_Turbo", prepared.transactions, "msisdn1"),
        ("Epargne courante_Turbo", prepared.current_savings, "msisdn"),
        ("DAT_Turbo", prepared.fixed_savings, "msisdn"),
        ("Credits_Turbo", prepared.loans, "msisdn1"),
        ("Clients_Turbo", prepared.customers, "msisdn1"),
    ]
    for label, frame, phone_col in source_map:
        if frame.empty or "customer_id" not in frame.columns:
            continue
        columns = ["customer_id"]
        if "Nom_client" in frame.columns:
            columns.append("Nom_client")
        if phone_col in frame.columns:
            columns.append(phone_col)
        tmp = frame[columns].copy().rename(columns={phone_col: "telephone"})
        if "Nom_client" not in tmp.columns:
            tmp["Nom_client"] = pd.NA
        if "telephone" not in tmp.columns:
            tmp["telephone"] = pd.NA
        tmp["source"] = label
        frames.append(tmp)
    if not frames:
        return pd.DataFrame(columns=["customer_id", "Nom_client", "telephone", "source"])
    clients = concat_frames_stable(frames).drop_duplicates()
    mask = clients["customer_id"].astype("string").str.contains(text, case=False, regex=False, na=False)
    normalized_name = normalize_label(text)
    if normalized_name:
        mask = mask | clients["Nom_client"].apply(normalize_label).str.contains(normalized_name, regex=False, na=False)
    if not _is_empty_text(normalized_phone):
        mask = mask | clients["telephone"].astype("string").str.contains(str(normalized_phone), case=False, regex=False, na=False)
    else:
        mask = mask | clients["telephone"].astype("string").str.contains(text, case=False, regex=False, na=False)
    return clients.loc[mask].sort_values(["customer_id", "telephone", "source"]).reset_index(drop=True)


def resolve_customer_id(query: object, prepared: MpesaPreparedData) -> str | None:
    matches = search_customers(query, prepared)
    if matches.empty:
        return None
    unique_ids = matches["customer_id"].dropna().astype(str).unique().tolist()
    if len(unique_ids) == 1:
        return unique_ids[0]
    return None


def _build_dat_direct(transactions_client: pd.DataFrame) -> pd.DataFrame:
    if transactions_client.empty:
        return pd.DataFrame()
    dat_direct = transactions_client.loc[transactions_client["account_type"].eq("FIXED SAVINGS")].copy()
    if dat_direct.empty:
        return pd.DataFrame()
    dat_direct["variation_dat_operation"] = dat_direct["bal_after"] - dat_direct["bal_before"]
    return (
        dat_direct.groupby(["currency_code", "created_at"], as_index=False, dropna=False)
        .agg(
            reference_dat_direct=("reference_id", concat_unique),
            solde_dat_operation_avant=("bal_before", "sum"),
            solde_dat_operation_apres=("bal_after", "sum"),
            variation_dat_operation=("variation_dat_operation", "sum"),
        )
    )


CUSTOMER_POSITION_ACCOUNT_LABELS = {
    "NORMAL SAVINGS": "Epargne courante",
    "FIXED SAVINGS": "DAT",
    "PRINCIPLE": "Credit",
}


def _build_turbo_operation_events(
    transactions: pd.DataFrame | None,
    customer_id: object | None = None,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Regroupe les lignes Turbo au grain de l'evenement metier.

    ``ref_no`` reste prioritaire. Quand il manque, les lignes partageant le meme
    client, la meme devise et le meme horodatage sont reunies afin de conserver
    les ventilations credit et les transferts internes DAT qui utilisent des
    ``reference_id`` de comptes differents. ``customer_id`` permet de reutiliser
    exactement le meme moteur pour un extrait client ou pour tout le portefeuille.
    """
    if not isinstance(transactions, pd.DataFrame) or transactions.empty:
        return pd.DataFrame(), pd.DataFrame()
    required = {"customer_id", "currency_code", "created_at", "account_type", "description"}
    if not required.issubset(transactions.columns):
        return pd.DataFrame(), pd.DataFrame()

    frame = transactions.copy()
    if customer_id is not None:
        frame = frame.loc[
            clean_identifier(frame["customer_id"]).eq(str(customer_id).strip())
        ].copy()
    if frame.empty:
        return pd.DataFrame(), pd.DataFrame()

    for column in ["id", "customer_id", "msisdn1", "account_type", "reference_id", "currency_code", "ref_no", "description"]:
        if column not in frame.columns:
            frame[column] = ""
        frame[column] = clean_text(frame[column])
    frame["currency_code"] = frame["currency_code"].str.upper().replace("", "NON RENSEIGNEE")
    frame["created_at"] = pd.to_datetime(frame["created_at"], errors="coerce")
    frame = frame.dropna(subset=["created_at"]).copy()
    if frame.empty:
        return pd.DataFrame(), pd.DataFrame()
    for column in ["dr", "cr", "bal_before", "bal_after"]:
        frame[column] = numeric_column(frame, column)
    frame["__row_order"] = np.arange(len(frame))
    frame["__description_norm"] = frame["description"].apply(normalize_label)
    frame["__balance_delta_abs"] = (frame["bal_after"] - frame["bal_before"]).abs()
    frame["__ref_no_present"] = frame["ref_no"].replace("", pd.NA)
    frame["__reference_present"] = frame["reference_id"].replace("", pd.NA)
    frame["__retrait_vers_mpesa"] = frame["__description_norm"].str.contains(
        "retrait vers m-pesa", na=False
    )

    timestamp_keys = ["customer_id", "currency_code", "created_at"]
    reference_summary = (
        frame.groupby(timestamp_keys, as_index=False, dropna=False)
        .agg(
            ref_no_horodatage=("__ref_no_present", "first"),
            nombre_ref_no_horodatage=("__ref_no_present", "nunique"),
            retrait_vers_mpesa_horodatage=("__retrait_vers_mpesa", "max"),
            reference_horodatage=("__reference_present", "first"),
        )
    )
    reference_summary["ref_no_horodatage"] = reference_summary[
        "ref_no_horodatage"
    ].fillna("")
    reference_summary["reference_horodatage"] = reference_summary[
        "reference_horodatage"
    ].fillna("")
    frame = frame.merge(reference_summary, on=timestamp_keys, how="left")

    def event_reference(row: pd.Series) -> str:
        ref_count = int(row.get("nombre_ref_no_horodatage", 0) or 0)
        if ref_count == 1:
            return str(row.get("ref_no_horodatage", "")).strip()
        source_ref = str(row.get("ref_no", "")).strip()
        if ref_count > 1 and source_ref:
            return source_ref
        timestamp = pd.Timestamp(row["created_at"]).strftime("%Y%m%d%H%M%S%f")
        internal_references = str(row.get("reference_horodatage", "")).strip()
        if bool(row.get("retrait_vers_mpesa_horodatage", False)) and internal_references:
            return f"RETRAIT-{internal_references}-{timestamp}"
        prefix = "TURBO-MULTI" if ref_count > 1 else "TURBO"
        return f"{prefix}-{timestamp}"

    frame["event_reference"] = frame.apply(event_reference, axis=1)
    frame["event_key"] = (
        frame["customer_id"].astype(str)
        + "|" + frame["currency_code"].astype(str)
        + "|" + frame["created_at"].dt.strftime("%Y-%m-%d %H:%M:%S.%f")
        + "|" + frame["event_reference"].astype(str)
    )

    account = frame["account_type"]
    description = frame["__description_norm"]
    is_mpesa = account.eq("MPESA ACCOUNT")
    is_normal = account.eq("NORMAL SAVINGS")
    is_fixed = account.eq("FIXED SAVINGS")
    is_loan = account.eq("LOAN ACCOUNT")
    is_principle = account.eq("PRINCIPLE")

    component_rules: dict[str, pd.Series] = {
        "mpesa_debit_total": frame["dr"].where(is_mpesa, 0.0),
        "mpesa_credit_total": frame["cr"].where(is_mpesa, 0.0),
        "depot_normal_mpesa": frame["dr"].where(is_mpesa & description.str.contains("m-pesa depot", na=False), 0.0),
        "depot_normal_epargne": frame["cr"].where(is_normal & description.str.contains("epargne depot", na=False), 0.0),
        "depot_dat_mpesa": frame["dr"].where(is_mpesa & description.str.contains("m-pesa compte", na=False), 0.0),
        "depot_dat_compte": frame["cr"].where(is_fixed & description.str.contains("depot bloque", na=False), 0.0),
        "retrait_epargne_mpesa": frame["cr"].where(is_mpesa & description.str.contains("retrait vers m-pesa", na=False), 0.0),
        "retrait_epargne_compte": frame["dr"].where(is_normal & description.str.contains("retrait vers m-pesa", na=False), 0.0),
        "montant_decaisse_client": frame["cr"].where(is_mpesa & description.str.contains("montant pret", na=False), 0.0),
        "montant_decaisse_miroir": frame["cr"].where(account.eq("LOAN AMOUNT A/C") & description.str.contains("montant pret", na=False), 0.0),
        "dette_creee_observee": frame["cr"].where(is_loan & description.eq("compte de pret"), 0.0),
        "interet_observe": frame["cr"].where(account.eq("INTEREST EARNED") & description.str.contains("interet", na=False), 0.0),
        "principal_rembourse": frame["cr"].where(is_principle & description.str.contains("remboursement", na=False), 0.0),
        "remboursement_mpesa": frame["dr"].where(is_mpesa & description.str.contains("remboursement", na=False), 0.0),
        "penalite_observee": frame[["dr", "cr"]].max(axis=1).where(account.eq("LOAN PENALTY FEES"), 0.0),
        "transfert_dat_sortie": frame["__balance_delta_abs"].where(is_fixed & description.str.contains("retrait compte bloque", na=False), 0.0),
        "transfert_epargne_entree": frame["__balance_delta_abs"].where(is_normal & description.str.contains("retrait compte bloque", na=False), 0.0),
        "epargne_dat_remboursement": frame["__balance_delta_abs"].where(
            (is_normal | is_fixed) & description.str.contains("remboursement du compte bloque", na=False), 0.0
        ),
    }
    for column, values in component_rules.items():
        frame[column] = pd.to_numeric(values, errors="coerce").fillna(0.0)
    frame["__unknown_account_type"] = frame["account_type"].where(~frame["account_type"].isin(KNOWN_ACCOUNT_TYPES), "")
    frame["__line_zero"] = frame["dr"].eq(0) & frame["cr"].eq(0)
    frame["__negative_balance"] = frame["bal_after"].lt(0)

    event_keys = ["event_key", "customer_id", "currency_code", "created_at", "event_reference"]
    aggregations: dict[str, tuple[str, object]] = {
        "telephone": ("msisdn1", "first"),
        "ref_no": ("ref_no", "first"),
        "reference_ids": ("reference_id", concat_unique),
        "account_types": ("account_type", concat_unique),
        "descriptions": ("description", concat_unique),
        "nombre_lignes": ("id", "size"),
        "nombre_ref_no_horodatage": ("nombre_ref_no_horodatage", "max"),
        "total_debit_ecritures": ("dr", "sum"),
        "total_credit_ecritures": ("cr", "sum"),
        "lignes_mouvement_nul": ("__line_zero", "sum"),
        "lignes_solde_negatif": ("__negative_balance", "sum"),
        "types_comptes_inconnus": ("__unknown_account_type", concat_unique),
    }
    for column in component_rules:
        aggregation = "max" if column in {"transfert_dat_sortie", "transfert_epargne_entree", "epargne_dat_remboursement"} else "sum"
        aggregations[column] = (column, aggregation)
    if "Nom_client" in frame.columns:
        aggregations["Nom_client"] = ("Nom_client", "first")
    events = frame.groupby(event_keys, as_index=False, dropna=False).agg(**aggregations)
    if "Nom_client" not in events.columns:
        events["Nom_client"] = ""

    def classify_event(row: pd.Series) -> str:
        event_text = normalize_label(f"{row.get('descriptions', '')} {row.get('account_types', '')}")
        if "retrait compte bloque" in event_text and not any(
            token in event_text for token in ["loan account", "loan portfolio", "principle"]
        ):
            return "Transfert DAT vers epargne courante"
        movement_net = float(row.get("mpesa_credit_total", 0.0)) - float(row.get("mpesa_debit_total", 0.0))
        return classify_mpesa_operation(row.get("descriptions", ""), row.get("account_types", ""), movement_net)

    events["type_operation"] = events.apply(classify_event, axis=1)
    fallback_amount = events[["mpesa_debit_total", "mpesa_credit_total"]].max(axis=1)
    events["montant_operation"] = np.select(
        [
            events["type_operation"].eq("Sortie M-PESA_Turbo vers epargne"),
            events["type_operation"].eq("Sortie M-PESA_Turbo vers DAT"),
            events["type_operation"].eq("Entree M-PESA_Turbo depuis epargne"),
            events["type_operation"].eq("Decaissement de credit"),
            events["type_operation"].isin(["Remboursement de credit", "Remboursement avec penalite"]),
            events["type_operation"].eq("Transfert DAT vers epargne courante"),
        ],
        [
            events["depot_normal_mpesa"],
            events["depot_dat_mpesa"],
            events["retrait_epargne_mpesa"],
            events["montant_decaisse_client"],
            events[["remboursement_mpesa", "principal_rembourse"]].max(axis=1),
            events[["transfert_dat_sortie", "transfert_epargne_entree"]].max(axis=1),
        ],
        default=fallback_amount,
    )
    events["sens_metier"] = np.select(
        [
            events["type_operation"].isin(
                [
                    "Sortie M-PESA_Turbo vers epargne",
                    "Sortie M-PESA_Turbo vers DAT",
                    "Remboursement de credit",
                    "Remboursement avec penalite",
                ]
            ),
            events["type_operation"].isin(["Entree M-PESA_Turbo depuis epargne", "Decaissement de credit"]),
            events["type_operation"].eq("Transfert DAT vers epargne courante"),
        ],
        ["Entree Bisou Bisou", "Sortie Bisou Bisou", "Interne"],
        default="A verifier",
    )
    events["montant_entree_bisou"] = events["montant_operation"].where(events["sens_metier"].eq("Entree Bisou Bisou"), 0.0)
    events["montant_sortie_bisou"] = events["montant_operation"].where(events["sens_metier"].eq("Sortie Bisou Bisou"), 0.0)
    events["revenu_credit_observe"] = events["interet_observe"] + events["penalite_observee"]
    events["ecart_debit_credit_observe"] = events["total_debit_ecritures"] - events["total_credit_ecritures"]

    def repayment_mode(row: pd.Series) -> str:
        if row["type_operation"] not in {"Remboursement de credit", "Remboursement avec penalite"}:
            return ""
        sources: list[str] = []
        if float(row.get("remboursement_mpesa", 0.0)) > 0:
            sources.append("M-PESA_Turbo")
        if float(row.get("epargne_dat_remboursement", 0.0)) > 0:
            sources.append("Epargne/DAT")
        if float(row.get("penalite_observee", 0.0)) > 0:
            sources.append("Penalite")
        return " + ".join(sources) or "A verifier"

    events["mode_remboursement_observe"] = events.apply(repayment_mode, axis=1)

    def amount_control(row: pd.Series) -> tuple[str, str]:
        operation_type = row["type_operation"]
        pairs: list[tuple[float, float, str]] = []
        if operation_type == "Sortie M-PESA_Turbo vers epargne":
            pairs.append((row["depot_normal_mpesa"], row["depot_normal_epargne"], "Depot M-PESA contre epargne"))
        elif operation_type == "Sortie M-PESA_Turbo vers DAT":
            pairs.append((row["depot_dat_mpesa"], row["depot_dat_compte"], "Depot M-PESA contre DAT"))
        elif operation_type == "Entree M-PESA_Turbo depuis epargne":
            pairs.append((row["retrait_epargne_mpesa"], row["retrait_epargne_compte"], "Retrait M-PESA contre epargne"))
        elif operation_type == "Decaissement de credit":
            pairs.append((row["montant_decaisse_client"], row["montant_decaisse_miroir"], "Decaissement contre compte montant pret"))
            if row["dette_creee_observee"] > 0 and row["montant_decaisse_client"] > 0:
                pairs.append(
                    (
                        row["dette_creee_observee"],
                        row["montant_decaisse_client"] + row["interet_observe"],
                        "Dette creee contre capital verse plus interet",
                    )
                )
        elif operation_type in {"Remboursement de credit", "Remboursement avec penalite"}:
            pairs.append((row["remboursement_mpesa"], row["principal_rembourse"], "Remboursement M-PESA contre principal"))
        elif operation_type == "Transfert DAT vers epargne courante":
            pairs.append((row["transfert_dat_sortie"], row["transfert_epargne_entree"], "Sortie DAT contre entree epargne"))
        usable = [(float(left), float(right), label) for left, right, label in pairs if float(left) > 0 or float(right) > 0]
        if not usable:
            return "Non applicable", "Aucune paire metier stable a comparer"
        differences = []
        for left, right, label in usable:
            tolerance = max(0.01, max(abs(left), abs(right)) * 1e-6)
            if abs(left - right) > tolerance:
                differences.append(f"{label} : {left:.2f} contre {right:.2f}")
        if differences:
            return "A verifier", " | ".join(differences)
        return "Conforme", "Montants miroirs conformes"

    amount_checks = events.apply(amount_control, axis=1)
    events["controle_montant_operation"] = amount_checks.map(lambda value: value[0])
    events["detail_controle_montant"] = amount_checks.map(lambda value: value[1])

    def global_control(row: pd.Series) -> tuple[str, str]:
        reasons: list[str] = []
        operation_type = row["type_operation"]
        if row["controle_montant_operation"] == "A verifier":
            reasons.append(row["detail_controle_montant"])
        if int(row.get("nombre_ref_no_horodatage", 0)) > 1:
            reasons.append("Plusieurs ref_no au meme horodatage")
        if not str(row.get("ref_no", "")).strip() and operation_type in {
            "Sortie M-PESA_Turbo vers epargne",
            "Sortie M-PESA_Turbo vers DAT",
        }:
            reasons.append("ref_no absent pour un depot")
        if int(row.get("lignes_solde_negatif", 0)) > 0:
            reasons.append("Solde negatif observe")
        if str(row.get("types_comptes_inconnus", "")).strip():
            reasons.append(f"Type de compte inconnu : {row['types_comptes_inconnus']}")
        if float(row.get("montant_operation", 0.0)) == 0 and int(row.get("lignes_mouvement_nul", 0)) == int(row.get("nombre_lignes", 0)):
            reasons.append("Operation entierement nulle")
        return ("A verifier", " | ".join(reasons)) if reasons else ("Conforme", "Aucun ecart metier detecte")

    global_checks = events.apply(global_control, axis=1)
    events["statut_controle_turbo"] = global_checks.map(lambda value: value[0])
    events["observation_controle_turbo"] = global_checks.map(lambda value: value[1])
    return events.sort_values(["created_at", "event_reference"]).reset_index(drop=True), frame


def _build_customer_turbo_events(
    transactions: pd.DataFrame | None,
    customer_id: object,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Construit les evenements Turbo d'un client avec le moteur portefeuille."""
    return _build_turbo_operation_events(transactions, customer_id=customer_id)


def build_turbo_operation_events(
    transactions: pd.DataFrame | None,
) -> dict[str, pd.DataFrame]:
    """Expose le journal d'evenements Turbo reutilisable par les caches de l'UI."""
    events, lines = _build_turbo_operation_events(transactions)
    return {"events": events, "lines": lines}


def _turbo_period_bucket(values: pd.Series, frequency: str) -> pd.Series:
    """Ramene les dates Turbo au debut du jour, de la semaine ou du mois."""
    dates = pd.to_datetime(values, errors="coerce")
    normalized_frequency = normalize_label(frequency)
    if normalized_frequency in {"m", "mois", "mensuel", "monthly"}:
        return dates.dt.to_period("M").dt.start_time
    if normalized_frequency in {"w", "semaine", "hebdomadaire", "weekly"}:
        return dates.dt.to_period("W-SUN").dt.start_time
    return dates.dt.normalize()


def _turbo_amount_band(currency: object, amount: object) -> str:
    """Classe l'encours initial sans convertir ni additionner les devises."""
    code = str(currency).strip().upper()
    value = float(pd.to_numeric(pd.Series([amount]), errors="coerce").fillna(0).iloc[0])
    if code == "USD":
        if value < 100:
            return "Moins de 100"
        if value < 500:
            return "100 a 499,99"
        if value < 1_000:
            return "500 a 999,99"
        return "1 000 et plus"
    if code == "CDF":
        if value < 500_000:
            return "Moins de 500 000"
        if value < 2_000_000:
            return "500 000 a 1 999 999"
        if value < 5_000_000:
            return "2 000 000 a 4 999 999"
        return "5 000 000 et plus"
    return "Devise non parametree"


def _turbo_deposit_band(currency: object, amount: object) -> str:
    """Reprend les tranches de depots prioritaires de Perfect Vision."""
    code = str(currency).strip().upper()
    value = float(pd.to_numeric(pd.Series([amount]), errors="coerce").fillna(0).iloc[0])
    if code == "USD":
        if value < 10:
            return "Moins de 10"
        if value < 25:
            return "10 a 24,99"
        if value < 50:
            return "25 a 49,99"
        if value < 100:
            return "50 a 99,99"
        return "100 et plus"
    if code == "CDF":
        if value < 125_000:
            return "Moins de 125 000"
        if value < 200_000:
            return "125 000 a 199 999"
        return "200 000 et plus"
    return "Devise non parametree"


def build_mpesa_turbo_financial_analysis(
    prepared: MpesaPreparedData,
    *,
    date_start: Any | None = None,
    date_end: Any | None = None,
    frequency: str = "Jour",
    dat_annual_interest_rate_pct: float | None = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
    fractionation_thresholds: dict[str, float] | None = None,
    large_transaction_thresholds: dict[str, float] | None = None,
    turbo_events: pd.DataFrame | None = None,
    turbo_transaction_lines: pd.DataFrame | None = None,
) -> dict[str, Any]:
    """Construit le cockpit financier exclusivement depuis les sources Turbo.

    Les montants, soldes, credits, DAT et remboursements ne proviennent jamais de
    G2. G2 demeure disponible ailleurs pour enrichir l'identite et documenter le
    rapprochement des ecritures, sans modifier un seul indicateur de ce rapport.
    Les positions Credits/Epargne/DAT sont des instantanes du portail Turbo ; les
    flux de periode sont reconstruits depuis Transactions M-PESA_Turbo.
    """
    transactions = prepared.transactions if isinstance(prepared.transactions, pd.DataFrame) else pd.DataFrame()
    transaction_dates = (
        pd.to_datetime(transactions["created_at"], errors="coerce").dropna()
        if not transactions.empty and "created_at" in transactions.columns
        else pd.Series(dtype="datetime64[ns]")
    )
    fallback_end = _mpesa_analysis_date(prepared, date_end)
    end_date = pd.Timestamp(fallback_end).normalize()
    if date_start is not None and pd.notna(pd.to_datetime(date_start, errors="coerce")):
        start_date = pd.Timestamp(pd.to_datetime(date_start, errors="coerce")).normalize()
    elif not transaction_dates.empty:
        start_date = pd.Timestamp(transaction_dates.min()).normalize()
    else:
        start_date = end_date
    if start_date > end_date:
        start_date, end_date = end_date, start_date
    period_end_exclusive = end_date + pd.Timedelta(days=1)

    fractionation_limits = {"CDF": 14_000_000.0, "USD": 5_000.0}
    if fractionation_thresholds:
        fractionation_limits.update(
            {str(key).upper(): float(value) for key, value in fractionation_thresholds.items()}
        )
    large_limits = {"CDF": 28_000_000.0, "USD": 10_000.0}
    if large_transaction_thresholds:
        large_limits.update(
            {str(key).upper(): float(value) for key, value in large_transaction_thresholds.items()}
        )

    if isinstance(turbo_events, pd.DataFrame) and isinstance(
        turbo_transaction_lines, pd.DataFrame
    ):
        all_events = turbo_events.copy()
        transaction_lines = turbo_transaction_lines.copy()
    else:
        all_events, transaction_lines = _build_turbo_operation_events(transactions)
    if all_events.empty:
        period_events = pd.DataFrame()
    else:
        all_events = all_events.copy()
        all_events["created_at"] = pd.to_datetime(all_events["created_at"], errors="coerce")
        all_events["date_operation"] = all_events["created_at"].dt.normalize()
        period_events = all_events.loc[
            all_events["created_at"].ge(start_date)
            & all_events["created_at"].lt(period_end_exclusive)
        ].copy()
        period_events["periode_analyse"] = _turbo_period_bucket(
            period_events["created_at"], frequency
        )

    flow_numeric_columns = [
        "montant_entree_bisou",
        "montant_sortie_bisou",
        "depot_normal_mpesa",
        "depot_dat_mpesa",
        "retrait_epargne_mpesa",
        "montant_decaisse_client",
        "remboursement_mpesa",
        "principal_rembourse",
        "interet_observe",
        "penalite_observee",
        "revenu_credit_observe",
    ]
    if not period_events.empty:
        for column in flow_numeric_columns:
            period_events[column] = numeric_column(period_events, column)
        period_events["flux_net_bisou"] = (
            period_events["montant_entree_bisou"] - period_events["montant_sortie_bisou"]
        )
        flow_aggregations: dict[str, tuple[str, object]] = {
            "nombre_operations": ("event_key", "nunique"),
            "nombre_clients": (
                "customer_id",
                lambda values: clean_identifier(values).replace("", pd.NA).nunique(),
            ),
            "montant_entrees": ("montant_entree_bisou", "sum"),
            "montant_sorties": ("montant_sortie_bisou", "sum"),
            "flux_net": ("flux_net_bisou", "sum"),
            "depots_epargne_courante": ("depot_normal_mpesa", "sum"),
            "depots_dat": ("depot_dat_mpesa", "sum"),
            "retraits_epargne": ("retrait_epargne_mpesa", "sum"),
            "nouveaux_credits_decaissements": ("montant_decaisse_client", "sum"),
            "remboursements_observes": ("remboursement_mpesa", "sum"),
            "principal_rembourse_observe": ("principal_rembourse", "sum"),
            "interets_credit_observes": ("interet_observe", "sum"),
            "penalites_observees": ("penalite_observee", "sum"),
        }
        flow_summary = (
            period_events.groupby("currency_code", as_index=False, dropna=False)
            .agg(**flow_aggregations)
            .sort_values("currency_code")
            .reset_index(drop=True)
        )
        flow_evolution = (
            period_events.groupby(
                ["periode_analyse", "currency_code"], as_index=False, dropna=False
            )
            .agg(**flow_aggregations)
            .sort_values(["currency_code", "periode_analyse"])
            .reset_index(drop=True)
        )
    else:
        flow_summary = pd.DataFrame()
        flow_evolution = pd.DataFrame()

    repayment_types = {"Remboursement de credit", "Remboursement avec penalite"}
    repayments = (
        period_events.loc[period_events["type_operation"].isin(repayment_types)].copy()
        if not period_events.empty
        else pd.DataFrame()
    )
    if not repayments.empty:
        repayment_summary = (
            repayments.groupby("currency_code", as_index=False, dropna=False)
            .agg(
                nombre_remboursements=("event_key", "nunique"),
                nombre_clients=("customer_id", "nunique"),
                montant_rembourse=("montant_operation", "sum"),
                principal_observe=("principal_rembourse", "sum"),
                interets_observes=("interet_observe", "sum"),
                penalites_observees=("penalite_observee", "sum"),
                remboursements_a_verifier=(
                    "statut_controle_turbo",
                    lambda values: int(pd.Series(values).eq("A verifier").sum()),
                ),
            )
            .sort_values("currency_code")
            .reset_index(drop=True)
        )
    else:
        repayment_summary = pd.DataFrame()

    new_credit_events = (
        period_events.loc[period_events["type_operation"].eq("Decaissement de credit")].copy()
        if not period_events.empty
        else pd.DataFrame()
    )

    loans_source = prepared.loans.copy() if isinstance(prepared.loans, pd.DataFrame) else pd.DataFrame()
    if not loans_source.empty:
        loans_source["created_at"] = pd.to_datetime(
            loans_source.get("created_at", pd.Series(pd.NaT, index=loans_source.index)),
            errors="coerce",
        )
        loans_for_position = loans_source.loc[
            loans_source["created_at"].isna() | loans_source["created_at"].lt(period_end_exclusive)
        ].copy()
    else:
        loans_for_position = pd.DataFrame()
    credit_risk = build_mpesa_credit_risk_analysis(loans_for_position, as_of_date=end_date)
    credit_detail = credit_risk["detail"].copy()
    credit_summary = credit_risk["synthese"].copy()

    new_loan_accounts = pd.DataFrame()
    if not loans_source.empty:
        new_loan_accounts = loans_source.loc[
            loans_source["created_at"].ge(start_date)
            & loans_source["created_at"].lt(period_end_exclusive)
        ].copy()
        if not new_loan_accounts.empty:
            new_loan_accounts["loan_id"] = clean_identifier(
                new_loan_accounts.get("loan_id", pd.Series("", index=new_loan_accounts.index))
            )
            new_loan_accounts["currency_code"] = clean_text(
                new_loan_accounts.get("currency_code", pd.Series("", index=new_loan_accounts.index))
            ).str.upper().replace("", "NON RENSEIGNEE")
            new_loan_accounts["loan_amount"] = numeric_column(new_loan_accounts, "loan_amount")
            new_loan_accounts = new_loan_accounts.drop_duplicates("loan_id", keep="last")

    event_new_summary = pd.DataFrame()
    if not new_credit_events.empty:
        event_new_summary = new_credit_events.groupby(
            "currency_code", as_index=False, dropna=False
        ).agg(
            nombre_decaissements_turbo=("event_key", "nunique"),
            clients_decaissements_turbo=("customer_id", "nunique"),
            montant_decaisse_turbo=("montant_operation", "sum"),
        )
    account_new_summary = pd.DataFrame()
    if not new_loan_accounts.empty:
        account_new_summary = new_loan_accounts.groupby(
            "currency_code", as_index=False, dropna=False
        ).agg(
            nouveaux_comptes_credit=("loan_id", "nunique"),
            clients_nouveaux_comptes=("customer_id", "nunique"),
            montant_initial_comptes=("loan_amount", "sum"),
        )
    if event_new_summary.empty:
        new_credit_summary = account_new_summary
    elif account_new_summary.empty:
        new_credit_summary = event_new_summary
    else:
        new_credit_summary = event_new_summary.merge(
            account_new_summary, on="currency_code", how="outer"
        )
    if not new_credit_summary.empty:
        for column in [
            "nombre_decaissements_turbo",
            "clients_decaissements_turbo",
            "montant_decaisse_turbo",
            "nouveaux_comptes_credit",
            "clients_nouveaux_comptes",
            "montant_initial_comptes",
        ]:
            if column not in new_credit_summary.columns:
                new_credit_summary[column] = 0.0
            new_credit_summary[column] = pd.to_numeric(
                new_credit_summary[column], errors="coerce"
            ).fillna(0.0)
        new_credit_summary["ecart_decaissement_comptes"] = (
            new_credit_summary["montant_decaisse_turbo"]
            - new_credit_summary["montant_initial_comptes"]
        )
        new_credit_summary["statut_rapprochement"] = np.where(
            new_credit_summary["ecart_decaissement_comptes"].abs().le(
                np.maximum(
                    0.01,
                    new_credit_summary[
                        ["montant_decaisse_turbo", "montant_initial_comptes"]
                    ].abs().max(axis=1)
                    * 1e-6,
                )
            ),
            "Conforme",
            "A expliquer",
        )
        new_credit_summary = new_credit_summary.sort_values("currency_code").reset_index(drop=True)

    credit_concentration_clients = pd.DataFrame()
    credit_concentration_summary = pd.DataFrame()
    par_by_amount = pd.DataFrame()
    if not credit_detail.empty:
        active_credit = credit_detail.loc[credit_detail["encours_total"].gt(0)].copy()
        if not active_credit.empty:
            credit_concentration_clients = (
                active_credit.groupby(
                    ["currency_code", "customer_id"], as_index=False, dropna=False
                )
                .agg(
                    nombre_credits=("loan_id", "nunique"),
                    encours_total=("encours_total", "sum"),
                    encours_retard_1j=(
                        "encours_total",
                        lambda values: float(
                            values.loc[active_credit.loc[values.index, "jours_retard"].ge(1)].sum()
                        ),
                    ),
                )
            )
            credit_concentration_clients["rang_encours"] = (
                credit_concentration_clients.groupby("currency_code")["encours_total"]
                .rank(method="first", ascending=False)
                .astype(int)
            )
            concentration_rows: list[dict[str, Any]] = []
            for currency, group in credit_concentration_clients.groupby("currency_code"):
                ordered = group.sort_values("encours_total", ascending=False)
                total = float(ordered["encours_total"].sum())
                top_ten_percent_count = max(1, int(np.ceil(len(ordered) * 0.10)))
                concentration_rows.append(
                    {
                        "currency_code": currency,
                        "nombre_clients": int(len(ordered)),
                        "encours_total": total,
                        "part_top_1_client_pct": float(ordered.head(1)["encours_total"].sum() / total * 100) if total else np.nan,
                        "part_top_5_clients_pct": float(ordered.head(5)["encours_total"].sum() / total * 100) if total else np.nan,
                        "part_top_10_clients_pct": float(ordered.head(10)["encours_total"].sum() / total * 100) if total else np.nan,
                        "part_top_10_pct_clients_pct": float(ordered.head(top_ten_percent_count)["encours_total"].sum() / total * 100) if total else np.nan,
                    }
                )
            credit_concentration_summary = pd.DataFrame(concentration_rows)

        band_detail = credit_detail.copy()
        band_detail["tranche_montant_credit"] = band_detail.apply(
            lambda row: _turbo_amount_band(row["currency_code"], row["loan_amount"]), axis=1
        )
        par_by_amount = (
            band_detail.groupby(
                ["currency_code", "tranche_montant_credit"], as_index=False, dropna=False
            )
            .agg(
                nombre_credits=("loan_id", "nunique"),
                montant_initial=("loan_amount", "sum"),
                encours_total=("encours_total", "sum"),
                encours_retard_1j=(
                    "encours_total",
                    lambda values: float(
                        values.loc[band_detail.loc[values.index, "jours_retard"].ge(1)].sum()
                    ),
                ),
                encours_retard_7j=(
                    "encours_total",
                    lambda values: float(
                        values.loc[band_detail.loc[values.index, "jours_retard"].ge(7)].sum()
                    ),
                ),
                encours_retard_30j=(
                    "encours_total",
                    lambda values: float(
                        values.loc[band_detail.loc[values.index, "jours_retard"].ge(30)].sum()
                    ),
                ),
            )
        )
        denominator = par_by_amount["encours_total"].replace(0, pd.NA)
        for days in [1, 7, 30]:
            par_by_amount[f"par_{days}j_pct"] = (
                par_by_amount[f"encours_retard_{days}j"].div(denominator).mul(100)
            )

    savings_types = {
        "Sortie M-PESA_Turbo vers epargne",
        "Sortie M-PESA_Turbo vers DAT",
        "Entree M-PESA_Turbo depuis epargne",
    }
    savings_events = (
        period_events.loc[period_events["type_operation"].isin(savings_types)].copy()
        if not period_events.empty
        else pd.DataFrame()
    )
    savings_activity = pd.DataFrame()
    weekly_deposits = pd.DataFrame()
    deposit_bands = pd.DataFrame()
    if not savings_events.empty:
        savings_events["depot_courant"] = savings_events["montant_operation"].where(
            savings_events["type_operation"].eq("Sortie M-PESA_Turbo vers epargne"), 0.0
        )
        savings_events["depot_dat"] = savings_events["montant_operation"].where(
            savings_events["type_operation"].eq("Sortie M-PESA_Turbo vers DAT"), 0.0
        )
        savings_events["retrait_epargne"] = savings_events["montant_operation"].where(
            savings_events["type_operation"].eq("Entree M-PESA_Turbo depuis epargne"), 0.0
        )
        savings_activity = (
            savings_events.groupby(
                ["currency_code", "customer_id"], as_index=False, dropna=False
            )
            .agg(
                telephone=("telephone", concat_unique),
                Nom_client=("Nom_client", concat_unique),
                nombre_operations=("event_key", "nunique"),
                depots_epargne_courante=("depot_courant", "sum"),
                depots_dat=("depot_dat", "sum"),
                retraits_epargne=("retrait_epargne", "sum"),
                premiere_operation=("created_at", "min"),
                derniere_operation=("created_at", "max"),
            )
        )
        savings_activity["flux_net_epargne"] = (
            savings_activity["depots_epargne_courante"]
            + savings_activity["depots_dat"]
            - savings_activity["retraits_epargne"]
        )
        deposits = savings_events.loc[
            savings_events["type_operation"].isin(
                ["Sortie M-PESA_Turbo vers epargne", "Sortie M-PESA_Turbo vers DAT"]
            )
        ].copy()
        if not deposits.empty:
            deposits["semaine"] = deposits["created_at"].dt.to_period("W-SUN").dt.start_time
            weekly_deposits = (
                deposits.groupby(
                    ["semaine", "currency_code", "customer_id"], as_index=False, dropna=False
                )
                .agg(
                    telephone=("telephone", concat_unique),
                    nombre_depots=("event_key", "nunique"),
                    montant_depose=("montant_operation", "sum"),
                )
            )
            weekly_deposits["deposant_frequent_3_plus"] = weekly_deposits["nombre_depots"].ge(3)
            deposits["tranche_depot"] = deposits.apply(
                lambda row: _turbo_deposit_band(row["currency_code"], row["montant_operation"]), axis=1
            )
            deposit_bands = (
                deposits.groupby(
                    ["currency_code", "tranche_depot"], as_index=False, dropna=False
                )
                .agg(
                    nombre_depots=("event_key", "nunique"),
                    nombre_clients=("customer_id", "nunique"),
                    montant_depose=("montant_operation", "sum"),
                )
            )

    transaction_concentration_clients = pd.DataFrame()
    transaction_concentration_summary = pd.DataFrame()
    if not period_events.empty:
        transaction_concentration_clients = (
            period_events.groupby(
                ["currency_code", "customer_id"], as_index=False, dropna=False
            )
            .agg(
                telephone=("telephone", concat_unique),
                nombre_operations=("event_key", "nunique"),
                volume_total=("montant_operation", "sum"),
                montant_entrees=("montant_entree_bisou", "sum"),
                montant_sorties=("montant_sortie_bisou", "sum"),
            )
        )
        transaction_concentration_clients["rang_volume"] = (
            transaction_concentration_clients.groupby("currency_code")["volume_total"]
            .rank(method="first", ascending=False)
            .astype(int)
        )
        concentration_rows = []
        for currency, group in transaction_concentration_clients.groupby("currency_code"):
            ordered = group.sort_values("volume_total", ascending=False)
            total = float(ordered["volume_total"].sum())
            top_count = max(1, int(np.ceil(len(ordered) * 0.10)))
            concentration_rows.append(
                {
                    "currency_code": currency,
                    "nombre_clients": int(len(ordered)),
                    "volume_total": total,
                    "part_top_5_clients_pct": float(ordered.head(5)["volume_total"].sum() / total * 100) if total else np.nan,
                    "part_top_10_clients_pct": float(ordered.head(10)["volume_total"].sum() / total * 100) if total else np.nan,
                    "part_top_10_pct_clients_pct": float(ordered.head(top_count)["volume_total"].sum() / total * 100) if total else np.nan,
                }
            )
        transaction_concentration_summary = pd.DataFrame(concentration_rows)

    alerts: list[pd.DataFrame] = []
    if not period_events.empty:
        controls = period_events.loc[period_events["statut_controle_turbo"].eq("A verifier")].copy()
        if not controls.empty:
            alerts.append(
                pd.DataFrame(
                    {
                        "alerte": "Controle comptable Turbo",
                        "customer_id": controls["customer_id"],
                        "currency_code": controls["currency_code"],
                        "date_reference": controls["created_at"],
                        "montant": controls["montant_operation"],
                        "nombre_operations": 1,
                        "ratio_activite": np.nan,
                        "detail": controls["observation_controle_turbo"],
                        "event_reference": controls["event_reference"],
                    }
                )
            )
        thresholds = period_events["currency_code"].map(large_limits)
        large_mask = thresholds.notna() & period_events["montant_operation"].ge(thresholds)
        large_events = period_events.loc[large_mask].copy()
        if not large_events.empty:
            alerts.append(
                pd.DataFrame(
                    {
                        "alerte": "Transaction importante",
                        "customer_id": large_events["customer_id"],
                        "currency_code": large_events["currency_code"],
                        "date_reference": large_events["created_at"],
                        "montant": large_events["montant_operation"],
                        "nombre_operations": 1,
                        "ratio_activite": np.nan,
                        "detail": "Seuil de transaction importante atteint",
                        "event_reference": large_events["event_reference"],
                    }
                )
            )
        fractionation = period_events.loc[period_events["montant_operation"].gt(0)].copy()
        fractionation["seuil_fractionnement"] = fractionation["currency_code"].map(fractionation_limits)
        fractionation = fractionation.loc[
            fractionation["seuil_fractionnement"].notna()
            & fractionation["montant_operation"].lt(fractionation["seuil_fractionnement"])
        ]
        if not fractionation.empty:
            fractionation_daily = (
                fractionation.groupby(
                    ["date_operation", "currency_code", "customer_id"],
                    as_index=False,
                    dropna=False,
                )
                .agg(
                    montant_cumule=("montant_operation", "sum"),
                    nombre_operations=("event_key", "nunique"),
                    seuil_fractionnement=("seuil_fractionnement", "max"),
                    references=("event_reference", concat_unique),
                )
            )
            fractionation_daily = fractionation_daily.loc[
                fractionation_daily["nombre_operations"].ge(2)
                & fractionation_daily["montant_cumule"].ge(
                    fractionation_daily["seuil_fractionnement"]
                )
            ]
            if not fractionation_daily.empty:
                alerts.append(
                    pd.DataFrame(
                        {
                            "alerte": "Fractionnement potentiel",
                            "customer_id": fractionation_daily["customer_id"],
                            "currency_code": fractionation_daily["currency_code"],
                            "date_reference": fractionation_daily["date_operation"],
                            "montant": fractionation_daily["montant_cumule"],
                            "nombre_operations": fractionation_daily["nombre_operations"],
                            "ratio_activite": np.nan,
                            "detail": "Cumul journalier au-dessus du seuil avec operations unitaires sous le seuil",
                            "event_reference": fractionation_daily["references"],
                        }
                    )
                )

        baseline_start = start_date - pd.Timedelta(days=90)
        baseline = all_events.loc[
            all_events["created_at"].ge(baseline_start)
            & all_events["created_at"].lt(start_date)
        ].copy()
        current_activity = (
            period_events.groupby(["currency_code", "customer_id"], as_index=False)
            .agg(
                montant_periode=("montant_operation", "sum"),
                operations_periode=("event_key", "nunique"),
            )
        )
        period_days = max(1, int((end_date - start_date).days) + 1)
        current_activity["moyenne_journaliere_periode"] = current_activity["montant_periode"] / period_days
        if baseline.empty:
            current_activity["moyenne_journaliere_90j"] = 0.0
        else:
            baseline_activity = (
                baseline.groupby(["currency_code", "customer_id"], as_index=False)
                .agg(montant_90j=("montant_operation", "sum"))
            )
            baseline_activity["moyenne_journaliere_90j"] = baseline_activity["montant_90j"] / 90.0
            current_activity = current_activity.merge(
                baseline_activity[["currency_code", "customer_id", "moyenne_journaliere_90j"]],
                on=["currency_code", "customer_id"],
                how="left",
            )
            current_activity["moyenne_journaliere_90j"] = current_activity[
                "moyenne_journaliere_90j"
            ].fillna(0.0)
        current_activity["ratio_activite"] = current_activity["moyenne_journaliere_periode"].div(
            current_activity["moyenne_journaliere_90j"].replace(0, pd.NA)
        )
        minimum_alert = current_activity["currency_code"].map(large_limits).fillna(np.inf) * 0.5
        unusual_mask = (
            current_activity["ratio_activite"].ge(3)
            | current_activity["moyenne_journaliere_90j"].eq(0)
        ) & (
            current_activity["operations_periode"].ge(3)
            | current_activity["montant_periode"].ge(minimum_alert)
        )
        unusual = current_activity.loc[unusual_mask].copy()
        if not unusual.empty:
            alerts.append(
                pd.DataFrame(
                    {
                        "alerte": "Activite inhabituelle vs 90 jours",
                        "customer_id": unusual["customer_id"],
                        "currency_code": unusual["currency_code"],
                        "date_reference": end_date,
                        "montant": unusual["montant_periode"],
                        "nombre_operations": unusual["operations_periode"],
                        "ratio_activite": unusual["ratio_activite"],
                        "detail": "Moyenne journaliere de la periode comparee aux 90 jours precedents",
                        "event_reference": "",
                    }
                )
            )
    alerts_frame = (
        pd.concat(alerts, ignore_index=True, sort=False)
        .sort_values(["currency_code", "montant"], ascending=[True, False])
        .reset_index(drop=True)
        if alerts
        else pd.DataFrame()
    )

    inactive_movements = pd.DataFrame()
    if not transaction_lines.empty:
        scoped_lines = transaction_lines.loc[
            transaction_lines["created_at"].ge(start_date)
            & transaction_lines["created_at"].lt(period_end_exclusive)
        ].copy()
        account_rows: list[pd.DataFrame] = []
        for family, source in [
            ("Epargne courante", prepared.current_savings),
            ("DAT", prepared.fixed_savings),
        ]:
            if not isinstance(source, pd.DataFrame) or source.empty:
                continue
            accounts = source.copy()
            accounts["customer_id"] = clean_identifier(
                accounts.get("customer_id", pd.Series("", index=accounts.index))
            )
            accounts["currency_code"] = clean_text(
                accounts.get("currency_code", pd.Series("", index=accounts.index))
            ).str.upper()
            accounts["status_compte"] = clean_text(
                accounts.get("status", pd.Series("", index=accounts.index))
            )
            accounts["date_closed"] = pd.to_datetime(
                accounts.get("date_closed", pd.Series(pd.NaT, index=accounts.index)),
                errors="coerce",
            )
            for reference_column in ["id", "savings_id"]:
                if reference_column not in accounts.columns:
                    continue
                references = accounts[
                    ["customer_id", "currency_code", "status_compte", "date_closed", reference_column]
                ].copy()
                references["reference_id"] = clean_identifier(references[reference_column])
                references["famille_compte"] = family
                references = references.loc[references["reference_id"].ne("")]
                account_rows.append(
                    references.drop(columns=[reference_column]).drop_duplicates(
                        ["customer_id", "currency_code", "reference_id"]
                    )
                )
        if account_rows and not scoped_lines.empty:
            account_refs = pd.concat(account_rows, ignore_index=True).drop_duplicates(
                ["customer_id", "currency_code", "reference_id"], keep="last"
            )
            scoped_lines["reference_id"] = clean_identifier(scoped_lines["reference_id"])
            scoped_lines["status_normalise"] = ""
            merged_lines = scoped_lines.merge(
                account_refs,
                on=["customer_id", "currency_code", "reference_id"],
                how="inner",
            )
            merged_lines["status_normalise"] = merged_lines["status_compte"].apply(normalize_label)
            active_status = merged_lines["status_normalise"].isin(
                ["", "open", "active", "approved", "activated"]
            )
            inactive_mask = (~active_status) | (
                merged_lines["date_closed"].notna()
                & merged_lines["created_at"].ge(merged_lines["date_closed"])
            )
            inactive_movements = merged_lines.loc[inactive_mask].copy()
            if not inactive_movements.empty:
                inactive_movements["montant_mouvement"] = inactive_movements[["dr", "cr"]].max(axis=1)
                inactive_movements = inactive_movements[
                    [
                        "created_at", "customer_id", "msisdn1", "currency_code",
                        "reference_id", "famille_compte", "status_compte", "date_closed",
                        "description", "dr", "cr", "montant_mouvement", "event_reference",
                    ]
                ].sort_values(["currency_code", "created_at"])

    dat_without_credit = pd.DataFrame()
    fixed = prepared.fixed_savings.copy() if isinstance(prepared.fixed_savings, pd.DataFrame) else pd.DataFrame()
    if not fixed.empty:
        fixed["customer_id"] = clean_identifier(
            fixed.get("customer_id", pd.Series("", index=fixed.index))
        )
        fixed["currency_code"] = clean_text(
            fixed.get("currency_code", pd.Series("", index=fixed.index))
        ).str.upper().replace("", "NON RENSEIGNEE")
        fixed["balance"] = numeric_column(fixed, "balance")
        fixed["maturity_date"] = pd.to_datetime(
            fixed.get("maturity_date", pd.Series(pd.NaT, index=fixed.index)), errors="coerce"
        )
        positive_dat = fixed.loc[fixed["balance"].gt(0) & fixed["customer_id"].ne("")].copy()
        if not positive_dat.empty:
            active_credit_keys = (
                credit_detail.loc[credit_detail["encours_total"].gt(0), ["customer_id", "currency_code"]]
                .drop_duplicates()
                .assign(credit_actif=True)
                if not credit_detail.empty
                else pd.DataFrame(columns=["customer_id", "currency_code", "credit_actif"])
            )
            positive_dat = positive_dat.merge(
                active_credit_keys,
                on=["customer_id", "currency_code"],
                how="left",
            )
            dat_without_credit = positive_dat.loc[positive_dat["credit_actif"].isna()].drop(
                columns=["credit_actif"]
            )

    savings_reconciliation = build_loan_savings_reconciliation(
        loans_for_position,
        prepared.current_savings,
        prepared.fixed_savings,
    )
    credit_savings_clients = savings_reconciliation.get("clients", pd.DataFrame()).copy()
    if not credit_savings_clients.empty:
        outstanding_column = next(
            (
                column
                for column in ["encours_credit", "encours_total", "loan_balance"]
                if column in credit_savings_clients.columns
            ),
            None,
        )
        if outstanding_column:
            credit_savings_clients = credit_savings_clients.loc[
                pd.to_numeric(credit_savings_clients[outstanding_column], errors="coerce").fillna(0).gt(0)
            ].copy()

    dat_maturity = build_mpesa_dat_maturity_analysis(
        prepared.fixed_savings,
        as_of_date=end_date,
        annual_interest_rate_pct=dat_annual_interest_rate_pct,
    )

    customer_quality = pd.DataFrame()
    customer_quality_detail = pd.DataFrame()
    if isinstance(prepared.customers, pd.DataFrame) and not prepared.customers.empty:
        customer_quality_detail = prepared.customers.copy()
        source_phone = customer_quality_detail.get(
            "msisdn1", pd.Series("", index=customer_quality_detail.index)
        )
        customer_quality_detail["telephone_normalise"] = normalize_phone(source_phone)
        customer_quality_detail["telephone_valide"] = customer_quality_detail[
            "telephone_normalise"
        ].str.fullmatch(r"243\d{9}", na=False)
        customer_quality = pd.DataFrame(
            [
                {
                    "source": "Clients_Turbo",
                    "nombre_lignes": int(len(customer_quality_detail)),
                    "telephones_valides": int(customer_quality_detail["telephone_valide"].sum()),
                    "telephones_invalides_ou_absents": int((~customer_quality_detail["telephone_valide"]).sum()),
                    "telephones_distincts_valides": int(
                        customer_quality_detail.loc[
                            customer_quality_detail["telephone_valide"], "telephone_normalise"
                        ].nunique()
                    ),
                }
            ]
        )
        customer_quality_detail = customer_quality_detail.loc[
            ~customer_quality_detail["telephone_valide"]
        ].reset_index(drop=True)

    source_rows: list[dict[str, Any]] = []
    for source_name, frame, date_column, role in [
        ("Transactions M-PESA_Turbo", prepared.transactions, "created_at", "Flux et ecritures"),
        ("Savings Account_Turbo", pd.concat([prepared.current_savings, prepared.fixed_savings], ignore_index=True, sort=False), "updated_at", "Positions epargne et DAT"),
        ("Loans Account_Turbo", prepared.loans, "updated_at", "Positions de credit"),
        ("Customers_Turbo", prepared.customers, "created_at", "Referentiel client"),
        ("Transactions M-PESA_G2", prepared.g2_transactions, "completion_time", "Identite et preuve de rapprochement uniquement"),
    ]:
        available = isinstance(frame, pd.DataFrame) and not frame.empty
        dates = (
            pd.to_datetime(frame[date_column], errors="coerce").dropna()
            if available and date_column in frame.columns
            else pd.Series(dtype="datetime64[ns]")
        )
        source_rows.append(
            {
                "source": source_name,
                "role": role,
                "intervient_dans_les_montants": not source_name.endswith("G2"),
                "disponible": available,
                "nombre_lignes": int(len(frame)) if available else 0,
                "date_min": dates.min() if not dates.empty else pd.NaT,
                "date_max": dates.max() if not dates.empty else pd.NaT,
            }
        )

    definitions = pd.DataFrame(
        [
            {"indicateur": "Flux de periode", "source": "Transactions M-PESA_Turbo", "definition": "Evenements regroupes par ref_no, sinon client + devise + horodatage."},
            {"indicateur": "Remboursements observes", "source": "Transactions M-PESA_Turbo", "definition": "Ecritures classees remboursement ; principal, interet et penalite restent observes separement."},
            {"indicateur": "Nouveaux credits", "source": "Transactions et Loans Account_Turbo", "definition": "Decaissements observes rapproches des comptes de credit crees dans la periode."},
            {"indicateur": "Encours et PAR", "source": "Loans Account_Turbo", "definition": "Position instantanee ; retard simplifie depuis due_date faute de plan d'amortissement detaille."},
            {"indicateur": "Epargne et DAT", "source": "Savings Account_Turbo", "definition": "Position instantanee, sans compensation avec l'encours de credit."},
            {"indicateur": "G2", "source": "Transactions M-PESA_G2", "definition": "Enrichissement du nom et preuve de rapprochement uniquement ; aucun calcul financier."},
        ]
    )

    return {
        "date_debut": start_date,
        "date_fin": end_date,
        "frequence": frequency,
        "sources": pd.DataFrame(source_rows),
        "definitions": definitions,
        "operations_turbo": period_events.reset_index(drop=True),
        "lignes_turbo_periode": transaction_lines.loc[
            transaction_lines["created_at"].ge(start_date)
            & transaction_lines["created_at"].lt(period_end_exclusive)
        ].reset_index(drop=True) if not transaction_lines.empty else pd.DataFrame(),
        "flux_synthese": flow_summary,
        "flux_evolution": flow_evolution,
        "remboursements_synthese": repayment_summary,
        "remboursements_detail": repayments.reset_index(drop=True),
        "nouveaux_credits_synthese": new_credit_summary,
        "nouveaux_credits_detail": new_credit_events.reset_index(drop=True),
        "nouveaux_comptes_credit": new_loan_accounts.reset_index(drop=True),
        "credit_synthese": credit_summary,
        "credit_detail": credit_detail,
        "par_tranches_montant": par_by_amount,
        "concentration_credit_synthese": credit_concentration_summary,
        "concentration_credit_clients": credit_concentration_clients,
        "activite_epargne_clients": savings_activity,
        "depots_frequents_hebdo": weekly_deposits,
        "tranches_depots": deposit_bands,
        "concentration_transactions_synthese": transaction_concentration_summary,
        "concentration_transactions_clients": transaction_concentration_clients,
        "alertes_transactions": alerts_frame,
        "controles_operations": period_events.loc[
            period_events["statut_controle_turbo"].eq("A verifier")
        ].reset_index(drop=True) if not period_events.empty else pd.DataFrame(),
        "mouvements_comptes_inactifs": inactive_movements.reset_index(drop=True),
        "dat_sans_credit_actif": dat_without_credit.reset_index(drop=True),
        "credits_epargne_disponible": credit_savings_clients.reset_index(drop=True),
        "rapprochement_credit_epargne_synthese": savings_reconciliation.get("synthese", pd.DataFrame()),
        "rapprochement_credit_epargne_controles": savings_reconciliation.get("controles", pd.DataFrame()),
        "dat_echeances_synthese": dat_maturity.get("synthese", pd.DataFrame()),
        "dat_echeances_detail": dat_maturity.get("detail", pd.DataFrame()),
        "qualite_clients_synthese": customer_quality,
        "qualite_clients_detail": customer_quality_detail,
    }


def _build_customer_observed_positions(
    prepared: MpesaPreparedData,
    transaction_lines: pd.DataFrame,
    customer_id: object,
    *,
    comparison_allowed: bool,
) -> pd.DataFrame:
    position_columns = [
        "customer_id", "currency_code", "famille_position", "source_observation",
        "solde_transactions_observe", "date_derniere_ecriture", "nombre_comptes_observes",
        "references_ambigues", "source_solde_reference", "solde_reference",
        "nombre_comptes_reference", "ecart_reference", "statut_rapprochement_solde",
    ]
    lines = transaction_lines.loc[
        transaction_lines.get("account_type", pd.Series("", index=transaction_lines.index)).isin(
            CUSTOMER_POSITION_ACCOUNT_LABELS
        )
    ].copy()
    if not lines.empty:
        normalized_description = lines.get(
            "description", pd.Series("", index=lines.index)
        ).apply(normalize_label)
        principle_lines = lines["account_type"].eq("PRINCIPLE")
        valid_principle_position = (
            normalized_description.str.contains("montant principal", na=False)
            & numeric_column(lines, "dr").gt(0)
        ) | (
            normalized_description.str.contains("remboursement du principal", na=False)
            & numeric_column(lines, "cr").gt(0)
        )
        lines = lines.loc[~principle_lines | valid_principle_position].copy()
    observed = pd.DataFrame()
    if not lines.empty:
        lines["reference_id"] = clean_text(lines.get("reference_id", pd.Series("", index=lines.index)))
        lines["customer_id"] = clean_identifier(lines["customer_id"])
        lines["currency_code"] = clean_text(lines["currency_code"]).str.upper().replace("", "NON RENSEIGNEE")
        lines["created_at"] = pd.to_datetime(lines["created_at"], errors="coerce")
        lines["bal_after"] = numeric_column(lines, "bal_after")
        grouping = ["customer_id", "currency_code", "account_type"]
        known_references = {
            key: [value for value in values if value]
            for key, values in lines.groupby(grouping, dropna=False)["reference_id"].apply(
                lambda values: list(dict.fromkeys(clean_text(values).replace("", pd.NA).dropna().tolist()))
            ).items()
        }

        def resolved_account(row: pd.Series) -> str:
            reference = str(row.get("reference_id", "")).strip()
            if reference:
                return reference
            known = known_references.get((row["customer_id"], row["currency_code"], row["account_type"]), [])
            if len(known) == 1:
                return known[0]
            if not known:
                return "SANS_REFERENCE"
            return "REFERENCE_AMBIGUE"

        lines["__account_key"] = lines.apply(resolved_account, axis=1)
        ambiguous = (
            lines.loc[lines["__account_key"].eq("REFERENCE_AMBIGUE")]
            .groupby(["currency_code", "account_type"], as_index=False)
            .size()
            .rename(columns={"size": "references_ambigues"})
        )
        usable = lines.loc[~lines["__account_key"].eq("REFERENCE_AMBIGUE")].copy()
        latest = (
            usable.sort_values(["created_at", "__row_order"], na_position="first")
            .drop_duplicates(["currency_code", "account_type", "__account_key"], keep="last")
        )
        if not latest.empty:
            observed = (
                latest.groupby(["currency_code", "account_type"], as_index=False, dropna=False)
                .agg(
                    solde_transactions_observe=("bal_after", "sum"),
                    date_derniere_ecriture=("created_at", "max"),
                    nombre_comptes_observes=("__account_key", "nunique"),
                )
            )
            observed = observed.merge(ambiguous, on=["currency_code", "account_type"], how="left")
            observed["references_ambigues"] = numeric_column(observed, "references_ambigues").astype(int)
            observed["famille_position"] = observed["account_type"].map(CUSTOMER_POSITION_ACCOUNT_LABELS)
            observed["source_observation"] = "Transactions M-PESA_Turbo - dernier bal_after observe"
            observed = observed.drop(columns="account_type")

    customer_key = str(customer_id).strip()
    reference_blocks: list[pd.DataFrame] = []

    def reference_block(
        source: pd.DataFrame,
        *,
        family: str,
        source_label: str,
        balance_column: str,
        account_column: str | None = None,
    ) -> None:
        if not isinstance(source, pd.DataFrame) or source.empty or "customer_id" not in source.columns:
            return
        block = source.loc[clean_identifier(source["customer_id"]).eq(customer_key)].copy()
        if block.empty or "currency_code" not in block.columns or balance_column not in block.columns:
            return
        block["currency_code"] = clean_text(block["currency_code"]).str.upper().replace("", "NON RENSEIGNEE")
        block[balance_column] = numeric_column(block, balance_column)
        if account_column and account_column in block.columns:
            account_values = clean_identifier(block[account_column])
            account_count = lambda values: int(clean_identifier(values).replace("", pd.NA).nunique())
            account_source = account_column
        else:
            block["__account_row"] = np.arange(len(block))
            account_count = "size"
            account_source = "__account_row"
        aggregation: dict[str, tuple[str, object]] = {
            "solde_reference": (balance_column, "sum"),
            "nombre_comptes_reference": (account_source, account_count),
        }
        reference = block.groupby("currency_code", as_index=False, dropna=False).agg(**aggregation)
        reference["famille_position"] = family
        reference["source_solde_reference"] = source_label
        reference_blocks.append(reference)

    reference_block(
        prepared.current_savings,
        family="Epargne courante",
        source_label="Epargne courante_Turbo (Current Savings)",
        balance_column="balance",
    )
    reference_block(
        prepared.fixed_savings,
        family="DAT",
        source_label="DAT_Turbo (Fixed Savings)",
        balance_column="balance",
    )
    reference_block(
        prepared.loans,
        family="Credit",
        source_label="Credits_Turbo (Loans)",
        balance_column="loan_balance",
        account_column="loan_id",
    )
    reference = concat_frames_stable(reference_blocks)
    if observed.empty and reference.empty:
        return pd.DataFrame(columns=position_columns)
    if observed.empty:
        positions = reference.copy()
        positions["source_observation"] = pd.NA
        positions["solde_transactions_observe"] = np.nan
        positions["date_derniere_ecriture"] = pd.NaT
        positions["nombre_comptes_observes"] = 0
        positions["references_ambigues"] = 0
    elif reference.empty:
        positions = observed.copy()
        positions["source_solde_reference"] = pd.NA
        positions["solde_reference"] = np.nan
        positions["nombre_comptes_reference"] = 0
    else:
        positions = observed.merge(
            reference,
            on=["currency_code", "famille_position"],
            how="outer",
        )
    positions["customer_id"] = customer_key
    for column in ["nombre_comptes_observes", "references_ambigues", "nombre_comptes_reference"]:
        if column not in positions.columns:
            positions[column] = 0
        positions[column] = pd.to_numeric(positions[column], errors="coerce").fillna(0).astype(int)
    positions["ecart_reference"] = np.where(
        positions["solde_transactions_observe"].notna() & positions["solde_reference"].notna(),
        positions["solde_transactions_observe"] - positions["solde_reference"],
        np.nan,
    )

    def position_status(row: pd.Series) -> str:
        observed_value = row.get("solde_transactions_observe")
        reference_value = row.get("solde_reference")
        if pd.isna(observed_value):
            return "Non observe dans Transactions M-PESA_Turbo"
        if pd.isna(reference_value):
            return "Solde observe uniquement - fichier de reference absent"
        if not comparison_allowed:
            return "Non comparable - derniere operation hors perimetre"
        tolerance = max(0.01, max(abs(float(observed_value)), abs(float(reference_value))) * 1e-6)
        return "Conforme" if abs(float(observed_value) - float(reference_value)) <= tolerance else "Ecart a expliquer"

    positions["statut_rapprochement_solde"] = positions.apply(position_status, axis=1)
    positions["source_observation"] = positions.get("source_observation", pd.Series(pd.NA, index=positions.index)).fillna(
        "Transactions M-PESA_Turbo - aucune ecriture observee"
    )
    positions["source_solde_reference"] = positions.get(
        "source_solde_reference", pd.Series(pd.NA, index=positions.index)
    ).fillna("Fichier de reference non charge")
    return positions[position_columns].sort_values(["currency_code", "famille_position"]).reset_index(drop=True)


def _build_customer_behavior(events: pd.DataFrame) -> pd.DataFrame:
    columns = [
        "currency_code", "nombre_operations", "jours_actifs", "premiere_operation", "derniere_operation",
        "operations_par_jour_actif", "montant_moyen", "montant_median", "plus_forte_operation",
        "jour_semaine_frequent", "heure_frequente", "type_operation_frequent",
        "intervalle_median_heures", "plus_longue_inactivite_jours",
    ]
    if not isinstance(events, pd.DataFrame) or events.empty:
        return pd.DataFrame(columns=columns)
    french_weekdays = {
        0: "Lundi", 1: "Mardi", 2: "Mercredi", 3: "Jeudi",
        4: "Vendredi", 5: "Samedi", 6: "Dimanche",
    }
    rows: list[dict[str, object]] = []
    for currency, group in events.groupby("currency_code", sort=True, dropna=False):
        group = group.sort_values("created_at").copy()
        dates = pd.to_datetime(group["created_at"], errors="coerce").dropna()
        amounts = pd.to_numeric(group["montant_operation"], errors="coerce").dropna()
        active_days = int(dates.dt.date.nunique()) if not dates.empty else 0
        intervals = dates.diff().dropna().dt.total_seconds().div(3600)
        weekday_counts = dates.dt.dayofweek.value_counts()
        hour_counts = dates.dt.hour.value_counts()
        type_counts = clean_text(group["type_operation"]).replace("", pd.NA).dropna().value_counts()
        rows.append(
            {
                "currency_code": currency,
                "nombre_operations": int(len(group)),
                "jours_actifs": active_days,
                "premiere_operation": dates.min() if not dates.empty else pd.NaT,
                "derniere_operation": dates.max() if not dates.empty else pd.NaT,
                "operations_par_jour_actif": len(group) / active_days if active_days else np.nan,
                "montant_moyen": float(amounts.mean()) if not amounts.empty else np.nan,
                "montant_median": float(amounts.median()) if not amounts.empty else np.nan,
                "plus_forte_operation": float(amounts.max()) if not amounts.empty else np.nan,
                "jour_semaine_frequent": french_weekdays.get(int(weekday_counts.index[0]), "-") if not weekday_counts.empty else "-",
                "heure_frequente": f"{int(hour_counts.index[0]):02d}h" if not hour_counts.empty else "-",
                "type_operation_frequent": str(type_counts.index[0]) if not type_counts.empty else "-",
                "intervalle_median_heures": float(intervals.median()) if not intervals.empty else np.nan,
                "plus_longue_inactivite_jours": float(intervals.max() / 24) if not intervals.empty else np.nan,
            }
        )
    return pd.DataFrame(rows, columns=columns)


def build_customer_matured_dat_interest_entries(
    fixed_savings: pd.DataFrame | None,
    transactions: pd.DataFrame | None,
    customer_id: object,
    *,
    as_of_date: object | None = None,
    date_start: object | None = None,
    date_end: object | None = None,
    currency: object | None = None,
    reference_query: object = "",
    annual_interest_rate_pct: float = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
) -> pd.DataFrame:
    """Restitue les interets constates des DAT echus sans modifier le solde M-PESA.

    ``interest_earned`` provient de la position ``Savings Account``. Une ligne
    n'est retenue que si le DAT est arrive a echeance, porte un interet positif
    et presente un statut de denouement. La presence d'une ecriture de sortie
    portant le meme ``savings_id`` sert uniquement a qualifier la tracabilite.
    """
    empty = pd.DataFrame(columns=CUSTOMER_DAT_INTEREST_COLUMNS)
    if not isinstance(fixed_savings, pd.DataFrame) or fixed_savings.empty:
        return empty
    if "customer_id" not in fixed_savings.columns:
        return empty

    customer_key = str(customer_id).strip()
    frame = fixed_savings.loc[
        clean_identifier(fixed_savings["customer_id"]).eq(customer_key)
    ].copy()
    if frame.empty:
        return empty

    for column in [
        "savings_id",
        "msisdn",
        "msisdn1",
        "currency_code",
        "product_name",
        "status",
        "balance",
        "date_approved",
        "maturity_date",
        "interest_earned",
        "voda_interest",
    ]:
        if column not in frame.columns:
            frame[column] = pd.NA
    frame["savings_id"] = clean_identifier(frame["savings_id"])
    frame["customer_id"] = clean_identifier(frame["customer_id"])
    frame["currency_code"] = clean_text(frame["currency_code"]).str.upper()
    frame["msisdn"] = normalize_phone(frame["msisdn"].where(frame["msisdn"].notna(), frame["msisdn1"]))
    frame["date_approved"] = pd.to_datetime(frame["date_approved"], errors="coerce")
    frame["maturity_date"] = pd.to_datetime(frame["maturity_date"], errors="coerce")
    frame["balance"] = pd.to_numeric(frame["balance"], errors="coerce").fillna(0.0)
    frame["interest_earned"] = pd.to_numeric(frame["interest_earned"], errors="coerce").fillna(0.0)
    frame["voda_interest"] = pd.to_numeric(frame["voda_interest"], errors="coerce").fillna(0.0)

    transaction_dates = pd.Series(dtype="datetime64[ns]")
    if isinstance(transactions, pd.DataFrame) and not transactions.empty and "created_at" in transactions.columns:
        transaction_dates = pd.to_datetime(transactions["created_at"], errors="coerce").dropna()
    effective_end = pd.to_datetime(
        date_end if date_end is not None else as_of_date,
        errors="coerce",
    )
    if pd.isna(effective_end):
        effective_end = transaction_dates.max() if not transaction_dates.empty else pd.Timestamp.now()
    if pd.Timestamp(effective_end) == pd.Timestamp(effective_end).normalize():
        effective_end = _timestamp_plus(effective_end, days=1, microseconds=-1)

    normalized_status = frame["status"].apply(normalize_label)
    settled_status = normalized_status.str.contains(
        r"withdraw|mature|closed|clotur|retir",
        regex=True,
        na=False,
    )
    legacy_settled = normalized_status.eq("") & frame["balance"].le(0)
    frame = frame.loc[
        frame["maturity_date"].notna()
        & frame["maturity_date"].le(effective_end)
        & frame["interest_earned"].gt(0)
        & (settled_status | legacy_settled)
    ].copy()
    if frame.empty:
        return empty

    start = pd.to_datetime(date_start, errors="coerce") if date_start is not None else pd.NaT
    if pd.notna(start):
        frame = frame.loc[frame["maturity_date"].ge(start)].copy()
    end = pd.to_datetime(date_end, errors="coerce") if date_end is not None else pd.NaT
    if pd.notna(end):
        if pd.Timestamp(end) == pd.Timestamp(end).normalize():
            end = _timestamp_plus(end, days=1, microseconds=-1)
        frame = frame.loc[frame["maturity_date"].le(end)].copy()
    currency_text = str(currency).strip().upper() if currency is not None else ""
    if currency_text and currency_text not in {"TOUTES", "ALL"}:
        frame = frame.loc[frame["currency_code"].eq(currency_text)].copy()
    if frame.empty:
        return empty

    transaction_summary = pd.DataFrame(
        columns=[
            "savings_id",
            "capital_place",
            "reference_transaction_turbo",
            "date_ecriture_turbo",
        ]
    )
    if isinstance(transactions, pd.DataFrame) and not transactions.empty:
        if "customer_id" in transactions.columns:
            tx = transactions.loc[
                clean_identifier(transactions["customer_id"]).eq(customer_key)
            ].copy()
        else:
            tx = pd.DataFrame()
        for column in [
            "customer_id",
            "currency_code",
            "account_type",
            "reference_id",
            "ref_no",
            "description",
            "created_at",
            "dr",
            "cr",
        ]:
            if column not in tx.columns:
                tx[column] = pd.NA
        tx["customer_id"] = clean_identifier(tx["customer_id"])
        tx["currency_code"] = clean_text(tx["currency_code"]).str.upper()
        tx["account_type"] = clean_text(tx["account_type"]).str.upper()
        tx["reference_id"] = clean_identifier(tx["reference_id"])
        tx["ref_no"] = clean_identifier(tx["ref_no"])
        tx["created_at"] = pd.to_datetime(tx["created_at"], errors="coerce")
        tx["dr"] = pd.to_numeric(tx["dr"], errors="coerce").fillna(0.0)
        tx["cr"] = pd.to_numeric(tx["cr"], errors="coerce").fillna(0.0)
        savings_ids = set(frame["savings_id"].loc[frame["savings_id"].ne("")].tolist())
        fixed_tx = tx.loc[
            tx["customer_id"].eq(customer_key)
            & tx["account_type"].eq("FIXED SAVINGS")
            & tx["reference_id"].isin(savings_ids)
            & (tx["created_at"].isna() | tx["created_at"].le(effective_end))
        ].copy()
        if not fixed_tx.empty:
            capital = (
                fixed_tx.groupby("reference_id", as_index=False, dropna=False)["cr"]
                .sum()
                .rename(columns={"reference_id": "savings_id", "cr": "capital_place"})
            )
            settlement_description = fixed_tx["description"].apply(normalize_label)
            settlement = fixed_tx.loc[
                fixed_tx["dr"].gt(0)
                & settlement_description.str.contains(
                    r"retrait|withdraw|remboursement.*compte|fixed savings",
                    regex=True,
                    na=False,
                )
            ].copy()
            if not settlement.empty:
                settlement["reference_transaction_turbo"] = settlement["ref_no"].where(
                    settlement["ref_no"].ne(""), settlement["reference_id"]
                )
                traced = settlement.groupby("reference_id", as_index=False, dropna=False).agg(
                    reference_transaction_turbo=("reference_transaction_turbo", concat_unique),
                    date_ecriture_turbo=("created_at", "max"),
                ).rename(columns={"reference_id": "savings_id"})
                transaction_summary = capital.merge(traced, on="savings_id", how="left")
            else:
                transaction_summary = capital
                transaction_summary["reference_transaction_turbo"] = ""
                transaction_summary["date_ecriture_turbo"] = pd.NaT

    frame = frame.merge(transaction_summary, on="savings_id", how="left")
    frame["capital_place"] = pd.to_numeric(frame["capital_place"], errors="coerce")
    frame.loc[frame["capital_place"].le(0), "capital_place"] = np.nan
    frame["taux_interet_annuel_pct"] = float(annual_interest_rate_pct)
    frame["interet_client_constate"] = frame["interest_earned"]
    frame["montant_echeance_client"] = frame["capital_place"] + frame["interet_client_constate"]
    frame["reference_transaction_turbo"] = clean_text(
        frame.get("reference_transaction_turbo", pd.Series("", index=frame.index))
    )
    frame["date_ecriture_turbo"] = pd.to_datetime(
        frame.get("date_ecriture_turbo", pd.Series(pd.NaT, index=frame.index)),
        errors="coerce",
    )
    traced = frame["date_ecriture_turbo"].notna()
    frame["statut_tracabilite"] = np.where(
        traced,
        "Comptabilise et trace dans Transactions M-PESA_Turbo",
        "Constate dans Savings Account - ecriture detaillee absente",
    )
    frame["source_interet"] = "Savings Account [Turbo] - interest_earned"
    frame["impact_solde_mpesa"] = "Hors solde M-PESA - versement en epargne courante"

    query = str(reference_query).strip()
    if query:
        query_mask = pd.Series(False, index=frame.index)
        for column in ["savings_id", "reference_transaction_turbo"]:
            query_mask |= frame[column].astype("string").str.contains(
                query, case=False, regex=False, na=False
            )
        frame = frame.loc[query_mask].copy()
    if frame.empty:
        return empty
    return (
        frame[CUSTOMER_DAT_INTEREST_COLUMNS]
        .sort_values(["maturity_date", "currency_code", "savings_id"])
        .reset_index(drop=True)
    )


def build_customer_active_dat_positions(
    fixed_savings: pd.DataFrame | None,
    transactions: pd.DataFrame | None,
    customer_id: object,
    *,
    currency: object | None = None,
    annual_interest_rate_pct: float = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
) -> pd.DataFrame:
    """Construit la position courante des DAT positifs d'un client depuis Turbo.

    La date de situation vient d'abord de ``Savings Account`` puis, à défaut,
    de la dernière transaction Turbo du client. G2 n'intervient ni dans le
    périmètre, ni dans les montants, ni dans l'estimation des intérêts.
    """
    empty = pd.DataFrame(columns=CUSTOMER_ACTIVE_DAT_COLUMNS)
    if not isinstance(fixed_savings, pd.DataFrame) or fixed_savings.empty:
        return empty
    if "customer_id" not in fixed_savings.columns:
        return empty

    customer_key = clean_identifier(pd.Series([customer_id])).iloc[0]
    customer_dat = fixed_savings.loc[
        clean_identifier(fixed_savings["customer_id"]).eq(customer_key)
    ].copy()
    if customer_dat.empty:
        return empty

    snapshot_candidates: list[pd.Series] = []
    for column in ["updated_at", "date_locked"]:
        if column in customer_dat.columns:
            values = pd.to_datetime(customer_dat[column], errors="coerce").dropna()
            if not values.empty:
                snapshot_candidates.append(values)
                break
    if (
        not snapshot_candidates
        and isinstance(transactions, pd.DataFrame)
        and not transactions.empty
        and {"customer_id", "created_at"}.issubset(transactions.columns)
    ):
        tx_dates = pd.to_datetime(
            transactions.loc[
                clean_identifier(transactions["customer_id"]).eq(customer_key),
                "created_at",
            ],
            errors="coerce",
        ).dropna()
        if not tx_dates.empty:
            snapshot_candidates.append(tx_dates)
    if not snapshot_candidates:
        for column in ["created_at", "date_approved"]:
            if column in customer_dat.columns:
                values = pd.to_datetime(customer_dat[column], errors="coerce").dropna()
                if not values.empty:
                    snapshot_candidates.append(values)
                    break
    situation_date = (
        pd.Timestamp(pd.concat(snapshot_candidates, ignore_index=True).max()).normalize()
        if snapshot_candidates
        else pd.Timestamp.now().normalize()
    )

    maturity = build_mpesa_dat_maturity_analysis(
        customer_dat,
        as_of_date=situation_date,
        annual_interest_rate_pct=annual_interest_rate_pct,
        preparation_horizon_days=DEFAULT_DAT_REPAYMENT_PREPARATION_HORIZON_DAYS,
    )
    detail = maturity.get("detail", pd.DataFrame()).copy()
    if detail.empty:
        return empty

    currency_text = str(currency).strip().upper() if currency is not None else ""
    if currency_text and currency_text not in {"TOUTES", "ALL"}:
        detail = detail.loc[detail["currency_code"].eq(currency_text)].copy()
    if detail.empty:
        return empty

    days = pd.to_numeric(detail["jours_avant_echeance"], errors="coerce")
    detail["situation_dat_client"] = np.select(
        [
            days.isna(),
            days.lt(0),
            days.eq(0),
            days.le(DEFAULT_DAT_REPAYMENT_PREPARATION_HORIZON_DAYS),
        ],
        [
            "Date d'échéance à compléter",
            "Échu à rembourser",
            "Échéance aujourd'hui",
            "Échéance proche",
        ],
        default="En cours",
    )
    detail["date_situation"] = situation_date
    detail["savings_id"] = clean_identifier(detail["savings_id"])
    detail["customer_id"] = clean_identifier(detail["customer_id"])
    detail["msisdn"] = normalize_phone(detail["msisdn"])
    detail["balance"] = pd.to_numeric(detail["balance"], errors="coerce")
    detail["interet_estime_echeance"] = pd.to_numeric(
        detail["interet_estime_echeance"], errors="coerce"
    )
    detail["capital_plus_interet_estime"] = pd.to_numeric(
        detail["capital_plus_interet_estime"], errors="coerce"
    )
    detail["__ordre_situation"] = np.select(
        [days.lt(0), days.eq(0), days.le(DEFAULT_DAT_REPAYMENT_PREPARATION_HORIZON_DAYS)],
        [0, 1, 2],
        default=3,
    )
    return (
        detail.sort_values(
            ["currency_code", "__ordre_situation", "maturity_date", "savings_id"],
            na_position="last",
        )[CUSTOMER_ACTIVE_DAT_COLUMNS]
        .reset_index(drop=True)
    )


def build_customer_transaction_analysis(
    prepared: MpesaPreparedData,
    customer_id: object,
    *,
    currency: object | None = None,
    operation_types: Iterable[object] | None = None,
    date_start: object | None = None,
    date_end: object | None = None,
    reference_query: object = "",
    annual_interest_rate_pct: float = DEFAULT_DAT_ANNUAL_INTEREST_RATE_PCT,
) -> dict[str, pd.DataFrame]:
    """Construit les analyses client disponibles dans Transactions M-PESA_Turbo."""
    empty = {
        "parcours_turbo": pd.DataFrame(),
        "jalons_turbo": pd.DataFrame(),
        "credit_turbo_synthese_client": pd.DataFrame(),
        "credit_turbo_detail_client": pd.DataFrame(),
        "remboursements_turbo_synthese_client": pd.DataFrame(),
        "remboursements_turbo_detail_client": pd.DataFrame(),
        "positions_turbo": pd.DataFrame(),
        "comportement_turbo": pd.DataFrame(),
        "mouvements_internes_turbo": pd.DataFrame(),
        "controles_client_turbo": pd.DataFrame(),
        "dat_en_cours_client": build_customer_active_dat_positions(
            prepared.fixed_savings,
            prepared.transactions,
            customer_id,
            currency=currency,
            annual_interest_rate_pct=annual_interest_rate_pct,
        ),
    }
    events, lines = _build_customer_turbo_events(prepared.transactions, customer_id)
    if events.empty:
        return empty
    full_customer_latest = pd.to_datetime(events["created_at"], errors="coerce").max()
    base = events.copy()
    currency_text = str(currency).strip().upper() if currency is not None else ""
    if currency_text and currency_text not in {"TOUTES", "ALL"}:
        base = base.loc[base["currency_code"].eq(currency_text)].copy()
    start = pd.to_datetime(date_start, errors="coerce") if date_start is not None else pd.NaT
    end = pd.to_datetime(date_end, errors="coerce") if date_end is not None else pd.NaT
    if pd.notna(start):
        base = base.loc[pd.to_datetime(base["created_at"], errors="coerce").ge(start)].copy()
    if pd.notna(end):
        if not hasattr(date_end, "hour"):
            end = _timestamp_plus(end, days=1, microseconds=-1)
        base = base.loc[pd.to_datetime(base["created_at"], errors="coerce").le(end)].copy()
    query = str(reference_query).strip()
    if query:
        reference_mask = pd.Series(False, index=base.index)
        for column in ["event_reference", "ref_no", "reference_ids"]:
            reference_mask |= base[column].astype("string").str.contains(query, case=False, regex=False, na=False)
        base = base.loc[reference_mask].copy()
    if base.empty:
        return empty

    base_event_keys = set(base["event_key"].astype(str))
    base_lines = lines.loc[lines["event_key"].astype(str).isin(base_event_keys)].copy()
    internal = base.loc[base["type_operation"].eq("Transfert DAT vers epargne courante")].copy()

    selected_types = [str(value) for value in (operation_types or []) if str(value).strip()]
    scoped = base.loc[base["type_operation"].isin(selected_types)].copy() if selected_types else base.copy()
    if scoped.empty:
        positions = _build_customer_observed_positions(
            prepared,
            base_lines,
            customer_id,
            comparison_allowed=False,
        )
        result = dict(empty)
        result["positions_turbo"] = positions
        result["mouvements_internes_turbo"] = internal
        return result

    latest_in_scope = pd.to_datetime(base["created_at"], errors="coerce").max()
    comparison_allowed = (
        not query
        and pd.notna(latest_in_scope)
        and pd.notna(full_customer_latest)
        and pd.Timestamp(latest_in_scope) >= pd.Timestamp(full_customer_latest)
    )
    positions = _build_customer_observed_positions(
        prepared,
        base_lines,
        customer_id,
        comparison_allowed=comparison_allowed,
    )
    behavior = _build_customer_behavior(scoped)

    milestones = (
        scoped.groupby(["currency_code", "type_operation"], as_index=False, dropna=False)
        .agg(
            nombre_operations=("event_key", "nunique"),
            montant_total_observe=("montant_operation", "sum"),
            premiere_operation=("created_at", "min"),
            derniere_operation=("created_at", "max"),
        )
        .sort_values(["currency_code", "premiere_operation", "type_operation"])
        .reset_index(drop=True)
    )
    credit_mask = (
        scoped["type_operation"].isin(["Decaissement de credit", "Remboursement de credit", "Remboursement avec penalite"])
        | scoped[["montant_decaisse_client", "interet_observe", "principal_rembourse", "penalite_observee"]].gt(0).any(axis=1)
    )
    credit_detail = scoped.loc[credit_mask].copy()
    credit_summary = pd.DataFrame()
    repayment_detail = pd.DataFrame()
    repayment_summary = pd.DataFrame()
    if not credit_detail.empty:
        credit_detail["est_decaissement"] = credit_detail["type_operation"].eq("Decaissement de credit")
        credit_detail["est_remboursement"] = credit_detail["type_operation"].isin(
            ["Remboursement de credit", "Remboursement avec penalite"]
        )
        credit_detail["remboursement_avec_penalite"] = credit_detail["penalite_observee"].gt(0)
        credit_detail["remboursement_avec_epargne_dat"] = credit_detail["epargne_dat_remboursement"].gt(0)
        credit_summary = (
            credit_detail.groupby("currency_code", as_index=False, dropna=False)
            .agg(
                nombre_decaissements=("est_decaissement", "sum"),
                montant_decaisse_client=("montant_decaisse_client", "sum"),
                dette_creee_observee=("dette_creee_observee", "sum"),
                interet_observe=("interet_observe", "sum"),
                nombre_remboursements=("est_remboursement", "sum"),
                principal_rembourse=("principal_rembourse", "sum"),
                remboursements_avec_penalite=("remboursement_avec_penalite", "sum"),
                penalite_observee=("penalite_observee", "sum"),
                remboursements_avec_epargne_dat=("remboursement_avec_epargne_dat", "sum"),
                revenu_credit_observe=("revenu_credit_observe", "sum"),
            )
        )
        credit_summary["ratio_interet_decaissement_pct"] = (
            credit_summary["interet_observe"]
            .div(credit_summary["montant_decaisse_client"].replace(0, pd.NA))
            .mul(100)
        )
        repayment_detail = credit_detail.loc[credit_detail["est_remboursement"]].copy()
        if not repayment_detail.empty:
            repayment_detail["montant_paye_observe"] = repayment_detail[
                "remboursement_mpesa"
            ].where(
                repayment_detail["remboursement_mpesa"].gt(0),
                repayment_detail["montant_operation"],
            )
            repayment_summary = (
                repayment_detail.groupby(
                    ["customer_id", "currency_code"],
                    as_index=False,
                    dropna=False,
                )
                .agg(
                    nombre_remboursements=("event_key", "nunique"),
                    montant_paye_observe=("montant_paye_observe", "sum"),
                    principal_rembourse=("principal_rembourse", "sum"),
                    interet_observe=("interet_observe", "sum"),
                    penalite_observee=("penalite_observee", "sum"),
                )
            )

    path_columns = [
        "created_at", "currency_code", "event_reference", "ref_no", "type_operation", "sens_metier",
        "montant_operation", "montant_entree_bisou", "montant_sortie_bisou", "mode_remboursement_observe",
        "account_types", "descriptions", "nombre_lignes", "statut_controle_turbo",
    ]
    credit_columns = [
        "created_at", "currency_code", "event_reference", "type_operation", "montant_decaisse_client",
        "dette_creee_observee", "interet_observe", "remboursement_mpesa", "principal_rembourse",
        "penalite_observee", "epargne_dat_remboursement", "mode_remboursement_observe",
        "statut_controle_turbo", "observation_controle_turbo",
    ]
    repayment_columns = [
        "customer_id", "created_at", "event_reference", "ref_no", "currency_code",
        "montant_paye_observe", "principal_rembourse", "interet_observe",
        "penalite_observee", "mode_remboursement_observe",
    ]
    control_columns = [
        "created_at", "currency_code", "event_reference", "type_operation", "montant_operation",
        "controle_montant_operation", "detail_controle_montant", "nombre_lignes",
        "lignes_mouvement_nul", "lignes_solde_negatif", "types_comptes_inconnus",
        "total_debit_ecritures", "total_credit_ecritures", "ecart_debit_credit_observe",
        "statut_controle_turbo", "observation_controle_turbo",
    ]
    internal_columns = [
        "created_at", "currency_code", "event_reference", "type_operation", "montant_operation",
        "transfert_dat_sortie", "transfert_epargne_entree", "account_types", "descriptions",
        "statut_controle_turbo", "observation_controle_turbo",
    ]
    return {
        "parcours_turbo": scoped[path_columns].sort_values("created_at").reset_index(drop=True),
        "jalons_turbo": milestones,
        "credit_turbo_synthese_client": credit_summary,
        "credit_turbo_detail_client": credit_detail[credit_columns].sort_values("created_at").reset_index(drop=True),
        "remboursements_turbo_synthese_client": repayment_summary,
        "remboursements_turbo_detail_client": (
            repayment_detail[repayment_columns].sort_values("created_at").reset_index(drop=True)
            if not repayment_detail.empty
            else pd.DataFrame(columns=repayment_columns)
        ),
        "positions_turbo": positions,
        "comportement_turbo": behavior,
        "mouvements_internes_turbo": internal[internal_columns].sort_values("created_at").reset_index(drop=True),
        "controles_client_turbo": scoped[control_columns].sort_values("created_at").reset_index(drop=True),
        "dat_en_cours_client": empty["dat_en_cours_client"],
    }


def build_mpesa_statement(
    prepared: MpesaPreparedData,
    customer_id: str,
    opening_balances: dict[str, float | None] | None = None,
    date_start: object | None = None,
    date_end: object | None = None,
) -> dict[str, Any]:
    opening_balances = opening_balances or {}
    transactions = prepared.transactions
    tx_client = transactions.loc[transactions["customer_id"].eq(str(customer_id))].copy()
    dat_client = filter_customer_frame(prepared.fixed_savings, str(customer_id))
    savings_client = filter_customer_frame(prepared.current_savings, str(customer_id))
    loans_client = filter_customer_frame(prepared.loans, str(customer_id))

    if tx_client.empty:
        raise ValueError(f"Aucune transaction trouvee pour le client {customer_id}.")

    if not dat_client.empty:
        situation_date = pd.to_datetime(transactions["created_at"], errors="coerce").max()
        dat_client["statut_dat"] = np.select(
            [
                numeric_column(dat_client, "balance").le(0),
                dat_client.get("maturity_date", pd.Series(pd.NaT, index=dat_client.index)).notna()
                & dat_client.get("maturity_date", pd.Series(pd.NaT, index=dat_client.index)).lt(situation_date),
            ],
            ["Solde nul", "Echu"],
            default="Actif",
        )

    tx_client["operation_reference"] = tx_client["ref_no"].where(
        ~tx_client["ref_no"].apply(_is_empty_text), tx_client["reference_id"]
    )
    missing_reference = tx_client["operation_reference"].apply(_is_empty_text)
    tx_client.loc[missing_reference, "operation_reference"] = (
        "LIGNE-" + tx_client.loc[missing_reference, "id"].astype("string")
    )
    operation_metadata = (
        tx_client.groupby(
            ["customer_id", "currency_code", "created_at", "operation_reference"],
            as_index=False,
            dropna=False,
        )
        .agg(
            description_turbo=("description", concat_unique),
            account_types_operation=("account_type", concat_unique),
        )
    )

    mpesa = tx_client.loc[tx_client["account_type"].eq("MPESA ACCOUNT")].copy()
    if mpesa.empty:
        raise ValueError(f"Aucun mouvement MPESA ACCOUNT trouve pour le client {customer_id}.")

    statement_aggregations: dict[str, tuple[str, object]] = {
        "telephone": ("msisdn1", concat_unique),
        "debit_mpesa": ("dr", "sum"),
        "credit_mpesa": ("cr", "sum"),
        "references_internes": ("reference_id", concat_unique),
        "descriptions": ("description", concat_unique),
        "account_types": ("account_type", concat_unique),
        "nombre_lignes_comptables": ("id", "count"),
    }
    if "Nom_client" in mpesa.columns:
        statement_aggregations["Nom_client"] = ("Nom_client", concat_unique)
    if "mode_rapprochement_nom_client" in mpesa.columns:
        statement_aggregations["mode_rapprochement_nom_client"] = ("mode_rapprochement_nom_client", concat_unique)
    statement = (
        mpesa.groupby(["customer_id", "currency_code", "created_at", "operation_reference"], as_index=False, dropna=False)
        .agg(**statement_aggregations)
        .sort_values(["currency_code", "created_at", "operation_reference"])
    )
    statement = statement.merge(
        operation_metadata,
        on=["customer_id", "currency_code", "created_at", "operation_reference"],
        how="left",
    )
    has_turbo_description = ~statement["description_turbo"].apply(_is_empty_text)
    statement["descriptions"] = statement["description_turbo"].where(
        has_turbo_description, statement["descriptions"]
    )
    has_operation_accounts = ~statement["account_types_operation"].apply(_is_empty_text)
    statement["account_types"] = statement["account_types_operation"].where(
        has_operation_accounts, statement["account_types"]
    )
    statement = statement.drop(columns=["account_types_operation"])
    if "Nom_client" not in statement.columns:
        statement["Nom_client"] = pd.NA
    if "mode_rapprochement_nom_client" not in statement.columns:
        statement["mode_rapprochement_nom_client"] = "Nom client non disponible"

    statement["reference_dat_operation"] = statement["references_internes"].apply(lambda value: extract_prefixed_reference(value, "FA"))
    statement["reference_credit_operation"] = statement["references_internes"].apply(lambda value: extract_prefixed_reference(value, "LN"))

    dat_direct = _build_dat_direct(tx_client)
    if not dat_direct.empty:
        statement = statement.merge(dat_direct, on=["currency_code", "created_at"], how="left")
        statement["reference_dat_operation"] = statement["reference_dat_direct"].fillna(statement["reference_dat_operation"])
        statement = statement.drop(columns=["reference_dat_direct"])
    else:
        statement["solde_dat_operation_avant"] = np.nan
        statement["solde_dat_operation_apres"] = np.nan
        statement["variation_dat_operation"] = np.nan

    statement["mouvement_net_mpesa"] = statement["credit_mpesa"] - statement["debit_mpesa"]
    statement["entree_mpesa"] = statement["mouvement_net_mpesa"].clip(lower=0)
    statement["sortie_mpesa"] = (-statement["mouvement_net_mpesa"].clip(upper=0))
    statement["type_operation"] = statement.apply(
        lambda row: classify_mpesa_operation(row["descriptions"], row["account_types"], row["mouvement_net_mpesa"]),
        axis=1,
    )
    statement["cumul_net_depuis_debut_fichier"] = statement.groupby("currency_code")["mouvement_net_mpesa"].cumsum()
    statement["solde_mpesa_avant"] = np.nan
    statement["solde_mpesa_apres"] = np.nan
    statement["solde_mpesa_disponible"] = False

    for currency, index in statement.groupby("currency_code").groups.items():
        opening = opening_balances.get(str(currency))
        if opening is None or pd.isna(opening):
            continue
        opening = float(opening)
        cumulative = statement.loc[index, "cumul_net_depuis_debut_fichier"]
        statement.loc[index, "solde_mpesa_apres"] = opening + cumulative
        statement.loc[index, "solde_mpesa_avant"] = statement.loc[index, "solde_mpesa_apres"] - statement.loc[index, "mouvement_net_mpesa"]
        statement.loc[index, "solde_mpesa_disponible"] = True

    dat_final = build_dat_final(prepared.fixed_savings, str(customer_id))
    savings_final = build_savings_final(prepared.current_savings, str(customer_id))
    savings_events = build_account_events(tx_client, "NORMAL SAVINGS")
    dat_events = build_account_events(tx_client, "FIXED SAVINGS")
    statement = add_reconstructed_balance(statement, savings_events, savings_final, "solde_epargne_au_moment")
    statement = add_reconstructed_balance(statement, dat_events, dat_final, "solde_dat_total_au_moment")
    statement["epargne_courante_finale"] = statement["currency_code"].map(lambda currency: float(savings_final.get(currency, 0.0)))
    statement["dat_final_client"] = statement["currency_code"].map(lambda currency: float(dat_final.get(currency, 0.0)))

    statement = enrich_with_loans(statement, prepared.loans)

    if date_start is not None:
        start = pd.to_datetime(date_start, errors="coerce")
        if pd.notna(start):
            statement = statement.loc[statement["created_at"].ge(start)]
    if date_end is not None:
        end = pd.to_datetime(date_end, errors="coerce")
        if pd.notna(end):
            if not hasattr(date_end, "hour"):
                end = _timestamp_plus(end, days=1, microseconds=-1)
            statement = statement.loc[statement["created_at"].le(end)]

    statement["controle_mouvement"] = np.select(
        [
            statement["mouvement_net_mpesa"].eq(0),
            statement["operation_reference"].astype("string").str.startswith("LIGNE-", na=False),
        ],
        ["Mouvement nul a verifier", "Reference absente"],
        default="OK",
    )
    statement = statement.sort_values(["created_at", "currency_code", "operation_reference"]).reset_index(drop=True)

    summary = build_customer_summary(statement, dat_client, loans_client, str(customer_id))
    diagnostics = build_diagnostics(prepared, str(customer_id), statement)
    g2_dat = build_g2_dat_crosscheck(prepared, str(customer_id))
    customer_transaction_analysis = build_customer_transaction_analysis(prepared, str(customer_id))
    client_visible_analysis = dict(customer_transaction_analysis)
    for internal_credit_key in [
        "credit_turbo_synthese_client",
        "credit_turbo_detail_client",
        "positions_turbo",
    ]:
        client_visible_analysis.pop(internal_credit_key, None)
    g2_available = not prepared.g2_transactions.empty
    name_enriched_by_g2 = bool(
        "Nom_client" in statement.columns
        and clean_text(statement["Nom_client"]).ne("").any()
        and g2_available
    )
    return {
        "customer_id": str(customer_id),
        "mode_source_extrait": "Turbo principal + verification G2" if g2_available else "Turbo seul",
        "controle_g2_disponible": g2_available,
        "nom_client_enrichi_g2": name_enriched_by_g2,
        "extrait": format_statement_columns(statement),
        "synthese": summary,
        "dat_final": dat_client,
        "mouvements_dat": dat_events,
        "mouvements_epargne": savings_events,
        "credits": loans_client,
        "g2_dat": g2_dat,
        "diagnostics": diagnostics,
        **client_visible_analysis,
    }


def enrich_with_loans(statement: pd.DataFrame, loans: pd.DataFrame) -> pd.DataFrame:
    result = statement.copy()
    if loans.empty or "loan_id" not in loans.columns:
        for column in [
            "loan_id",
            "loan_amount",
            "loan_balance",
            "amount_paid",
            "outstanding_principle",
            "outstanding_interest",
            "outstanding_penalty_fees",
            "status_name",
            "due_date",
        ]:
            result[column] = np.nan
        return result
    loan_columns = [
        column
        for column in [
            "loan_id",
            "loan_amount",
            "loan_balance",
            "amount_paid",
            "outstanding_principle",
            "outstanding_interest",
            "outstanding_penalty_fees",
            "status_name",
            "due_date",
        ]
        if column in loans.columns
    ]
    loan_frame = loans[loan_columns].drop_duplicates("loan_id").copy()
    return result.merge(loan_frame, left_on="reference_credit_operation", right_on="loan_id", how="left")


def build_customer_summary(statement: pd.DataFrame, dat_client: pd.DataFrame, loans_client: pd.DataFrame, customer_id: str) -> pd.DataFrame:
    if statement.empty:
        return pd.DataFrame()
    rows: list[dict[str, object]] = []
    for currency, group in statement.groupby("currency_code", sort=True, dropna=False):
        real_balance = bool(group["solde_mpesa_disponible"].fillna(False).all())
        loan_balance = 0.0
        loan_count = 0
        if not loans_client.empty and "currency_code" in loans_client.columns:
            loans_currency = loans_client.loc[loans_client["currency_code"].eq(currency)]
            loan_count = int(loans_currency["loan_id"].nunique()) if "loan_id" in loans_currency.columns else len(loans_currency)
            if "loan_balance" in loans_currency.columns:
                loan_balance = float(pd.to_numeric(loans_currency["loan_balance"], errors="coerce").fillna(0).sum())
        rows.append(
            {
                "customer_id": customer_id,
                "Nom_client": concat_unique(group["Nom_client"]) if "Nom_client" in group.columns else "",
                "telephone": concat_unique(group["telephone"]),
                "devise": currency,
                "nombre_operations_mpesa": int(len(group)),
                "premiere_transaction": group["created_at"].min(),
                "derniere_transaction": group["created_at"].max(),
                "total_entrees_mpesa": float(group["entree_mpesa"].sum()),
                "total_sorties_mpesa": float(group["sortie_mpesa"].sum()),
                "mouvement_net": float(group["mouvement_net_mpesa"].sum()),
                "solde_mpesa_final": group["solde_mpesa_apres"].iloc[-1] if real_balance else np.nan,
                "solde_mpesa_est_reel": real_balance,
                "epargne_courante_finale": float(group["epargne_courante_finale"].iloc[-1]),
                "dat_final": float(group["dat_final_client"].iloc[-1]),
                "nombre_dat": int(dat_client["currency_code"].eq(currency).sum()) if not dat_client.empty and "currency_code" in dat_client.columns else 0,
                "nombre_credits": loan_count,
                "solde_credit_total": loan_balance,
            }
        )
    return pd.DataFrame(rows)


def build_diagnostics(prepared: MpesaPreparedData, customer_id: str | None = None, statement: pd.DataFrame | None = None) -> pd.DataFrame:
    tx = prepared.transactions
    diagnostics: list[dict[str, object]] = []

    def add(control: str, value: int | str, status: str, detail: str = "") -> None:
        diagnostics.append({"controle": control, "valeur": value, "statut": status, "detail": detail})

    def add_dat_controls() -> None:
        fixed = prepared.fixed_savings
        if not isinstance(fixed, pd.DataFrame) or fixed.empty:
            return
        fixed = fixed.copy()
        if customer_id is not None and "customer_id" in fixed.columns:
            fixed = fixed.loc[fixed["customer_id"].astype("string").eq(str(customer_id))].copy()
        if fixed.empty or not {"date_approved", "maturity_date"}.issubset(fixed.columns):
            return
        approved = pd.to_datetime(fixed["date_approved"], errors="coerce")
        maturity = pd.to_datetime(fixed["maturity_date"], errors="coerce")
        invalid_order = approved.notna() & maturity.notna() & maturity.lt(approved)
        invalid_count = int(invalid_order.sum())
        add(
            "DAT - echeance anterieure a l'approbation",
            invalid_count,
            "OK" if invalid_count == 0 else "A surveiller",
            "Verifier les dates d'approbation et d'echeance des DAT concernes.",
        )

        analysis_date = pd.to_datetime(_mpesa_analysis_date(prepared), errors="coerce")
        if pd.isna(analysis_date):
            analysis_date = pd.Timestamp.now()
        positive_balance = numeric_column(fixed, "balance").gt(0)
        expired_positive = maturity.notna() & maturity.lt(pd.Timestamp(analysis_date).normalize()) & positive_balance
        expired_count = int(expired_positive.sum())
        rate = 100 * expired_count / len(fixed) if len(fixed) else 0.0
        add(
            "DAT echus avec solde positif",
            expired_count,
            "OK" if expired_count == 0 else "Controle metier",
            (
                f"{rate:.1f}% des DAT du perimetre. Peut correspondre a un renouvellement ou a un DAT non cloture; "
                "ne pas traiter automatiquement comme une erreur de donnees."
            ),
        )

    if tx.empty:
        add("Transactions chargees", 0, "A verifier", "Aucun fichier Transactions exploitable.")
        if not prepared.g2_transactions.empty:
            add("Transactions G2 chargees", int(len(prepared.g2_transactions)), "Information")
        add_dat_controls()
        return pd.DataFrame(diagnostics)

    missing_customer = int(tx["customer_id"].apply(_is_empty_text).sum()) if "customer_id" in tx.columns else len(tx)
    add("Lignes sans customer_id", missing_customer, "OK" if missing_customer == 0 else "A verifier")
    reference_missing = 0
    if "reference_id" in tx.columns:
        reference_missing += int(tx["reference_id"].apply(_is_empty_text).sum())
    add("Lignes sans reference_id", reference_missing, "OK" if reference_missing == 0 else "A verifier")
    ref_no_missing = int(tx["ref_no"].apply(_is_empty_text).sum()) if "ref_no" in tx.columns else len(tx)
    add("Lignes sans ref_no", ref_no_missing, "OK" if ref_no_missing == 0 else "Information")
    invalid_dates = int(pd.to_datetime(tx["created_at"], errors="coerce").isna().sum()) if "created_at" in tx.columns else len(tx)
    add("Dates invalides", invalid_dates, "OK" if invalid_dates == 0 else "A verifier")
    dr_values = numeric_column(tx, "dr")
    cr_values = numeric_column(tx, "cr")
    zero_moves = int((dr_values.eq(0) & cr_values.eq(0)).sum())
    add("Mouvements dr = 0 et cr = 0", zero_moves, "OK" if zero_moves == 0 else "A verifier")
    both_dr_cr = int((dr_values.gt(0) & cr_values.gt(0)).sum())
    add("Lignes avec dr > 0 et cr > 0", both_dr_cr, "OK" if both_dr_cr == 0 else "A verifier")
    negative_balance = int((numeric_column(tx, "bal_before").lt(0) | numeric_column(tx, "bal_after").lt(0)).sum())
    add("Soldes bal_before/bal_after negatifs", negative_balance, "OK" if negative_balance == 0 else "A verifier")
    empty_currency = int(tx["currency_code"].apply(_is_empty_text).sum()) if "currency_code" in tx.columns else len(tx)
    add("currency_code vide", empty_currency, "OK" if empty_currency == 0 else "A verifier")
    if "account_type" in tx.columns:
        account_types = clean_text(tx["account_type"]).str.upper()
        empty_account = int(account_types.eq("").sum())
        unknown_mask = account_types.ne("") & ~account_types.isin(KNOWN_ACCOUNT_TYPES)
        unknown_account = int(unknown_mask.sum())
        unknown_values = concat_unique(account_types.loc[unknown_mask])
    else:
        empty_account = len(tx)
        unknown_account = 0
        unknown_values = ""
    add("account_type vide", empty_account, "OK" if empty_account == 0 else "A verifier")
    add(
        "Types de comptes a classifier",
        unknown_account,
        "OK" if unknown_account == 0 else "A verifier",
        f"Valeurs non referencees : {unknown_values}" if unknown_values else "Tous les types charges sont references.",
    )

    exact_duplicates = int(tx.duplicated().sum())
    add("Doublons exacts", exact_duplicates, "OK" if exact_duplicates == 0 else "A verifier")
    duplicate_key_columns = [
        column
        for column in ["customer_id", "created_at", "ref_no", "reference_id", "dr", "cr"]
        if column in tx.columns
    ]
    linked_groups = 0
    linked_rows = 0
    repeated_to_review = 0
    if duplicate_key_columns:
        duplicate_work = tx.copy()
        duplicate_work["__account_type_control"] = clean_text(
            duplicate_work.get("account_type", pd.Series("", index=duplicate_work.index))
        ).str.upper()
        duplicate_work["__id_control"] = clean_identifier(
            duplicate_work.get("id", pd.Series("", index=duplicate_work.index))
        )
        duplicate_groups = (
            duplicate_work.groupby(duplicate_key_columns, dropna=False, as_index=False)
            .agg(
                nombre_lignes=("__account_type_control", "size"),
                nombre_types_comptes=("__account_type_control", "nunique"),
                nombre_ids=("__id_control", "nunique"),
            )
        )
        repeated = duplicate_groups["nombre_lignes"].gt(1)
        linked = repeated & duplicate_groups["nombre_types_comptes"].gt(1)
        linked_groups = int(linked.sum())
        linked_rows = int(duplicate_groups.loc[linked, "nombre_lignes"].sum())
        repeated_to_review = int((repeated & ~linked).sum())
    add(
        "Ecritures comptables liees",
        linked_groups,
        "Information",
        f"{linked_rows} lignes reparties dans {linked_groups} groupe(s) multi-comptes; elles ne sont pas des doublons.",
    )
    add(
        "Groupes d'ecritures repetees a verifier",
        repeated_to_review,
        "OK" if repeated_to_review == 0 else "A verifier",
        "Memes attributs de controle sans changement de type de compte.",
    )
    add_dat_controls()
    if not prepared.g2_transactions.empty:
        g2_missing_phone = (
            int(prepared.g2_transactions["phone_prefixe"].isna().sum())
            if "phone_prefixe" in prepared.g2_transactions.columns
            else int(len(prepared.g2_transactions))
        )
        add("Transactions G2 chargees", int(len(prepared.g2_transactions)), "Information")
        add("Transactions G2 sans telephone exploitable", g2_missing_phone, "OK" if g2_missing_phone == 0 else "A verifier")

    if customer_id:
        clients_dat = set(prepared.fixed_savings["customer_id"].dropna().astype(str)) if not prepared.fixed_savings.empty and "customer_id" in prepared.fixed_savings.columns else set()
        clients_savings = set(prepared.current_savings["customer_id"].dropna().astype(str)) if not prepared.current_savings.empty and "customer_id" in prepared.current_savings.columns else set()
        add("Client sans fichier DAT", int(str(customer_id) not in clients_dat), "OK" if str(customer_id) in clients_dat else "Information")
        add("Client sans compte epargne", int(str(customer_id) not in clients_savings), "OK" if str(customer_id) in clients_savings else "Information")
        if statement is not None and not statement.empty:
            dat_refs = statement["reference_dat_operation"].astype("string").fillna("")
            credit_refs = statement["reference_credit_operation"].astype("string").fillna("")
            known_loans = set(prepared.loans["loan_id"].dropna().astype(str)) if not prepared.loans.empty and "loan_id" in prepared.loans.columns else set()
            unknown_credit_refs = int(credit_refs.ne("").sum()) if not known_loans else int((credit_refs.ne("") & ~credit_refs.isin(known_loans)).sum())
            add("References DAT detectees", int(dat_refs.ne("").sum()), "Information")
            add("References credit non reconnues", unknown_credit_refs, "OK" if unknown_credit_refs == 0 else "A verifier")

    return pd.DataFrame(diagnostics)


def format_statement_columns(statement: pd.DataFrame) -> pd.DataFrame:
    ordered = [
        "created_at",
        "customer_id",
        "Nom_client",
        "mode_rapprochement_nom_client",
        "telephone",
        "currency_code",
        "type_operation",
        "operation_reference",
        "references_internes",
        "reference_dat_operation",
        "reference_credit_operation",
        "description_turbo",
        "descriptions",
        "entree_mpesa",
        "sortie_mpesa",
        "mouvement_net_mpesa",
        "solde_mpesa_avant",
        "solde_mpesa_apres",
        "cumul_net_depuis_debut_fichier",
        "solde_epargne_au_moment",
        "solde_dat_operation_avant",
        "solde_dat_operation_apres",
        "variation_dat_operation",
        "solde_dat_total_au_moment",
        "dat_final_client",
        "epargne_courante_finale",
        "loan_id",
        "loan_amount",
        "loan_balance",
        "amount_paid",
        "outstanding_principle",
        "outstanding_interest",
        "outstanding_penalty_fees",
        "status_name",
        "due_date",
        "controle_mouvement",
        "nombre_lignes_comptables",
    ]
    present = [column for column in ordered if column in statement.columns]
    rest = [column for column in statement.columns if column not in present]
    return statement[present + rest].copy()


def build_customer_statement_view(
    statement: pd.DataFrame,
    *,
    account_number: object = "",
    entry_account_number: object | None = None,
    output_account_number: object | None = None,
    allow_multiple_currencies: bool = False,
) -> dict[str, Any]:
    """Construit l'extrait officiel selon le sens metier Bisou Bisou.

    Le compte Turbo ``MPESA ACCOUNT`` porte le mouvement inverse du flux de
    l'organisation : un debit Turbo correspond a une entree Bisou Bisou et un
    credit Turbo a une sortie. Les cumuls restent calcules separement par devise.
    """
    legacy_account = str(account_number).strip()
    entry_account = (
        legacy_account if entry_account_number is None else str(entry_account_number).strip()
    )
    output_account = (
        legacy_account if output_account_number is None else str(output_account_number).strip()
    )
    empty_transactions = pd.DataFrame(columns=CUSTOMER_STATEMENT_COLUMNS)
    empty_result: dict[str, Any] = {
        "transactions": empty_transactions,
        "currency": "",
        "customer_id": "",
        "customer_name": "",
        "telephone": "",
        "account_number": legacy_account,
        "entry_account_number": entry_account,
        "output_account_number": output_account,
        "total_entries": 0.0,
        "total_outputs": 0.0,
        "opening_amount": np.nan,
        "closing_amount": np.nan,
        "balance_is_real": False,
        "balance_label": "Cumul net",
        "summary_by_currency": pd.DataFrame(),
    }
    if not isinstance(statement, pd.DataFrame) or statement.empty:
        return empty_result

    frame = statement.copy()
    frame["currency_code"] = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper()
    currencies = [value for value in frame["currency_code"].dropna().unique().tolist() if value]
    if len(currencies) > 1 and not allow_multiple_currencies:
        raise ValueError("L'extrait client officiel doit contenir une seule devise.")
    currency = "ALL" if len(currencies) > 1 else currencies[0] if currencies else "SANS DEVISE"

    frame["created_at"] = pd.to_datetime(
        frame.get("created_at", pd.Series(pd.NaT, index=frame.index)), errors="coerce"
    )
    frame = frame.sort_values(["created_at", "operation_reference"], na_position="last").reset_index(drop=True)
    turbo_entries = pd.to_numeric(
        frame.get("entree_mpesa", pd.Series(0.0, index=frame.index)), errors="coerce"
    ).fillna(0.0)
    turbo_outputs = pd.to_numeric(
        frame.get("sortie_mpesa", pd.Series(0.0, index=frame.index)), errors="coerce"
    ).fillna(0.0)
    # Le sens du compte MPESA Turbo est inverse du flux de l'organisation.
    entries = turbo_outputs
    outputs = turbo_entries
    movements = entries - outputs

    real_before = pd.to_numeric(
        frame.get("solde_mpesa_avant", pd.Series(np.nan, index=frame.index)), errors="coerce"
    )
    if "solde_mpesa_disponible" in frame.columns:
        available = frame["solde_mpesa_disponible"].astype("boolean").fillna(False)
    else:
        available = real_before.notna()

    displayed_balance = pd.Series(np.nan, index=frame.index, dtype="float64")
    summary_rows: list[dict[str, object]] = []
    for currency_code, index in frame.groupby("currency_code", sort=True, dropna=False).groups.items():
        group_movements = movements.loc[index]
        group_real = bool(available.loc[index].all() and real_before.loc[index].notna().all())
        if group_real:
            opening = float(real_before.loc[index].iloc[0])
            group_balance = opening + group_movements.cumsum()
        else:
            opening = 0.0
            group_balance = group_movements.cumsum()
        displayed_balance.loc[index] = group_balance
        summary_rows.append(
            {
                "currency_code": currency_code or "SANS DEVISE",
                "total_entries": float(entries.loc[index].sum()),
                "total_outputs": float(outputs.loc[index].sum()),
                "opening_amount": opening,
                "closing_amount": float(group_balance.iloc[-1]),
                "balance_is_real": group_real,
                "balance_label": "Solde" if group_real else "Cumul net",
            }
        )
    summary_by_currency = pd.DataFrame(summary_rows)
    balance_is_real = bool(
        not summary_by_currency.empty
        and summary_by_currency["balance_is_real"].astype(bool).all()
    )

    names = frame.get("Nom_client", pd.Series("", index=frame.index))
    phones = frame.get("telephone", pd.Series("", index=frame.index))

    def statement_description(row: pd.Series) -> str:
        turbo_description = row.get("description_turbo")
        if _is_empty_text(turbo_description):
            turbo_description = row.get("descriptions")
        parts = [turbo_description, row.get("telephone"), row.get("Nom_client")]
        return " - ".join(str(value).strip() for value in parts if not _is_empty_text(value))

    account_values = np.select(
        [entries.gt(0), outputs.gt(0)],
        [entry_account, output_account],
        default=entry_account or output_account,
    )
    transactions = pd.DataFrame(
        {
            "date": frame["created_at"],
            "compte": account_values,
            "receipt_no": frame.get("operation_reference", pd.Series("", index=frame.index)),
            "devise": frame["currency_code"],
            "description": frame.apply(statement_description, axis=1),
            "entree": entries,
            "sortie": outputs,
            "solde": displayed_balance,
        }
    )[CUSTOMER_STATEMENT_COLUMNS]

    if len(summary_by_currency) == 1:
        summary_row = summary_by_currency.iloc[0]
        opening_amount = summary_row["opening_amount"]
        closing_amount = summary_row["closing_amount"]
        total_entries = summary_row["total_entries"]
        total_outputs = summary_row["total_outputs"]
        balance_label = summary_row["balance_label"]
    else:
        opening_amount = np.nan
        closing_amount = np.nan
        total_entries = np.nan
        total_outputs = np.nan
        balance_label = "Solde par devise" if balance_is_real else "Cumul net par devise"

    result = dict(empty_result)
    result.update(
        {
            "transactions": transactions,
            "currency": currency,
            "customer_id": concat_unique(frame.get("customer_id", pd.Series("", index=frame.index))),
            "customer_name": concat_unique(names),
            "telephone": concat_unique(phones),
            "total_entries": float(total_entries) if pd.notna(total_entries) else np.nan,
            "total_outputs": float(total_outputs) if pd.notna(total_outputs) else np.nan,
            "opening_amount": float(opening_amount) if pd.notna(opening_amount) else np.nan,
            "closing_amount": float(closing_amount) if pd.notna(closing_amount) else np.nan,
            "balance_is_real": balance_is_real,
            "balance_label": balance_label,
            "entry_account_number": entry_account,
            "output_account_number": output_account,
            "summary_by_currency": summary_by_currency,
        }
    )
    return result


def _pdf_number(value: Any, *, decimals: int = 0) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "-"
    if pd.isna(number):
        return "-"
    return f"{number:,.{decimals}f}".replace(",", " ")


def _retention_svg(monthly: pd.DataFrame) -> str:
    if monthly.empty:
        return ""
    charts: list[str] = []
    colors = {"retention_m1_pct": "#17805c", "retention_90j_pct": "#173f73"}
    labels = {"retention_m1_pct": "Retention M+1", "retention_90j_pct": "Retention 90 jours"}
    for currency, frame in monthly.groupby("currency_code", dropna=False):
        frame = frame.sort_values("periode").reset_index(drop=True)
        if frame[["retention_m1_pct", "retention_90j_pct"]].notna().sum().sum() == 0:
            continue
        width, height = 760, 270
        left, right, top, bottom = 55, 22, 35, 48
        plot_width = width - left - right
        plot_height = height - top - bottom
        count = max(len(frame), 1)

        def x_position(index: int) -> float:
            return left + (plot_width / max(count - 1, 1)) * index

        def y_position(value: float) -> float:
            return top + plot_height * (1 - max(0.0, min(100.0, value)) / 100)

        svg_parts = [
            f'<svg class="retention-chart" viewBox="0 0 {width} {height}" role="img" '
            f'aria-label="Evolution de la retention {escape(str(currency))}">',
            f'<text x="{left}" y="20" class="chart-title">Devise {escape(str(currency))}</text>',
        ]
        for tick in [0, 25, 50, 75, 100]:
            y = y_position(float(tick))
            svg_parts.append(f'<line x1="{left}" y1="{y:.1f}" x2="{width-right}" y2="{y:.1f}" class="grid"/>')
            svg_parts.append(f'<text x="{left-10}" y="{y+4:.1f}" text-anchor="end" class="axis-label">{tick}%</text>')
        label_step = max(1, int(np.ceil(count / 6)))
        for index, month in enumerate(frame["mois"].astype(str)):
            if index % label_step == 0 or index == count - 1:
                x = x_position(index)
                svg_parts.append(f'<text x="{x:.1f}" y="{height-18}" text-anchor="middle" class="axis-label">{escape(month)}</text>')
        for column in ["retention_m1_pct", "retention_90j_pct"]:
            points: list[tuple[float, float, float]] = []
            for index, value in enumerate(pd.to_numeric(frame[column], errors="coerce")):
                if pd.notna(value):
                    points.append((x_position(index), y_position(float(value)), float(value)))
            if points:
                polyline = " ".join(f"{x:.1f},{y:.1f}" for x, y, _ in points)
                svg_parts.append(
                    f'<polyline points="{polyline}" fill="none" stroke="{colors[column]}" stroke-width="3"/>'
                )
                for x, y, value in points:
                    svg_parts.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="4" fill="{colors[column]}"/>')
                    svg_parts.append(
                        f'<text x="{x:.1f}" y="{y-8:.1f}" text-anchor="middle" class="point-label">{value:.0f}%</text>'
                    )
        legend_x = width - 300
        for offset, column in enumerate(["retention_m1_pct", "retention_90j_pct"]):
            x = legend_x + offset * 150
            svg_parts.append(f'<line x1="{x}" y1="18" x2="{x+24}" y2="18" stroke="{colors[column]}" stroke-width="3"/>')
            svg_parts.append(f'<text x="{x+30}" y="22" class="legend">{labels[column]}</text>')
        svg_parts.append("</svg>")
        charts.append("".join(svg_parts))
    return "".join(charts)


def _html_table(frame: pd.DataFrame, columns: list[str], labels: dict[str, str], *, limit: int = 80) -> str:
    present = [column for column in columns if column in frame.columns]
    if frame.empty or not present:
        return '<p class="empty">Aucune donnee disponible.</p>'
    display = frame[present].head(limit).rename(columns=labels).copy()
    for column in display.columns:
        source_name = next((key for key, label in labels.items() if label == column), column)
        if source_name.endswith("_pct"):
            display[column] = pd.to_numeric(display[column], errors="coerce").map(
                lambda value: f"{value:.1f}%" if pd.notna(value) else "-"
            )
        elif any(token in source_name for token in ["montant", "solde", "volume"]):
            display[column] = pd.to_numeric(display[column], errors="coerce").map(
                lambda value: _pdf_number(value, decimals=2)
            )
        else:
            display[column] = display[column].where(display[column].notna(), "-")
    return display.to_html(index=False, escape=True, border=0, classes="report-table")


def _g2_executive_context(report: dict[str, Any]) -> dict[str, Any]:
    source_label = str(report.get("analysis_source_label", "G2") or "G2")
    turbo_only = source_label.casefold() == "turbo"
    daily_pivot = report.get("rapport_journalier_pivot", pd.DataFrame())
    g2_dat = report.get("g2_dat", pd.DataFrame())
    monthly = report.get("retention_mensuelle", pd.DataFrame())
    transaction_detail = report.get("rapport_journalier_detail", pd.DataFrame())
    transactions_by_day = report.get("transactions_par_jour", pd.DataFrame())
    transactions_by_weekday = report.get("transactions_par_jour_semaine", pd.DataFrame())
    transactions_by_hour = report.get("transactions_par_heure", pd.DataFrame())
    status_summary = report.get("statuts_g2", pd.DataFrame())

    if not isinstance(transaction_detail, pd.DataFrame):
        transaction_detail = pd.DataFrame()
    eligible_transaction_detail = transaction_detail.copy()
    if (
        not eligible_transaction_detail.empty
        and "incluse_synthese" in eligible_transaction_detail.columns
    ):
        eligible_mask = (
            eligible_transaction_detail["incluse_synthese"]
            .astype("boolean")
            .fillna(False)
            .astype(bool)
        )
        eligible_transaction_detail = eligible_transaction_detail.loc[eligible_mask].copy()
    if not isinstance(transactions_by_day, pd.DataFrame):
        transactions_by_day = pd.DataFrame()
    if not isinstance(transactions_by_weekday, pd.DataFrame):
        transactions_by_weekday = pd.DataFrame()
    if not isinstance(transactions_by_hour, pd.DataFrame):
        transactions_by_hour = pd.DataFrame()
    if not isinstance(status_summary, pd.DataFrame):
        status_summary = pd.DataFrame()
    if status_summary.empty and not transaction_detail.empty:
        status_summary = build_g2_transaction_status_summary(transaction_detail)
    if (
        transactions_by_day.empty
        or transactions_by_weekday.empty
        or transactions_by_hour.empty
    ) and not transaction_detail.empty:
        time_report = build_g2_transaction_time_analysis(transaction_detail)
        transactions_by_day = time_report["par_jour"]
        transactions_by_weekday = time_report["par_jour_semaine"]
        transactions_by_hour = time_report["par_heure"]

    active_items: list[str] = []
    retention_items: list[str] = []
    latest_retention_rows: list[dict[str, object]] = []
    chart_blocks: list[pd.DataFrame] = []
    if not eligible_transaction_detail.empty:
        activity = eligible_transaction_detail.copy()
        activity["currency_code"] = clean_text(
            activity.get("currency_code", pd.Series("", index=activity.index))
        ).str.upper().replace("", "SANS DEVISE")
        phone_values = normalize_phone(
            activity.get("phone_prefixe", pd.Series(pd.NA, index=activity.index, dtype="string"))
        )
        if phone_values.notna().sum() == 0 and "opposite_party" in activity.columns:
            phone_values = normalize_phone(
                _extract_phone_from_opposite_party(activity["opposite_party"])
            )
        activity["__client_phone"] = phone_values
        for currency, frame in activity.groupby("currency_code", dropna=False, sort=True):
            transaction_count = int(len(frame))
            client_count = int(frame["__client_phone"].dropna().nunique())
            operation_status = "operation(s) comptabilisee(s) dans Turbo" if turbo_only else "transaction(s) Completed"
            active_items.append(
                f"{currency} : {transaction_count} {operation_status}, "
                f"{client_count} client(s) distinct(s)"
            )
    if not monthly.empty:
        for currency, frame in monthly.groupby("currency_code", dropna=False):
            frame = frame.sort_values("periode")
            eligible_m1 = frame.dropna(subset=["retention_m1_pct"])
            eligible_90 = frame.dropna(subset=["retention_90j_pct"])
            m1_text = f"{eligible_m1.iloc[-1]['retention_m1_pct']:.1f}%" if not eligible_m1.empty else "non calculable"
            day90_text = f"{eligible_90.iloc[-1]['retention_90j_pct']:.1f}%" if not eligible_90.empty else "non calculable"
            retention_items.append(f"{currency} : M+1 {m1_text}, 90 jours {day90_text}")
            latest_retention_rows.append(
                {
                    "Devise": currency,
                    "Retention M+1": m1_text,
                    "Retention 90 jours": day90_text,
                }
            )
            eligible_periods = int(frame[["retention_m1_pct", "retention_90j_pct"]].notna().any(axis=1).sum())
            if eligible_periods >= 6:
                chart_blocks.append(frame)

    flow_items: list[str] = []
    if not daily_pivot.empty:
        for _, row in daily_pivot.iterrows():
            currency = row.get("currency_code", "-")
            flow_items.append(
                f"{currency} : entrees {_pdf_number(row.get('montant_total_entrees'), decimals=2)}, "
                f"sorties {_pdf_number(row.get('montant_total_sorties'), decimals=2)}, "
                f"net {_pdf_number(row.get('solde_net_flux'), decimals=2)}"
            )

    status_text = ""
    if not status_summary.empty and "nombre_transactions" in status_summary.columns:
        status_counts = pd.to_numeric(status_summary["nombre_transactions"], errors="coerce").fillna(0)
        included = clean_text(
            status_summary.get("prise_en_compte_analyse", pd.Series("", index=status_summary.index))
        ).eq("Oui")
        completed_count = int(status_counts.loc[included].sum())
        control_only_count = int(status_counts.loc[~included].sum())
        if turbo_only:
            status_text = (
                f"{completed_count} operation(s) comptabilisee(s) dans Turbo incluse(s) dans les analyses; "
                f"{control_only_count} operation(s) exclue(s) du perimetre analytique. "
                "Le statut de transaction G2 n'est pas disponible."
            )
        else:
            status_text = (
                f"{completed_count} transaction(s) Completed incluse(s) dans les analyses; "
                f"{control_only_count} transaction(s) d'autres statuts conservee(s) pour controle uniquement."
            )

    time_items: list[str] = []
    total_transactions = 0
    if not transactions_by_hour.empty and "nombre_transactions" in transactions_by_hour.columns:
        hourly_totals = (
            transactions_by_hour.groupby(["heure_num", "heure"], as_index=False, dropna=False)["nombre_transactions"]
            .sum()
            .sort_values("heure_num")
            .reset_index(drop=True)
        )
        hourly_totals["nombre_transactions"] = pd.to_numeric(
            hourly_totals["nombre_transactions"], errors="coerce"
        ).fillna(0)
        total_transactions = int(hourly_totals["nombre_transactions"].sum())
        if total_transactions > 0:
            busiest_hour = hourly_totals.loc[hourly_totals["nombre_transactions"].idxmax()]
            busiest_hour_count = int(busiest_hour["nombre_transactions"])
            busiest_hour_share = 100 * busiest_hour_count / total_transactions
            time_items.append(
                f"Heure la plus fréquente : {busiest_hour.get('heure', '-')}, "
                f"avec {busiest_hour_count} transaction(s), soit {busiest_hour_share:.1f}% du volume"
            )

    analysis_start = pd.to_datetime(report.get("analysis_date_start"), errors="coerce")
    analysis_end = pd.to_datetime(report.get("analysis_date_end"), errors="coerce")
    if pd.notna(analysis_start) and pd.notna(analysis_end):
        spans_multiple_days = analysis_end.normalize() > analysis_start.normalize()
    else:
        detail_dates = pd.to_datetime(
            eligible_transaction_detail.get("date", pd.Series(dtype="datetime64[ns]")), errors="coerce"
        ).dropna()
        spans_multiple_days = detail_dates.dt.normalize().nunique() > 1
        if not spans_multiple_days and not transactions_by_day.empty:
            reported_dates = pd.to_datetime(
                transactions_by_day.get(
                    "date_transaction", pd.Series(dtype="datetime64[ns]")
                ),
                errors="coerce",
            ).dropna()
            spans_multiple_days = reported_dates.dt.normalize().nunique() > 1

    if (
        spans_multiple_days
        and not transactions_by_weekday.empty
        and "nombre_transactions" in transactions_by_weekday.columns
    ):
        weekday_totals = (
            transactions_by_weekday.groupby(
                ["jour_semaine_num", "jour_semaine"], as_index=False, dropna=False
            )["nombre_transactions"]
            .sum()
            .sort_values("jour_semaine_num")
            .reset_index(drop=True)
        )
        weekday_totals["nombre_transactions"] = pd.to_numeric(
            weekday_totals["nombre_transactions"], errors="coerce"
        ).fillna(0)
        weekday_total = int(weekday_totals["nombre_transactions"].sum())
        if weekday_total > 0:
            busiest_weekday = weekday_totals.loc[weekday_totals["nombre_transactions"].idxmax()]
            busiest_weekday_count = int(busiest_weekday["nombre_transactions"])
            busiest_weekday_share = 100 * busiest_weekday_count / weekday_total
            time_items.append(
                f"Jour de semaine le plus actif : {busiest_weekday.get('jour_semaine', '-')}, "
                f"avec {busiest_weekday_count} transaction(s), soit {busiest_weekday_share:.1f}% du volume"
            )

    control_text = ""
    if turbo_only:
        control_text = (
            "Mode Turbo seul : les controles croises G2/Turbo sont non applicables. "
            "Les depots sont regroupes par ref_no et les retraits M-PESA par reference_id + created_at."
        )
    elif not g2_dat.empty:
        if "statut_rapprochement" in g2_dat.columns:
            reference_status = g2_dat["statut_rapprochement"].astype("string").fillna("")
            exact = int(reference_status.eq("Rapproche exact").sum())
            with_gap = int(reference_status.eq("Rapproche avec ecart").sum())
            unmatched = int(reference_status.eq("Non rapproche").sum())
            anomalies = int(
                g2_dat.get("est_anomalie", pd.Series(False, index=g2_dat.index))
                .fillna(False)
                .astype(bool)
                .sum()
            )
            total = int(len(g2_dat))
            rate = 100 * exact / total if total else 0
            control_text = (
                f"Rapprochement Receipt No/ref_no : {exact}/{total} exact(s), soit {rate:.1f}%; "
                f"{with_gap} avec ecart, {unmatched} non rapproche(s), {anomalies} anomalie(s) au total."
            )
        else:
            status = g2_dat.get("statut_rapprochement_dat", pd.Series("", index=g2_dat.index)).astype("string")
            dat_absent = bool(status.str.contains("Fichier DAT absent", case=False, na=False).all())
            if dat_absent:
                control_text = "Rapprochement G2/DAT non evalue : fichier DAT non charge."
            else:
                matched = int(
                    g2_dat.get("customer_id_dat", pd.Series(dtype="string"))
                    .astype("string").fillna("").ne("").sum()
                )
                total = int(len(g2_dat))
                rate = 100 * matched / total if total else 0
                control_text = f"Rapprochement G2/DAT : {matched}/{total} entree(s), soit {rate:.1f}%."

    has_retention = bool(
        not monthly.empty
        and monthly[["retention_m1_pct", "retention_90j_pct"]].notna().any().any()
    )
    attention_text = (
        "Les mois les plus recents restent provisoires tant que leur fenetre de suivi n'est pas terminee."
        if has_retention
        else "La fidelisation M+1 et a 90 jours exige un historique couvrant les mois suivants."
    )
    return {
        "active_text": "; ".join(active_items) or "Aucune transaction Completed dans le perimetre filtre.",
        "flow_text": "; ".join(flow_items) or "Aucun flux disponible.",
        "status_text": status_text,
        "time_text": "; ".join(time_items),
        "retention_text": "; ".join(retention_items) or "Non calculable sur la periode chargee.",
        "control_text": control_text,
        "attention_text": attention_text,
        "has_retention": has_retention,
        "latest_retention": pd.DataFrame(latest_retention_rows),
        "chart_monthly": concat_frames_stable(chart_blocks) if chart_blocks else pd.DataFrame(),
    }


def build_g2_dat_pdf_html(
    report: dict[str, Any],
    *,
    period_text: str,
    direction_label: str,
    generated_at: pd.Timestamp | None = None,
) -> str:
    """Construit la version courte destinee a la Direction generale."""
    daily_pivot = report.get("rapport_journalier_pivot", pd.DataFrame())
    daily_synthese = report.get("rapport_journalier_synthese", pd.DataFrame())
    generated_at = generated_at if generated_at is not None else pd.Timestamp.now()
    context = _g2_executive_context(report)

    flow_table = _html_table(
        daily_pivot,
        ["currency_code", "nombre_entrees", "montant_total_entrees", "nombre_sorties", "montant_total_sorties", "solde_net_flux"],
        {
            "currency_code": "Devise", "nombre_entrees": "Nb entrees", "montant_total_entrees": "Montant entrees",
            "nombre_sorties": "Nb sorties", "montant_total_sorties": "Montant sorties", "solde_net_flux": "Solde net",
        },
    )
    classified = daily_synthese.copy()
    if not classified.empty and "details_rapport" in classified.columns:
        classified = classified.loc[~classified["details_rapport"].astype("string").str.startswith("Total ", na=False)]
    operation_table = _html_table(
        classified,
        ["currency_code", "sens_flux", "details_rapport", "nombre", "montant"],
        {"currency_code": "Devise", "sens_flux": "Sens", "details_rapport": "Operation", "nombre": "Nombre", "montant": "Montant"},
    )
    retention_table = _html_table(
        context["latest_retention"],
        ["Devise", "Retention M+1", "Retention 90 jours"],
        {},
    )
    chart_html = _retention_svg(context["chart_monthly"])
    control_bullet = (
        f'<li><strong>Controle.</strong> {escape(context["control_text"])}</li>'
        if context["control_text"] else ""
    )
    retention_section = (
        f'<h2>Fidelisation</h2><p>{escape(context["retention_text"])}</p>{chart_html}{retention_table}'
        if context["has_retention"]
        else f'<div class="note"><strong>Point de vigilance.</strong> {escape(context["attention_text"])}</div>'
    )

    return f"""<!doctype html>
<html lang="fr"><head><meta charset="utf-8"><title>Rapport M-PESA - G2/DAT</title>
<style>
@page {{ size: A4; margin: 12mm; }}
* {{ box-sizing: border-box; }}
body {{ font-family: Arial, Helvetica, sans-serif; color: #18314f; margin: 0; font-size: 9.5pt; line-height: 1.3; }}
h1 {{ font-size: 22pt; color: #12385f; margin: 0 0 3px; }}
h2 {{ font-size: 13.5pt; color: #12385f; margin: 14px 0 6px; border-bottom: 2px solid #dbe6f1; padding-bottom: 3px; page-break-after: avoid; }}
h3 {{ font-size: 11pt; color: #173f73; margin: 10px 0 5px; page-break-after: avoid; }}
p {{ margin: 4px 0 7px; }}
.meta {{ color: #5d7187; margin-bottom: 10px; }}
.summary {{ background: #edf5fb; border-left: 5px solid #2b6ea6; padding: 8px 12px; border-radius: 4px; }}
.summary h2 {{ border: 0; margin: 0 0 5px; font-size: 12.5pt; }}
.summary ul {{ margin: 2px 0; padding-left: 18px; }}
.report-table {{ width: 100%; border-collapse: collapse; margin: 4px 0 9px; font-size: 8.2pt; }}
.report-table th {{ background: #1f4e78; color: white; padding: 5px; text-align: left; }}
.report-table td {{ border-bottom: 1px solid #dce4eb; padding: 4px 5px; }}
.report-table tr {{ page-break-inside: avoid; }}
.retention-chart {{ width: 100%; height: auto; display: block; margin: 3px 0 8px; page-break-inside: avoid; }}
.chart-title {{ font-size: 13px; font-weight: bold; fill: #173f73; }}
.grid {{ stroke: #dfe7ee; stroke-width: 1; }}
.axis-label, .legend {{ font-size: 10px; fill: #5d7187; }}
.point-label {{ font-size: 9px; font-weight: bold; fill: #20364d; }}
.note {{ background: #fff7e8; border-left: 4px solid #d69a25; padding: 7px 9px; margin: 9px 0; }}
.empty {{ color: #6c7d8c; font-style: italic; }}
.footer-note {{ color: #60758a; font-size: 8pt; margin-top: 10px; }}
</style></head><body>
<h1>Rapport M-PESA - G2/DAT</h1>
<p class="meta">{escape(period_text)} | Sens : {escape(direction_label)} | {generated_at:%d/%m/%Y}</p>
<section class="summary"><h2>Synthese executive</h2><ul>
<li><strong>Activite.</strong> {escape(context["active_text"])}</li>
{f'<li><strong>Perimetre des statuts.</strong> {escape(context["status_text"])}</li>' if context["status_text"] else ''}
{f'<li><strong>Fréquence temporelle.</strong> {escape(context["time_text"])}</li>' if context["time_text"] else ''}
<li><strong>Flux financiers.</strong> {escape(context["flow_text"])}</li>
{control_bullet}
</ul></section>
<h2>Flux par devise</h2>{flow_table}
<h3>Principales operations</h3>{operation_table}
{retention_section}
</body></html>"""


def _find_pdf_browser() -> Path | None:
    candidates = [
        shutil.which("msedge"),
        shutil.which("chrome"),
        shutil.which("chromium"),
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/usr/bin/google-chrome",
        "/usr/bin/chromium",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    return None


def create_g2_dat_pdf(
    report: dict[str, Any],
    *,
    period_text: str,
    direction_label: str,
) -> bytes:
    """Convertit le rapport G2/DAT statique en PDF via Edge/Chrome headless."""
    browser = _find_pdf_browser()
    if browser is None:
        raise RuntimeError("Microsoft Edge, Google Chrome ou Chromium est requis pour generer le PDF.")
    html_content = build_g2_dat_pdf_html(
        report,
        period_text=period_text,
        direction_label=direction_label,
    )
    with tempfile.TemporaryDirectory(prefix="mpesa_g2_pdf_") as temp_dir:
        temp_path = Path(temp_dir)
        html_path = temp_path / "rapport_g2_dat.html"
        pdf_path = temp_path / "rapport_g2_dat.pdf"
        html_path.write_text(html_content, encoding="utf-8")
        command = [
            str(browser),
            "--headless=new",
            "--disable-gpu",
            "--no-first-run",
            "--no-default-browser-check",
            "--no-pdf-header-footer",
            "--run-all-compositor-stages-before-draw",
            f"--user-data-dir={temp_path / 'browser_profile'}",
            f"--print-to-pdf={pdf_path}",
            html_path.as_uri(),
        ]
        creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=90,
            creationflags=creation_flags,
            check=False,
        )
        if completed.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size == 0:
            message = completed.stderr.strip() or completed.stdout.strip() or "conversion sans fichier de sortie"
            raise RuntimeError(f"Impossible de generer le PDF G2/DAT : {message[:300]}")
        return pdf_path.read_bytes()


def create_customer_statement_word(
    statement: pd.DataFrame,
    *,
    analysis_report: dict[str, pd.DataFrame] | None = None,
    customer_id: object,
    customer_name: object = "",
    telephone: object = "",
    currency: object,
    account_number: object = "",
    entry_account_number: object | None = None,
    output_account_number: object | None = None,
    period_start: object | None = None,
    period_end: object | None = None,
    generated_at: pd.Timestamp | None = None,
) -> bytes:
    """Genere un extrait Word CDF, USD ou ALL sans sommer les devises."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("La dependance python-docx est requise pour generer l'extrait Word.") from exc

    if not isinstance(statement, pd.DataFrame) or statement.empty:
        raise ValueError("Aucune operation filtree n'est disponible pour l'extrait Word.")
    currency_text = str(currency).strip().upper()
    all_currencies = currency_text in {"ALL", "TOUTES", "TOUS", "CDF + USD"}
    frame = statement.copy()
    frame_currency = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper()
    if not all_currencies:
        frame = frame.loc[frame_currency.eq(currency_text)].copy()
    if frame.empty:
        raise ValueError(f"Aucune operation {currency_text or 'sans devise'} n'est disponible pour l'extrait Word.")

    view = build_customer_statement_view(
        frame,
        account_number=account_number,
        entry_account_number=entry_account_number,
        output_account_number=output_account_number,
        allow_multiple_currencies=all_currencies,
    )
    transactions = view["transactions"]
    customer_id_text = str(customer_id).strip() or view["customer_id"] or "Non disponible"
    customer_name_text = str(customer_name).strip() or view["customer_name"] or "Nom non disponible"
    telephone_text = str(telephone).strip() or view["telephone"] or "Non disponible"
    observed_currencies = sorted(
        value for value in transactions["devise"].dropna().astype(str).unique().tolist() if value
    )
    currency_label = (
        f"ALL ({', '.join(observed_currencies)})"
        if all_currencies
        else currency_text or view["currency"]
    )
    generated_at = generated_at if generated_at is not None else pd.Timestamp.now()

    observed_dates = pd.to_datetime(transactions["date"], errors="coerce").dropna()

    def report_date(value: object | None, fallback: pd.Timestamp | None) -> str:
        parsed = pd.to_datetime(value, errors="coerce") if value is not None else pd.NaT
        if pd.isna(parsed):
            parsed = fallback
        return f"{parsed:%d/%m/%Y}" if parsed is not None and pd.notna(parsed) else "Non disponible"

    first_observed = observed_dates.min() if not observed_dates.empty else None
    last_observed = observed_dates.max() if not observed_dates.empty else None
    start_text = report_date(period_start, first_observed)
    end_text = report_date(period_end, last_observed)
    decimals = 0 if currency_text == "CDF" else 2

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(16)
    styles["Title"].font.color.rgb = RGBColor(24, 41, 58)

    def set_cell_shading(cell: Any, fill: str) -> None:
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_properties.append(shading)

    def set_repeat_header(row: Any) -> None:
        row_properties = row._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        row_properties.append(repeat_header)

    header = document.add_table(rows=2, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    brand_cell = header.cell(0, 0).merge(header.cell(1, 0))
    brand_cell.width = Cm(17.5)
    brand_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    brand_paragraph = brand_cell.paragraphs[0]
    brand_paragraph.paragraph_format.space_after = Pt(0)
    if CUSTOMER_STATEMENT_LOGO_PATH.is_file():
        try:
            brand_paragraph.add_run().add_picture(
                str(CUSTOMER_STATEMENT_LOGO_PATH), width=Cm(4.6)
            )
        except (OSError, ValueError):
            brand_run = brand_paragraph.add_run("IMF Microfinance Bisou Bisou")
            brand_run.bold = True
            brand_run.font.size = Pt(16)
    else:
        brand_run = brand_paragraph.add_run("IMF Microfinance Bisou Bisou")
        brand_run.bold = True
        brand_run.font.size = Pt(16)

    criteria_title = header.cell(0, 1)
    criteria_title.width = Cm(9.0)
    criteria_title.text = "Critères"
    criteria_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(criteria_title, "1F2937")
    for run in criteria_title.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    criteria_cell = header.cell(1, 1)
    criteria_cell.width = Cm(9.0)
    criteria = criteria_cell.add_table(rows=5, cols=2)
    criteria.autofit = False
    criteria_rows = [
        ("Date du :", start_text),
        ("Au :", end_text),
        ("Numéro du client :", customer_id_text),
        ("Téléphone :", telephone_text),
        ("Devise :", currency_label),
    ]
    for row_index, (label, value) in enumerate(criteria_rows):
        criteria.cell(row_index, 0).text = label
        criteria.cell(row_index, 1).text = value
        criteria.cell(row_index, 0).width = Cm(4.1)
        criteria.cell(row_index, 1).width = Cm(4.9)
        for run in criteria.cell(row_index, 0).paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        for run in criteria.cell(row_index, 1).paragraphs[0].runs:
            run.font.size = Pt(8)
    criteria_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(5)
    title.paragraph_format.space_after = Pt(6)
    customer_name_available = (
        not _is_empty_text(customer_name_text)
        and normalize_label(customer_name_text)
        not in {"non disponible", "nom non disponible"}
    )
    title_parts = ["Extrait de compte", telephone_text]
    if customer_name_available:
        title_parts.append(customer_name_text.upper())
    title_parts.append(currency_label)
    title.add_run(" - ".join(title_parts))

    summary_by_currency = view["summary_by_currency"]
    if all_currencies:
        summary_headers = ["Devise", "Ouverture", "Entrees", "Sorties", "Cloture"]
        summary = document.add_table(rows=1, cols=len(summary_headers))
        summary.alignment = WD_TABLE_ALIGNMENT.RIGHT
        summary.autofit = False
        for index, label in enumerate(summary_headers):
            summary.cell(0, index).text = label
            summary.cell(0, index).width = Cm(3.3)
            set_cell_shading(summary.cell(0, index), "E8EEF4")
            summary.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in summary.cell(0, index).paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
        for _, summary_row in summary_by_currency.iterrows():
            cells = summary.add_row().cells
            row_currency = str(summary_row.get("currency_code", ""))
            row_decimals = 0 if row_currency == "CDF" else 2
            balance_kind = "Solde" if bool(summary_row.get("balance_is_real", False)) else "Cumul"
            values = [
                f"{row_currency} ({balance_kind})",
                _pdf_number(summary_row.get("opening_amount"), decimals=row_decimals),
                _pdf_number(summary_row.get("total_entries"), decimals=row_decimals),
                _pdf_number(summary_row.get("total_outputs"), decimals=row_decimals),
                _pdf_number(summary_row.get("closing_amount"), decimals=row_decimals),
            ]
            for index, value in enumerate(values):
                cells[index].text = value
                cells[index].width = Cm(3.3)
                cells[index].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                )
                for run in cells[index].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(8.5)
    else:
        balance_prefix = "Solde" if view["balance_is_real"] else "Cumul"
        summary = document.add_table(rows=2, cols=4)
        summary.alignment = WD_TABLE_ALIGNMENT.RIGHT
        summary.autofit = False
        summary_labels = [
            f"{balance_prefix} initial",
            "Total entrees",
            "Total sorties",
            f"{balance_prefix} final",
        ]
        summary_values = [
            view["opening_amount"],
            view["total_entries"],
            view["total_outputs"],
            view["closing_amount"],
        ]
        for index, label in enumerate(summary_labels):
            summary.cell(0, index).text = label
            summary.cell(1, index).text = _pdf_number(summary_values[index], decimals=decimals)
            summary.cell(0, index).width = Cm(3.8)
            summary.cell(1, index).width = Cm(3.8)
            set_cell_shading(summary.cell(0, index), "E8EEF4")
            summary.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            summary.cell(1, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in summary.cell(0, index).paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            for run in summary.cell(1, index).paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

    analysis_report = analysis_report or {}

    def analysis_frame(key: str) -> pd.DataFrame:
        source = analysis_report.get(key, pd.DataFrame())
        if not isinstance(source, pd.DataFrame) or source.empty:
            return pd.DataFrame()
        scoped = source.copy()
        customer_scope = clean_identifier(pd.Series([customer_id_text])).iloc[0]
        if customer_scope and "customer_id" in scoped.columns:
            scoped = scoped.loc[
                clean_identifier(scoped["customer_id"]).eq(customer_scope)
            ].copy()
        if not all_currencies and "currency_code" in scoped.columns:
            scoped = scoped.loc[
                clean_text(scoped["currency_code"]).str.upper().eq(currency_text)
            ].copy()
        return scoped

    def add_analysis_title(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(24, 63, 91)

    def add_analysis_table(
        source: pd.DataFrame,
        columns: list[str],
        labels: dict[str, str],
    ) -> None:
        present = [column for column in columns if column in source.columns]
        if source.empty or not present:
            return
        analysis_table = document.add_table(rows=1, cols=len(present))
        analysis_table.style = "Table Grid"
        analysis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        analysis_table.autofit = True
        set_repeat_header(analysis_table.rows[0])
        for index, column in enumerate(present):
            cell = analysis_table.rows[0].cells[index]
            cell.text = labels.get(column, column)
            set_cell_shading(cell, "DCE6F1")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(7.2)
        for _, analysis_row in source.iterrows():
            cells = analysis_table.add_row().cells
            row_currency = str(analysis_row.get("currency_code", currency_text)).upper()
            row_decimals = 0 if row_currency == "CDF" else 2
            for index, column in enumerate(present):
                value = analysis_row.get(column)
                if column in {
                    "created_at",
                    "premiere_operation",
                    "derniere_operation",
                    "date_derniere_ecriture",
                    "date_approved",
                    "maturity_date",
                    "date_situation",
                    "date_ecriture_turbo",
                }:
                    parsed = pd.to_datetime(value, errors="coerce")
                    text = f"{parsed:%d/%m/%Y %H:%M}" if pd.notna(parsed) else "-"
                elif column.endswith("_pct"):
                    numeric_value = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
                    text = f"{numeric_value:.1f}%" if pd.notna(numeric_value) else "-"
                elif any(
                    token in column
                    for token in [
                        "montant", "solde", "balance", "ecart", "interet", "principal",
                        "penalite", "dette", "revenu", "intervalle", "inactivite",
                    ]
                ):
                    text = _pdf_number(value, decimals=row_decimals)
                else:
                    text = "-" if _is_empty_text(value) else str(value)
                cells[index].text = text
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[index].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    if any(
                        token in column
                        for token in ["montant", "solde", "balance", "ecart", "interet", "principal", "penalite", "dette", "revenu"]
                    ):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(6.8)

    active_dat = analysis_frame("dat_en_cours_client")
    if not active_dat.empty:
        situation_dates = pd.to_datetime(active_dat.get("date_situation"), errors="coerce").dropna()
        dat_title = "DAT en cours"
        if not situation_dates.empty:
            dat_title += f" - situation au {situation_dates.max():%d/%m/%Y}"
        add_analysis_title(dat_title)
        add_analysis_table(
            active_dat,
            [
                "savings_id",
                "product_name",
                "date_approved",
                "maturity_date",
                "jours_avant_echeance",
                "currency_code",
                "balance",
                "taux_interet_annuel_pct",
                "interet_estime_echeance",
                "capital_plus_interet_estime",
                "situation_dat_client",
            ],
            {
                "savings_id": "DAT",
                "product_name": "Durée",
                "date_approved": "Souscription",
                "maturity_date": "Échéance",
                "jours_avant_echeance": "Jours restants",
                "currency_code": "Devise",
                "balance": "Capital bloqué",
                "taux_interet_annuel_pct": "Taux annuel",
                "interet_estime_echeance": "Intérêt estimé",
                "capital_plus_interet_estime": "Capital + intérêt estimé",
                "situation_dat_client": "Situation",
            },
        )

    repayments = analysis_frame("remboursements_turbo_detail_client")
    if not repayments.empty:
        add_analysis_title("Remboursements observés")
        add_analysis_table(
            repayments,
            [
                "created_at", "event_reference", "currency_code", "montant_paye_observe",
                "principal_rembourse", "interet_observe", "penalite_observee",
                "mode_remboursement_observe",
            ],
            {
                "created_at": "Date",
                "event_reference": "Référence",
                "currency_code": "Devise",
                "montant_paye_observe": "Montant payé",
                "principal_rembourse": "Principal remboursé",
                "interet_observe": "Intérêts",
                "penalite_observee": "Pénalités",
                "mode_remboursement_observe": "Mode observé",
            },
        )

    internal = analysis_frame("mouvements_internes_turbo")
    if not internal.empty:
        add_analysis_title("Mouvements internes epargne / DAT")
        add_analysis_table(
            internal,
            ["created_at", "currency_code", "type_operation", "montant_operation", "descriptions"],
            {
                "created_at": "Date",
                "currency_code": "Devise",
                "type_operation": "Operation interne",
                "montant_operation": "Montant",
                "descriptions": "Description",
            },
        )

    add_analysis_title("Detail des transactions")
    table = document.add_table(rows=1, cols=len(CUSTOMER_STATEMENT_COLUMNS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = {
        "date": "Date",
        "compte": "Compte",
        "receipt_no": "Receipt No",
        "devise": "Devise",
        "description": "Description",
        "entree": "Entrée",
        "sortie": "Sortie",
        "solde": view["balance_label"],
    }
    widths = {
        "date": 2.3,
        "compte": 1.7,
        "receipt_no": 2.8,
        "devise": 1.5,
        "description": 8.9,
        "entree": 2.4,
        "sortie": 2.4,
        "solde": 2.6,
    }
    set_repeat_header(table.rows[0])
    for index, column in enumerate(CUSTOMER_STATEMENT_COLUMNS):
        cell = table.rows[0].cells[index]
        cell.text = labels[column]
        cell.width = Cm(widths[column])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1F2937")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)

    for _, row in transactions.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(CUSTOMER_STATEMENT_COLUMNS):
            value = row.get(column)
            row_currency = str(row.get("devise", currency_text)).upper()
            row_decimals = 0 if row_currency == "CDF" else 2
            if column == "date":
                parsed = pd.to_datetime(value, errors="coerce")
                text = f"{parsed:%d/%m/%Y}" if pd.notna(parsed) else "-"
            elif column in {"entree", "sortie"}:
                numeric_value = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
                text = "" if pd.isna(numeric_value) or float(numeric_value) == 0 else _pdf_number(numeric_value, decimals=row_decimals)
            elif column == "solde":
                text = _pdf_number(value, decimals=row_decimals)
            else:
                text = "-" if _is_empty_text(value) else str(value)
            cells[index].text = text
            cells[index].width = Cm(widths[column])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if column in {"entree", "sortie", "solde"}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)

    if view["balance_is_real"]:
        note = document.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note.paragraph_format.space_after = Pt(0)
        note_text = "Le solde est calcule a partir du solde d'ouverture renseigne et des mouvements du fichier charge."
        note_run = note.add_run(note_text)
        note_run.italic = True
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor(70, 90, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"Extrait genere le {generated_at:%d/%m/%Y %H:%M} - Solution Bisou Bisou Digital")
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor(110, 125, 140)

    document.core_properties.title = f"Extrait de compte Turbo {customer_id_text} - {currency_label}"
    document.core_properties.subject = "Extrait client M-PESA_Turbo"
    document.core_properties.author = "Solution Controle Interne"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_customer_statement_pdf(
    statement: pd.DataFrame,
    *,
    analysis_report: dict[str, pd.DataFrame] | None = None,
    customer_id: object,
    customer_name: object = "",
    telephone: object = "",
    currency: object,
    account_number: object = "",
    entry_account_number: object | None = None,
    output_account_number: object | None = None,
    period_start: object | None = None,
    period_end: object | None = None,
    generated_at: pd.Timestamp | None = None,
) -> bytes:
    """Genere la version PDF de l'extrait officiel, sans conversion externe."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Image,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError("La dependance reportlab est requise pour generer l'extrait PDF.") from exc

    if not isinstance(statement, pd.DataFrame) or statement.empty:
        raise ValueError("Aucune operation filtree n'est disponible pour l'extrait PDF.")
    currency_text = str(currency).strip().upper()
    all_currencies = currency_text in {"ALL", "TOUTES", "TOUS", "CDF + USD"}
    frame = statement.copy()
    frame_currency = clean_text(
        frame.get("currency_code", pd.Series("", index=frame.index))
    ).str.upper()
    if not all_currencies:
        frame = frame.loc[frame_currency.eq(currency_text)].copy()
    if frame.empty:
        raise ValueError(f"Aucune operation {currency_text or 'sans devise'} n'est disponible pour l'extrait PDF.")

    view = build_customer_statement_view(
        frame,
        account_number=account_number,
        entry_account_number=entry_account_number,
        output_account_number=output_account_number,
        allow_multiple_currencies=all_currencies,
    )
    transactions = view["transactions"]
    customer_id_text = str(customer_id).strip() or view["customer_id"] or "Non disponible"
    customer_name_text = str(customer_name).strip() or view["customer_name"] or "Nom non disponible"
    telephone_text = str(telephone).strip() or view["telephone"] or "Non disponible"
    observed_currencies = sorted(
        value for value in transactions["devise"].dropna().astype(str).unique() if value
    )
    currency_label = (
        f"ALL ({', '.join(observed_currencies)})"
        if all_currencies
        else currency_text or view["currency"]
    )
    observed_dates = pd.to_datetime(transactions["date"], errors="coerce").dropna()

    def report_date(value: object | None, fallback: pd.Timestamp | None) -> str:
        parsed = pd.to_datetime(value, errors="coerce") if value is not None else pd.NaT
        if pd.isna(parsed):
            parsed = fallback
        return f"{parsed:%d/%m/%Y}" if parsed is not None and pd.notna(parsed) else "Non disponible"

    start_text = report_date(period_start, observed_dates.min() if not observed_dates.empty else None)
    end_text = report_date(period_end, observed_dates.max() if not observed_dates.empty else None)
    generated_at = generated_at if generated_at is not None else pd.Timestamp.now()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomerStatementTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=15, leading=18, textColor=colors.HexColor("#18293A"), alignment=TA_CENTER,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "CustomerStatementBody", parent=styles["BodyText"], fontName="Helvetica",
        fontSize=6.8, leading=8.2,
    )
    header_style = ParagraphStyle(
        "CustomerStatementHeader", parent=body_style, fontName="Helvetica-Bold",
        textColor=colors.white, alignment=TA_CENTER,
    )
    amount_style = ParagraphStyle(
        "CustomerStatementAmount", parent=body_style, alignment=TA_RIGHT,
    )
    section_style = ParagraphStyle(
        "CustomerStatementSection", parent=styles["Heading3"], fontName="Helvetica-Bold",
        fontSize=10, textColor=colors.HexColor("#183F5B"), spaceBefore=7, spaceAfter=4,
    )
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer, pagesize=landscape(A4), leftMargin=1 * cm, rightMargin=1 * cm,
        topMargin=0.8 * cm, bottomMargin=1.1 * cm,
        title=f"Extrait de compte {customer_id_text} - {currency_label}",
        author="Solution Controle Interne",
    )
    story: list[Any] = []
    if CUSTOMER_STATEMENT_LOGO_PATH.is_file():
        logo: Any = Image(str(CUSTOMER_STATEMENT_LOGO_PATH), width=3.1 * cm, height=3.1 * cm)
    else:
        logo = Paragraph("<b>IMF Microfinance Bisou Bisou</b>", styles["Heading2"])
    criteria = Table(
        [
            [Paragraph("<b>Critères</b>", header_style), ""],
            ["Date du :", start_text], ["Au :", end_text],
            ["Numéro du client :", customer_id_text], ["Téléphone :", telephone_text],
            ["Devise :", currency_label],
        ],
        colWidths=[4.2 * cm, 4.8 * cm],
    )
    criteria.setStyle(TableStyle([
        ("SPAN", (0, 0), (1, 0)), ("BACKGROUND", (0, 0), (1, 0), colors.HexColor("#1F2937")),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 1), (-1, -1), 7.5),
        ("GRID", (0, 1), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4), ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    masthead = Table([[logo, criteria]], colWidths=[17.5 * cm, 9 * cm])
    masthead.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.extend([masthead, Spacer(1, 0.15 * cm)])
    name_available = normalize_label(customer_name_text) not in {"", "non disponible", "nom non disponible"}
    title_parts = ["Extrait de compte", telephone_text]
    if name_available:
        title_parts.append(customer_name_text.upper())
    title_parts.append(currency_label)
    story.append(Paragraph(" - ".join(escape(part) for part in title_parts), title_style))

    summary_rows = [["Devise", "Ouverture", "Entrées", "Sorties", "Clôture"]]
    for _, row in view["summary_by_currency"].iterrows():
        row_currency = str(row.get("currency_code", ""))
        decimals = 0 if row_currency == "CDF" else 2
        balance_kind = "Solde" if bool(row.get("balance_is_real", False)) else "Cumul"
        summary_rows.append([
            f"{row_currency} ({balance_kind})",
            _pdf_number(row.get("opening_amount"), decimals=decimals),
            _pdf_number(row.get("total_entries"), decimals=decimals),
            _pdf_number(row.get("total_outputs"), decimals=decimals),
            _pdf_number(row.get("closing_amount"), decimals=decimals),
        ])
    summary_table = Table(summary_rows, colWidths=[3.5 * cm] * 5, hAlign="RIGHT", repeatRows=1)
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"), ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CBD5E1")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(summary_table)

    analysis_report = analysis_report or {}

    def analysis_frame(key: str) -> pd.DataFrame:
        source = analysis_report.get(key, pd.DataFrame())
        if not isinstance(source, pd.DataFrame) or source.empty:
            return pd.DataFrame()
        scoped = source.copy()
        customer_scope = clean_identifier(pd.Series([customer_id_text])).iloc[0]
        if customer_scope and "customer_id" in scoped.columns:
            scoped = scoped.loc[
                clean_identifier(scoped["customer_id"]).eq(customer_scope)
            ].copy()
        if not all_currencies and "currency_code" in scoped.columns:
            scoped = scoped.loc[
                clean_text(scoped["currency_code"]).str.upper().eq(currency_text)
            ].copy()
        return scoped

    active_dat = analysis_frame("dat_en_cours_client")
    if not active_dat.empty:
        situation_dates = pd.to_datetime(active_dat.get("date_situation"), errors="coerce").dropna()
        dat_title = "DAT en cours"
        if not situation_dates.empty:
            dat_title += f" - situation au {situation_dates.max():%d/%m/%Y}"
        story.append(Paragraph(dat_title, section_style))
        dat_headers = [
            "DAT",
            "Durée",
            "Souscription",
            "Échéance",
            "Jours",
            "Devise",
            "Capital bloqué",
            "Taux",
            "Intérêt estimé",
            "Capital + intérêt estimé",
            "Situation",
        ]
        dat_rows: list[list[Any]] = [
            [Paragraph(label, header_style) for label in dat_headers]
        ]
        for _, dat_row in active_dat.iterrows():
            row_currency = str(dat_row.get("currency_code", currency_text)).upper()
            row_decimals = 0 if row_currency == "CDF" else 2
            approved = pd.to_datetime(dat_row.get("date_approved"), errors="coerce")
            maturity = pd.to_datetime(dat_row.get("maturity_date"), errors="coerce")
            values = [
                dat_row.get("savings_id", "-"),
                dat_row.get("product_name", "-"),
                f"{approved:%d/%m/%Y}" if pd.notna(approved) else "-",
                f"{maturity:%d/%m/%Y}" if pd.notna(maturity) else "-",
                _pdf_number(dat_row.get("jours_avant_echeance"), decimals=0),
                row_currency,
                _pdf_number(dat_row.get("balance"), decimals=row_decimals),
                (
                    f"{float(dat_row.get('taux_interet_annuel_pct')):.1f}%"
                    if pd.notna(dat_row.get("taux_interet_annuel_pct"))
                    else "-"
                ),
                _pdf_number(dat_row.get("interet_estime_echeance"), decimals=row_decimals),
                _pdf_number(dat_row.get("capital_plus_interet_estime"), decimals=row_decimals),
                dat_row.get("situation_dat_client", "-"),
            ]
            dat_rows.append(
                [
                    Paragraph(escape("-" if _is_empty_text(value) else str(value)), body_style)
                    for value in values
                ]
            )
        dat_table = Table(
            dat_rows,
            colWidths=[
                2.2 * cm,
                1.8 * cm,
                2.0 * cm,
                2.0 * cm,
                1.3 * cm,
                1.2 * cm,
                2.3 * cm,
                1.3 * cm,
                2.3 * cm,
                2.5 * cm,
                3.0 * cm,
            ],
            repeatRows=1,
        )
        dat_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C1CC")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (4, 1), (9, -1), "RIGHT"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2.2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2.2),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        story.append(dat_table)

    repayments = analysis_frame("remboursements_turbo_detail_client")
    if not repayments.empty:
        story.append(Paragraph("Remboursements observés", section_style))
        repayment_headers = [
            "Date",
            "Référence",
            "Devise",
            "Montant payé",
            "Principal remboursé",
            "Intérêts",
            "Pénalités",
            "Mode observé",
        ]
        repayment_rows: list[list[Any]] = [
            [Paragraph(label, header_style) for label in repayment_headers]
        ]
        for _, repayment_row in repayments.iterrows():
            row_currency = str(repayment_row.get("currency_code", currency_text)).upper()
            row_decimals = 0 if row_currency == "CDF" else 2
            repayment_date = pd.to_datetime(repayment_row.get("created_at"), errors="coerce")
            values = [
                f"{repayment_date:%d/%m/%Y %H:%M}" if pd.notna(repayment_date) else "-",
                repayment_row.get("event_reference", "-"),
                row_currency,
                _pdf_number(repayment_row.get("montant_paye_observe"), decimals=row_decimals),
                _pdf_number(repayment_row.get("principal_rembourse"), decimals=row_decimals),
                _pdf_number(repayment_row.get("interet_observe"), decimals=row_decimals),
                _pdf_number(repayment_row.get("penalite_observee"), decimals=row_decimals),
                repayment_row.get("mode_remboursement_observe", "-"),
            ]
            repayment_rows.append(
                [
                    Paragraph(escape("-" if _is_empty_text(value) else str(value)), body_style)
                    for value in values
                ]
            )
        repayment_table = Table(
            repayment_rows,
            colWidths=[2.2 * cm, 3.5 * cm, 1.4 * cm, 3.0 * cm, 3.0 * cm, 2.7 * cm, 2.7 * cm, 5.0 * cm],
            repeatRows=1,
        )
        repayment_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C1CC")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (3, 1), (6, -1), "RIGHT"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2.2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2.2),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        story.append(repayment_table)

    story.append(Paragraph("Détail des transactions", section_style))

    labels = ["Date", "Compte", "Receipt No", "Devise", "Description", "Entrée", "Sortie", view["balance_label"]]
    detail_rows: list[list[Any]] = [[Paragraph(label, header_style) for label in labels]]
    for _, row in transactions.iterrows():
        row_currency = str(row.get("devise", currency_text)).upper()
        decimals = 0 if row_currency == "CDF" else 2
        date_value = pd.to_datetime(row.get("date"), errors="coerce")
        values = [
            f"{date_value:%d/%m/%Y}" if pd.notna(date_value) else "-",
            row.get("compte", "-"), row.get("receipt_no", "-"), row_currency,
            row.get("description", "-"), row.get("entree"), row.get("sortie"), row.get("solde"),
        ]
        formatted: list[Any] = []
        for index, value in enumerate(values):
            if index in {5, 6}:
                numeric = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
                text_value = "" if pd.isna(numeric) or float(numeric) == 0 else _pdf_number(numeric, decimals=decimals)
                formatted.append(Paragraph(text_value, amount_style))
            elif index == 7:
                formatted.append(Paragraph(_pdf_number(value, decimals=decimals), amount_style))
            else:
                formatted.append(Paragraph(escape("-" if _is_empty_text(value) else str(value)), body_style))
        detail_rows.append(formatted)
    detail = Table(
        detail_rows,
        colWidths=[2.1 * cm, 1.5 * cm, 2.7 * cm, 1.4 * cm, 9.1 * cm, 2.3 * cm, 2.3 * cm, 2.6 * cm],
        repeatRows=1,
    )
    detail.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C1CC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.5), ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
        ("TOPPADDING", (0, 0), (-1, -1), 2.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
    ]))
    story.append(detail)

    def draw_footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#6E7D8C"))
        footer = f"Extrait généré le {generated_at:%d/%m/%Y %H:%M} - Solution Bisou Bisou Digital - Page {doc.page}"
        canvas.drawCentredString(landscape(A4)[0] / 2, 0.45 * cm, footer)
        canvas.restoreState()

    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return buffer.getvalue()


def create_g2_dat_word(
    report: dict[str, Any],
    *,
    period_text: str,
    direction_label: str,
    generated_at: pd.Timestamp | None = None,
) -> bytes:
    """Genere la version Word courte et editable destinee a la Direction generale."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT, WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("La dependance python-docx est requise pour generer le rapport Word.") from exc

    source_label = str(report.get("analysis_source_label", "G2") or "G2")
    report_scope = "Turbo/DAT" if source_label.casefold() == "turbo" else "G2/DAT"
    daily_pivot = report.get("rapport_journalier_pivot", pd.DataFrame())
    daily_synthese = report.get("rapport_journalier_synthese", pd.DataFrame())
    transaction_detail = report.get("rapport_journalier_detail", pd.DataFrame())
    eligible_transaction_detail = transaction_detail
    if isinstance(transaction_detail, pd.DataFrame):
        eligible_transaction_detail = transaction_detail.copy()
        if (
            not eligible_transaction_detail.empty
            and "incluse_synthese" in eligible_transaction_detail.columns
        ):
            eligible_word_detail = (
                eligible_transaction_detail["incluse_synthese"]
                .astype("boolean")
                .fillna(False)
                .astype(bool)
            )
            eligible_transaction_detail = eligible_transaction_detail.loc[eligible_word_detail].copy()
    if not isinstance(daily_pivot, pd.DataFrame) or daily_pivot.empty:
        daily_pivot = build_entry_pivot(eligible_transaction_detail)
    generated_at = generated_at if generated_at is not None else pd.Timestamp.now()
    word_report = dict(report)
    word_report["rapport_journalier_pivot"] = daily_pivot
    context = _g2_executive_context(word_report)
    transaction_detail = eligible_transaction_detail

    classified = daily_synthese.copy()
    if not classified.empty and "details_rapport" in classified.columns:
        classified = classified.loc[
            ~classified["details_rapport"].astype("string").str.startswith("Total ", na=False)
        ]

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor(18, 56, 95)
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(11)

    title = document.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(2)
    title.add_run(f"Rapport M-PESA - {report_scope}")
    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(7)
    meta_run = meta.add_run(f"{period_text} | Sens : {direction_label} | {generated_at:%d/%m/%Y}")
    meta_run.font.size = Pt(8.5)
    meta_run.font.color.rgb = RGBColor(93, 113, 135)

    document.add_heading("Synthese executive", level=1)

    def add_summary_bullet(label: str, text: str) -> None:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        bold_run = paragraph.add_run(f"{label}. ")
        bold_run.bold = True
        paragraph.add_run(text)

    add_summary_bullet("Activite", context["active_text"])
    if context["status_text"]:
        add_summary_bullet("Perimetre des statuts", context["status_text"])
    if context["time_text"]:
        add_summary_bullet("Fréquence temporelle", context["time_text"])
    add_summary_bullet("Flux financiers", context["flow_text"])
    if context["control_text"]:
        add_summary_bullet("Controle", context["control_text"])

    def set_cell_shading(cell: Any, fill: str) -> None:
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_properties.append(shading)

    def add_table(
        frame: pd.DataFrame,
        columns: list[str],
        labels: dict[str, str],
        *,
        font_size: float = 8,
        amount_decimals: int = 2,
        column_widths_cm: dict[str, float] | None = None,
    ) -> Any | None:
        present = [column for column in columns if column in frame.columns]
        if frame.empty or not present:
            document.add_paragraph("Aucune donnee disponible.")
            return None
        table = document.add_table(rows=1, cols=len(present))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = column_widths_cm is None
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)
        for index, column in enumerate(present):
            cell = table.rows[0].cells[index]
            cell.text = labels.get(column, column)
            if column_widths_cm and column in column_widths_cm:
                cell.width = Cm(column_widths_cm[column])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "1F4E78")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(font_size)
        for _, row in frame[present].iterrows():
            cells = table.add_row().cells
            for index, column in enumerate(present):
                value = row.get(column)
                if any(token in column for token in ["montant", "solde", "volume"]):
                    text = _pdf_number(value, decimals=amount_decimals)
                elif str(column).endswith("_pct"):
                    text = f"{float(value):.1f}%" if pd.notna(value) else "-"
                elif column in {"retenu_m1", "retenu_90j"}:
                    text = "-" if pd.isna(value) else "Oui" if bool(value) else "Non"
                elif column in {"date", "compte_cree"}:
                    date_value = pd.to_datetime(value, errors="coerce")
                    text = f"{date_value:%d/%m/%Y %H:%M:%S}" if pd.notna(date_value) else "-"
                elif column == "premier_retour":
                    date_value = pd.to_datetime(value, errors="coerce")
                    text = f"{date_value:%d/%m/%Y}" if pd.notna(date_value) else "-"
                else:
                    text = "-" if pd.isna(value) else str(value)
                cells[index].text = text
                if column_widths_cm and column in column_widths_cm:
                    cells[index].width = Cm(column_widths_cm[column])
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[index].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
        document.add_paragraph().paragraph_format.space_after = Pt(0)
        return table

    document.add_heading(f"Synthese des flux {source_label} par devise", level=1)
    add_table(
        daily_pivot,
        ["currency_code", "nombre_entrees", "montant_total_entrees", "nombre_sorties", "montant_total_sorties", "solde_net_flux"],
        {
            "currency_code": "Devise", "nombre_entrees": "Nb entrees", "montant_total_entrees": "Montant entrees",
            "nombre_sorties": "Nb sorties", "montant_total_sorties": "Montant sorties", "solde_net_flux": "Solde net",
        },
    )
    document.add_heading("Principales operations", level=2)
    add_table(
        classified,
        ["currency_code", "sens_flux", "details_rapport", "nombre", "montant"],
        {"currency_code": "Devise", "sens_flux": "Sens", "details_rapport": "Operation", "nombre": "Nombre", "montant": "Montant"},
    )

    if context["has_retention"]:
        document.add_heading("Fidelisation", level=1)
        paragraph = document.add_paragraph(context["retention_text"])
        paragraph.paragraph_format.space_after = Pt(4)
        add_table(
            context["latest_retention"],
            ["Devise", "Retention M+1", "Retention 90 jours"],
            {},
        )
    else:
        document.add_heading("Point de vigilance", level=1)
        warning = document.add_paragraph(context["attention_text"])
        warning.paragraph_format.space_after = Pt(5)
        warning.runs[0].font.color.rgb = RGBColor(156, 103, 10)

    if isinstance(transaction_detail, pd.DataFrame) and not transaction_detail.empty:
        detail = transaction_detail.copy()
        detail["currency_code"] = clean_text(
            detail.get("currency_code", pd.Series("", index=detail.index))
        ).str.upper().replace("", "SANS DEVISE")
        detail["date"] = pd.to_datetime(
            detail.get("date", pd.Series(pd.NaT, index=detail.index)), errors="coerce"
        )
        detail = detail.sort_values(
            ["currency_code", "date"], ascending=[True, False], na_position="last"
        ).reset_index(drop=True)
        detail_section = document.add_section(WD_SECTION.NEW_PAGE)
        detail_section.orientation = WD_ORIENT.LANDSCAPE
        detail_section.page_width, detail_section.page_height = (
            detail_section.page_height,
            detail_section.page_width,
        )
        detail_section.top_margin = Cm(1.0)
        detail_section.bottom_margin = Cm(1.0)
        detail_section.left_margin = Cm(0.5)
        detail_section.right_margin = Cm(0.5)
        document.add_heading("Transactions", level=1)
        # detail_caption.paragraph_format.space_after = Pt(4)
        add_table(
            detail,
            G2_CLASSIFIED_TRANSACTION_COLUMNS,
            {column: column for column in G2_CLASSIFIED_TRANSACTION_COLUMNS},
            font_size=5.5,
            amount_decimals=2,
            column_widths_cm={
                "date": 2.7,
                "receipt_no": 2.1,
                "currency_code": 1.4,
                "details_rapport": 2.3,
                "opposite_party": 6.2,
                "duree": 1.6,
                "compte_cree": 2.7,
                "montant": 1.7,
                "montant_entree": 1.7,
                "montant_sortie": 1.7,
                "balance_numeric": 1.9,
            },
        )

    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    # source_run.italic = True
    # source_run.font.size = Pt(8)
    # source_run.font.color.rgb = RGBColor(96, 117, 138)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Solution Bisou Bisou Digital ")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(110, 125, 140)

    document.core_properties.title = f"Rapport M-PESA - {report_scope}"
    document.core_properties.subject = "Synthese destinee a la Direction generale"
    document.core_properties.author = "Solution Controle Interne"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_excel_export(report: dict[str, Any]) -> bytes:
    sheet_contract = [
        ("synthese", "Synthese"),
        ("extrait", "Extrait_Turbo"),
        ("parcours_turbo", "Parcours_Turbo"),
        ("dat_en_cours_client", "DAT_En_Cours"),
        ("remboursements_turbo_detail_client", "Remboursements_Turbo"),
        ("credit_turbo_detail_client", "Credit_Client_Turbo"),
        ("positions_turbo", "Positions_Turbo"),
        ("comportement_turbo", "Comportement_Turbo"),
        ("mouvements_internes_turbo", "Mouvements_Internes"),
        ("controles_client_turbo", "Controles_Client_Turbo"),
        ("interets_dat_echus", "Interets_DAT_Echus"),
        ("dat_final", "DAT_Final"),
        ("forts_dat", "Forts_DAT"),
        ("portefeuille_dat", "Portefeuille_DAT"),
        ("mouvements_dat", "Mouvements_DAT"),
        ("mouvements_epargne", "Mouvements_Epargne"),
        ("credits", "Credits"),
        ("g2_dat", "G2_DAT"),
        ("turbo_dat", "Turbo_DAT"),
        ("rapport_g2_pivot", "Rapport_G2_Pivot"),
        ("rapport_g2_comptages", "Rapport_G2_Comptages"),
        ("rapport_g2_vertical", "Rapport_G2_Vertical"),
        ("rapport_g2_synthese", "Rapport_G2_Synthese"),
        ("rapport_g2_detail", "Rapport_G2_Detail"),
        ("rapport_journalier_pivot", "Rapport_Journalier_Pivot"),
        ("rapport_journalier_comptages", "Rapport_Journalier_Comptages"),
        ("rapport_journalier_vertical", "Rapport_Journalier_Vertical"),
        ("rapport_journalier_synthese", "Rapport_Journalier_Synthese"),
        ("statuts_g2", "Statuts_G2"),
        ("statuts_turbo", "Statuts_Turbo"),
        ("rapport_journalier_detail", "Rapport_Journalier_Detail"),
        ("rapport_journalier_anomalies", "Anomalies_G2"),
        ("rapport_turbo_anomalies", "Anomalies_Turbo"),
        ("transactions_par_jour", "Transactions_Jour"),
        ("transactions_par_jour_semaine", "Transactions_Jour_Semaine"),
        ("transactions_par_heure", "Transactions_Heure"),
        ("transactions_jour_heure", "Transactions_Jour_Heure"),
        ("retention_mensuelle", "Retention_Mensuelle"),
        ("retention_operations", "Retention_Operations"),
        ("retention_detail", "Retention_Detail"),
        ("retention_definitions", "Retention_Definitions"),
        ("perfect_clients", "Clients_Perfect"),
        ("perfect_operations", "Operations_Turbo_G2"),
        ("clients_perfect_dans_mpesa", "Clients_Perfect_G2"),
        ("clients_perfect_dans_turbo", "Clients_Perfect_Turbo"),
        ("clients_perfect_dans_turbo_et_mpesa", "Clients_Perfect_Turbo_G2"),
        ("clients_3_systemes", "Clients_Perfect_3_Systemes"),
        ("credit_synthese", "Pilotage_Credit_Turbo"),
        ("credit_detail", "Credits_Risque_Turbo"),
        ("flux_synthese", "Flux_Synthese_Turbo"),
        ("flux_evolution", "Flux_Evolution_Turbo"),
        ("remboursements_synthese", "Remboursements_Synthese"),
        ("remboursements_detail", "Remboursements_Pilotage"),
        ("nouveaux_credits_synthese", "Nouveaux_Credits_Synthese"),
        ("nouveaux_credits_detail", "Nouveaux_Credits_Turbo"),
        ("par_tranches_montant", "PAR_Tranches_Turbo"),
        ("concentration_credit_synthese", "Concentration_Credit"),
        ("activite_epargne_clients", "Activite_Epargne_Clients"),
        ("depots_frequents_hebdo", "Depots_Frequents_Hebdo"),
        ("tranches_depots", "Tranches_Depots_Turbo"),
        ("concentration_transactions_synthese", "Concentration_Transactions"),
        ("mouvements_comptes_inactifs", "Mouvements_Comptes_Inactifs"),
        ("dat_sans_credit_actif", "DAT_Sans_Credit_Actif"),
        ("credits_epargne_disponible", "Credit_Epargne_Disponible"),
        ("qualite_clients_detail", "Qualite_Clients_Turbo"),
        ("definitions", "Definitions_Pilotage"),
        ("sources", "Sources_Pilotage"),
        ("loan_savings_summary", "Credit_Epargne_Synthese"),
        ("loan_savings_clients", "Credit_Epargne_Clients"),
        ("loan_savings_detail", "Credit_Epargne_Detail"),
        ("loan_savings_controls", "Controle_Credit_Epargne"),
        ("liquidite_synthese", "Liquidite_Turbo"),
        ("liquidite_journaliere", "Liquidite_Jour_Turbo"),
        ("activite_clients", "Activite_Turbo"),
        ("conversion_clients", "Conversion_DAT_Turbo"),
        ("concentration_clients", "Concentration_Turbo"),
        ("qualite_synthese", "Qualite_Turbo"),
        ("alertes_transactions", "Alertes_Turbo"),
        ("dat_echeances_detail", "Echeances_DAT_Turbo"),
        ("perfect_adoption_detail", "Adoption_Turbo_G2"),
        ("accounting_summary", "Compta_Synthese_Turbo"),
        ("accounting_client_balances", "Balance_Clients_Turbo"),
        ("accounting_client_positions", "Positions_Clients_Turbo"),
        ("accounting_account_balance", "Balance_Comptes_Turbo"),
        ("accounting_operation_journal", "Journal_Operations_Turbo"),
        ("accounting_entry_journal", "Journal_Ecritures_Turbo"),
        ("accounting_operation_controls", "Controles_Operations_Turbo"),
        ("accounting_balance_controls", "Controles_Soldes_Turbo"),
        ("accounting_cash_flow", "Flux_MPESA_Turbo"),
        ("accounting_financial_products", "Produits_Financiers_Turbo"),
        ("accounting_portfolio_positions", "Positions_Portefeuille_Turbo"),
        ("accounting_g2_controls", "Controle_G2_Turbo"),
        ("diagnostics", "Diagnostics"),
    ]
    sheets = {
        sheet_name: report[report_key]
        for report_key, sheet_name in sheet_contract
        if report_key in report
    }
    if not sheets:
        sheets = {"Information": pd.DataFrame({"message": ["Aucune donnee a exporter."]})}
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet_name, frame in sheets.items():
            safe_frame = frame if isinstance(frame, pd.DataFrame) else pd.DataFrame()
            safe_frame.to_excel(writer, sheet_name=sheet_name[:31], index=False)
            worksheet = writer.sheets[sheet_name[:31]]
            worksheet.freeze_panes = "A2"
            worksheet.auto_filter.ref = worksheet.dimensions
            for cell in worksheet[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(fill_type="solid", fgColor="1F4E78")
            for index, column in enumerate(safe_frame.columns, start=1):
                values = [str(column)] + ["" if pd.isna(value) else str(value) for value in safe_frame[column].head(300)]
                width = min(max(len(value) for value in values) + 2, 45)
                worksheet.column_dimensions[worksheet.cell(row=1, column=index).column_letter].width = width
                name = str(column).lower()
                if "date" in name or "created_at" in name:
                    number_format = "dd/mm/yyyy hh:mm:ss"
                elif any(token in name for token in ["solde", "montant", "entree", "sortie", "mouvement", "balance", "variation", "amount", "loan"]):
                    number_format = '#,##0.00'
                else:
                    number_format = None
                if number_format:
                    for row in range(2, len(safe_frame) + 2):
                        worksheet.cell(row=row, column=index).number_format = number_format
    buffer.seek(0)
    return buffer.getvalue()


def load_excel_file(file_like: Any, sheet_name: str | int | None = 0) -> pd.DataFrame:
    if file_like is None:
        return pd.DataFrame()
    return pd.read_excel(file_like, sheet_name=sheet_name)

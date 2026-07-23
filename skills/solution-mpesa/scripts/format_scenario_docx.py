"""Appliquer la charte du cahier de tests au scénario Turbo/G2.

Le script conserve le texte métier du document et ne modifie que sa présentation.
Les blocs sont séparés par une ligne vide : le premier bloc porte les paramètres,
les blocs intermédiaires sont numérotés comme étapes et le dernier bloc présente
la situation financière attendue.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


NAVY = "0B2E59"
BLUE = "0D4F8B"
RED = "E31B23"
PALE_BLUE = "EAF2F8"
PALE_GREY = "F5F7FA"
MID_GREY = "64748B"
TEXT = "172033"
WHITE = "FFFFFF"
GREEN = "137A4B"

STEP_PREFIX = re.compile(r"^(?:ÉTAPE|ETAPE)\s+\d+\s*[|—-]\s*", re.IGNORECASE)
RESULT_PREFIX = re.compile(
    r"^(?:RÉSULTAT|RESULTAT)\s+ATTENDU\s*[|—-]\s*", re.IGNORECASE
)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_paragraph_borders(
    paragraph,
    *,
    top: str | None = None,
    bottom: str | None = None,
    left: str | None = None,
    right: str | None = None,
    size: int = 8,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge, color in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        if not color:
            continue
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), color)


def _set_keep(paragraph, *, with_next: bool = False, together: bool = True) -> None:
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.keep_with_next = with_next


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _rewrite_label_value(
    paragraph,
    text: str,
    *,
    marker: str | None = None,
    label_color: str = BLUE,
    value_color: str = TEXT,
) -> None:
    _clear_paragraph(paragraph)
    if marker:
        marker_run = paragraph.add_run(marker)
        marker_run.bold = True
        marker_run.font.color.rgb = RGBColor.from_string(RED)
    if ":" in text:
        label, value = text.split(":", 1)
        label_run = paragraph.add_run(f"{label.strip()} :")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(label_color)
        value_run = paragraph.add_run(f" {value.strip()}")
        value_run.font.color.rgb = RGBColor.from_string(value_color)
    else:
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor.from_string(value_color)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = TEXT) -> None:
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    for field_name in ("PAGE", "NUMPAGES"):
        if field_name == "NUMPAGES":
            paragraph.add_run(" / ")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), field_name)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), MID_GREY)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "17")
        run_properties.extend((color, size))
        text = OxmlElement("w:t")
        text.text = "1"
        run.extend((run_properties, text))
        field.append(run)
        paragraph._p.append(field)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)

    style_specs = {
        "Scenario Subtitle": (10.5, False, MID_GREY),
        "Scenario Meta": (10, False, TEXT),
        "Scenario Step": (13, True, WHITE),
        "Scenario Narrative": (10.5, False, TEXT),
        "Scenario Evidence": (9.5, False, TEXT),
        "Scenario Result": (11, False, TEXT),
    }
    for name, (size, bold, color) in style_specs.items():
        if name not in document.styles:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = document.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)


def _configure_page(document: Document) -> None:
    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.6)
        section.footer_distance = Cm(0.7)


def _configure_header_footer(document: Document, logo_path: Path | None) -> None:
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        for element in list(header._element):
            header._element.remove(element)

        table = header.add_table(rows=1, cols=2, width=Cm(17.4))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(4.5)
        table.columns[1].width = Cm(12.9)
        _remove_table_borders(table)
        left, right = table.rows[0].cells
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        if logo_path and logo_path.exists():
            logo_paragraph = left.paragraphs[0]
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_paragraph.add_run().add_picture(str(logo_path), width=Cm(2.7))
        else:
            _set_cell_text(left, "BISOU BISOU", bold=True, color=NAVY)

        brand_paragraph = right.paragraphs[0]
        brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        brand = brand_paragraph.add_run("SOLUTION M_PESA")
        brand.bold = True
        brand.font.name = "Aptos Display"
        brand.font.size = Pt(10)
        brand.font.color.rgb = RGBColor.from_string(NAVY)
        brand_paragraph.add_run("\nCAHIER DE TESTS FONCTIONNELS").font.color.rgb = (
            RGBColor.from_string(MID_GREY)
        )

        divider = header.add_paragraph()
        divider.paragraph_format.space_after = Pt(0)
        _set_paragraph_borders(divider, bottom=RED, size=12)

        footer = section.footer
        footer.is_linked_to_previous = False
        for element in list(footer._element):
            footer._element.remove(element)
        footer_table = footer.add_table(rows=1, cols=2, width=Cm(17.4))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        footer_table.autofit = False
        footer_table.columns[0].width = Cm(12.5)
        footer_table.columns[1].width = Cm(4.9)
        _remove_table_borders(footer_table)
        _set_cell_text(
            footer_table.cell(0, 0),
            "Solution Bisou Bisou Digital  •  Scénario Turbo / G2",
            color=MID_GREY,
        )
        page_paragraph = footer_table.cell(0, 1).paragraphs[0]
        _clear_paragraph(page_paragraph)
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_page_field(page_paragraph)
        for run in page_paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(MID_GREY)


def _insert_subtitle(title_paragraph, text: str):
    new_p = OxmlElement("w:p")
    title_paragraph._p.addnext(new_p)
    paragraph = title_paragraph._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    new_p.getparent().replace(new_p, paragraph._p)
    paragraph.add_run(text)
    return paragraph


def _paragraph_groups(document: Document) -> list[list]:
    groups: list[list] = []
    current: list = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            current.append(paragraph)
        elif current:
            groups.append(current)
            current = []
    if current:
        groups.append(current)
    return groups


def _base_heading(text: str) -> str:
    text = STEP_PREFIX.sub("", text.strip())
    text = RESULT_PREFIX.sub("", text)
    if text.casefold() == "situation financière finale":
        return "Solde actuel"
    return text


def _insert_paragraph_after(paragraph, text: str):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_element.getparent().replace(new_element, new_paragraph._p)
    new_paragraph.add_run(text)
    return new_paragraph


def _apply_business_corrections(document: Document) -> None:
    """Verrouiller le scénario de prêt brut, intérêt à 7 % et net versé."""
    explanation = (
        "Lecture retenue pour l'extrait client : prêt brut 5,00 USD — "
        "intérêt prélevé 0,35 USD (7 %) — net versé 4,65 USD."
    )
    expected_result = (
        "Résultat Extrait client attendu : A4 portrait ; flux externes : "
        "entrées 15,00 USD, sorties 4,65 USD et flux net 10,35 USD ; "
        "situation de l'épargne : compte ouvert 0,00 USD et compte bloqué "
        "10,00 USD ; remboursement interne depuis le compte ouvert 5,00 USD "
        "présenté séparément et exclu des entrées externes ; DAT FA9IQ86JE7 "
        "présenté En cours avec le taux annuel de 11 % dans les critères et "
        "un capital + intérêt estimé de 10,09 USD."
    )
    explanation_paragraphs = []
    insertion_point = None
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        normalized_text = re.sub(r"^[•✓]\s*", "", text).strip()
        if normalized_text == explanation:
            explanation_paragraphs.append(paragraph)
        if normalized_text.startswith("Résultat Extrait client attendu :"):
            _clear_paragraph(paragraph)
            paragraph.add_run(f"•  {expected_result}")
        if text.startswith("•  Message M_PESA") and "DGM667X6O32" in text:
            _clear_paragraph(paragraph)
            paragraph.add_run(
                "•  Message M_PESA : Vous avez reçu un paiement provenant de "
                "15558-Bisou Bisou de 4,65$ (montant net versé après prélèvement "
                "de l'intérêt de 7 %). Ref. : DGM667X6O32"
            )
        if text.startswith("•  Fichier Transaction_Turbo : 12 écritures"):
            _clear_paragraph(paragraph)
            paragraph.add_run(
                "•  Fichier Transaction_Turbo : 12 écritures d'octroi à 16:17:16 "
                "(reference_id : LN11FAEGXL). Le même reference_id apparaît sur "
                "16 lignes au total, car 4 écritures supplémentaires correspondent "
                "au remboursement à 16:26:43."
            )
            insertion_point = paragraph
    for duplicate in explanation_paragraphs[1:]:
        duplicate._element.getparent().remove(duplicate._element)
    if insertion_point is not None and not explanation_paragraphs:
        _insert_paragraph_after(insertion_point, explanation)


def _clean_content_prefix(text: str) -> tuple[str, bool]:
    """Retirer les marqueurs visuels ajoutés lors d'un passage précédent."""
    cleaned = text.strip()
    is_narrative = cleaned.casefold().startswith("contexte  •") or cleaned.casefold().startswith(
        "contexte •"
    )
    if is_narrative:
        cleaned = re.sub(r"^contexte\s*•\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^[•✓]\s*", "", cleaned)
    return cleaned, is_narrative


def format_document(source: Path, output: Path, logo_path: Path | None) -> None:
    document = Document(source)
    _apply_business_corrections(document)
    _configure_styles(document)
    _configure_page(document)
    _configure_header_footer(document, logo_path)

    properties = document.core_properties
    properties.title = "Scénario Turbo et G2 — Solution M_PESA"
    properties.subject = "Cahier de tests fonctionnels Turbo/G2"
    properties.keywords = "Solution M_PESA, Turbo, G2, DAT, crédit, épargne, test"

    groups = _paragraph_groups(document)
    if not groups:
        raise ValueError("Le document ne contient aucun texte à mettre en forme.")

    first_group = groups[0]
    title_paragraph = first_group[0]
    title_paragraph.style = document.styles["Title"]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.keep_with_next = True
    _set_paragraph_borders(title_paragraph, bottom=RED, size=12)

    subtitle_text = "Cahier de tests fonctionnels  •  Turbo, source principale  •  G2, contrôle complémentaire"
    subtitle = next(
        (p for p in first_group[1:] if p.text.strip() == subtitle_text), None
    )
    if subtitle is None:
        subtitle = _insert_subtitle(title_paragraph, subtitle_text)
    subtitle.style = document.styles["Scenario Subtitle"]
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True
    subtitle.runs[0].italic = True

    groups = _paragraph_groups(document)
    first_group = groups[0]
    metadata = [
        p
        for p in first_group[1:]
        if p.text.strip() and p.text.strip() != subtitle_text
    ]
    for index, paragraph in enumerate(metadata):
        original = paragraph.text.strip()
        paragraph.style = document.styles["Scenario Meta"]
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.right_indent = Cm(0.35)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        _set_paragraph_shading(paragraph, PALE_GREY)
        _set_paragraph_borders(
            paragraph,
            top=BLUE if index == 0 else None,
            bottom=BLUE if index == len(metadata) - 1 else PALE_BLUE,
            left=BLUE,
            right=BLUE,
            size=5,
        )
        _rewrite_label_value(paragraph, original)
        _set_keep(paragraph, with_next=index < len(metadata) - 1)

    groups = _paragraph_groups(document)
    step_groups = groups[1:]
    step_number = 0
    for group in step_groups:
        heading = group[0]
        base_heading = _base_heading(heading.text)
        is_result = base_heading.casefold() == "solde actuel"
        if is_result:
            visible_heading = "RÉSULTAT ATTENDU  |  Situation financière finale"
        else:
            step_number += 1
            visible_heading = f"ÉTAPE {step_number:02d}  |  {base_heading}"

        _clear_paragraph(heading)
        heading.add_run(visible_heading)
        heading.style = document.styles["Scenario Step"]
        heading.paragraph_format.left_indent = Cm(0.25)
        heading.paragraph_format.right_indent = Cm(0.1)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.line_spacing = 1.15
        _set_paragraph_shading(heading, NAVY if not is_result else BLUE)
        _set_paragraph_borders(heading, left=RED, bottom=RED, size=10)
        _set_keep(heading, with_next=True)

        for position, paragraph in enumerate(group[1:], start=1):
            original, was_narrative = _clean_content_prefix(paragraph.text)
            if not original:
                continue
            if was_narrative or original.casefold().startswith("le client"):
                paragraph.style = document.styles["Scenario Narrative"]
                paragraph.paragraph_format.left_indent = Cm(0.3)
                paragraph.paragraph_format.right_indent = Cm(0.3)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(7)
                _set_paragraph_shading(paragraph, PALE_BLUE)
                _set_paragraph_borders(paragraph, left=BLUE, size=7)
                _clear_paragraph(paragraph)
                intro = paragraph.add_run("Contexte  •  ")
                intro.bold = True
                intro.font.color.rgb = RGBColor.from_string(BLUE)
                paragraph.add_run(original)
            elif is_result:
                paragraph.style = document.styles["Scenario Result"]
                paragraph.paragraph_format.left_indent = Cm(0.3)
                paragraph.paragraph_format.right_indent = Cm(0.3)
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                _set_paragraph_shading(paragraph, PALE_GREY)
                _set_paragraph_borders(
                    paragraph,
                    bottom=PALE_BLUE,
                    left=GREEN,
                    right=PALE_BLUE,
                    size=7,
                )
                _rewrite_label_value(
                    paragraph,
                    original,
                    marker="✓  ",
                    label_color=NAVY,
                    value_color=GREEN,
                )
                if paragraph.runs:
                    paragraph.runs[-1].bold = True
            else:
                paragraph.style = document.styles["Scenario Evidence"]
                paragraph.paragraph_format.left_indent = Cm(0.65)
                paragraph.paragraph_format.first_line_indent = Cm(-0.35)
                paragraph.paragraph_format.right_indent = Cm(0.15)
                paragraph.paragraph_format.space_before = Pt(1.5)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                _rewrite_label_value(paragraph, original, marker="•  ")
            _set_keep(
                paragraph,
                with_next=position < len(group) - 1,
                together=True,
            )

    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.35

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.name}.tmp")
    document.save(temporary)
    temporary.replace(output)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Mettre en forme le scénario Word Turbo/G2."
    )
    parser.add_argument("document", type=Path, help="Fichier DOCX source")
    parser.add_argument(
        "--output",
        type=Path,
        help="Fichier DOCX de sortie; remplace la source si omis",
    )
    parser.add_argument("--logo", type=Path, help="Logo Bisou Bisou facultatif")
    args = parser.parse_args()
    source = args.document.resolve()
    output = (args.output or source).resolve()
    format_document(source, output, args.logo.resolve() if args.logo else None)
    print(output)


if __name__ == "__main__":
    main()

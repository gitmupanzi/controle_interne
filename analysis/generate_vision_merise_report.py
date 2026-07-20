from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\Benjamin-mupanzi\Documents\GitHub\controle_interne")
VISION_DIR = ROOT / "data" / "vision"
SCHEMA_PATH = VISION_DIR / "BB_VISION_PRO.sql"
DIAGRAM_PATH = VISION_DIR / "modelisation_visuelle_bb_vision_controle_interne.png"
OUTPUT_PATH = VISION_DIR / "analyse_normalisation_merise_bb_vision_pro.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.right_indent = Cm(0.7)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def add_issue(
    document: Document,
    number: int,
    title: str,
    severity: str,
    observation: str,
    evidence: str,
    impact: str,
    recommendation: str,
) -> None:
    heading = document.add_heading(f"{number}. {title}", level=2)
    heading.paragraph_format.keep_with_next = True
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Niveau : {severity}")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 57, 43) if severity == "Haute" else RGBColor(211, 84, 0)
    document.add_paragraph(observation)
    evidence_p = document.add_paragraph()
    evidence_p.add_run("Éléments observés : ").bold = True
    evidence_p.add_run(evidence)
    impact_p = document.add_paragraph()
    impact_p.add_run("Impact : ").bold = True
    impact_p.add_run(impact)
    recommendation_p = document.add_paragraph()
    recommendation_p.add_run("Recommandation : ").bold = True
    recommendation_p.add_run(recommendation)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(13)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Analyse des faiblesses de normalisation MERISE")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BB_VISION_PRO — périmètre client, comptes, opérations et crédit")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(89, 89, 89)
    source = document.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.add_run("Sources analysées : BB_VISION_PRO.sql et modelisation_visuelle_bb_vision_controle_interne.png")
    source.runs[0].italic = True
    source.runs[0].font.size = Pt(9)

    document.add_paragraph()
    summary_box = document.add_table(rows=1, cols=1)
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_cell = summary_box.cell(0, 0)
    shade_cell(summary_cell, "EAF2F8")
    summary_cell.text = ""
    p = summary_cell.paragraphs[0]
    p.add_run("Conclusion générale\n").bold = True
    p.add_run(
        "Le modèle possède une décomposition fonctionnelle solide, notamment pour la chaîne crédit. "
        "Ses principales faiblesses concernent les cardinalités non imposées, les coordonnées multivaluées, "
        "les attributs redondants, les historiques susceptibles de multiplier les lignes et l'intégrité "
        "entre les sous-modèles métier et comptable."
    )

    document.add_heading("1. Périmètre et méthode", level=1)
    document.add_paragraph(
        "L'analyse confronte le modèle logique simplifié à la définition physique SQL Server. Elle examine "
        "les entités, associations, identifiants, dépendances fonctionnelles, cardinalités minimales et "
        "maximales, attributs multivalués, redondances et contraintes d'intégrité."
    )
    document.add_paragraph(
        "Cette analyse est structurelle : elle identifie ce que le schéma autorise. Elle ne mesure pas la "
        "fréquence réelle des anomalies, faute de connexion aux données de production."
    )

    document.add_heading("2. Modélisation visuelle examinée", level=1)
    if DIAGRAM_PATH.exists():
        document.add_picture(str(DIAGRAM_PATH), width=Cm(17.0))
        caption = document.add_paragraph("Figure 1 — MLD simplifié BB_VISION_PRO / Contrôle interne")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(8.5)

    document.add_heading("3. Points positifs du modèle", level=1)
    positives = [
        "Séparation claire des grands domaines : référentiels, clients, comptes, opérations et crédit.",
        "Chaîne crédit correctement décomposée conceptuellement : demande, dossier, prêt, cycle, échéance et remboursement.",
        "Utilisation de clés primaires et de nombreuses clés étrangères.",
        "Spécialisation d'ADHERENTS en INDIVIDUS, GROUPES, ENTREPRISES et IMFS avec clé primaire partagée.",
        "Séparation entre le compte générique COMPTES et son extension métier COMPTES_ADHERENT.",
        "Référentiels dédiés pour les devises, produits, types de clients et points de service.",
    ]
    for item in positives:
        add_bullet(document, item)

    document.add_heading("4. Synthèse des faiblesses", level=1)
    issues_summary = [
        ("Haute", "Téléphones et emails stockés dans ADRESSES", "Coordonnées multiples difficiles à gérer et risque de non-respect de la 1NF."),
        ("Haute", "Cardinalités obligatoires non imposées", "Chaîne crédit physiquement incomplète malgré des cardinalités (1,1) dans le MLD."),
        ("Haute", "Double relation client-compte", "Incohérence possible entre le compte principal et les comptes rattachés."),
        ("Haute", "Montants financiers en float", "Erreurs d'arrondi et comparaisons imprécises."),
        ("Haute", "HDPM.ID_OPERATION sans clé étrangère", "Écritures comptables potentiellement orphelines."),
        ("Moyenne", "Informations client dupliquées", "Sources de vérité concurrentes pour le nom et certains attributs."),
        ("Moyenne", "Classification des comptes redondante", "Codes de type et catégorie potentiellement contradictoires."),
        ("Moyenne", "Historiques joints comme des attributs uniques", "Multiplication des clients, comptes, crédits ou échéances."),
        ("Moyenne", "Généralisation client non exclusive", "Un adhérent peut théoriquement avoir zéro ou plusieurs sous-types."),
        ("Faible à moyenne", "Clés et domaines trop larges", "Contrôle sémantique faible et index plus volumineux."),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Priorité", "Faiblesse", "Risque principal")
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[idx], "1F4E79")
    for severity, weakness, risk in issues_summary:
        cells = table.add_row().cells
        set_cell_text(cells[0], severity, bold=True)
        set_cell_text(cells[1], weakness)
        set_cell_text(cells[2], risk)

    document.add_page_break()
    document.add_heading("5. Analyse détaillée", level=1)

    add_issue(
        document,
        1,
        "Coordonnées client insuffisamment normalisées",
        "Haute",
        "Le téléphone et l'email sont portés par ADRESSES, alors qu'une adresse postale, un téléphone et un email sont des concepts distincts.",
        "ADRESSES.TEL et ADRESSES.EMAIL ; PERSONNES possède ID_ADRESSE et ID_ADRESSE_CONTACT. Le modèle ne contient pas d'entité TELEPHONE ou CONTACT.",
        "Un client ayant plusieurs numéros ne peut pas être représenté proprement. Si plusieurs numéros sont saisis dans une seule chaîne TEL, l'atomicité de la première forme normale est perdue.",
        "Créer une entité CONTACT reliée à ADHERENT avec type, valeur, indicateur principal, date de début et date de fin. Séparer les adresses postales dans une association ADHERENT_ADRESSE avec un rôle.",
    )
    add_code_block(
        document,
        "CONTACT(id_contact, id_adherent, type_contact, valeur, principal, date_debut, date_fin)\n"
        "ADHERENT_ADRESSE(id_adherent, id_adresse, role_adresse, date_debut, date_fin)",
    )

    add_issue(
        document,
        2,
        "Double relation entre l'adhérent et ses comptes",
        "Haute",
        "Le schéma permet d'atteindre les comptes d'un client par deux chemins différents.",
        "ADHERENTS.ID_COMPTE_ADHERENT pointe vers un compte, tandis que COMPTES_ADHERENT.ID_ADHERENT rattache tous les comptes à l'adhérent.",
        "Le compte indiqué directement dans ADHERENTS peut ne pas appartenir au même client selon COMPTES_ADHERENT. Le nom ID_COMPTE_ADHERENT ne précise pas s'il s'agit du compte principal.",
        "Utiliser une seule association de détention. Si un compte principal doit être conservé, le modéliser comme un rôle et contrôler qu'il appartient bien au client.",
    )
    add_code_block(
        document,
        "DETENTION_COMPTE(id_adherent, id_compte, role, date_debut, date_fin)\n"
        "Rôles possibles : PRINCIPAL, EPARGNE, REMBOURSEMENT, GARANTIE",
    )

    add_issue(
        document,
        3,
        "Cardinalités MERISE non imposées dans le schéma physique",
        "Haute",
        "Plusieurs associations affichées comme obligatoires dans le diagramme sont traduites par des clés étrangères acceptant NULL.",
        "DOSSIERS_CREDIT.ID_DEMANDE, PRETS.ID_DOSSIER_CREDIT, CYCLES_PRET.ID_PRET, TABAMOR.ID_CYCLE_PRET, REMBOURS_CRD.ID_TABAMORT et REMBOURS_CRD.ID_OPERATION_CRD sont facultatifs.",
        "Des dossiers, prêts, cycles, échéances ou remboursements peuvent être enregistrés sans parent métier. Une clé étrangère empêche une référence invalide, mais n'empêche pas une référence absente.",
        "Rendre NOT NULL les relations réellement obligatoires après avoir contrôlé et corrigé les enregistrements orphelins ou incomplets.",
    )

    add_issue(
        document,
        4,
        "Généralisation ADHERENTS non exclusive et non totale",
        "Moyenne",
        "La clé primaire partagée traduit correctement les sous-types, mais aucune contrainte inter-table ne garantit la règle de spécialisation.",
        "INDIVIDUS.ID, GROUPES.ID, ENTREPRISES.ID et IMFS.ID référencent ADHERENTS.ID indépendamment du discriminant ID_CATEGORIE_ADHERENT.",
        "Un adhérent peut théoriquement ne figurer dans aucun sous-type ou figurer dans plusieurs sous-types. Le sous-type peut aussi être incompatible avec sa catégorie déclarée.",
        "Ajouter des contrôles d'exclusivité et de complétude dans le processus d'écriture, ou centraliser la nature juridique dans une structure qui impose exactement un sous-type.",
    )

    add_issue(
        document,
        5,
        "Informations client redondantes",
        "Moyenne",
        "Le nom du client est stocké dans plusieurs entités qui peuvent devenir des sources concurrentes.",
        "ADHERENTS.NOM_ADHERENT coexiste avec PERSONNES.NOM, GROUPES.NOM, ENTREPRISES.NOM et IMFS.NOM.",
        "Une correction effectuée dans une table peut ne pas être propagée dans les autres, entraînant des noms différents selon la requête utilisée.",
        "Définir une source de vérité unique. Les vues d'export doivent calculer le nom d'affichage à partir du sous-type, sans entretenir plusieurs valeurs modifiables du même attribut.",
    )

    add_issue(
        document,
        6,
        "Classification des comptes redondante",
        "Moyenne",
        "La nature fonctionnelle et commerciale du compte est représentée dans plusieurs colonnes sans référentiel unique commun.",
        "COMPTES.TYPE_COMPTE, COMPTES_ADHERENT.TYPE_CPTE_ADH, COMPTES_ADHERENT.CATEG_CPTE_ADH et COMPTES_ADHERENT_INFO.ID_PRODUIT_EPG.",
        "Un même compte peut recevoir des codes incompatibles, par exemple un type technique différent du produit ou de la catégorie métier.",
        "Séparer explicitement classe technique, catégorie métier et produit commercial. Remplacer les codes libres par des clés étrangères ou des contraintes CHECK.",
    )

    add_issue(
        document,
        7,
        "Montants et taux financiers stockés en float",
        "Haute",
        "La plupart des valeurs financières sont définies avec un type approximatif.",
        "PRETS.MONTANT, TABAMOR.CAPITAL, TABAMOR.INTERET, REMBOURS_CRD.CAPITAL, HDPM.MONTANT_OPERATION et de nombreux taux utilisent float.",
        "Les totaux, soldes et comparaisons peuvent subir des écarts binaires. Les tolérances de type 0,01 dans les requêtes masquent partiellement ce risque.",
        "Migrer progressivement vers decimal(19,4), ou une précision validée par devise et par règle comptable.",
    )

    add_issue(
        document,
        8,
        "Redondances et attributs dérivés dans OPERATIONS_CRD",
        "Moyenne",
        "La table mélange des références, montants détaillés et valeurs pouvant être recalculées.",
        "OPERATIONS_CRD contient ID_DEVISE et devise, ainsi que MONTANT, MONTANT_TOTAL et IMPAYE.",
        "Deux colonnes de devise peuvent diverger. Un impayé enregistré peut ne plus correspondre au calcul fondé sur TABAMOR et REMBOURS_CRD.",
        "Supprimer ou documenter les doublons, imposer une seule devise de référence et distinguer clairement données saisies, instantanés historiques et données calculées.",
    )

    add_issue(
        document,
        9,
        "Lien comptable HDPM vers OPERATIONS non sécurisé",
        "Haute",
        "HDPM contient un identifiant d'opération, mais le lien reste logique et non référentiel.",
        "HDPM.ID_OPERATION existe, mais aucune clé étrangère ne le relie à OPERATIONS.ID. Le diagramme le matérialise en pointillés.",
        "Une écriture peut pointer vers une opération inexistante ou devenir impossible à rapprocher avec le fait générateur métier.",
        "Mesurer d'abord les écritures orphelines, corriger les références, puis ajouter une contrainte de clé étrangère si le cycle de conservation des données le permet.",
    )

    add_issue(
        document,
        10,
        "Historiques utilisés comme attributs uniques",
        "Moyenne",
        "Les tables de départ et de clôture sont des historiques 1,N, mais certaines vues les joignent directement sans sélectionner l'événement applicable.",
        "DEPARTS_ADHERENT peut contenir plusieurs lignes par client et CLOTURE_COMPTE plusieurs lignes par compte.",
        "Les jointures peuvent multiplier un client, un compte, un crédit ou une échéance, puis gonfler les montants agrégés.",
        "Dans chaque requête de reporting, isoler l'événement applicable avec ROW_NUMBER(), TOP (1) ou une règle temporelle fondée sur la date de situation.",
    )

    add_issue(
        document,
        11,
        "Règles de cohérence des devises insuffisantes",
        "Haute",
        "La devise est portée à plusieurs niveaux, ce qui est parfois justifié, mais aucune règle transversale visible ne garantit leur compatibilité.",
        "PRODUITS_CRD.ID_DEVISE, PRETS.ID_DEVISE, COMPTES.ID_DEVISE, OPERATIONS_CRD.ID_DEVISE et HDPM.ID_DEVISE peuvent être renseignés indépendamment.",
        "Un prêt peut être associé à un compte de remboursement d'une autre devise sans règle explicite de conversion, ce qui fragilise les contrôles de solde et d'impayé.",
        "Définir les combinaisons autorisées, stocker la devise source et la devise cible, ainsi que le taux et la date du taux pour toute conversion.",
    )

    add_issue(
        document,
        12,
        "Domaines et identifiants insuffisamment contraints",
        "Faible à moyenne",
        "De nombreux états, types, sens et modes sont stockés dans varchar(255), tandis que beaucoup de clés techniques utilisent également varchar(255).",
        "Exemples : ETAT_DOSSIER, SENS_OPERATION, TYPE_DUREE_CREDIT, MODE_PAIEMENT et plusieurs identifiants techniques.",
        "Les variantes orthographiques sont possibles et les index deviennent plus volumineux. Les règles de domaine restent dispersées dans le code applicatif.",
        "Utiliser des référentiels ou des contraintes CHECK pour les domaines stables, et dimensionner les clés selon leur format réel.",
    )

    document.add_heading("6. Écarts entre le diagramme et le schéma physique", level=1)
    mismatch_table = document.add_table(rows=1, cols=3)
    mismatch_table.style = "Table Grid"
    mismatch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(("Association", "Lecture du diagramme", "Réalité physique")):
        set_cell_text(mismatch_table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(mismatch_table.rows[0].cells[idx], "1F4E79")
    mismatches = [
        ("DEMANDE → DOSSIER", "Dossier rattaché obligatoirement", "DOSSIERS_CREDIT.ID_DEMANDE accepte NULL."),
        ("DOSSIER → PRÊT", "Prêt rattaché à un dossier", "PRETS.ID_DOSSIER_CREDIT accepte NULL."),
        ("PRÊT → CYCLE", "Cycle rattaché à un prêt", "CYCLES_PRET.ID_PRET accepte NULL."),
        ("CYCLE → ÉCHÉANCE", "Échéance rattachée au cycle", "TABAMOR.ID_CYCLE_PRET accepte NULL."),
        ("OPÉRATION → HDPM", "Lien logique visible", "HDPM.ID_OPERATION n'est pas une clé étrangère."),
        ("ADHÉRENT → COMPTE", "Un compte principal et plusieurs comptes", "Deux chemins de rattachement sans contrôle de cohérence croisée."),
    ]
    for association, conceptual, physical in mismatches:
        cells = mismatch_table.add_row().cells
        set_cell_text(cells[0], association, bold=True)
        set_cell_text(cells[1], conceptual)
        set_cell_text(cells[2], physical)

    document.add_heading("7. Modèle cible recommandé", level=1)
    add_code_block(
        document,
        "ADHERENT\n"
        " ├── exactement un sous-type : INDIVIDU | GROUPE | ENTREPRISE | IMF\n"
        " ├── 0,N CONTACTS\n"
        " ├── 0,N ADRESSES via ADHERENT_ADRESSE\n"
        " └── 0,N COMPTES via DETENTION_COMPTE\n\n"
        "DEMANDE_CREDIT\n"
        " └── 1,N DOSSIERS\n"
        "      └── 0,N PRETS\n"
        "           ├── 1,N AFFECTATIONS_COMPTE avec rôle\n"
        "           └── 1,N CYCLES\n"
        "                └── 1,N ECHEANCES\n"
        "                     └── 0,N REGLEMENTS\n\n"
        "OPERATION\n"
        " └── 1,N ECRITURES_COMPTABLES",
    )

    document.add_heading("8. Plan d'action recommandé", level=1)
    actions = [
        "Sécuriser immédiatement le grain « une ligne par client » et « une ligne par compte » dans les requêtes de contrôle.",
        "Ajouter des requêtes d'audit des doublons, des clés étrangères NULL et des écritures HDPM orphelines.",
        "Définir une source de vérité pour le nom, le téléphone principal et l'état courant du client.",
        "Documenter la signification exacte d'ADHERENTS.ID_COMPTE_ADHERENT et contrôler son appartenance au client.",
        "Normaliser les téléphones et emails dans une future évolution du modèle.",
        "Contrôler les incohérences de devise entre produit, prêt, comptes, opérations et écritures.",
        "Préparer une migration des montants float vers decimal après étude des impacts applicatifs.",
        "Rendre NOT NULL les relations obligatoires après correction des données existantes.",
        "Ajouter la clé étrangère HDPM.ID_OPERATION seulement après traitement des références orphelines.",
        "Mettre à jour le MLD visuel afin que les cardinalités reflètent exactement les contraintes physiques réellement appliquées.",
    ]
    for action in actions:
        add_number(document, action)

    document.add_heading("9. Contrôles de données à prévoir", level=1)
    checks = [
        "Clients apparaissant plusieurs fois dans extra_clients_view.",
        "Clients possédant plusieurs valeurs distinctes de téléphone.",
        "Comptes sans adhérent ou adhérents dont le compte principal ne leur appartient pas.",
        "Adhérents sans sous-type ou présents dans plusieurs sous-types.",
        "Dossiers, prêts, cycles, échéances et remboursements dont la référence parent est NULL.",
        "Écritures HDPM dont ID_OPERATION ne correspond à aucune opération.",
        "Divergences entre OPERATIONS_CRD.ID_DEVISE et OPERATIONS_CRD.devise.",
        "Prêts dont la devise diffère de celle du compte crédit ou du compte de remboursement.",
        "Comptes répétés par plusieurs événements CLOTURE_COMPTE.",
        "Écarts entre les nombres de membres stockés dans GROUPES et les lignes de MEMBRES_GPE.",
    ]
    for check in checks:
        add_bullet(document, check)

    document.add_heading("10. Sources et repères", level=1)
    source_table = document.add_table(rows=1, cols=2)
    source_table.style = "Table Grid"
    source_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(("Source", "Repères principaux")):
        set_cell_text(source_table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(source_table.rows[0].cells[idx], "1F4E79")
    source_rows = [
        (str(SCHEMA_PATH), "ADHERENTS ligne 1807 ; ADRESSES 1852 ; CYCLES_PRET 2245 ; TABAMOR 2300 ; HDPM 2553 ; COMPTES_ADHERENT 3506 ; OPERATIONS_CRD 14754 ; REMBOURS_CRD 17853."),
        (str(DIAGRAM_PATH), "MLD simplifié des domaines référentiels, client, épargne/mouvements et crédit."),
    ]
    for source_path, reference in source_rows:
        cells = source_table.add_row().cells
        set_cell_text(cells[0], source_path)
        set_cell_text(cells[1], reference)

    footer = document.sections[0].footer.paragraphs[0]
    footer.add_run("Analyse MERISE — BB_VISION_PRO | ")
    add_page_number(footer)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()

from __future__ import annotations

from datetime import time
from io import BytesIO
import unittest

import pandas as pd

from credit_app.services.mpesa_analysis import (
    MpesaPreparedData,
    CUSTOMER_STATEMENT_COLUMNS,
    CUSTOMER_STATEMENT_FOCUS_OPERATION_TYPES,
    G2_CLASSIFIED_TRANSACTION_COLUMNS,
    build_customer_statement_view,
    build_customer_client_statement_view,
    build_customer_statement_filename,
    build_customer_matured_dat_interest_entries,
    build_customer_statement_elements,
    build_customer_transaction_analysis,
    build_filtered_turbo_deposit_withdrawal_pivot_report,
    build_filtered_turbo_balance_report,
    build_g2_daily_savings_report,
    build_g2_dat_crosscheck,
    build_g2_retention_report,
    build_g2_transaction_time_analysis,
    build_turbo_only_g2_transactions,
    build_mpesa_dat_maturity_analysis,
    build_mpesa_accounting_analysis,
    build_mpesa_management_dashboard,
    build_mpesa_credit_cockpit,
    build_mpesa_savings_cockpit,
    build_mpesa_comparison_windows,
    build_mpesa_forecast_report,
    build_mpesa_g2_statistics_quality,
    build_mpesa_statistics_report,
    build_mpesa_weekly_comparison,
    build_mpesa_year_over_year_comparison,
    build_mpesa_turbo_financial_analysis,
    build_turbo_operation_events,
    build_loan_savings_reconciliation,
    build_g2_dat_pdf_html,
    create_g2_dat_word,
    create_mpesa_statistics_word,
    create_customer_client_statement_pdf,
    create_customer_client_statement_word,
    create_customer_statement_pdf,
    create_customer_statement_word,
    create_turbo_balance_pdf,
    create_turbo_balance_word,
    build_g2_entry_report,
    build_entry_count_summary,
    build_entry_pivot,
    build_large_dat_summary,
    build_diagnostics,
    build_transaction_anomalies,
    build_load_report,
    build_savings_accounts_reconciliation,
    build_savings_final,
    build_mpesa_statement,
    build_perfect_client_crosscheck,
    create_excel_export,
    enrich_transactions_with_g2_customer_names,
    enrich_turbo_with_g2_customer_names,
    filter_g2_transactions_by_completion_time,
    filter_g2_transactions_by_direction,
    numeric_column,
    prepare_current_savings,
    prepare_customers,
    prepare_fixed_savings,
    prepare_g2_transactions,
    prepare_loans,
    prepare_perfect_clients,
    prepare_fixed_savings_from_accounts,
    prepare_savings_accounts,
    prepare_transactions,
    scope_mpesa_prepared_data_by_year,
    search_customers,
    validate_required_columns,
    TRANSACTION_REQUIRED_COLUMNS,
)


def _sample_prepared_data() -> MpesaPreparedData:
    transactions = pd.DataFrame(
        [
            {
                "id": 1,
                "customer_id": 1001.0,
                "msisdn1": "0812345678",
                "account_type": "MPESA ACCOUNT",
                "reference_id": "FA001",
                "currency_code": "cdf",
                "dr": 1000,
                "cr": 0,
                "bal_before": 9000,
                "bal_after": 8000,
                "ref_no": "TX001",
                "description": "M-Pesa Compte",
                "created_at": "2026-07-01 10:00:00",
            },
            {
                "id": 2,
                "customer_id": 1001.0,
                "msisdn1": "0812345678",
                "account_type": "FIXED SAVINGS",
                "reference_id": "FA001",
                "currency_code": "cdf",
                "dr": 0,
                "cr": 1000,
                "bal_before": 4000,
                "bal_after": 5000,
                "ref_no": "TX001",
                "description": "Depot Bloque",
                "created_at": "2026-07-01 10:00:00",
            },
            {
                "id": 3,
                "customer_id": 1001.0,
                "msisdn1": "0812345678",
                "account_type": "MPESA ACCOUNT",
                "reference_id": "LN001",
                "currency_code": "cdf",
                "dr": 0,
                "cr": 2000,
                "bal_before": 8000,
                "bal_after": 10000,
                "ref_no": "TX002",
                "description": "Montant pret",
                "created_at": "2026-07-02 10:00:00",
            },
        ]
    )
    current = pd.DataFrame(
        [
            {
                "customer_id": 1001.0,
                "msisdn": "0812345678",
                "product_name": "Courant",
                "account_type": "NORMAL SAVINGS",
                "balance": 12500,
                "currency_code": "cdf",
                "created_at": "2026-01-01",
                "updated_at": "2026-07-02",
            }
        ]
    )
    fixed = pd.DataFrame(
        [
            {
                "customer_id": 1001.0,
                "msisdn": "0812345678",
                "product_name": "DAT",
                "account_type": "FIXED SAVINGS",
                "balance": 5000,
                "currency_code": "cdf",
                "date_approved": "2026-06-01",
                "maturity_date": "2026-09-01",
            }
        ]
    )
    loans = pd.DataFrame(
        [
            {
                "loan_id": "LN001",
                "customer_id": 1001.0,
                "customer": "Client Test",
                "msisdn1": "0812345678",
                "currency_code": "cdf",
                "loan_amount": 2000,
                "loan_balance": 1500,
                "amount_paid": 500,
                "outstanding_principle": 1500,
                "outstanding_interest": 0,
                "outstanding_penalty_fees": 0,
                "status_name": "Active",
                "due_date": "2026-08-01",
                "last_repayment_date": None,
                "created_at": "2026-07-02",
                "updated_at": "2026-07-02",
            }
        ]
    )
    return MpesaPreparedData(
        transactions=prepare_transactions(transactions),
        current_savings=prepare_current_savings(current),
        fixed_savings=prepare_fixed_savings(fixed),
        loans=prepare_loans(loans),
        load_report=build_load_report({}, {}),
    )


def _sample_customer_transaction_analysis_data() -> MpesaPreparedData:
    common = {
        "customer_id": "CLIENT-ANALYSE",
        "msisdn1": "0812345678",
        "currency_code": "CDF",
    }
    rows: list[dict[str, object]] = []

    def add(
        row_id: int,
        *,
        account_type: str,
        reference_id: str,
        description: str,
        created_at: str,
        dr: float = 0,
        cr: float = 0,
        bal_before: float = 0,
        bal_after: float = 0,
        ref_no: str = "",
    ) -> None:
        rows.append(
            {
                **common,
                "id": row_id,
                "account_type": account_type,
                "reference_id": reference_id,
                "dr": dr,
                "cr": cr,
                "bal_before": bal_before,
                "bal_after": bal_after,
                "ref_no": ref_no,
                "description": description,
                "created_at": created_at,
            }
        )

    add(1, account_type="MPESA ACCOUNT", reference_id="SAV-1", description="M-Pesa Depot", created_at="2026-07-01 08:00:00", dr=100, bal_after=100, ref_no="DEP-1")
    add(2, account_type="NORMAL SAVINGS", reference_id="SAV-1", description="Epargne depot", created_at="2026-07-01 08:00:00", cr=100, bal_after=100, ref_no="DEP-1")
    add(3, account_type="MPESA ACCOUNT", reference_id="DAT-1", description="M-Pesa Compte", created_at="2026-07-02 08:00:00", dr=50, bal_after=50, ref_no="DAT-DEP-1")
    add(4, account_type="FIXED SAVINGS", reference_id="DAT-1", description="Depot Bloque", created_at="2026-07-02 08:00:00", cr=50, bal_after=50, ref_no="DAT-DEP-1")

    loan_time = "2026-07-03 08:00:00"
    add(5, account_type="PRINCIPLE", reference_id="LOAN-1", description="Montant principal", created_at=loan_time, dr=100, bal_after=100)
    add(6, account_type="LOAN ACCOUNT", reference_id="LOAN-1", description="Compte de pret", created_at=loan_time, cr=110, bal_after=110)
    add(7, account_type="PRINCIPLE", reference_id="LOAN-1", description="Montant principal", created_at=loan_time, cr=10, bal_after=10)
    add(8, account_type="LOAN ACCOUNT", reference_id="LOAN-1", description="Revenu du interets", created_at=loan_time, dr=10, bal_after=10)
    add(9, account_type="INTEREST EARNED", reference_id="LOAN-1", description="Revenu du interets", created_at=loan_time, cr=10, bal_after=10)
    add(10, account_type="LOAN AMOUNT A/C", reference_id="LOAN-1", description="Montant pret", created_at=loan_time, cr=100, bal_after=100)
    add(11, account_type="MPESA ACCOUNT", reference_id="LOAN-1", description="Montant pret", created_at=loan_time, cr=100, bal_after=100)
    add(12, account_type="MPESA ACCOUNT", reference_id="LOAN-1", description="Compte du M-Pesa", created_at=loan_time, dr=10, bal_after=10)

    repayment_time = "2026-07-04 08:00:00"
    add(13, account_type="PRINCIPLE", reference_id="LOAN-1", description="Remboursement du principal", created_at=repayment_time, cr=40, bal_before=100, bal_after=60)
    add(14, account_type="LOAN ACCOUNT", reference_id="LOAN-1", description="Remboursement du Pret", created_at=repayment_time, dr=40, bal_before=100, bal_after=60)
    add(15, account_type="LOAN PORTFOLIO", reference_id="LOAN-1", description="Portefeuille Pret Remboursement", created_at=repayment_time, dr=40, bal_before=100, bal_after=60)
    add(16, account_type="MPESA ACCOUNT", reference_id="LOAN-1", description="Remboursement du M-Pesa", created_at=repayment_time, dr=40, bal_after=40)
    add(17, account_type="LOAN PENALTY FEES", reference_id="LOAN-1", description="Compte de penalite de pret", created_at=repayment_time, cr=5, bal_after=5)
    add(18, account_type="CUSTOMER USD WALLET PENALTY", reference_id="LOAN-1", description="Compte de penalite de pret", created_at=repayment_time, cr=5, bal_after=5)

    transfer_time = "2026-07-05 08:00:00"
    add(19, account_type="FIXED SAVINGS", reference_id="DAT-1", description="Retrait Compte Bloque", created_at=transfer_time, dr=20, bal_before=50, bal_after=30)
    add(20, account_type="NORMAL SAVINGS", reference_id="SAV-1", description="Retrait Compte Bloque", created_at=transfer_time, dr=20, bal_before=100, bal_after=120)

    current = pd.DataFrame(
        [
            {
                "customer_id": "CLIENT-ANALYSE",
                "msisdn": "0812345678",
                "product_name": "Courant",
                "account_type": "NORMAL SAVINGS",
                "balance": 120,
                "currency_code": "CDF",
                "created_at": "2026-01-01",
                "updated_at": "2026-07-05",
            }
        ]
    )
    fixed = pd.DataFrame(
        [
            {
                "savings_id": "DAT-1",
                "customer_id": "CLIENT-ANALYSE",
                "msisdn": "0812345678",
                "product_name": "DAT",
                "account_type": "FIXED SAVINGS",
                "balance": 30,
                "currency_code": "CDF",
                "date_approved": "2026-07-02",
                "maturity_date": "2026-10-02",
                "created_at": "2026-07-02 07:45:00",
            }
        ]
    )
    loans = pd.DataFrame(
        [
            {
                "loan_id": "LOAN-1",
                "customer_id": "CLIENT-ANALYSE",
                "currency_code": "CDF",
                "loan_amount": 100,
                "loan_balance": 60,
                "amount_paid": 40,
                "outstanding_principle": 60,
                "outstanding_interest": 0,
                "outstanding_penalty_fees": 5,
                "status_name": "Active",
                "created_at": "2026-07-03",
                "updated_at": "2026-07-05",
            }
        ]
    )
    return MpesaPreparedData(
        transactions=prepare_transactions(pd.DataFrame(rows)),
        current_savings=prepare_current_savings(current),
        fixed_savings=prepare_fixed_savings(fixed),
        loans=prepare_loans(loans),
        load_report=pd.DataFrame(),
    )


def _sample_seven_percent_loan_data() -> MpesaPreparedData:
    """Reproduit le pret de 5 USD du scenario Benjamin au grain Turbo."""
    common = {
        "customer_id": "37370",
        "msisdn1": "243000000000",
        "currency_code": "USD",
        "reference_id": "PRET_TEST_001",
    }
    rows: list[dict[str, object]] = []

    def add(
        row_id: int,
        account_type: str,
        description: str,
        created_at: str,
        *,
        dr: float = 0.0,
        cr: float = 0.0,
        ref_no: str = "",
        reference_id: str | None = None,
    ) -> None:
        rows.append(
            {
                **common,
                "reference_id": common["reference_id"] if reference_id is None else reference_id,
                "id": row_id,
                "account_type": account_type,
                "dr": dr,
                "cr": cr,
                "bal_before": 0.0,
                "bal_after": max(dr, cr),
                "ref_no": ref_no,
                "description": description,
                "created_at": created_at,
            }
        )

    origination = "2026-07-22 16:17:16"
    add(98365, "PRINCIPLE", "Montant principal", origination, dr=5.0)
    add(98366, "LOAN ACCOUNT", "Compte de pret", origination, cr=5.35)
    add(98367, "PRINCIPLE", "Montant principal", origination, cr=0.35)
    add(98368, "LOAN ACCOUNT", "Revenu du interets", origination, dr=0.35)
    add(98369, "INTEREST EARNED", "Revenu du interets", origination, cr=0.35)
    add(98370, "LOAN PORTFOLIO", "Portefeuille Pret", origination, cr=5.35)
    add(98371, "LOAN PORTFOLIO", "Revenu du interets", origination, dr=0.35)
    add(98372, "LOAN AMOUNT A/C", "Montant pret", origination, cr=5.0)
    add(98373, "MPESA ACCOUNT", "Montant pret", origination, cr=5.0)
    add(98374, "MPESA ACCOUNT", "Compte du M-Pesa", origination, dr=0.35)
    add(98375, "BISOU COLLECTION", "Part Bisou", origination, cr=2.50)
    add(98376, "VODA COLLECTION A/C", "Part Voda", origination, cr=1.00)

    repayment = "2026-07-22 16:26:43"
    repayment_reference = "RETRAIT_TEST_001"
    add(98385, "PRINCIPLE", "Remboursement du principal", repayment, cr=5.0, ref_no=repayment_reference)
    add(98386, "LOAN ACCOUNT", "Remboursement du Pret", repayment, dr=5.0, ref_no=repayment_reference)
    add(98387, "LOAN PORTFOLIO", "Portefeuille Pret Remboursement", repayment, dr=5.0, ref_no=repayment_reference)
    add(98388, "MPESA ACCOUNT", "Remboursement du M-Pesa", repayment, dr=5.0, ref_no=repayment_reference)
    add(98389, "NORMAL SAVINGS", "Remboursement de compte epargne", repayment, dr=5.0, ref_no=repayment_reference, reference_id="")
    add(98390, "FIXED SAVINGS", "Remboursement de compte epargne", repayment, dr=5.0, ref_no=repayment_reference, reference_id="")

    loans = prepare_loans(
        pd.DataFrame(
            [
                {
                    "loan_id": "PRET_TEST_001",
                    "customer_id": "37370",
                    "customer": "MUPANZI KITSHI BENJAMIN",
                    "currency_code": "USD",
                    "loan_amount": 5.0,
                    "loan_balance": 0.0,
                    "amount_paid": 5.35,
                    "outstanding_principle": 0.0,
                    "outstanding_interest": 0.0,
                    "interest_earned": 0.35,
                    "status_name": "Matured",
                    "created_at": origination,
                    "updated_at": repayment,
                    "msisdn1": "243000000000",
                }
            ]
        )
    )
    return MpesaPreparedData(
        transactions=prepare_transactions(pd.DataFrame(rows)),
        current_savings=pd.DataFrame(),
        fixed_savings=pd.DataFrame(),
        loans=loans,
        load_report=build_load_report({}, {}),
    )


class MpesaAnalysisTests(unittest.TestCase):
    def test_multi_file_turbo_and_perfect_snapshots_are_deduplicated_by_business_key(self) -> None:
        transaction_rows = pd.DataFrame(
            [
                {
                    "id": 1,
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "REF-1",
                    "currency_code": "CDF",
                    "dr": 100,
                    "cr": 0,
                    "bal_before": 500,
                    "bal_after": 400,
                    "ref_no": "TX-1",
                    "description": "M-Pesa Depot",
                    "created_at": "2026-07-15 08:00:00",
                    "fichier_source_transactions_turbo": "transactions_a.xlsx",
                    "ordre_fichier_import": 0,
                },
                {
                    "id": 1,
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "REF-1",
                    "currency_code": "CDF",
                    "dr": 100,
                    "cr": 0,
                    "bal_before": 500,
                    "bal_after": 400,
                    "ref_no": "TX-1",
                    "description": "M-Pesa Depot",
                    "created_at": "2026-07-15 08:00:00",
                    "fichier_source_transactions_turbo": "transactions_b.xlsx",
                    "ordre_fichier_import": 1,
                },
                {
                    "id": 2,
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "REF-2",
                    "currency_code": "CDF",
                    "dr": 0,
                    "cr": 50,
                    "bal_before": 400,
                    "bal_after": 450,
                    "ref_no": "TX-2",
                    "description": "Retrait vers M-Pesa",
                    "created_at": "2026-07-15 09:00:00",
                    "fichier_source_transactions_turbo": "transactions_b.xlsx",
                    "ordre_fichier_import": 1,
                },
            ]
        )
        transactions = prepare_transactions(transaction_rows).set_index("id")

        self.assertEqual(len(transactions), 2)
        self.assertEqual(float(transactions.loc["1", "dr"]), 100.0)
        self.assertEqual(float(transactions.loc["2", "cr"]), 50.0)
        self.assertEqual(
            transactions.loc["1", "fichiers_sources_transactions_turbo"],
            "transactions_a.xlsx | transactions_b.xlsx",
        )
        self.assertNotIn("ordre_fichier_import", transactions.columns)

        current_rows = pd.DataFrame(
            [
                {
                    "customer_id": 10,
                    "msisdn": "0811111111",
                    "product_name": "Open Savings",
                    "account_type": "Current account",
                    "balance": balance,
                    "currency_code": "CDF",
                    "created_at": "2026-07-01 08:00:00",
                    "updated_at": updated_at,
                    "fichier_source_epargne_turbo": source,
                    "ordre_fichier_import": order,
                }
                for balance, updated_at, source, order in [
                    (100, "2026-07-14 08:00:00", "epargne_a.xlsx", 0),
                    (150, "2026-07-15 08:00:00", "epargne_b.xlsx", 1),
                ]
            ]
        )
        current = prepare_current_savings(current_rows)
        self.assertEqual(len(current), 1)
        self.assertEqual(float(current.iloc[0]["balance"]), 150.0)
        self.assertEqual(current.iloc[0]["fichiers_sources_epargne_turbo"], "epargne_a.xlsx | epargne_b.xlsx")

        fixed_rows = pd.DataFrame(
            [
                {
                    "customer_id": 10,
                    "msisdn": "0811111111",
                    "product_name": "1 Month",
                    "account_type": "1 Month Fixed Account",
                    "balance": balance,
                    "currency_code": "CDF",
                    "date_approved": "2026-07-01 08:00:00",
                    "maturity_date": "2026-08-01 08:00:00",
                    "fichier_source_dat_turbo": source,
                    "ordre_fichier_import": order,
                }
                for balance, source, order in [
                    (1000, "dat_a.xlsx", 0),
                    (1100, "dat_b.xlsx", 1),
                ]
            ]
        )
        fixed = prepare_fixed_savings(fixed_rows)
        self.assertEqual(len(fixed), 1)
        self.assertEqual(float(fixed.iloc[0]["balance"]), 1100.0)

        loan_rows = pd.DataFrame(
            [
                {
                    "id": 1,
                    "loan_id": "LN-1",
                    "customer_id": 10,
                    "loan_balance": balance,
                    "updated_at": updated_at,
                    "fichier_source_credits_turbo": source,
                    "ordre_fichier_import": order,
                }
                for balance, updated_at, source, order in [
                    (500, "2026-07-14", "credits_a.xlsx", 0),
                    (400, "2026-07-15", "credits_b.xlsx", 1),
                ]
            ]
        )
        loans = prepare_loans(loan_rows)
        self.assertEqual(len(loans), 1)
        self.assertEqual(float(loans.iloc[0]["loan_balance"]), 400.0)

        customer_rows = pd.DataFrame(
            [
                {
                    "msisdn1": "0811111111",
                    "created_at": "2026-07-01",
                    "fichier_source_clients_turbo": source,
                    "ordre_fichier_import": order,
                }
                for source, order in [("clients_a.xlsx", 0), ("clients_b.xlsx", 1)]
            ]
        )
        customers = prepare_customers(customer_rows)
        self.assertEqual(len(customers), 1)
        self.assertEqual(customers.iloc[0]["fichier_source_clients_turbo"], "clients_b.xlsx")

        perfect_rows = pd.DataFrame(
            [
                {
                    "id_client": "P-1",
                    "code_client": "C-1",
                    "nom_complet": name,
                    "Phone_Prefixe": "243811111111",
                    "fichier_source_clients_perfect": source,
                    "ordre_fichier_import": order,
                }
                for name, source, order in [
                    ("ANCIEN NOM", "perfect_a.xlsx", 0),
                    ("NOUVEAU NOM", "perfect_b.xlsx", 1),
                ]
            ]
        )
        perfect = prepare_perfect_clients(perfect_rows)
        self.assertEqual(len(perfect), 1)
        self.assertEqual(perfect.iloc[0]["nom_complet"], "NOUVEAU NOM")
        self.assertEqual(perfect.iloc[0]["fichiers_sources_clients_perfect"], "perfect_a.xlsx | perfect_b.xlsx")

    def test_prepare_current_savings_accepts_mixed_savings_account_export(self) -> None:
        raw = pd.DataFrame(
            [
                {
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "balance": 150,
                    "currency_code": "CDF",
                    "created_at": "2026-07-01 08:00:00",
                    "updated_at": "2026-07-17 08:00:00",
                },
                {
                    "customer_id": 11,
                    "msisdn1": "0822222222",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "balance": 25,
                    "currency_code": "USD",
                    "created_at": "2026-07-02 08:00:00",
                    "updated_at": "2026-07-17 08:00:00",
                },
                {
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "balance": 500,
                    "currency_code": "CDF",
                    "created_at": "2026-07-03 08:00:00",
                    "updated_at": "2026-07-17 08:00:00",
                },
            ]
        )

        result = prepare_current_savings(raw)

        self.assertEqual(len(result), 2)
        self.assertEqual(set(result["account_type"]), {"NORMAL SAVINGS"})
        self.assertEqual(set(result["msisdn"]), {"243811111111", "243822222222"})
        self.assertEqual(float(result["balance"].sum()), 175.0)

    def test_complete_savings_account_source_retains_and_reconciles_fixed_accounts(self) -> None:
        raw = pd.DataFrame(
            [
                {
                    "savings_id": "SA-1",
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "balance": 150,
                    "currency_code": "CDF",
                    "created_at": "2026-07-01",
                    "fichier_source_epargne_turbo": "Savings Account.xlsx",
                },
                {
                    "savings_id": "FA-1",
                    "customer_id": 10,
                    "msisdn1": "0811111111",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "balance": 500,
                    "currency_code": "CDF",
                    "date_approved": "2026-07-02",
                    "maturity_date": "2026-08-02",
                    "created_at": "2026-07-02",
                    "fichier_source_epargne_turbo": "Savings Account.xlsx",
                },
                {
                    "savings_id": "FA-2",
                    "customer_id": 11,
                    "msisdn1": "0822222222",
                    "product_name": "3 Months",
                    "product_description": "3 MONTH Fixed Account",
                    "balance": 0,
                    "currency_code": "USD",
                    "date_approved": "2026-04-02",
                    "maturity_date": "2026-07-02",
                    "created_at": "2026-04-02",
                    "fichier_source_epargne_turbo": "Savings Account.xlsx",
                },
            ]
        )
        control = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": 10,
                        "msisdn": "0811111111",
                        "product_name": "1 Month",
                        "account_type": "1 Month Fixed Account",
                        "balance": 500,
                        "currency_code": "CDF",
                        "date_approved": "2026-07-02",
                        "maturity_date": "2026-08-02",
                    }
                ]
            )
        )
        all_accounts = prepare_savings_accounts(raw)
        fixed = prepare_fixed_savings_from_accounts(raw)
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=all_accounts.loc[
                all_accounts["account_type"].eq("NORMAL SAVINGS")
            ],
            fixed_savings=fixed,
            fixed_savings_control=control,
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        report = build_savings_accounts_reconciliation(prepared)
        summary = report["synthese"].iloc[0]

        self.assertEqual(len(all_accounts), 3)
        self.assertEqual(len(fixed), 2)
        self.assertEqual(int(summary["dat_solde_positif"]), 1)
        self.assertEqual(int(summary["dat_solde_nul"]), 1)
        self.assertEqual(int(summary["dat_export_retrouves"]), 1)
        self.assertEqual(summary["statut_rapprochement"], "Concordance exacte")
        self.assertTrue(report["ecarts"].empty)

    def test_savings_account_is_autonomous_without_summary_exports(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(
                [
                    {"customer_id": 10, "balance": 150},
                    {"customer_id": 11, "balance": 0},
                ]
            ),
            fixed_savings=pd.DataFrame(
                [
                    {
                        "customer_id": 10,
                        "currency_code": "CDF",
                        "balance": 500,
                        "fichier_source_epargne_turbo": "Savings Account.xlsx",
                    },
                    {
                        "customer_id": 11,
                        "currency_code": "USD",
                        "balance": 0,
                        "fichier_source_epargne_turbo": "Savings Account.xlsx",
                    },
                ]
            ),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        report = build_savings_accounts_reconciliation(prepared)
        summary = report["synthese"].iloc[0]

        self.assertEqual(summary["statut_rapprochement"], "Source autonome")
        self.assertEqual(int(summary["comptes_courants"]), 2)
        self.assertEqual(int(summary["dat_total_source_complete"]), 2)
        self.assertEqual(int(summary["dat_solde_positif"]), 1)
        self.assertEqual(int(summary["dat_solde_nul"]), 1)
        self.assertEqual(int(summary["dat_positifs_absents_export_resume"]), 0)
        self.assertTrue(report["ecarts"].empty)

    def test_prepare_g2_transactions_promotes_two_organization_statement_headers(self) -> None:
        def statement_frame(
            account_name: str,
            source_name: str,
            receipt_no: str,
            *,
            paid_in: object = None,
            withdrawn: object = None,
        ) -> pd.DataFrame:
            columns = ["Account Holder:", account_name] + [
                f"Unnamed: {index}" for index in range(2, 12)
            ]
            rows = [
                ["Short Code:", "1441", *([None] * 10)],
                ["Account:", "All Account", *([None] * 10)],
                ["Time Period:", "From", "01-07-2026", "To", "17-07-2026", *([None] * 7)],
                ["Operator:", "GOMA", "Organization:", account_name, *([None] * 8)],
                [
                    "Receipt No.",
                    "Completion Time",
                    "Initiation Time",
                    "Details",
                    "Transaction Status",
                    "Currency",
                    "Paid In",
                    "Withdrawn",
                    "Balance",
                    "Reason Type",
                    "Opposite Party",
                    "Linked Transaction ID",
                ],
                [
                    receipt_no,
                    "\t16-07-2026 10:00:00",
                    "\t16-07-2026 09:59:00",
                    "BisouBisouC2B" if paid_in is not None else "Bisou Bisou B2C payment",
                    "Completed",
                    "CDF",
                    paid_in,
                    withdrawn,
                    "1000",
                    "BisouBisouC2B" if paid_in is not None else "BisouBisouB2C",
                    "\t243811111111 - CLIENT TEST",
                    None,
                ],
            ]
            frame = pd.DataFrame(rows, columns=columns)
            frame["fichier_source_g2"] = source_name
            frame["ordre_fichier_import"] = 0
            return frame

        raw = pd.concat(
            [
                statement_frame(
                    "IMF Bisou Bisou SA",
                    "ORG_1441.xlsx",
                    "ENTRY-001",
                    paid_in="100",
                ),
                statement_frame(
                    "IMF Bisou Bisou  SA",
                    "ORG_15558.xlsx",
                    "OUTPUT-001",
                    withdrawn="-75",
                ),
            ],
            ignore_index=True,
            sort=False,
        )

        result = prepare_g2_transactions(raw).set_index("receipt_no")

        self.assertEqual(set(result.index), {"ENTRY-001", "OUTPUT-001"})
        self.assertEqual(result.loc["ENTRY-001", "sens_flux"], "Entree")
        self.assertEqual(result.loc["OUTPUT-001", "sens_flux"], "Sortie")
        self.assertEqual(float(result.loc["OUTPUT-001", "montant_sortie"]), 75.0)
        self.assertEqual(result.loc["ENTRY-001", "fichier_source_g2"], "ORG_1441.xlsx")
        self.assertEqual(result.loc["OUTPUT-001", "fichier_source_g2"], "ORG_15558.xlsx")

    def test_g2_count_summary_keeps_currencies_separate_and_fills_missing_categories(self) -> None:
        detail = pd.DataFrame(
            [
                {"currency_code": "CDF", "details_rapport": "DAT"},
                {"currency_code": "CDF", "details_rapport": "DAT"},
                {"currency_code": "CDF", "details_rapport": "Depot normal"},
                {"currency_code": "USD", "details_rapport": "Depot normal"},
                {"currency_code": "USD", "details_rapport": "Remboursement prets"},
            ]
        )

        result = build_entry_count_summary(detail).set_index("currency_code")

        self.assertEqual(int(result.loc["CDF", "Nombre de DAT"]), 2)
        self.assertEqual(int(result.loc["CDF", "Nombre de remboursement de pret"]), 0)
        self.assertEqual(int(result.loc["CDF", "Nombre total"]), 3)
        self.assertEqual(int(result.loc["USD", "Nombre de DAT"]), 0)
        self.assertEqual(int(result.loc["USD", "Nombre de depot normal"]), 1)
        self.assertEqual(int(result.loc["USD", "Nombre de remboursement de pret"]), 1)
        self.assertEqual(int(result.loc["USD", "Nombre total"]), 2)

    def test_filter_g2_completion_time_uses_inclusive_dates(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {"Receipt No.": "D1", "Completion Time": "10-07-2026 23:59:59", "Currency": "CDF", "Opposite Party": "0811111111 - A"},
                    {"Receipt No.": "D2", "Completion Time": "11-07-2026 00:00:00", "Currency": "CDF", "Opposite Party": "0822222222 - B"},
                    {"Receipt No.": "D3", "Completion Time": "11-07-2026 23:59:59", "Currency": "USD", "Opposite Party": "0833333333 - C"},
                    {"Receipt No.": "D4", "Completion Time": "12-07-2026 00:00:00", "Currency": "USD", "Opposite Party": "0844444444 - D"},
                ]
            )
        )

        result = filter_g2_transactions_by_completion_time(g2, "2026-07-11", "2026-07-11")

        self.assertEqual(result["receipt_no"].tolist(), ["D2", "D3"])

    def test_filter_g2_completion_time_uses_inclusive_hours(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {"Receipt No.": "T1", "Completion Time": "11-07-2026 09:59:59", "Currency": "CDF", "Opposite Party": "0811111111 - A"},
                    {"Receipt No.": "T2", "Completion Time": "11-07-2026 10:00:00", "Currency": "CDF", "Opposite Party": "0822222222 - B"},
                    {"Receipt No.": "T3", "Completion Time": "11-07-2026 15:30:00", "Currency": "USD", "Opposite Party": "0833333333 - C"},
                    {"Receipt No.": "T4", "Completion Time": "11-07-2026 15:30:01", "Currency": "USD", "Opposite Party": "0844444444 - D"},
                ]
            )
        )

        result = filter_g2_transactions_by_completion_time(
            g2,
            "2026-07-11",
            "2026-07-11",
            time(10, 0),
            time(15, 30),
        )

        self.assertEqual(result["receipt_no"].tolist(), ["T2", "T3"])

    def test_filter_g2_direction_supports_entries_outputs_and_both(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {"Receipt No.": "IN", "Currency": "CDF", "Paid In": 100, "Opposite Party": "0811111111 - A"},
                    {"Receipt No.": "OUT", "Currency": "CDF", "Withdrawn": -50, "Opposite Party": "0822222222 - B"},
                    {"Receipt No.": "CHECK", "Currency": "CDF", "Opposite Party": "0833333333 - C"},
                ]
            )
        )

        entries = filter_g2_transactions_by_direction(g2, "Entrees")
        outputs = filter_g2_transactions_by_direction(g2, "Sorties")
        both = filter_g2_transactions_by_direction(g2, None)

        self.assertEqual(entries["receipt_no"].tolist(), ["IN"])
        self.assertEqual(outputs["receipt_no"].tolist(), ["OUT"])
        self.assertEqual(set(both["receipt_no"]), {"IN", "OUT", "CHECK"})

    def test_prepare_g2_transactions_accepts_paid_in_and_withdrawn_format(self) -> None:
        raw = pd.DataFrame(
            [
                {
                    "Receipt No.": "ORG-IN",
                    "Completion Time": "11-07-2026 12:47:24",
                    "Initiation Time": "11-07-2026 12:46:00",
                    "Details": "BisouBisouC2B",
                    "Transaction Status": "Completed",
                    "Currency": "CDF",
                    "Paid In": 1000.0,
                    "Withdrawn": None,
                    "Balance": 5000.0,
                    "Opposite Party": "0812345678 - CLIENT DEPOT",
                },
                {
                    "Receipt No.": "ORG-REPAY",
                    "Completion Time": "11-07-2026 13:00:00",
                    "Initiation Time": "11-07-2026 12:59:00",
                    "Details": "BisouBisouRepayment",
                    "Transaction Status": "Completed",
                    "Currency": "CDF",
                    "Paid In": 250.0,
                    "Withdrawn": None,
                    "Balance": 5250.0,
                    "Opposite Party": "0812345678 - CLIENT DEPOT",
                },
                {
                    "Receipt No.": "ORG-OUT",
                    "Completion Time": "11-07-2026 14:00:00",
                    "Initiation Time": "11-07-2026 13:59:00",
                    "Details": "Super Transaction",
                    "Transaction Status": "Completed",
                    "Currency": "CDF",
                    "Paid In": None,
                    "Withdrawn": -500.0,
                    "Balance": 4750.0,
                    "Opposite Party": "0812345678 - CLIENT DEPOT",
                },
            ]
        )

        result = prepare_g2_transactions(raw).set_index("receipt_no")

        self.assertEqual(float(result.loc["ORG-IN", "transaction_amount_numeric"]), 1000.0)
        self.assertEqual(result.loc["ORG-IN", "transaction_amount_source"], "Paid In")
        self.assertEqual(float(result.loc["ORG-REPAY", "transaction_amount_numeric"]), 250.0)
        self.assertEqual(result.loc["ORG-REPAY", "transaction_amount_source"], "Paid In")
        self.assertEqual(float(result.loc["ORG-OUT", "transaction_amount_numeric"]), -500.0)
        self.assertEqual(result.loc["ORG-OUT", "transaction_amount_source"], "Withdrawn")
        self.assertEqual(result.loc["ORG-IN", "sens_flux"], "Entree")
        self.assertEqual(result.loc["ORG-OUT", "sens_flux"], "Sortie")
        self.assertEqual(float(result.loc["ORG-IN", "montant_entree"]), 1000.0)
        self.assertEqual(float(result.loc["ORG-OUT", "montant_sortie"]), 500.0)
        self.assertEqual(result.loc["ORG-OUT", "type_operation_g2"], "Operation interne Bisou")
        self.assertEqual(float(result.loc["ORG-OUT", "balance_numeric"]), 4750.0)
        self.assertEqual(result.loc["ORG-IN", "completion_time"], pd.Timestamp("2026-07-11 12:47:24"))
        self.assertEqual(result.loc["ORG-IN", "Nom_client"], "CLIENT DEPOT")

    def test_g2_report_classifies_b2c_and_loan_request_as_outflows(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "ENTRY",
                        "Completion Time": "13-07-2026 08:00:00",
                        "Details": "BisouBisouC2B",
                        "Reason Type": "BisouBisouC2B",
                        "Currency": "CDF",
                        "Paid In": 1000,
                        "Opposite Party": "0811111111 - CLIENT ENTREE",
                    },
                    {
                        "Receipt No.": "REPAY",
                        "Completion Time": "13-07-2026 08:10:00",
                        "Details": "BisouBisouC2BRepayment",
                        "Reason Type": "BisouBisouC2BRepayment",
                        "Currency": "CDF",
                        "Paid In": 250,
                        "Opposite Party": "0811111111 - CLIENT ENTREE",
                    },
                    {
                        "Receipt No.": "B2C",
                        "Completion Time": "13-07-2026 08:20:00",
                        "Details": "Bisou Bisou B2C payment to client",
                        "Reason Type": "BisouBisouB2C",
                        "Currency": "CDF",
                        "Withdrawn": -400,
                        "Opposite Party": "0822222222 - CLIENT SORTIE",
                    },
                    {
                        "Receipt No.": "LOAN",
                        "Completion Time": "13-07-2026 08:30:00",
                        "Details": "Bisou Bisou Loan Request payment",
                        "Reason Type": "BisouBisouLoanRequest",
                        "Currency": "CDF",
                        "Withdrawn": -1500,
                        "Opposite Party": "0833333333 - CLIENT CREDIT",
                    },
                    {
                        "Receipt No.": "SUPER",
                        "Completion Time": "13-07-2026 08:40:00",
                        "Details": "Super Transaction",
                        "Reason Type": "Super Transaction",
                        "Currency": "CDF",
                        "Withdrawn": -100,
                        "Opposite Party": "15558 - IMF BISOU",
                    },
                ]
            )
        )
        fixed = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": "3003",
                        "msisdn": "0833333333",
                        "product_name": "1 Month",
                        "account_type": "FIXED SAVINGS",
                        "balance": 1500,
                        "currency_code": "CDF",
                        "date_approved": "2026-07-13 08:30:00",
                        "maturity_date": "2026-08-13",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=fixed,
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_g2_daily_savings_report(prepared)
        detail = report["detail"].set_index("receipt_no")
        pivot = report["pivot"].set_index("currency_code")

        self.assertEqual(detail.loc["ENTRY", "sens_flux"], "Entree")
        self.assertEqual(detail.loc["ENTRY", "details_rapport"], "Depot normal")
        self.assertEqual(detail.loc["REPAY", "details_rapport"], "Remboursement prets")
        self.assertEqual(detail.loc["B2C", "sens_flux"], "Sortie")
        self.assertEqual(detail.loc["B2C", "details_rapport"], "Paiement client B2C")
        self.assertEqual(detail.loc["LOAN", "details_rapport"], "Demande de credit")
        self.assertTrue(pd.isna(detail.loc["LOAN", "dat_customer_id"]))
        self.assertEqual(detail.loc["SUPER", "details_rapport"], "Operation interne Bisou")
        self.assertEqual(
            detail.loc["SUPER", "statut_rapprochement"],
            "Non applicable - operation interne",
        )
        self.assertEqual(detail.loc["SUPER", "motif_anomalie"], "")
        self.assertEqual(int(pivot.loc["CDF", "nombre_entrees"]), 2)
        self.assertEqual(int(pivot.loc["CDF", "nombre_sorties"]), 3)
        self.assertEqual(float(pivot.loc["CDF", "montant_Demande de credit"]), 1500.0)
        self.assertEqual(float(pivot.loc["CDF", "montant_total_sorties"]), 2000.0)
        self.assertEqual(float(pivot.loc["CDF", "solde_net_flux"]), -750.0)

    def test_large_dat_summary_ranks_clients_by_currency_without_mixing_totals(self) -> None:
        fixed = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": "1001",
                        "msisdn": "0811111111",
                        "Nom_client": "CLIENT A",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 7000,
                        "currency_code": "CDF",
                        "date_approved": "2026-06-01",
                        "maturity_date": "2026-07-20",
                    },
                    {
                        "customer_id": "1001",
                        "msisdn": "0811111111",
                        "Nom_client": "CLIENT A",
                        "product_name": "6 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 3000,
                        "currency_code": "CDF",
                        "date_approved": "2026-05-01",
                        "maturity_date": "2026-11-01",
                    },
                    {
                        "customer_id": "1002",
                        "msisdn": "0822222222",
                        "Nom_client": "CLIENT B",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 2000,
                        "currency_code": "CDF",
                        "date_approved": "2026-06-10",
                        "maturity_date": "2026-09-10",
                    },
                    {
                        "customer_id": "2001",
                        "msisdn": "0833333333",
                        "Nom_client": "CLIENT USD",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 500,
                        "currency_code": "USD",
                        "date_approved": "2026-04-01",
                        "maturity_date": "2026-07-01",
                    },
                ]
            )
        )

        result = build_large_dat_summary(fixed, percentile=0.90, as_of_date="2026-07-13")
        clients = result["clients"]
        portfolio = result["portefeuille"].set_index("currency_code")
        cdf = clients.loc[clients["currency_code"].eq("CDF")].set_index("customer_id")

        self.assertEqual(float(cdf.loc["1001", "solde_dat_total"]), 10000.0)
        self.assertEqual(int(cdf.loc["1001", "nb_comptes_dat"]), 2)
        self.assertEqual(int(cdf.loc["1001", "rang_devise"]), 1)
        self.assertTrue(bool(cdf.loc["1001", "est_fort_dat"]))
        self.assertFalse(bool(cdf.loc["1002", "est_fort_dat"]))
        self.assertEqual(float(portfolio.loc["CDF", "total_dat"]), 12000.0)
        self.assertEqual(float(portfolio.loc["USD", "total_dat"]), 500.0)
        self.assertEqual(float(portfolio.loc["CDF", "solde_echeance_30j"]), 7000.0)
        self.assertEqual(float(portfolio.loc["USD", "solde_dat_echu"]), 500.0)

    def test_g2_customer_name_enrichment_prioritizes_phone_then_reference(self) -> None:
        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": "1001",
                        "msisdn1": "0812345678",
                        "ref_no": "REF-PHONE",
                    },
                    {
                        "id": 2,
                        "customer_id": "1002",
                        "msisdn1": "0999999999",
                        "ref_no": "REF-FALLBACK",
                    },
                    {
                        "id": 3,
                        "customer_id": "1003",
                        "msisdn1": "0977777777",
                        "ref_no": "REF-ABSENT",
                    },
                ]
            )
        )
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {"Receipt No.": "REF-PHONE", "Opposite Party": "0812345678 - NOM PAR TELEPHONE", "Currency": "CDF"},
                    {"Receipt No.": "REF-FALLBACK", "Opposite Party": "0822222222 - NOM PAR REFERENCE", "Currency": "CDF"},
                ]
            )
        )

        result = enrich_transactions_with_g2_customer_names(transactions, g2)

        self.assertEqual(result.loc[0, "Nom_client"], "NOM PAR TELEPHONE")
        self.assertEqual(result.loc[0, "mode_rapprochement_nom_client"], "Telephone G2 = msisdn1 Solution Numérique")
        self.assertEqual(result.loc[1, "Nom_client"], "NOM PAR REFERENCE")
        self.assertEqual(result.loc[1, "mode_rapprochement_nom_client"], "Receipt No G2 = ref_no Solution Numérique")
        self.assertTrue(pd.isna(result.loc[2, "Nom_client"]))
        self.assertEqual(result.loc[2, "mode_rapprochement_nom_client"], "Nom G2 non rapproche")

    def test_g2_customer_name_enrichment_handles_missing_optional_file(self) -> None:
        transactions = prepare_transactions(pd.DataFrame([{"id": 1, "msisdn1": "0812345678", "ref_no": "REF-1"}]))

        result = enrich_transactions_with_g2_customer_names(transactions, pd.DataFrame())

        self.assertIn("Nom_client", result.columns)
        self.assertTrue(pd.isna(result.loc[0, "Nom_client"]))
        self.assertEqual(result.loc[0, "mode_rapprochement_nom_client"], "Fichier G2 absent")

    def test_g2_customer_name_is_propagated_to_turbo_reports(self) -> None:
        prepared = _sample_prepared_data()
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {"Receipt No.": "TX001", "Opposite Party": "0812345678 - CLIENT TEST", "Currency": "CDF"},
                ]
            )
        )
        transactions = enrich_transactions_with_g2_customer_names(prepared.transactions, g2)
        current = enrich_turbo_with_g2_customer_names(prepared.current_savings, g2, phone_column="msisdn")
        prepared = MpesaPreparedData(
            transactions,
            current,
            prepared.fixed_savings,
            prepared.loans,
            prepared.load_report,
            g2,
        )

        report = build_mpesa_statement(prepared, "1001", {"CDF": None})
        matches = search_customers("1001", prepared)
        matches_by_name = search_customers("client test", prepared)

        self.assertEqual(report["extrait"].iloc[0]["Nom_client"], "CLIENT TEST")
        self.assertEqual(report["synthese"].iloc[0]["Nom_client"], "CLIENT TEST")
        self.assertEqual(report["mouvements_dat"].iloc[0]["Nom_client"], "CLIENT TEST")
        self.assertIn("CLIENT TEST", matches["Nom_client"].dropna().tolist())
        self.assertEqual(matches_by_name["customer_id"].dropna().unique().tolist(), ["1001"])

    def test_prepare_perfect_clients_normalizes_only_valid_phone_keys(self) -> None:
        result = prepare_perfect_clients(
            pd.DataFrame(
                [
                    {"id_client": 1.0, "Phone_Prefixe": "0812345678", "nom_complet": "CLIENT VALIDE"},
                    {"id_client": 2.0, "Phone_Prefixe": "A VERIFIER", "nom_complet": "CLIENT INVALIDE"},
                    {"id_client": 3.0, "Phone_Prefixe": "243701234567", "nom_complet": "PREFIXE INVALIDE"},
                ]
            )
        )

        self.assertEqual(result.loc[0, "phone_prefixe"], "243812345678")
        self.assertTrue(pd.isna(result.loc[1, "phone_prefixe"]))
        self.assertTrue(pd.isna(result.loc[2, "phone_prefixe"]))
        self.assertEqual(result.loc[1, "phone_prefixe_source"], "A VERIFIER")
        self.assertEqual(result.loc[0, "id_client"], "1")

    def test_perfect_crosscheck_aggregates_shared_phone_without_multiplying_operations(self) -> None:
        prepared = _sample_prepared_data()
        perfect = prepare_perfect_clients(
            pd.DataFrame(
                [
                    {
                        "id_client": 10,
                        "code_client": "P001",
                        "nom_complet": "NOM PERFECT A",
                        "Phone_Prefixe": "243812345678",
                        "Statut_phone": "OK",
                    },
                    {
                        "id_client": 11,
                        "code_client": "P002",
                        "nom_complet": "NOM PERFECT B",
                        "Phone_Prefixe": "0812345678",
                        "Statut_phone": "OK",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=prepared.transactions,
            current_savings=prepared.current_savings,
            fixed_savings=prepared.fixed_savings,
            loans=prepared.loans,
            load_report=prepared.load_report,
            perfect_clients=perfect,
        )

        report = build_perfect_client_crosscheck(prepared)
        summary = report["synthese"]
        operations = report["operations"]

        self.assertEqual(len(summary), 1)
        self.assertEqual(int(summary.loc[0, "nb_clients_perfect"]), 2)
        self.assertEqual(
            summary.loc[0, "statut_rapprochement_perfect"],
            "Trouve dans Perfect - plusieurs clients",
        )
        self.assertEqual(summary.loc[0, "noms_clients_perfect"], "NOM PERFECT A | NOM PERFECT B")
        self.assertEqual(len(operations), int(summary.loc[0, "nombre_operations_turbo"]))
        self.assertTrue(operations["nb_clients_perfect"].eq(2).all())
        self.assertTrue(operations["noms_clients_perfect"].eq("NOM PERFECT A | NOM PERFECT B").all())
        self.assertTrue(report["clients_trois_systemes"].empty)

    def test_perfect_crosscheck_identifies_clients_present_in_g2_turbo_and_perfect(self) -> None:
        prepared = _sample_prepared_data()
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "TX001",
                        "Completion Time": "2026-07-01 10:00:00",
                        "Opposite Party": "0812345678 - CLIENT G2",
                        "Currency": "CDF",
                        "Transaction Amount": 1000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    }
                ]
            )
        )
        perfect = prepare_perfect_clients(
            pd.DataFrame(
                [
                    {
                        "id_client": 10,
                        "code_client": "P001",
                        "nom_complet": "CLIENT PERFECT",
                        "Phone_Prefixe": "243812345678",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=prepared.transactions,
            current_savings=prepared.current_savings,
            fixed_savings=prepared.fixed_savings,
            loans=prepared.loans,
            load_report=prepared.load_report,
            g2_transactions=g2,
            perfect_clients=perfect,
        )

        report = build_perfect_client_crosscheck(prepared)
        summary = report["synthese"]
        clients_trois_systemes = report["clients_trois_systemes"]

        self.assertEqual(len(clients_trois_systemes), 1)
        self.assertTrue(bool(summary.loc[0, "present_dans_turbo"]))
        self.assertTrue(bool(summary.loc[0, "present_dans_g2"]))
        self.assertTrue(bool(summary.loc[0, "present_dans_perfect"]))
        self.assertTrue(bool(summary.loc[0, "present_dans_les_3_systemes"]))
        self.assertEqual(
            summary.loc[0, "statut_presence_systemes"],
            "Present dans G2, Solution Numérique et Perfect",
        )
        self.assertEqual(clients_trois_systemes.loc[0, "noms_clients_perfect"], "CLIENT PERFECT")
        self.assertEqual(len(report["clients_perfect_dans_mpesa"]), 1)
        self.assertEqual(len(report["clients_perfect_dans_turbo"]), 1)
        self.assertEqual(len(report["clients_perfect_dans_turbo_et_mpesa"]), 1)
        export = create_excel_export({"clients_3_systemes": clients_trois_systemes})
        exported = pd.read_excel(BytesIO(export), sheet_name="Clients_Perfect_3_Systemes")
        self.assertEqual(len(exported), 1)
        self.assertEqual(exported.loc[0, "phone_prefixe"], 243812345678)

    def test_perfect_crosscheck_builds_the_three_requested_populations(self) -> None:
        prepared = _sample_prepared_data()
        turbo_extra = prepared.transactions.iloc[[0]].copy()
        turbo_extra["id"] = "TURBO-B"
        turbo_extra["customer_id"] = "2002"
        turbo_extra["msisdn1"] = "243822222222"
        turbo_extra["ref_no"] = "TURBO-B"
        turbo_extra["reference_id"] = "TURBO-B"
        transactions = pd.concat([prepared.transactions, turbo_extra], ignore_index=True)
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "G2-A",
                        "Completion Time": "2026-07-01 11:00:00",
                        "Opposite Party": "0833333333 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 500,
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "TX001",
                        "Completion Time": "2026-07-01 10:00:00",
                        "Opposite Party": "0812345678 - CLIENT C",
                        "Currency": "CDF",
                        "Transaction Amount": 1000,
                        "Details": "BisouBisouC2B",
                    },
                ]
            )
        )
        perfect = prepare_perfect_clients(
            pd.DataFrame(
                [
                    {"id_client": 1, "nom_complet": "PERFECT A", "Phone_Prefixe": "243833333333"},
                    {"id_client": 2, "nom_complet": "PERFECT B", "Phone_Prefixe": "243822222222"},
                    {"id_client": 3, "nom_complet": "PERFECT C", "Phone_Prefixe": "243812345678"},
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=prepared.current_savings,
            fixed_savings=prepared.fixed_savings,
            loans=prepared.loans,
            load_report=prepared.load_report,
            g2_transactions=g2,
            perfect_clients=perfect,
        )

        report = build_perfect_client_crosscheck(prepared)

        self.assertEqual(
            set(report["clients_perfect_dans_mpesa"]["phone_prefixe"]),
            {"243833333333", "243812345678"},
        )
        self.assertEqual(
            set(report["clients_perfect_dans_turbo"]["phone_prefixe"]),
            {"243822222222", "243812345678"},
        )
        self.assertEqual(
            report["clients_perfect_dans_turbo_et_mpesa"]["phone_prefixe"].tolist(),
            ["243812345678"],
        )
        export = create_excel_export(
            {
                "clients_perfect_dans_mpesa": report["clients_perfect_dans_mpesa"],
                "clients_perfect_dans_turbo": report["clients_perfect_dans_turbo"],
                "clients_perfect_dans_turbo_et_mpesa": report["clients_perfect_dans_turbo_et_mpesa"],
            }
        )
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertEqual(
            workbook.sheet_names,
            ["Clients_Perfect_G2", "Clients_Perfect_Turbo", "Clients_Perfect_Turbo_G2"],
        )

    def test_perfect_crosscheck_keeps_unmatched_mpesa_client(self) -> None:
        prepared = _sample_prepared_data()
        perfect = prepare_perfect_clients(
            pd.DataFrame([{"id_client": 10, "Phone_Prefixe": "243999999999", "nom_complet": "AUTRE CLIENT"}])
        )
        prepared = MpesaPreparedData(
            transactions=prepared.transactions,
            current_savings=prepared.current_savings,
            fixed_savings=prepared.fixed_savings,
            loans=prepared.loans,
            load_report=prepared.load_report,
            perfect_clients=perfect,
        )

        summary = build_perfect_client_crosscheck(prepared)["synthese"]

        self.assertEqual(len(summary), 1)
        self.assertEqual(int(summary.loc[0, "nb_clients_perfect"]), 0)
        self.assertEqual(summary.loc[0, "statut_rapprochement_perfect"], "Non trouve dans Perfect")

    def test_validate_required_columns_detects_missing_values(self) -> None:
        missing = validate_required_columns(pd.DataFrame({"id": [1]}), TRANSACTION_REQUIRED_COLUMNS, "Transactions")

        self.assertIn("customer_id", missing)
        self.assertIn("created_at", missing)

    def test_search_customer_by_normalized_phone(self) -> None:
        prepared = _sample_prepared_data()

        result = search_customers("243812345678", prepared)

        self.assertFalse(result.empty)
        self.assertEqual(result.iloc[0]["customer_id"], "1001")

    def test_search_customer_identifies_clients_turbo_as_its_own_source(self) -> None:
        customers = prepare_customers(
            pd.DataFrame(
                [
                    {
                        "customer_id": "CLIENT-TURBO-1",
                        "msisdn1": "0811111111",
                        "created_at": "2026-07-01 08:00:00",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            customers=customers,
        )

        result = search_customers("CLIENT-TURBO-1", prepared)

        self.assertEqual(result["customer_id"].tolist(), ["CLIENT-TURBO-1"])
        self.assertEqual(result["source"].tolist(), ["Clients_Turbo"])

    def test_build_statement_reconstructs_balances_and_loans(self) -> None:
        prepared = _sample_prepared_data()

        report = build_mpesa_statement(prepared, "1001", {"CDF": 10000})
        statement = report["extrait"]

        self.assertEqual(report["mode_source_extrait"], "Solution Numérique seule")
        self.assertFalse(report["controle_g2_disponible"])
        self.assertFalse(report["nom_client_enrichi_g2"])
        self.assertTrue(report["g2_dat"].empty)
        self.assertEqual(len(statement), 2)
        self.assertIn("solde_mpesa_apres", statement.columns)
        self.assertEqual(float(statement.iloc[-1]["solde_mpesa_apres"]), 11000.0)
        self.assertIn("loan_balance", statement.columns)
        self.assertEqual(float(statement["dat_final_client"].iloc[0]), 5000.0)
        summary = report["synthese"].iloc[0]
        self.assertEqual(float(summary["total_entrees_mpesa"]), 1000.0)
        self.assertEqual(float(summary["total_sorties_mpesa"]), 2000.0)
        self.assertEqual(float(summary["mouvement_net"]), -1000.0)
        self.assertEqual(float(summary["solde_mpesa_final"]), 9000.0)

    def test_customer_summary_uses_bisou_perspective_for_every_currency(self) -> None:
        rows: list[dict[str, object]] = []
        row_id = 0
        for currency, debit, credit in [("CDF", 100.0, 30.0), ("USD", 5.0, 2.0)]:
            for suffix, dr, cr in [("ENTREE", debit, 0.0), ("SORTIE", 0.0, credit)]:
                row_id += 1
                rows.append(
                    {
                        "id": row_id,
                        "customer_id": "CLIENT-SENS",
                        "msisdn1": "0812345678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": f"{currency}-{suffix}",
                        "currency_code": currency,
                        "dr": dr,
                        "cr": cr,
                        "bal_before": 0,
                        "bal_after": abs(dr - cr),
                        "ref_no": f"{currency}-{suffix}",
                        "description": "M-Pesa Depot" if dr else "Retrait Vers M-Pesa",
                        "created_at": f"2026-07-21 {8 + row_id:02d}:00:00",
                    }
                )
        prepared = MpesaPreparedData(
            transactions=prepare_transactions(pd.DataFrame(rows)),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        summary = build_mpesa_statement(prepared, "CLIENT-SENS")["synthese"].set_index("devise")

        self.assertEqual(float(summary.loc["CDF", "total_entrees_mpesa"]), 100.0)
        self.assertEqual(float(summary.loc["CDF", "total_sorties_mpesa"]), 30.0)
        self.assertEqual(float(summary.loc["CDF", "mouvement_net"]), 70.0)
        self.assertEqual(float(summary.loc["USD", "total_entrees_mpesa"]), 5.0)
        self.assertEqual(float(summary.loc["USD", "total_sorties_mpesa"]), 2.0)
        self.assertEqual(float(summary.loc["USD", "mouvement_net"]), 3.0)

    def test_customer_transaction_analysis_reconstructs_credit_internal_dat_and_positions(self) -> None:
        prepared = _sample_customer_transaction_analysis_data()

        analysis = build_customer_transaction_analysis(prepared, "CLIENT-ANALYSE")

        self.assertEqual(len(analysis["parcours_turbo"]), 5)
        credit = analysis["credit_turbo_synthese_client"].iloc[0]
        self.assertEqual(credit["currency_code"], "CDF")
        self.assertEqual(int(credit["nombre_decaissements"]), 1)
        self.assertEqual(float(credit["montant_decaisse_client"]), 100.0)
        self.assertEqual(float(credit["dette_creee_observee"]), 110.0)
        self.assertEqual(float(credit["interet_observe"]), 10.0)
        self.assertEqual(int(credit["nombre_remboursements"]), 1)
        self.assertEqual(float(credit["principal_rembourse"]), 40.0)
        self.assertEqual(int(credit["remboursements_avec_penalite"]), 1)
        self.assertEqual(float(credit["penalite_observee"]), 5.0)

        repayments = analysis["remboursements_turbo_synthese_client"].iloc[0]
        self.assertEqual(repayments["customer_id"], "CLIENT-ANALYSE")
        self.assertEqual(int(repayments["nombre_remboursements"]), 1)
        self.assertEqual(float(repayments["montant_paye_observe"]), 40.0)
        self.assertEqual(float(repayments["principal_rembourse"]), 40.0)
        self.assertEqual(float(repayments["penalite_observee"]), 5.0)
        self.assertEqual(
            analysis["remboursements_turbo_detail_client"]["event_reference"].tolist(),
            ["TURBO-20260704080000000000"],
        )

        active_dat = analysis["dat_en_cours_client"].iloc[0]
        self.assertEqual(active_dat["customer_id"], "CLIENT-ANALYSE")
        self.assertEqual(float(active_dat["balance"]), 30.0)
        self.assertEqual(float(active_dat["taux_interet_annuel_pct"]), 11.0)
        self.assertEqual(active_dat["situation_dat_client"], "En cours")

        internal = analysis["mouvements_internes_turbo"]
        self.assertEqual(len(internal), 1)
        self.assertEqual(internal.iloc[0]["type_operation"], "Transfert DAT vers epargne courante")
        self.assertEqual(float(internal.iloc[0]["montant_operation"]), 20.0)
        self.assertEqual(internal.iloc[0]["reference_dat"], "DAT-1")
        self.assertEqual(
            pd.Timestamp(internal.iloc[0]["date_creation_dat"]),
            pd.Timestamp("2026-07-02 07:45:00"),
        )
        self.assertEqual(
            pd.Timestamp(internal.iloc[0]["date_fin_dat"]),
            pd.Timestamp("2026-07-05 08:00:00"),
        )

        positions = analysis["positions_turbo"].set_index("famille_position")
        self.assertEqual(float(positions.loc["Epargne courante", "solde_transactions_observe"]), 120.0)
        self.assertEqual(float(positions.loc["DAT", "solde_transactions_observe"]), 30.0)
        self.assertEqual(float(positions.loc["Credit", "solde_transactions_observe"]), 60.0)
        self.assertTrue(positions["statut_rapprochement_solde"].eq("Conforme").all())
        self.assertTrue(
            analysis["controles_client_turbo"]["statut_controle_turbo"].eq("Conforme").all()
        )

    def test_customer_transaction_analysis_applies_currency_type_date_and_reference_filters(self) -> None:
        prepared = _sample_customer_transaction_analysis_data()

        analysis = build_customer_transaction_analysis(
            prepared,
            "CLIENT-ANALYSE",
            currency="CDF",
            operation_types=["Sortie M-PESA_Turbo vers DAT"],
            date_start=pd.Timestamp("2026-07-02").date(),
            date_end=pd.Timestamp("2026-07-02").date(),
            reference_query="DAT-DEP-1",
        )

        self.assertEqual(len(analysis["parcours_turbo"]), 1)
        self.assertEqual(
            analysis["parcours_turbo"].iloc[0]["type_operation"],
            "Sortie M-PESA_Turbo vers DAT",
        )
        self.assertTrue(analysis["credit_turbo_detail_client"].empty)
        self.assertTrue(analysis["remboursements_turbo_detail_client"].empty)
        self.assertEqual(len(analysis["dat_en_cours_client"]), 1)
        self.assertEqual(analysis["jalons_turbo"]["nombre_operations"].tolist(), [1])
        self.assertTrue(
            analysis["positions_turbo"]["statut_rapprochement_solde"]
            .astype(str)
            .str.contains("Non comparable")
            .any()
        )

    def test_customer_statement_elements_cover_seven_turbo_focused_families(self) -> None:
        events = pd.DataFrame(
            [
                {
                    "created_at": "2026-07-21 08:00:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "DEP-1",
                    "type_operation": "Sortie M-PESA_Turbo vers epargne",
                    "montant_operation": 100.0,
                    "remboursement_compte_ouvert": 0.0,
                    "remboursement_mpesa": 0.0,
                },
                {
                    "created_at": "2026-07-21 08:30:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "DAT-DEP-1",
                    "type_operation": "Sortie M-PESA_Turbo vers DAT",
                    "montant_operation": 80.0,
                    "remboursement_compte_ouvert": 0.0,
                    "remboursement_mpesa": 0.0,
                },
                {
                    "created_at": "2026-07-21 09:00:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "RET-1",
                    "type_operation": "Entree M-PESA_Turbo depuis epargne",
                    "montant_operation": 40.0,
                    "remboursement_compte_ouvert": 0.0,
                    "remboursement_mpesa": 0.0,
                },
                {
                    "created_at": "2026-07-21 10:00:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "REM-MPESA",
                    "type_operation": "Remboursement de credit",
                    "montant_operation": 25.0,
                    "remboursement_compte_ouvert": 0.0,
                    "remboursement_mpesa": 25.0,
                },
                {
                    "created_at": "2026-07-21 11:00:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "REM-OUVERT",
                    "type_operation": "Remboursement de credit",
                    "montant_operation": 30.0,
                    "remboursement_compte_ouvert": 30.0,
                    "remboursement_mpesa": 0.0,
                },
                {
                    "created_at": "2026-07-21 12:00:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "event_reference": "DAT-RETOUR",
                    "type_operation": "Transfert DAT vers epargne courante",
                    "montant_operation": 200.0,
                    "remboursement_compte_ouvert": 0.0,
                    "remboursement_mpesa": 0.0,
                },
            ]
        )
        interests = pd.DataFrame(
            [
                {
                    "maturity_date": "2026-07-21",
                    "date_ecriture_turbo": "2026-07-21 12:01:00",
                    "customer_id": "CLIENT-1",
                    "currency_code": "CDF",
                    "savings_id": "DAT-1",
                    "reference_transaction_turbo": "DAT-RETOUR",
                    "interet_client_constate": 11.0,
                }
            ]
        )

        result = build_customer_statement_elements(events, interests)
        detail = result["detail"].set_index("type_element_extrait")
        summary = result["synthese"]

        self.assertEqual(len(detail), 7)
        self.assertEqual(len(summary), 7)
        self.assertEqual(
            float(detail.loc["Dépôt à terme (DAT)", "montant_observe"]),
            80.0,
        )
        self.assertEqual(
            detail.loc[
                "Remboursement d'un credit depuis le compte ouvert",
                "origine_operation",
            ],
            "Compte ouvert",
        )
        self.assertEqual(
            float(
                detail.loc[
                    "Remboursement d'un credit depuis le compte ouvert",
                    "montant_observe",
                ]
            ),
            30.0,
        )
        self.assertEqual(
            detail.loc[
                "Entree des interets du capital mis en DAT",
                "source_turbo",
            ],
            "Savings Account - interest_earned",
        )

    def test_customer_statement_exports_move_dat_deposit_out_of_transaction_detail(self) -> None:
        from docx import Document
        from pypdf import PdfReader

        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": "37370",
                        "msisdn1": "243000000000",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "DAT_TEST_001",
                        "currency_code": "USD",
                        "dr": 10.0,
                        "cr": 0.0,
                        "bal_before": 0.0,
                        "bal_after": 10.0,
                        "ref_no": "DAT_MOUVEMENT_TEST_001",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-18 14:52:41",
                    },
                    {
                        "id": 2,
                        "customer_id": "37370",
                        "msisdn1": "243000000000",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "DAT_TEST_001",
                        "currency_code": "USD",
                        "dr": 0.0,
                        "cr": 10.0,
                        "bal_before": 0.0,
                        "bal_after": 10.0,
                        "ref_no": "DAT_MOUVEMENT_TEST_001",
                        "description": "Depot Bloque",
                        "created_at": "2026-07-18 14:52:41",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )
        report = build_mpesa_statement(
            prepared,
            "37370",
            date_start="2026-07-18",
            date_end="2026-07-18",
        )
        dat_summary = report["elements_extrait_client_synthese"].loc[
            report["elements_extrait_client_synthese"]["type_element_extrait"].eq(
                "Dépôt à terme (DAT)"
            )
        ].iloc[0]
        self.assertEqual(int(dat_summary["nombre_operations"]), 1)
        self.assertEqual(float(dat_summary["montant_total_observe"]), 10.0)

        view = build_customer_statement_view(report["extrait"])
        self.assertEqual(float(view["total_entries"]), 10.0)

        export_kwargs = {
            "analysis_report": report,
            "customer_id": "37370",
            "customer_name": "MUPANZI KITSHI BENJAMIN",
            "telephone": "243000000000",
            "currency": "USD",
            "period_start": "2026-07-18",
            "period_end": "2026-07-18",
        }
        word = create_customer_statement_word(report["extrait"], **export_kwargs)
        document = Document(BytesIO(word))
        word_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
        )
        self.assertIn("Dépôt à terme (DAT)", word_text)
        transaction_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "Date",
                "Compte",
                "Référence",
                "Devise",
                "Description",
                "Entrées",
                "Sorties",
                "Solde",
            ]
        )
        self.assertEqual(len(transaction_table.rows), 1)

        pdf = create_customer_statement_pdf(report["extrait"], **export_kwargs)
        pdf_text = "\n".join(
            page.extract_text() or "" for page in PdfReader(BytesIO(pdf)).pages
        )
        self.assertIn("Dépôt à terme (DAT)", pdf_text)
        self.assertNotIn("DAT_MOUVEMENT_TEST_001", pdf_text)

    def test_customer_bisou_statement_matches_benjamin_scenario(self) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from pypdf import PdfReader

        rows: list[dict[str, object]] = []

        def add(
            created_at: str,
            account_type: str,
            *,
            dr: float = 0.0,
            cr: float = 0.0,
            bal_before: float = 0.0,
            bal_after: float = 0.0,
            ref_no: str = "",
            reference_id: str = "",
            description: str,
        ) -> None:
            rows.append(
                {
                    "id": len(rows) + 1,
                    "customer_id": "37370",
                    "msisdn1": "243000000000",
                    "account_type": account_type,
                    "reference_id": reference_id,
                    "currency_code": "USD",
                    "dr": dr,
                    "cr": cr,
                    "bal_before": bal_before,
                    "bal_after": bal_after,
                    "ref_no": ref_no,
                    "description": description,
                    "created_at": created_at,
                }
            )

        add(
            "2026-07-18 14:52:41",
            "MPESA ACCOUNT",
            dr=10.0,
            bal_after=10.0,
            ref_no="DAT_MOUVEMENT_TEST_001",
            reference_id="DAT_TEST_001",
            description="M-Pesa Compte",
        )
        add(
            "2026-07-18 14:52:41",
            "FIXED SAVINGS",
            cr=10.0,
            bal_after=10.0,
            ref_no="DAT_MOUVEMENT_TEST_001",
            reference_id="DAT_TEST_001",
            description="Depot Bloque",
        )
        add(
            "2026-07-22 16:17:16",
            "PRINCIPLE",
            dr=5.0,
            bal_after=5.0,
            reference_id="PRET_TEST_001",
            description="Montant principal",
        )
        add(
            "2026-07-22 16:17:16",
            "LOAN ACCOUNT",
            cr=5.35,
            bal_after=5.35,
            reference_id="PRET_TEST_001",
            description="Compte de pret",
        )
        add(
            "2026-07-22 16:17:16",
            "MPESA ACCOUNT",
            cr=5.0,
            bal_after=5.0,
            reference_id="PRET_TEST_001",
            description="Montant pret",
        )
        add(
            "2026-07-22 16:17:16",
            "MPESA ACCOUNT",
            dr=0.35,
            bal_after=0.35,
            reference_id="PRET_TEST_001",
            description="Compte du M-Pesa",
        )
        add(
            "2026-07-22 16:24:50",
            "MPESA ACCOUNT",
            dr=5.0,
            bal_after=5.0,
            ref_no="DEPOT_TEST_001",
            description="M-Pesa Depot",
        )
        add(
            "2026-07-22 16:24:50",
            "NORMAL SAVINGS",
            cr=5.0,
            bal_after=5.0,
            ref_no="DEPOT_TEST_001",
            description="Epargne depot",
        )
        for account_type, dr, cr, before, after, description in [
            ("PRINCIPLE", 0.0, 5.0, 5.0, 0.0, "Remboursement du principal"),
            ("LOAN ACCOUNT", 5.0, 0.0, 5.0, 0.0, "Remboursement du pret"),
            ("MPESA ACCOUNT", 5.0, 0.0, 0.0, 5.0, "Remboursement du M-Pesa"),
            ("NORMAL SAVINGS", 5.0, 0.0, 5.0, 0.0, "Remboursement de compte epargne"),
            # Ligne technique présente dans le scénario réel : elle ne doit
            # pas réduire le compte bloqué, car le paiement vient du compte ouvert.
            ("FIXED SAVINGS", 5.0, 0.0, 5.0, 0.0, "Remboursement de compte epargne"),
        ]:
            add(
                "2026-07-22 16:26:43",
                account_type,
                dr=dr,
                cr=cr,
                bal_before=before,
                bal_after=after,
                ref_no="RETRAIT_TEST_001",
                reference_id=(
                    "PRET_TEST_001"
                    if account_type in {"PRINCIPLE", "LOAN ACCOUNT", "MPESA ACCOUNT"}
                    else ""
                ),
                description=description,
            )

        savings_accounts = pd.DataFrame(
            [
                {
                    "savings_id": "SATXW2I7Y1",
                    "customer_id": "37370",
                    "msisdn1": "243000000000",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "currency_code": "CDF",
                    "balance": 0.0,
                    "status": "active",
                    "created_at": "2026-07-10 07:26:33",
                    "updated_at": "2026-07-10 07:26:33",
                },
                {
                    "savings_id": "SADW5C1Z50",
                    "customer_id": "37370",
                    "msisdn1": "243000000000",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "currency_code": "USD",
                    "balance": 0.0,
                    "status": "active",
                    "created_at": "2026-07-10 07:26:33",
                    "updated_at": "2026-07-22 16:26:43",
                },
                {
                    "savings_id": "DAT_TEST_001",
                    "customer_id": "37370",
                    "msisdn1": "243000000000",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "currency_code": "USD",
                    "balance": 10.0,
                    "status": "active",
                    "date_approved": "2026-07-18 14:52:41",
                    "date_activated": "2026-07-18 14:52:41",
                    "maturity_date": "2026-08-18 14:52:41",
                    "interest_earned": 0.0,
                    "created_at": "2026-07-18 14:52:41",
                    "updated_at": "2026-07-18 14:52:41",
                },
            ]
        )
        prepared = MpesaPreparedData(
            transactions=prepare_transactions(pd.DataFrame(rows)),
            current_savings=prepare_current_savings(savings_accounts),
            fixed_savings=prepare_fixed_savings_from_accounts(savings_accounts),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )
        report = build_mpesa_statement(
            prepared,
            "37370",
            {"USD": 13.34},
        )
        situation = report["synthese"].iloc[0]

        self.assertAlmostEqual(float(situation["solde_mpesa_client_final"]), 2.99)
        self.assertEqual(float(situation["total_entrees_mpesa"]), 15.0)
        self.assertEqual(float(situation["total_sorties_mpesa"]), 4.65)
        self.assertEqual(float(situation["mouvement_net"]), 10.35)
        self.assertEqual(float(situation["epargne_courante_finale"]), 0.0)
        self.assertEqual(float(situation["dat_final"]), 10.0)
        self.assertIn("Savings Account", situation["source_epargne_courante_finale"])
        self.assertIn("Savings Account", situation["source_dat_final"])

        analysis = build_customer_transaction_analysis(prepared, "37370")
        active_dat = analysis["dat_en_cours_client"].iloc[0]
        self.assertEqual(active_dat["savings_id"], "DAT_TEST_001")
        self.assertEqual(float(active_dat["balance"]), 10.0)
        self.assertAlmostEqual(float(active_dat["interet_estime_echeance"]), 0.09, places=2)
        self.assertEqual(active_dat["situation_dat_client"], "En cours")
        positions = analysis["positions_turbo"].set_index(
            ["famille_position", "currency_code"]
        )
        self.assertEqual(
            float(
                positions.loc[
                    ("Epargne courante", "USD"),
                    "solde_transactions_observe",
                ]
            ),
            0.0,
        )
        self.assertEqual(
            float(positions.loc[("DAT", "USD"), "solde_transactions_observe"]),
            10.0,
        )

        export_report = dict(report)
        export_report.update(analysis)
        export_kwargs = {
            "analysis_report": export_report,
            "customer_id": "37370",
            "customer_name": "MUPANZI KITSHI BENJAMIN",
            "telephone": "243000000000",
            "currency": "USD",
            "period_start": "2026-07-18",
            "period_end": "2026-07-22",
        }
        document = Document(
            BytesIO(create_customer_statement_word(report["extrait"], **export_kwargs))
        )
        word_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
        )
        self.assertNotIn("Situation financière actuelle du client", word_text)
        self.assertNotIn("Solde M-PESA | Compte ouvert", word_text)
        self.assertIn("Synthèse financière par devise", word_text)
        self.assertNotIn("Synthèse des flux", word_text)
        self.assertNotIn("Situation de l'épargne", word_text)
        self.assertNotIn("Synthèse des flux — point de vue Bisou Bisou", word_text)
        self.assertNotIn(
            "Les entrées, les sorties, le flux net et le cumul sont présentés du point de vue de Bisou Bisou.",
            word_text,
        )
        self.assertEqual(document.sections[0].orientation, WD_ORIENT.PORTRAIT)
        self.assertIn("DAT_TEST_001", word_text)

        financial_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "Devise",
                "Ouverture",
                "Entrée",
                "Sorties",
                "Cloture",
                "Compte ouvert",
                "Compte bloqué",
            ]
        )
        self.assertEqual(
            [cell.text for cell in financial_table.rows[1].cells],
            ["USD", "0.00", "5.00", "5.00", "0.00", "0.00", "10.00"],
        )
        self.assertEqual(financial_table.style.name, "Table Grid")
        self.assertFalse(
            any(
                bool(run.bold)
                for cell in financial_table.rows[1].cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
            )
        )
        criteria_table = document.tables[0].cell(1, 1).tables[0]
        criteria_text = "\n".join(
            cell.text for row in criteria_table.rows for cell in row.cells
        )
        self.assertIn("Taux annuel DAT :", criteria_text)
        self.assertIn("11,0 %", criteria_text)
        transaction_table = next(
            table
            for table in document.tables
            if table.rows
            and [cell.text for cell in table.rows[0].cells[:2]] == ["Date", "Compte"]
        )
        self.assertEqual(transaction_table.rows[0].cells[2].text, "Référence")
        self.assertEqual(
            transaction_table.rows[0].cells[-1].text,
            "Solde",
        )
        transaction_table_text = "\n".join(
            cell.text for row in transaction_table.rows for cell in row.cells
        )
        self.assertIn("DEPOT_TEST_001", transaction_table_text)
        self.assertIn("RETRAIT_TEST_001", transaction_table_text)
        self.assertNotIn("PRET_TEST_001", transaction_table_text)
        self.assertNotIn("DAT_MOUVEMENT_TEST_001", transaction_table_text)

        pdf_reader = PdfReader(
            BytesIO(create_customer_statement_pdf(report["extrait"], **export_kwargs))
        )
        self.assertLess(
            float(pdf_reader.pages[0].mediabox.width),
            float(pdf_reader.pages[0].mediabox.height),
        )
        pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
        self.assertNotIn("Situation financière actuelle du client", pdf_text)
        self.assertIn("Synthèse financière par devise", pdf_text)
        self.assertNotIn("Synthèse des flux", pdf_text)
        self.assertNotIn("Synthèse des flux — point de vue Bisou Bisou", pdf_text)
        self.assertIn("10.00", pdf_text)
        self.assertNotIn("Situation de l'épargne", pdf_text)
        self.assertIn("Compte ouvert", pdf_text)
        self.assertIn("Compte bloqué", pdf_text)
        self.assertIn("Ouverture", pdf_text)
        self.assertIn("Cloture", pdf_text)
        self.assertIn("Taux annuel DAT", pdf_text)
        self.assertIn("Référence", pdf_text)
        self.assertIn("DAT_TEST_001", pdf_text)
        self.assertNotIn("sont présentés du point de vue de Bisou Bisou", pdf_text)
        minimal_pdf_reader = PdfReader(
            BytesIO(create_customer_statement_pdf(report["extrait"], **export_kwargs, minimal=True))
        )
        minimal_pdf_text = "\n".join(
            page.extract_text() or "" for page in minimal_pdf_reader.pages
        )
        self.assertIn("Extrait minimal de compte", minimal_pdf_text)
        self.assertIn("Synthèse financière par devise", minimal_pdf_text)
        self.assertIn("Ouverture", minimal_pdf_text)
        self.assertIn("Cloture", minimal_pdf_text)
        self.assertIn("Détail des transactions", minimal_pdf_text)
        self.assertNotIn("Éléments couverts par l'extrait client", minimal_pdf_text)
        self.assertNotIn("DAT en cours", minimal_pdf_text)
        self.assertNotIn("Remboursements observés", minimal_pdf_text)

        client_view = build_customer_client_statement_view(
            report["extrait"],
            analysis_report=export_report,
            customer_id="37370",
            currency="USD",
        )
        client_summary = client_view["summary"].iloc[0]
        self.assertEqual(float(client_summary["compte_ouvert"]), 0.0)
        self.assertEqual(float(client_summary["compte_bloque"]), 10.0)
        self.assertEqual(float(client_summary["position_epargne_finale"]), 10.0)
        self.assertEqual(float(client_summary["pret_net_recu"]), 4.65)
        self.assertEqual(float(client_summary["frais_interets_credit"]), 0.35)
        self.assertEqual(float(client_summary["remboursements_observes"]), 5.0)
        client_detail = client_view["detail"]
        self.assertEqual(float(client_detail.iloc[-1]["position_epargne"]), 0.0)
        self.assertIn("Dépôt compte ouvert", set(client_detail["operation"]))
        self.assertIn("Remboursement depuis compte ouvert", set(client_detail["operation"]))
        self.assertNotIn("Dépôt DAT / compte bloqué", set(client_detail["operation"]))

        client_word_complete = create_customer_client_statement_word(
            report["extrait"],
            **export_kwargs,
        )
        client_word_complete_document = Document(BytesIO(client_word_complete))
        client_word_complete_text = "\n".join(
            [paragraph.text for paragraph in client_word_complete_document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in client_word_complete_document.tables
                for row in table.rows
            ]
        )
        self.assertIn("DAT en cours", client_word_complete_text)
        self.assertIn("Jours restants", client_word_complete_text)
        self.assertIn("Extrait de compte", client_word_complete_text)
        self.assertIn("Synthèse financière par devise", client_word_complete_text)
        self.assertNotIn("vue client", client_word_complete_text)
        self.assertNotIn("Synthèse de la position client", client_word_complete_text)

        client_pdf_complete_reader = PdfReader(
            BytesIO(create_customer_client_statement_pdf(report["extrait"], **export_kwargs))
        )
        client_pdf_complete_text = "\n".join(
            page.extract_text() or "" for page in client_pdf_complete_reader.pages
        )
        self.assertIn("DAT en cours", client_pdf_complete_text)
        self.assertIn("Jours restants", client_pdf_complete_text)
        self.assertIn("Extrait de compte", client_pdf_complete_text)
        self.assertIn("Synthèse financière par devise", client_pdf_complete_text)
        self.assertNotIn("vue client", client_pdf_complete_text)
        self.assertNotIn("Synthèse de la position client", client_pdf_complete_text)

        client_word = create_customer_client_statement_word(
            report["extrait"],
            **export_kwargs,
            minimal=True,
        )
        client_word_document = Document(BytesIO(client_word))
        client_word_text = "\n".join(
            [paragraph.text for paragraph in client_word_document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in client_word_document.tables
                for row in table.rows
            ]
        )
        self.assertIn("Extrait minimal de compte", client_word_text)
        self.assertIn("Position épargne", client_word_text)
        self.assertIn("10.00", client_word_text)
        self.assertIn("0.35", client_word_text)
        self.assertNotIn("vue client", client_word_text)
        self.assertNotIn("Synthèse de la position client", client_word_text)
        self.assertNotIn("Cumul net des flux", client_word_text)

        client_pdf_reader = PdfReader(
            BytesIO(create_customer_client_statement_pdf(report["extrait"], **export_kwargs, minimal=True))
        )
        client_pdf_text = "\n".join(page.extract_text() or "" for page in client_pdf_reader.pages)
        self.assertIn("Extrait minimal de compte", client_pdf_text)
        self.assertIn("Position épargne", client_pdf_text)
        self.assertIn("10.00", client_pdf_text)
        self.assertIn("0.35", client_pdf_text)
        self.assertNotIn("vue client", client_pdf_text)
        self.assertNotIn("Synthèse de la position client", client_pdf_text)
        self.assertNotIn("Cumul net des flux", client_pdf_text)

    def test_customer_statement_uses_g2_only_for_name_and_selected_customer_control(self) -> None:
        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": 1001,
                        "msisdn1": "0811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-1001",
                        "currency_code": "CDF",
                        "dr": 100,
                        "cr": 0,
                        "bal_before": 500,
                        "bal_after": 400,
                        "ref_no": "G2-1001",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-15 08:00:00",
                    },
                    {
                        "id": 2,
                        "customer_id": 2002,
                        "msisdn1": "0822222222",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-2002",
                        "currency_code": "CDF",
                        "dr": 200,
                        "cr": 0,
                        "bal_before": 800,
                        "bal_after": 600,
                        "ref_no": "G2-2002",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-15 09:00:00",
                    },
                ]
            )
        )
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "G2-1001",
                        "Initiation Time": "2026-07-15 08:00:00",
                        "Completion Time": "2026-07-15 08:00:10",
                        "Details": "BisouBisouC2B",
                        "Transaction Status": "Completed",
                        "Currency": "CDF",
                        "Paid In": 999,
                        "Opposite Party": "0811111111 - CLIENT UN",
                    },
                    {
                        "Receipt No.": "G2-2002",
                        "Initiation Time": "2026-07-15 09:00:00",
                        "Completion Time": "2026-07-15 09:00:10",
                        "Details": "BisouBisouC2B",
                        "Transaction Status": "Completed",
                        "Currency": "CDF",
                        "Paid In": 200,
                        "Opposite Party": "0822222222 - CLIENT DEUX",
                    },
                ]
            )
        )
        transactions = enrich_transactions_with_g2_customer_names(transactions, g2)
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_mpesa_statement(prepared, "1001", {"CDF": None})

        self.assertEqual(report["mode_source_extrait"], "Solution Numérique + verification G2")
        self.assertTrue(report["controle_g2_disponible"])
        self.assertTrue(report["nom_client_enrichi_g2"])
        self.assertEqual(report["extrait"]["Nom_client"].dropna().unique().tolist(), ["CLIENT UN"])
        self.assertEqual(report["g2_dat"]["receipt_no"].tolist(), ["G2-1001"])
        self.assertEqual(float(report["extrait"].iloc[0]["debit_mpesa"]), 100.0)
        self.assertEqual(
            float(report["g2_dat"].iloc[0]["transaction_amount_numeric"]),
            999.0,
        )
        self.assertEqual(report["g2_dat"].iloc[0]["controle_montant"], "Ecart")
        official_view = build_customer_statement_view(report["extrait"], account_number="1441")
        official_description = official_view["transactions"].iloc[0]["description"]
        self.assertIn("M-Pesa Compte", official_description)
        self.assertIn("CLIENT UN", official_description)
        self.assertNotIn("BisouBisouC2B", official_description)

    def test_customer_statement_view_matches_the_short_statement_contract(self) -> None:
        prepared = _sample_prepared_data()
        report = build_mpesa_statement(prepared, "1001", {"CDF": 10000})
        statement = report["extrait"].copy()
        statement["Nom_client"] = "CLIENT TEST"

        view = build_customer_statement_view(statement, account_number="1441")

        self.assertEqual(list(view["transactions"].columns), CUSTOMER_STATEMENT_COLUMNS)
        self.assertEqual(view["currency"], "CDF")
        self.assertTrue(view["balance_is_real"])
        self.assertEqual(view["balance_label"], "Solde")
        self.assertEqual(float(view["opening_amount"]), 10000.0)
        self.assertEqual(float(view["total_entries"]), 1000.0)
        self.assertEqual(float(view["total_outputs"]), 2000.0)
        self.assertEqual(float(view["flow_net"]), -1000.0)
        self.assertEqual(float(view["closing_amount"]), 9000.0)
        self.assertTrue(view["transactions"]["compte"].eq("1441").all())
        self.assertTrue(view["transactions"]["devise"].eq("CDF").all())
        first_description = view["transactions"].iloc[0]["description"]
        self.assertIn("M-Pesa Compte", first_description)
        self.assertIn("Depot Bloque", first_description)
        self.assertIn("CLIENT TEST", first_description)

        relative_report = build_mpesa_statement(prepared, "1001", {"CDF": None})
        relative_view = build_customer_statement_view(relative_report["extrait"], account_number="1441")
        self.assertFalse(relative_view["balance_is_real"])
        self.assertEqual(relative_view["balance_label"], "Cumul net")
        self.assertEqual(float(relative_view["opening_amount"]), 0.0)
        self.assertEqual(float(relative_view["closing_amount"]), -1000.0)
        self.assertEqual(float(relative_view["flow_net"]), -1000.0)

    def test_seven_percent_loan_keeps_gross_interest_and_net_separate(self) -> None:
        prepared = _sample_seven_percent_loan_data()

        journal = build_turbo_operation_events(prepared.transactions)
        self.assertEqual(len(journal["events"]), 2)
        origination_event = journal["events"].loc[
            journal["events"]["type_operation"].eq("Decaissement de credit")
        ].iloc[0]
        repayment_event = journal["events"].loc[
            journal["events"]["type_operation"].eq("Remboursement de credit")
        ].iloc[0]
        self.assertEqual(int(origination_event["nombre_lignes"]), 12)
        self.assertEqual(int(repayment_event["nombre_lignes"]), 6)
        self.assertEqual(
            int(prepared.transactions["reference_id"].eq("PRET_TEST_001").sum()),
            16,
        )
        self.assertEqual(float(origination_event["pret_brut_decaisse"]), 5.0)
        self.assertEqual(float(origination_event["interet_pret_preleve"]), 0.35)
        self.assertAlmostEqual(float(origination_event["taux_interet_pret_pct"]), 7.0)
        self.assertEqual(float(origination_event["net_pret_verse"]), 4.65)
        self.assertAlmostEqual(float(origination_event["ecart_interet_pret_vs_7pct"]), 0.0)

        statement = build_mpesa_statement(prepared, "37370")["extrait"]
        loan_statement = statement.loc[
            statement["type_operation"].eq("Decaissement de credit")
        ].copy()
        self.assertEqual(len(loan_statement), 1)
        self.assertEqual(float(loan_statement.iloc[0]["pret_brut_decaisse"]), 5.0)
        self.assertEqual(float(loan_statement.iloc[0]["interet_pret_preleve"]), 0.35)
        self.assertAlmostEqual(float(loan_statement.iloc[0]["taux_interet_pret_pct"]), 7.0)
        self.assertEqual(float(loan_statement.iloc[0]["net_pret_verse"]), 4.65)

        view = build_customer_statement_view(
            loan_statement,
            entry_account_number="1441",
            output_account_number="15558",
        )
        self.assertEqual(float(view["total_outputs"]), 4.65)
        self.assertEqual(float(view["flow_net"]), -4.65)
        self.assertEqual(float(view["transactions"].iloc[0]["sortie"]), 4.65)
        self.assertIn(
            "Prêt brut : 5,00 USD — intérêt prélevé : 0,35 USD (7 %) — net versé : 4,65 USD",
            view["transactions"].iloc[0]["description"],
        )

        from docx import Document
        from pypdf import PdfReader

        export_kwargs = {
            "customer_id": "37370",
            "customer_name": "MUPANZI KITSHI BENJAMIN",
            "telephone": "243000000000",
            "currency": "USD",
            "entry_account_number": "1441",
            "output_account_number": "15558",
            "period_start": "2026-07-22",
            "period_end": "2026-07-22",
            "generated_at": pd.Timestamp("2026-07-22 18:00:00"),
        }
        word = create_customer_statement_word(loan_statement, **export_kwargs)
        document = Document(BytesIO(word))
        word_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertNotIn("Prêt brut : 5,00 USD", word_text)
        self.assertNotIn("net versé : 4,65 USD", word_text)

        pdf = create_customer_statement_pdf(loan_statement, **export_kwargs)
        pdf_text = "\n".join(
            page.extract_text() or "" for page in PdfReader(BytesIO(pdf)).pages
        )
        self.assertNotIn("Prêt brut : 5,00 USD", pdf_text)
        self.assertNotIn("net versé : 4,65", pdf_text)

        finance = build_mpesa_turbo_financial_analysis(
            prepared,
            date_start="2026-07-22",
            date_end="2026-07-22",
            turbo_events=journal["events"],
            turbo_transaction_lines=journal["lines"],
        )
        self.assertEqual(
            float(finance["nouveaux_credits_synthese"].iloc[0]["montant_decaisse_turbo"]),
            5.0,
        )
        accounting = build_mpesa_accounting_analysis(
            prepared,
            date_start="2026-07-22",
            date_end="2026-07-22",
        )
        interest_product = accounting["produits_financiers"].loc[
            accounting["produits_financiers"]["account_type"].eq("INTEREST EARNED")
        ].iloc[0]
        self.assertEqual(float(interest_product["montant_observe"]), 0.35)
        interest_detail = accounting["produits_financiers_detail"].loc[
            accounting["produits_financiers_detail"]["account_type"].eq("INTEREST EARNED")
        ].iloc[0]
        self.assertEqual(interest_detail["reference_id"], "PRET_TEST_001")
        self.assertEqual(float(interest_detail["montant_observe"]), 0.35)

    def test_customer_statement_filename_uses_turbo_identity_and_optional_g2_name(self) -> None:
        turbo_only = build_customer_statement_filename(
            customer_id="37335",
            customer_name="CE NOM DOIT ETRE IGNORE",
            telephone="243827972206",
            currency="USD",
            period_start=pd.Timestamp("2026-07-09"),
            period_end=pd.Timestamp("2026-07-15"),
            g2_available=False,
        )
        turbo_with_g2 = build_customer_statement_filename(
            customer_id="37335",
            customer_name="ELIANE LUAMBA MULEMBO",
            telephone="243827972206",
            currency="USD",
            period_start=pd.Timestamp("2026-07-09"),
            period_end=pd.Timestamp("2026-07-15"),
            g2_available=True,
        )

        self.assertEqual(
            turbo_only,
            "extrait_compte_37335_243827972206_USD_20260709_20260715.docx",
        )
        self.assertEqual(
            turbo_with_g2,
            "extrait_compte_37335_ELIANE LUAMBA MULEMBO_243827972206_USD_20260709_20260715.docx",
        )

    def test_turbo_withdrawals_are_unique_default_outputs_in_customer_statement(self) -> None:
        amounts = [20.0, 30.0, 35.0, 6.0, 50.0, 10.0, 40.0]
        dates = pd.to_datetime(
            [
                "2026-07-09 23:47:49",
                "2026-07-11 14:03:57",
                "2026-07-12 18:00:12",
                "2026-07-13 07:49:05",
                "2026-07-14 14:42:18",
                "2026-07-14 16:29:23",
                "2026-07-15 08:48:11",
            ]
        )
        rows: list[dict[str, object]] = []
        for index, (amount, created_at) in enumerate(zip(amounts, dates, strict=True), start=1):
            rows.extend(
                [
                    {
                        "id": index * 2 - 1,
                        "customer_id": "37301",
                        "msisdn1": "243814256725",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA8G57RHR4",
                        "currency_code": "USD",
                        "dr": 0.0,
                        "cr": amount,
                        "bal_before": 0.0,
                        "bal_after": amount,
                        "ref_no": "",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": created_at,
                    },
                    {
                        "id": index * 2,
                        "customer_id": "37301",
                        "msisdn1": "243814256725",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA8G57RHR4",
                        "currency_code": "USD",
                        "dr": amount,
                        "cr": 0.0,
                        "bal_before": amount,
                        "bal_after": 0.0,
                        "ref_no": "",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": created_at,
                    },
                ]
            )
        transactions = prepare_transactions(pd.DataFrame(rows))
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        statement = build_mpesa_statement(prepared, "37301")["extrait"]
        focused = statement.loc[
            statement["type_operation"].isin(CUSTOMER_STATEMENT_FOCUS_OPERATION_TYPES)
        ].copy()
        view = build_customer_statement_view(
            focused,
            entry_account_number="1441",
            output_account_number="15558",
        )

        self.assertEqual(len(focused), 7)
        self.assertTrue(focused["type_operation"].eq("Entree M-PESA_Turbo depuis epargne").all())
        self.assertEqual(float(view["total_entries"]), 0.0)
        self.assertEqual(float(view["total_outputs"]), 191.0)
        self.assertTrue(view["transactions"]["compte"].eq("15558").all())
        self.assertEqual(len(view["transactions"]), 7)

    def test_customer_statement_word_is_editable_filtered_and_single_currency(self) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT

        prepared = _sample_prepared_data()
        report = build_mpesa_statement(prepared, "1001", {"CDF": 10000})
        statement = report["extrait"].copy()
        statement["Nom_client"] = "CLIENT TEST"

        content = create_customer_statement_word(
            statement,
            customer_id="1001",
            customer_name="CLIENT TEST",
            telephone="243812345678",
            currency="CDF",
            entry_account_number="1441",
            output_account_number="15558",
            period_start=pd.Timestamp("2026-07-01"),
            period_end=pd.Timestamp("2026-07-02"),
            generated_at=pd.Timestamp("2026-07-15 10:30:00"),
        )

        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        statement_tables = [
            table
            for table in document.tables
            if table.rows and table.rows[0].cells[0].text == "Date"
        ]
        self.assertTrue(content.startswith(b"PK"))
        self.assertIn("Extrait de compte - 243812345678 - CLIENT TEST - CDF", text)
        self.assertNotIn("[Turbo]", text)
        self.assertEqual(len(statement_tables), 1)
        self.assertEqual(
            [cell.text for cell in statement_tables[0].rows[0].cells],
            [
                "Date",
                "Compte",
                "Référence",
                "Devise",
                "Description",
                "Entrées",
                "Sorties",
                "Solde",
            ],
        )
        self.assertEqual(len(statement_tables[0].rows), 1)
        self.assertNotIn(
            "TX002",
            [row.cells[2].text for row in statement_tables[0].rows[1:]],
        )
        self.assertNotIn(
            "Montant pret",
            "\n".join(cell.text for row in statement_tables[0].rows for cell in row.cells),
        )
        self.assertNotIn(
            "TX001",
            [row.cells[2].text for row in statement_tables[0].rows[1:]],
        )
        criteria_table = document.tables[0].cell(1, 1).tables[0]
        criteria_labels = [row.cells[0].text for row in criteria_table.rows]
        self.assertIn("Devise :", criteria_labels)
        self.assertNotIn("Compte :", criteria_labels)
        self.assertEqual(document.sections[0].orientation, WD_ORIENT.PORTRAIT)
        self.assertGreaterEqual(len(document.inline_shapes), 1)

        relative_statement = build_mpesa_statement(prepared, "1001", {"CDF": None})["extrait"]
        relative_content = create_customer_statement_word(
            relative_statement,
            customer_id="1001",
            customer_name="NON DISPONIBLE",
            telephone="243812345678",
            currency="CDF",
            entry_account_number="1441",
            output_account_number="15558",
        )
        relative_document = Document(BytesIO(relative_content))
        relative_tables = [
            table
            for table in relative_document.tables
            if table.rows and table.rows[0].cells[0].text == "Date"
        ]
        relative_text = "\n".join(paragraph.text for paragraph in relative_document.paragraphs)
        self.assertEqual(
            relative_tables[0].rows[0].cells[-1].text,
            "Solde",
        )
        self.assertNotIn("le solde d'ouverture n'a pas ete fourni", relative_text)
        self.assertIn("Extrait de compte - 243812345678 - CDF", relative_text)
        self.assertNotIn("Extrait de compte - 243812345678 - NON DISPONIBLE - CDF", relative_text)

    def test_customer_matured_dat_interest_entries_are_separate_and_traceable(self) -> None:
        fixed = prepare_savings_accounts(
            pd.DataFrame(
                [
                    {
                        "savings_id": "FA-ECHU-SANS-ECRITURE",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "product_name": "1 Month",
                        "product_description": "1 Month Fixed Account",
                        "currency_code": "USD",
                        "balance": 0,
                        "status": "Withdrawal",
                        "date_approved": "2026-06-16 21:53:03",
                        "maturity_date": "2026-07-16 21:53:03",
                        "interest_earned": "22.60",
                        "voda_interest": "6.16",
                        "created_at": "2026-06-16 21:53:03",
                        "updated_at": "2026-07-16 00:00:02",
                    },
                    {
                        "savings_id": "FA-ECHU-TRACE",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "product_name": "1 Month",
                        "product_description": "1 Month Fixed Account",
                        "currency_code": "USD",
                        "balance": 0,
                        "status": "Withdrawal",
                        "date_approved": "2026-06-16 08:00:00",
                        "maturity_date": "2026-07-16 08:00:00",
                        "interest_earned": 10.0,
                        "voda_interest": 2.73,
                        "created_at": "2026-06-16 08:00:00",
                        "updated_at": "2026-07-16 08:05:00",
                    },
                    {
                        "savings_id": "FA-ECHU-ACTIF",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "product_name": "1 Month",
                        "product_description": "1 Month Fixed Account",
                        "currency_code": "USD",
                        "balance": 500,
                        "status": "Active",
                        "date_approved": "2026-06-16",
                        "maturity_date": "2026-07-16",
                        "interest_earned": 4.52,
                        "voda_interest": 1.23,
                        "created_at": "2026-06-16",
                        "updated_at": "2026-07-16",
                    },
                ]
            )
        )
        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": "1",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "FA-ECHU-SANS-ECRITURE",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 2500,
                        "bal_before": 0,
                        "bal_after": 2500,
                        "ref_no": "DEPOT-1",
                        "description": "Depot Bloque",
                        "created_at": "2026-06-16 21:53:03",
                    },
                    {
                        "id": "2",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "FA-ECHU-TRACE",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 1000,
                        "bal_before": 0,
                        "bal_after": 1000,
                        "ref_no": "DEPOT-2",
                        "description": "Depot Bloque",
                        "created_at": "2026-06-16 08:00:00",
                    },
                    {
                        "id": "3",
                        "customer_id": "DAT-CLIENT",
                        "msisdn1": "243814445054",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "FA-ECHU-TRACE",
                        "currency_code": "USD",
                        "dr": 1000,
                        "cr": 0,
                        "bal_before": 1000,
                        "bal_after": 0,
                        "ref_no": "RETRAIT-2",
                        "description": "Retrait Compte Bloque",
                        "created_at": "2026-07-16 08:00:00",
                    },
                ]
            )
        )

        entries = build_customer_matured_dat_interest_entries(
            fixed,
            transactions,
            "DAT-CLIENT",
            date_start="2026-07-16",
            date_end="2026-07-16",
            currency="USD",
        )

        self.assertEqual(entries["savings_id"].tolist(), ["FA-ECHU-TRACE", "FA-ECHU-SANS-ECRITURE"])
        self.assertAlmostEqual(float(entries["interet_client_constate"].sum()), 32.60)
        self.assertEqual(int(entries["date_ecriture_turbo"].notna().sum()), 1)
        traced = entries.loc[entries["savings_id"].eq("FA-ECHU-TRACE")].iloc[0]
        untraced = entries.loc[entries["savings_id"].eq("FA-ECHU-SANS-ECRITURE")].iloc[0]
        self.assertEqual(float(traced["capital_place"]), 1000.0)
        self.assertEqual(traced["reference_transaction_turbo"], "RETRAIT-2")
        self.assertIn("Comptabilise et trace", traced["statut_tracabilite"])
        self.assertEqual(float(untraced["montant_echeance_client"]), 2522.60)
        self.assertIn("ecriture detaillee absente", untraced["statut_tracabilite"])

    def test_customer_statement_word_and_pdf_include_active_dat_and_repayments_only(self) -> None:
        from docx import Document
        from pypdf import PdfReader

        prepared = _sample_prepared_data()
        statement = build_mpesa_statement(prepared, "1001", {"CDF": None})["extrait"]
        active_dat = pd.DataFrame(
            [
                {
                    "date_situation": pd.Timestamp("2026-07-17"),
                    "savings_id": "FA9T2OLVUC",
                    "customer_id": "1001",
                    "msisdn": "243812345678",
                    "currency_code": "CDF",
                    "product_name": "6 Months",
                    "date_approved": pd.Timestamp("2026-06-16"),
                    "maturity_date": pd.Timestamp("2026-12-16"),
                    "jours_avant_echeance": 152,
                    "balance": 20_000.0,
                    "taux_interet_annuel_pct": 11.0,
                    "interet_estime_echeance": 1_103.01,
                    "capital_plus_interet_estime": 21_103.01,
                    "situation_dat_client": "En cours",
                    "status": "Active",
                }
            ]
        )
        foreign_dat = active_dat.iloc[[0]].copy()
        foreign_dat["customer_id"] = "OTHER-CUSTOMER"
        foreign_dat["savings_id"] = "FOREIGN-DAT"
        active_dat = pd.concat([active_dat, foreign_dat], ignore_index=True)
        repayments = pd.DataFrame(
            [
                {
                    "customer_id": "1001",
                    "created_at": pd.Timestamp("2026-07-16 10:30:00"),
                    "event_reference": "REM-CLIENT",
                    "currency_code": "CDF",
                    "montant_paye_observe": 50_500.0,
                    "principal_rembourse": 50_000.0,
                    "interet_observe": 0.0,
                    "penalite_observee": 500.0,
                    "origine_remboursement_observee": "Compte M-PESA",
                    "mode_remboursement_observe": "M-PESA_Turbo + Penalite",
                }
            ]
        )
        foreign_repayment = repayments.iloc[[0]].copy()
        foreign_repayment["customer_id"] = "OTHER-CUSTOMER"
        foreign_repayment["event_reference"] = "FOREIGN-REPAYMENT"
        repayments = pd.concat([repayments, foreign_repayment], ignore_index=True)
        upcoming_repayments = pd.DataFrame(
            [
                {
                    "customer_id": "1001",
                    "loan_id": "LN-NEXT",
                    "due_date": pd.Timestamp("2026-07-20 09:00:00"),
                    "currency_code": "CDF",
                    "montant_a_rembourser": 15_000.0,
                    "outstanding_principle": 14_000.0,
                    "outstanding_interest": 800.0,
                    "outstanding_penalty_fees": 200.0,
                    "status_name": "Active",
                }
            ]
        )
        dat_interest = pd.DataFrame(
            [
                {
                    "customer_id": "1001",
                    "maturity_date": pd.Timestamp("2026-07-16"),
                    "savings_id": "DAT-INTERET-CLIENT",
                    "currency_code": "CDF",
                    "capital_place": 10_000.0,
                    "interet_client_constate": 250.0,
                    "montant_echeance_client": 10_250.0,
                    "voda_interest": 60.0,
                    "statut_tracabilite": "Colonne technique à masquer",
                },
                {
                    "customer_id": "OTHER-CUSTOMER",
                    "maturity_date": pd.Timestamp("2026-07-16"),
                    "savings_id": "FOREIGN-INTEREST",
                    "currency_code": "CDF",
                    "capital_place": 2_000.0,
                    "interet_client_constate": 50.0,
                    "montant_echeance_client": 2_050.0,
                },
            ]
        )
        analysis = {
            "dat_en_cours_client": active_dat,
            "remboursements_turbo_detail_client": repayments,
            "prochains_remboursements_client": upcoming_repayments,
            "interets_dat_credites_client": dat_interest,
        }

        word = create_customer_statement_word(
            statement,
            analysis_report=analysis,
            customer_id="1001",
            customer_name="CLIENT TEST",
            telephone="243812345678",
            currency="CDF",
        )
        document = Document(BytesIO(word))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        dat_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "DAT",
                "Souscription",
                "Échéance",
                "Jours restants",
                "Devise",
                "Capital bloqué",
                "Situation",
                "Capital + intérêt estimé",
            ]
        )
        repayment_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "Date",
                "Référence",
                "Devise",
                "Montant payé",
                "Intérêts",
                "Origine du paiement",
                "Pénalités",
            ]
        )
        self.assertIn("DAT en cours - situation au 17/07/2026", text)
        self.assertIn("Remboursements observés", text)
        self.assertIn("Prochains remboursements sur la période", text)
        self.assertIn("Entrées des intérêts du capital mis en DAT", text)
        self.assertNotIn("Intérêts des DAT échus", text)
        self.assertNotIn("Crédit et remboursements observés", text)
        self.assertNotIn("Montant reçu", text)
        self.assertEqual(len(dat_table.rows), 2)
        self.assertEqual(len(repayment_table.rows), 2)
        self.assertEqual(dat_table.rows[1].cells[0].text, "FA9T2OLVUC")
        self.assertEqual(repayment_table.rows[1].cells[1].text, "REM-CLIENT")
        self.assertEqual(repayment_table.rows[1].cells[4].text, "0")
        self.assertEqual(repayment_table.rows[1].cells[6].text, "500")
        self.assertNotIn("Principal remboursé", text)
        self.assertNotIn("Mode observé", text)
        self.assertNotIn("Colonne technique à masquer", text)

        pdf = create_customer_statement_pdf(
            statement,
            analysis_report=analysis,
            customer_id="1001",
            customer_name="CLIENT TEST",
            telephone="243812345678",
            currency="CDF",
        )
        self.assertTrue(pdf.startswith(b"%PDF-"))
        self.assertGreater(len(pdf), 5_000)
        pdf_reader = PdfReader(BytesIO(pdf))
        self.assertLess(
            float(pdf_reader.pages[0].mediabox.width),
            float(pdf_reader.pages[0].mediabox.height),
        )
        pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
        self.assertIn("DAT en cours - situation au 17/07/2026", pdf_text)
        self.assertIn("Remboursements observés", pdf_text)
        self.assertIn("Prochains remboursements sur la période", pdf_text)
        self.assertIn("Entrées des intérêts du capital mis en DAT", pdf_text)
        self.assertNotIn("Intérêts des DAT échus", pdf_text)
        self.assertNotIn("Crédit et remboursements observés", pdf_text)
        self.assertNotIn("Montant reçu", pdf_text)
        self.assertNotIn("FOREIGN-DAT", pdf_text)
        self.assertNotIn("FOREIGN-REPAYMENT", pdf_text)
        self.assertNotIn("FOREIGN-INTEREST", pdf_text)
        self.assertNotIn("Colonne technique à masquer", pdf_text)
        self.assertNotIn("Principal remboursé", pdf_text)
        self.assertNotIn("Mode observé", pdf_text)

    def test_customer_statement_pdf_contains_logo_and_keeps_currency_totals_separate(self) -> None:
        prepared = _sample_prepared_data()
        cdf_statement = build_mpesa_statement(prepared, "1001", {"CDF": None})["extrait"]
        usd_statement = cdf_statement.copy()
        usd_statement["currency_code"] = "USD"
        usd_statement["operation_reference"] = "USD-" + usd_statement["operation_reference"].astype(str)
        combined = pd.concat([cdf_statement, usd_statement], ignore_index=True)

        content = create_customer_statement_pdf(
            combined,
            customer_id="1001",
            customer_name="CLIENT TEST",
            telephone="243812345678",
            currency="ALL",
            entry_account_number="1441",
            output_account_number="15558",
            generated_at=pd.Timestamp("2026-07-16 10:30:00"),
        )

        self.assertTrue(content.startswith(b"%PDF-"))
        self.assertGreater(len(content), 5_000)
        self.assertIn(b"/Subtype /Image", content)

    def test_customer_statement_word_includes_filtered_turbo_analyses(self) -> None:
        from docx import Document

        prepared = _sample_customer_transaction_analysis_data()
        report = build_mpesa_statement(prepared, "CLIENT-ANALYSE", {"CDF": None})
        analysis = build_customer_transaction_analysis(prepared, "CLIENT-ANALYSE")

        content = create_customer_statement_word(
            report["extrait"],
            analysis_report=analysis,
            customer_id="CLIENT-ANALYSE",
            customer_name="CLIENT ANALYSE",
            telephone="243812345678",
            currency="CDF",
            entry_account_number="1441",
            output_account_number="15558",
        )

        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        footer_text = "\n".join(
            paragraph.text
            for section in document.sections
            for paragraph in section.footer.paragraphs
        )
        self.assertNotIn("Synthese du comportement observe", text)
        self.assertIn("DAT en cours", text)
        self.assertIn("Remboursements observés", text)
        self.assertNotIn("Crédit et remboursements observés", text)
        self.assertNotIn("Positions observees et rapprochement des soldes", text)
        self.assertNotIn("Jalons du parcours financier", text)
        self.assertIn("Retours du capital mis en DAT", text)
        return_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "Création du DAT",
                "Fin du DAT",
                "Référence",
                "Devise",
                "Capital DAT restitué",
                "Entrée compte ouvert",
                "Description",
            ]
        )
        self.assertEqual(return_table.rows[1].cells[0].text, "02/07/2026 07:45")
        self.assertEqual(return_table.rows[1].cells[1].text, "05/07/2026 08:00")
        self.assertIn("Detail des transactions", text)
        self.assertNotIn("[Turbo]", text)
        self.assertIn("Solution Bisou Bisou Digital", footer_text)
        self.assertNotIn("Solution Controle Interne", footer_text)

        from pypdf import PdfReader

        pdf = create_customer_statement_pdf(
            report["extrait"],
            analysis_report=analysis,
            customer_id="CLIENT-ANALYSE",
            customer_name="CLIENT ANALYSE",
            telephone="243812345678",
            currency="CDF",
            entry_account_number="1441",
            output_account_number="15558",
        )
        pdf_text = "\n".join(
            page.extract_text() or "" for page in PdfReader(BytesIO(pdf)).pages
        )
        self.assertIn("Création du DAT", pdf_text)
        self.assertIn("Fin du DAT", pdf_text)
        self.assertIn("02/07/2026 07:45", pdf_text)
        self.assertIn("05/07/2026 08:00", pdf_text)

    def test_customer_statement_word_all_keeps_currency_totals_separate(self) -> None:
        from docx import Document

        prepared = _sample_prepared_data()
        cdf_statement = build_mpesa_statement(prepared, "1001", {"CDF": None})["extrait"]
        usd_statement = cdf_statement.copy()
        usd_statement["currency_code"] = "USD"
        usd_statement["operation_reference"] = "USD-" + usd_statement["operation_reference"].astype(str)
        combined = pd.concat([cdf_statement, usd_statement], ignore_index=True)

        view = build_customer_statement_view(
            combined,
            entry_account_number="1441",
            output_account_number="15558",
            allow_multiple_currencies=True,
        )
        self.assertEqual(view["currency"], "ALL")
        self.assertTrue(pd.isna(view["total_entries"]))
        self.assertEqual(set(view["summary_by_currency"]["currency_code"]), {"CDF", "USD"})

        content = create_customer_statement_word(
            combined,
            customer_id="1001",
            customer_name="CLIENT TEST",
            telephone="243812345678",
            currency="ALL",
            entry_account_number="1441",
            output_account_number="15558",
            generated_at=pd.Timestamp("2026-07-16 10:30:00"),
        )
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        statement_table = next(
            table for table in document.tables if table.rows[0].cells[0].text == "Date"
        )
        summary_table = next(
            table
            for table in document.tables
            if [cell.text for cell in table.rows[0].cells]
            == [
                "Devise",
                "Ouverture",
                "Entrée",
                "Sorties",
                "Cloture",
                "Compte ouvert",
                "Compte bloqué",
            ]
        )
        self.assertIn("ALL (CDF, USD)", text)
        self.assertEqual(
            {row.cells[0].text.split()[0] for row in summary_table.rows[1:]},
            {"CDF", "USD"},
        )
        self.assertEqual(len(statement_table.rows), 1)

    def test_excel_export_contains_content(self) -> None:
        prepared = _sample_prepared_data()
        report = build_mpesa_statement(prepared, "1001", {"CDF": None})

        export = create_excel_export(report)

        self.assertGreater(len(export), 5000)

    def test_excel_export_writes_only_requested_sheets(self) -> None:
        export = create_excel_export(
            {
                "synthese": pd.DataFrame([{"indicateur": "Clients", "valeur": 1}]),
                "extrait": pd.DataFrame([{"operation_reference": "REF-001"}]),
            }
        )

        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")

        self.assertEqual(workbook.sheet_names, ["Synthese", "Extrait_Turbo"])

    def test_customer_excel_export_preserves_filtered_extract_and_required_sheets(self) -> None:
        prepared = _sample_prepared_data()
        report = build_mpesa_statement(prepared, "1001", {"CDF": None})
        filtered_report = dict(report)
        filtered_report["extrait"] = report["extrait"].iloc[[0]].copy()
        customer_export = {
            key: filtered_report.get(key, pd.DataFrame())
            for key in [
                "synthese",
                "extrait",
                "parcours_turbo",
                "dat_en_cours_client",
                "remboursements_turbo_detail_client",
                "elements_extrait_client_turbo",
                "interets_dat_credites_client",
                "comportement_turbo",
                "mouvements_internes_turbo",
                "controles_client_turbo",
                "dat_final",
                "mouvements_dat",
                "mouvements_epargne",
                "g2_dat",
                "diagnostics",
            ]
        }
        customer_export["synthese"] = customer_export["synthese"].drop(
            columns=["nombre_credits", "solde_credit_total"],
            errors="ignore",
        )

        export = create_excel_export(customer_export)
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        required_sheets = {
            "Synthese",
            "Extrait_Turbo",
            "Parcours_Turbo",
            "DAT_En_Cours",
            "Remboursements_Turbo",
            "Elements_Extrait_Turbo",
            "Interets_DAT_Credites",
            "Comportement_Turbo",
            "Mouvements_Internes",
            "Controles_Client_Turbo",
            "DAT_Final",
            "Mouvements_DAT",
            "Mouvements_Epargne",
            "G2_DAT",
            "Diagnostics",
        }

        self.assertTrue(required_sheets.issubset(workbook.sheet_names))
        self.assertNotIn("Credit_Client_Turbo", workbook.sheet_names)
        self.assertNotIn("Positions_Turbo", workbook.sheet_names)
        self.assertNotIn("Interets_DAT_Echus", workbook.sheet_names)
        self.assertNotIn("Credits", workbook.sheet_names)
        exported_statement = pd.read_excel(workbook, sheet_name="Extrait_Turbo")
        exported_summary = pd.read_excel(workbook, sheet_name="Synthese")
        exported_dat = pd.read_excel(workbook, sheet_name="DAT_En_Cours")
        self.assertEqual(len(exported_statement), 1)
        self.assertEqual(len(exported_summary), len(report["synthese"]))
        self.assertNotIn("nombre_credits", exported_summary.columns)
        self.assertNotIn("solde_credit_total", exported_summary.columns)
        self.assertIn("operation_reference", exported_statement.columns)
        self.assertIn("Nom_client", exported_statement.columns)
        self.assertIn("mouvement_net_mpesa", exported_statement.columns)
        self.assertIn("Jours restants", exported_dat.columns)
        self.assertNotIn("jours_avant_echeance", exported_dat.columns)

        filtered_report["extrait"] = report["extrait"].iloc[0:0].copy()
        empty_export = create_excel_export(filtered_report)
        empty_statement = pd.read_excel(BytesIO(empty_export), sheet_name="Extrait_Turbo")
        self.assertEqual(len(empty_statement), 0)
        self.assertIn("operation_reference", empty_statement.columns)

    def test_build_statement_accepts_transactions_only(self) -> None:
        prepared = _sample_prepared_data()
        prepared = MpesaPreparedData(
            transactions=prepared.transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=prepared.load_report,
        )

        report = build_mpesa_statement(prepared, "1001", {"CDF": None})

        self.assertEqual(len(report["extrait"]), 2)
        self.assertTrue(report["dat_final"].empty)
        self.assertTrue(report["credits"].empty)
        self.assertEqual(float(report["synthese"].iloc[0]["dat_final"]), 5000.0)
        self.assertIn(
            "Transactions",
            report["synthese"].iloc[0]["source_dat_final"],
        )

    def test_g2_transactions_extract_phone_and_match_dat(self) -> None:
        prepared = _sample_prepared_data()
        g2 = pd.DataFrame(
            [
                {
                    "Receipt No.\xa0": "TX001",
                    "Completion Time\xa0": "2026-07-11 10:23:05",
                    "Opposite Party\xa0": "0999999999 - AUTRE TELEPHONE",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "CDF 1,000.00",
                    "Balance\xa0": "CDF 2,000.00",
                    "Transaction Status\xa0": "Completed",
                }
            ]
        )
        prepared = MpesaPreparedData(
            prepared.transactions,
            prepared.current_savings,
            prepared.fixed_savings,
            prepared.loans,
            prepared.load_report,
            prepare_g2_transactions(g2),
        )

        result = build_g2_dat_crosscheck(prepared)

        self.assertEqual(result.iloc[0]["phone_prefixe"], "243999999999")
        self.assertEqual(result.iloc[0]["customer_id_ref_no"], "1001")
        self.assertEqual(result.iloc[0]["customer_id_dat"], "1001")
        self.assertEqual(result.iloc[0]["reference_dat_operation"], "FA001")
        self.assertEqual(float(result.iloc[0]["solde_dat_operation_avant"]), 4000.0)
        self.assertEqual(float(result.iloc[0]["solde_dat_operation_apres"]), 5000.0)
        self.assertEqual(float(result.iloc[0]["solde_dat_operation"]), 5000.0)
        self.assertEqual(float(result.iloc[0]["dat_final"]), 5000.0)
        self.assertEqual(float(result.iloc[0]["variation_dat_operation"]), 1000.0)
        self.assertEqual(result.iloc[0]["mode_rapprochement"], "Receipt No = ref_no + DAT operation")
        self.assertEqual(float(result.iloc[0]["dat_final_client_devise"]), 5000.0)

    def test_g2_entry_report_builds_detail_and_summary(self) -> None:
        prepared = _sample_prepared_data()
        g2 = pd.DataFrame(
            [
                {
                    "Receipt No.\xa0": "TX001",
                    "Completion Time\xa0": "2026-07-11 10:23:05",
                    "Opposite Party\xa0": "0812345678 - CLIENT TEST",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "CDF 1,000.00",
                    "Balance\xa0": "CDF 2,000.00",
                    "Transaction Status\xa0": "Completed",
                }
            ]
        )
        prepared = MpesaPreparedData(
            prepared.transactions,
            prepared.current_savings,
            prepared.fixed_savings,
            prepared.loans,
            prepared.load_report,
            prepare_g2_transactions(g2),
        )

        report = build_g2_entry_report(prepared)

        self.assertEqual(report["detail"].iloc[0]["details_rapport"], "DAT")
        self.assertEqual(float(report["detail"].iloc[0]["montant"]), 1000.0)
        self.assertIn("Total CDF", report["synthese"]["details_rapport"].tolist())
        self.assertIn("montant_DAT", report["pivot"].columns)
        self.assertEqual(float(report["pivot"].iloc[0]["montant_DAT"]), 1000.0)

    def test_g2_repayment_detail_has_priority_over_dat_match(self) -> None:
        prepared = _sample_prepared_data()
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "TX001",
                        "Completion Time": "2026-07-11 10:23:05",
                        "Details": "BisouBisouRepayment",
                        "Opposite Party": "0812345678 - CLIENT TEST",
                        "Currency": "CDF",
                        "Transaction Amount": "CDF 1,000.00",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            prepared.transactions,
            prepared.current_savings,
            prepared.fixed_savings,
            prepared.loans,
            prepared.load_report,
            g2,
        )

        detail = build_g2_entry_report(prepared)["detail"]

        self.assertEqual(detail.iloc[0]["details_rapport"], "Remboursement prets")
        self.assertEqual(detail.iloc[0]["Nom_client"], "CLIENT TEST")

    def test_daily_g2_savings_report_matches_dat_by_phone_currency_and_amount(self) -> None:
        g2 = pd.DataFrame(
            [
                {
                    "Receipt No.\xa0": "DAT001",
                    "Completion Time\xa0": "2026-07-11 12:47:24",
                    "Opposite Party\xa0": "243826325569 - CLIENT DAT",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "Fc 5,000",
                    "Transaction Status\xa0": "Completed",
                },
                {
                    "Receipt No.\xa0": "DAT002",
                    "Completion Time\xa0": "2026-07-11 12:44:13",
                    "Opposite Party\xa0": "243826325569 - CLIENT DAT",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "Fc 80,000",
                    "Transaction Status\xa0": "Completed",
                },
                {
                    "Receipt No.\xa0": "LOAN001",
                    "Completion Time\xa0": "2026-07-11 12:43:23",
                    "Opposite Party\xa0": "243835549888 - CLIENT CREDIT",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "Fc 1,285",
                    "Transaction Status\xa0": "Completed",
                },
                {
                    "Receipt No.\xa0": "SAVE001",
                    "Completion Time\xa0": "2026-07-11 10:18:58",
                    "Opposite Party\xa0": "243822452403 - CLIENT EPARGNE",
                    "Currency\xa0": "CDF",
                    "Transaction Amount\xa0": "Fc 2,000",
                    "Transaction Status\xa0": "Completed",
                },
            ]
        )
        fixed = pd.DataFrame(
            [
                {
                    "customer_id": 37478,
                    "msisdn": "243826325569",
                    "product_name": "3 Months",
                    "account_type": "3 MONTH Fixed Account",
                    "balance": 85000,
                    "currency_code": "CDF",
                    "date_approved": "2026-07-11 11:44:14",
                    "maturity_date": "2026-10-11 11:44:14",
                },
                {
                    "customer_id": 26303,
                    "msisdn": "243835549888",
                    "product_name": "1 Month",
                    "account_type": "1 Month Fixed Account",
                    "balance": 1500,
                    "currency_code": "CDF",
                    "date_approved": "2026-07-11 11:46:11",
                    "maturity_date": "2026-08-11 11:46:11",
                },
            ]
        )
        current = pd.DataFrame(
            [
                {
                    "customer_id": 37478,
                    "msisdn": "243826325569",
                    "product_name": "Open Savings",
                    "account_type": "Current account",
                    "balance": 0,
                    "currency_code": "CDF",
                    "created_at": "2026-07-10 22:37:54",
                    "updated_at": "2026-07-11 08:00:00",
                }
            ]
        )
        customers = pd.DataFrame(
            [
                {
                    "msisdn1": "243826325569",
                    "created_at": "2026-07-09 09:30:00",
                }
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=prepare_current_savings(current),
            fixed_savings=prepare_fixed_savings(fixed),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(g2),
            customers=prepare_customers(customers),
        )

        report = build_g2_daily_savings_report(prepared)
        detail = report["detail"].set_index("receipt_no")

        self.assertEqual(detail.loc["DAT001", "details_rapport"], "DAT")
        self.assertEqual(detail.loc["DAT002", "details_rapport"], "DAT")
        self.assertEqual(detail.loc["LOAN001", "details_rapport"], "Remboursement prets")
        self.assertEqual(detail.loc["SAVE001", "details_rapport"], "Depot normal")
        self.assertEqual(pd.Timestamp(detail.loc["DAT001", "compte_cree"]), pd.Timestamp("2026-07-09 09:30:00"))
        self.assertEqual(pd.Timestamp(detail.loc["DAT001", "compte_cree_client"]), pd.Timestamp("2026-07-09 09:30:00"))
        self.assertEqual(pd.Timestamp(detail.loc["DAT001", "compte_cree_epargne_courante"]), pd.Timestamp("2026-07-10 22:37:54"))
        self.assertEqual(pd.Timestamp(detail.loc["DAT001", "compte_cree_dat"]), pd.Timestamp("2026-07-11 11:44:14"))
        pivot = report["pivot"].set_index("currency_code")
        self.assertEqual(float(pivot.loc["CDF", "montant_DAT"]), 85000.0)
        self.assertEqual(float(pivot.loc["CDF", "montant_Remboursement prets"]), 1285.0)

    def test_daily_g2_report_uses_portal_reference_to_classify_each_receipt(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "REF-DAT",
                        "Completion Time": "2026-07-13 10:00:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 5000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "REF-SAVE",
                        "Completion Time": "2026-07-13 10:05:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 80000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "REF-LOAN",
                        "Completion Time": "2026-07-13 10:10:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 1285,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                ]
            )
        )
        portal_rows = []
        for ref_no, amount, target_type, description, minute in [
            ("REF-DAT", 5000, "FIXED SAVINGS", "Depot Bloque", "00"),
            ("REF-SAVE", 80000, "NORMAL SAVINGS", "Epargne depot", "05"),
            ("REF-LOAN", 1285, "LOAN ACCOUNT", "Remboursement", "10"),
        ]:
            portal_rows.extend(
                [
                    {
                        "customer_id": 1001,
                        "msisdn1": "0811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": ref_no,
                        "currency_code": "CDF",
                        "dr": amount,
                        "cr": 0,
                        "bal_before": 100000,
                        "bal_after": 100000 - amount,
                        "ref_no": ref_no,
                        "description": "M-Pesa Compte",
                        "created_at": f"2026-07-13 10:{minute}:00",
                    },
                    {
                        "customer_id": 1001,
                        "msisdn1": "0811111111",
                        "account_type": target_type,
                        "reference_id": ref_no,
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": amount,
                        "bal_before": 0,
                        "bal_after": amount,
                        "ref_no": ref_no,
                        "description": description,
                        "created_at": f"2026-07-13 10:{minute}:00",
                    },
                ]
            )
        portal_frame = pd.DataFrame(portal_rows)
        portal_frame["id"] = range(1, len(portal_frame) + 1)
        prepared = MpesaPreparedData(
            transactions=prepare_transactions(portal_frame),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_g2_daily_savings_report(prepared)
        detail = report["detail"].set_index("receipt_no")

        self.assertEqual(detail.loc["REF-DAT", "details_rapport"], "DAT")
        self.assertEqual(detail.loc["REF-SAVE", "details_rapport"], "Depot normal")
        self.assertEqual(detail.loc["REF-LOAN", "details_rapport"], "Remboursement prets")
        self.assertTrue(detail["statut_rapprochement"].eq("Rapproche exact").all())
        self.assertTrue(report["anomalies"].empty)

    def test_g2_daily_report_falls_back_to_turbo_when_g2_is_absent(self) -> None:
        portal_rows = [
            {
                "id": 1,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "MPESA ACCOUNT",
                "reference_id": "NORMAL-1",
                "currency_code": "CDF",
                "dr": 100,
                "cr": 0,
                "bal_before": 500,
                "bal_after": 400,
                "ref_no": "REF-NORMAL",
                "description": "M-Pesa Depot",
                "created_at": "2026-07-15 08:00:00",
            },
            {
                "id": 2,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "NORMAL SAVINGS",
                "reference_id": "NORMAL-1",
                "currency_code": "CDF",
                "dr": 0,
                "cr": 100,
                "bal_before": 0,
                "bal_after": 100,
                "ref_no": "REF-NORMAL",
                "description": "Epargne depot",
                "created_at": "2026-07-15 08:00:00",
            },
            {
                "id": 3,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "MPESA ACCOUNT",
                "reference_id": "FIXED-1",
                "currency_code": "CDF",
                "dr": 200,
                "cr": 0,
                "bal_before": 400,
                "bal_after": 200,
                "ref_no": "REF-DAT",
                "description": "M-Pesa Compte",
                "created_at": "2026-07-15 09:00:00",
            },
            {
                "id": 4,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "FIXED SAVINGS",
                "reference_id": "FIXED-1",
                "currency_code": "CDF",
                "dr": 0,
                "cr": 200,
                "bal_before": 0,
                "bal_after": 200,
                "ref_no": "REF-DAT",
                "description": "Depot Bloque",
                "created_at": "2026-07-15 09:00:00",
            },
            {
                "id": 5,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "NORMAL SAVINGS",
                "reference_id": "SA-OUT-1",
                "currency_code": "CDF",
                "dr": 50,
                "cr": 0,
                "bal_before": 100,
                "bal_after": 50,
                "ref_no": "",
                "description": "Retrait Vers M-Pesa",
                "created_at": "2026-07-15 10:00:00",
            },
            {
                "id": 6,
                "customer_id": 1001,
                "msisdn1": "0811111111",
                "account_type": "MPESA ACCOUNT",
                "reference_id": "SA-OUT-1",
                "currency_code": "CDF",
                "dr": 0,
                "cr": 50,
                "bal_before": 200,
                "bal_after": 250,
                "ref_no": "",
                "description": "Retrait Vers M-Pesa",
                "created_at": "2026-07-15 10:00:00",
            },
        ]
        transactions = prepare_transactions(pd.DataFrame(portal_rows))
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        proxy = build_turbo_only_g2_transactions(transactions)
        report = build_g2_daily_savings_report(prepared)
        detail = report["detail"].set_index("details_rapport")

        self.assertEqual(len(proxy), 3)
        self.assertEqual(len(detail), 3)
        self.assertEqual(set(detail.index), {"Depot normal", "DAT", "Paiement client B2C"})
        self.assertEqual(int(detail["sens_flux"].eq("Entree").sum()), 2)
        self.assertEqual(int(detail["sens_flux"].eq("Sortie").sum()), 1)
        self.assertTrue(detail["source_analytique"].eq("Solution Numérique seule").all())
        self.assertTrue(detail["statut_transaction_g2"].eq("Comptabilisee Solution Numérique").all())
        self.assertTrue(detail["incluse_synthese"].all())
        self.assertTrue(detail["statut_rapprochement"].eq("Non applicable - Solution Numérique seule").all())
        self.assertTrue(detail["controle_date"].eq("Non applicable - Solution Numérique seule").all())
        self.assertTrue(report["anomalies"].empty)
        self.assertEqual(float(report["pivot"].iloc[0]["montant_total_entrees"]), 300.0)
        self.assertEqual(float(report["pivot"].iloc[0]["montant_total_sorties"]), 50.0)
        self.assertFalse(build_g2_retention_report(prepared, daily_detail=report["detail"])["mensuelle"].empty)

    def test_daily_g2_report_counts_duplicate_receipt_once_and_exports_anomaly(self) -> None:
        duplicate_rows = [
            {
                "Receipt No.": "DUP-001",
                "Completion Time": completion,
                "Opposite Party": "0811111111 - CLIENT A",
                "Currency": "CDF",
                "Transaction Amount": 5000,
                "Transaction Status": "Completed",
                "Details": "BisouBisouC2B",
            }
            for completion in ["2026-07-13 10:00:00", "2026-07-13 10:01:00"]
        ]
        portal_frame = pd.DataFrame(
                [
                    {
                        "customer_id": 1001,
                        "msisdn1": "0811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "DUP-001",
                        "currency_code": "CDF",
                        "dr": 5000,
                        "cr": 0,
                        "bal_before": 10000,
                        "bal_after": 5000,
                        "ref_no": "DUP-001",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-13 10:01:00",
                    },
                    {
                        "customer_id": 1001,
                        "msisdn1": "0811111111",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "FA-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 5000,
                        "bal_before": 0,
                        "bal_after": 5000,
                        "ref_no": "DUP-001",
                        "description": "Depot Bloque",
                        "created_at": "2026-07-13 10:01:00",
                    },
                ]
            )
        portal_frame["id"] = range(1, len(portal_frame) + 1)
        portal = prepare_transactions(portal_frame)
        prepared = MpesaPreparedData(
            transactions=portal,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(pd.DataFrame(duplicate_rows)),
        )

        report = build_g2_daily_savings_report(prepared)

        self.assertEqual(len(report["detail"]), 1)
        self.assertEqual(int(report["detail"].iloc[0]["nombre_lignes_g2_reference"]), 2)
        self.assertEqual(float(report["pivot"].iloc[0]["montant_DAT"]), 5000.0)
        self.assertEqual(len(report["anomalies"]), 1)
        self.assertIn("Receipt No duplique", report["anomalies"].iloc[0]["motif_anomalie"])

    def test_g2_statuses_are_traced_but_only_completed_transactions_feed_analyses(self) -> None:
        statuses = ["Completed", "Declined", "Cancelled", "Expired", "Pending", ""]
        raw = pd.DataFrame(
            [
                {
                    "Receipt No.": f"STATUS-{index}",
                    "Completion Time": f"2026-07-15 {8 + index:02d}:00:00",
                    "Details": "BisouBisouC2B",
                    "Transaction Status": status,
                    "Currency": "CDF",
                    "Paid In": 1000,
                    "Withdrawn": 0,
                    "Opposite Party": f"08111111{index:02d} - CLIENT {index}",
                    "Fichier source G2": "statuts_entrees.xlsx",
                }
                for index, status in enumerate(statuses)
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(raw),
        )

        report = build_g2_daily_savings_report(prepared)
        detail = report["detail"].set_index("transaction_status")

        self.assertEqual(len(detail), 6)
        self.assertTrue(bool(detail.loc["Completed", "incluse_synthese"]))
        for status in ["Declined", "Cancelled", "Expired", "Pending", ""]:
            self.assertFalse(bool(detail.loc[status, "incluse_synthese"]))
        self.assertEqual(float(report["pivot"].iloc[0]["montant_total_entrees"]), 1000.0)
        self.assertEqual(int(report["pivot"].iloc[0]["nombre_entrees"]), 1)
        self.assertEqual(
            int(build_g2_transaction_time_analysis(report["detail"])["par_heure"]["nombre_transactions"].sum()),
            1,
        )
        status_summary = report["statuts"].set_index("statut_transaction_g2")
        self.assertEqual(int(status_summary.loc["Completed", "nombre_transactions"]), 1)
        self.assertEqual(status_summary.loc["Completed", "prise_en_compte_analyse"], "Oui")
        self.assertEqual(status_summary.loc["Declined", "prise_en_compte_analyse"], "Non - controle uniquement")
        self.assertIn("Non renseigne", status_summary.index)
        self.assertEqual(set(report["detail"]["fichier_source_g2"]), {"statuts_entrees.xlsx"})

    def test_g2_completed_b2c_withdrawals_are_integrated_as_outputs(self) -> None:
        raw = pd.DataFrame(
            [
                {
                    "Receipt No.": "B2C-OUT-001",
                    "Completion Time": "2026-07-15 14:17:59",
                    "Details": "Bisou Bisou B2C payment from 15558 to 243811495678 - CLIENT SORTIE",
                    "Transaction Status": "Completed",
                    "Currency": "USD",
                    "Paid In": pd.NA,
                    "Withdrawn": -200,
                    "Balance": 81859.97,
                    "Opposite Party": "243811495678 - CLIENT SORTIE",
                    "Fichier source G2": "sorties_g2.xlsx",
                }
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(raw),
        )

        report = build_g2_daily_savings_report(prepared)
        row = report["detail"].iloc[0]

        self.assertEqual(row["sens_flux"], "Sortie")
        self.assertEqual(row["details_rapport"], "Paiement client B2C")
        self.assertEqual(float(row["montant_sortie"]), 200.0)
        self.assertEqual(int(report["pivot"].iloc[0]["nombre_sorties"]), 1)
        self.assertEqual(float(report["pivot"].iloc[0]["montant_total_sorties"]), 200.0)
        self.assertEqual(row["fichier_source_g2"], "sorties_g2.xlsx")

    def test_g2_b2c_output_matches_turbo_withdrawal_without_ref_no(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "G2-OUTPUT-001",
                        "Initiation Time": "2026-07-15 14:17:59",
                        "Completion Time": "2026-07-15 14:17:59",
                        "Details": "Bisou Bisou B2C payment to 243811495678 - CLIENT SORTIE",
                        "Reason Type": "BisouBisouB2C",
                        "Transaction Status": "Completed",
                        "Currency": "USD",
                        "Paid In": pd.NA,
                        "Withdrawn": -200,
                        "Opposite Party": "243811495678 - CLIENT SORTIE",
                    }
                ]
            )
        )
        turbo = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-OUTPUT-001",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 200,
                        "bal_before": 0,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-15 13:17:58",
                    },
                    {
                        "id": 2,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-OUTPUT-001",
                        "currency_code": "USD",
                        "dr": 200,
                        "cr": 0,
                        "bal_before": 400,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-15 13:17:58",
                    },
                    {
                        "id": 3,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-OUTPUT-001",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 200,
                        "bal_before": 0,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-14 10:00:00",
                    },
                    {
                        "id": 4,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-OUTPUT-001",
                        "currency_code": "USD",
                        "dr": 200,
                        "cr": 0,
                        "bal_before": 600,
                        "bal_after": 400,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-14 10:00:00",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=turbo,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_g2_daily_savings_report(prepared)
        row = report["detail"].iloc[0]

        self.assertTrue(pd.isna(row["ref_no_portal"]))
        self.assertEqual(row["reference_sortie_turbo"], "SA-OUTPUT-001")
        self.assertEqual(
            row["methode_rapprochement_turbo"],
            "Telephone + devise + montant + heure (sortie)",
        )
        self.assertEqual(row["operation_turbo_confirmee"], "Retrait epargne vers M-PESA")
        self.assertEqual(row["customer_id_portal"], "321")
        self.assertEqual(int(row["nombre_candidats_sortie_turbo"]), 1)
        self.assertEqual(row["controle_telephone"], "Conforme")
        self.assertEqual(row["controle_devise"], "Conforme")
        self.assertEqual(row["controle_montant"], "Conforme")
        self.assertAlmostEqual(float(row["ecart_creation_minutes"]), 60.0166667, places=4)
        self.assertEqual(row["controle_date_creation"], "Ecart de date")
        self.assertEqual(row["statut_rapprochement"], "Rapproche avec ecart")
        self.assertEqual(row["source_analytique"], "G2 + Solution Numérique")
        self.assertEqual(len(report["anomalies"]), 1)
        self.assertIn("Ecart de date de creation", row["motif_anomalie"])

    def test_g2_turbo_gap_of_exactly_60_minutes_is_accepted(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "G2-OUTPUT-060",
                        "Initiation Time": "2026-07-16 11:00:00",
                        "Completion Time": "2026-07-16 11:00:00",
                        "Details": "Bisou Bisou B2C payment to 243811495678 - CLIENT SORTIE",
                        "Reason Type": "BisouBisouB2C",
                        "Transaction Status": "Completed",
                        "Currency": "USD",
                        "Paid In": pd.NA,
                        "Withdrawn": -200,
                        "Opposite Party": "243811495678 - CLIENT SORTIE",
                    }
                ]
            )
        )
        turbo = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-OUTPUT-060",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 200,
                        "bal_before": 0,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-16 10:00:00",
                    },
                    {
                        "id": 2,
                        "customer_id": 321,
                        "msisdn1": "243811495678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-OUTPUT-060",
                        "currency_code": "USD",
                        "dr": 200,
                        "cr": 0,
                        "bal_before": 400,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-16 10:00:00",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=turbo,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        row = build_g2_daily_savings_report(prepared)["detail"].iloc[0]

        self.assertEqual(float(row["ecart_creation_minutes"]), 60.0)
        self.assertEqual(row["controle_date_creation"], "Conforme")
        self.assertEqual(row["statut_rapprochement"], "Rapproche exact")
        self.assertNotIn("Ecart de date de creation", row["motif_anomalie"])

    def test_daily_g2_report_retains_phone_currency_amount_and_date_gaps(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "GAP-001",
                        "Completion Time": "2026-07-13 10:00:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 5000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    }
                ]
            )
        )
        portal_frame = pd.DataFrame(
                [
                    {
                        "customer_id": 9999,
                        "msisdn1": "0899999999",
                        "account_type": "FIXED SAVINGS",
                        "reference_id": "FA-GAP",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 4000,
                        "bal_before": 0,
                        "bal_after": 4000,
                        "ref_no": "GAP-001",
                        "description": "Depot Bloque",
                        "created_at": "2026-07-14 10:00:00",
                    }
                ]
            )
        portal_frame["id"] = range(1, len(portal_frame) + 1)
        portal = prepare_transactions(portal_frame)
        prepared = MpesaPreparedData(
            transactions=portal,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_g2_daily_savings_report(prepared)
        row = report["detail"].iloc[0]

        self.assertEqual(row["statut_rapprochement"], "Rapproche avec ecart")
        self.assertEqual(row["controle_telephone"], "Ecart")
        self.assertEqual(row["controle_devise"], "Ecart")
        self.assertEqual(row["controle_montant"], "Ecart")
        self.assertEqual(row["controle_date"], "Ecart de date")
        self.assertEqual(row["controle_date_creation"], "Ecart de date")
        self.assertEqual(row["source_date_creation_g2"], "Completion Time (repli)")
        self.assertEqual(float(row["ecart_creation_minutes"]), -1440.0)
        self.assertIn("Creation G2 : 13/07/2026 10:00:00", row["Observation"])
        self.assertIn("Creation Solution Numérique : 14/07/2026 10:00:00", row["Observation"])
        self.assertIn("Decalage creation : 1440 minute(s)", row["Observation"])
        self.assertEqual(len(report["anomalies"]), 1)
        self.assertIn("Observation", report["anomalies"].columns)
        self.assertIn("Ecart de montant", row["motif_anomalie"])

    def test_g2_creation_date_uses_initiation_time_and_treats_midnight_completion_as_delay(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "MIDNIGHT-001",
                        "Initiation Time": "2026-07-13 23:59:00",
                        "Completion Time": "2026-07-14 00:01:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Transaction Amount": 5000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    }
                ]
            )
        )
        portal = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": 9999,
                        "msisdn1": "0811111111",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-MIDNIGHT",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 5000,
                        "bal_before": 0,
                        "bal_after": 5000,
                        "ref_no": "MIDNIGHT-001",
                        "description": "Epargne depot",
                        "created_at": "2026-07-13 23:59:00",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=portal,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_g2_daily_savings_report(prepared)
        row = report["detail"].iloc[0]

        self.assertEqual(row["source_date_creation_g2"], "Initiation Time")
        self.assertEqual(row["date_creation_g2"], pd.Timestamp("2026-07-13 23:59:00"))
        self.assertEqual(row["date_creation_turbo"], pd.Timestamp("2026-07-13 23:59:00"))
        self.assertEqual(row["date_finalisation_g2"], pd.Timestamp("2026-07-14 00:01:00"))
        self.assertEqual(float(row["ecart_creation_minutes"]), 0.0)
        self.assertEqual(float(row["delai_traitement_g2_minutes"]), 2.0)
        self.assertEqual(float(row["ecart_finalisation_minutes"]), 2.0)
        self.assertEqual(row["controle_date_creation"], "Conforme")
        self.assertEqual(row["controle_date"], "Conforme")
        self.assertEqual(row["controle_date_finalisation"], "Conforme - passage de date")
        self.assertEqual(row["statut_rapprochement"], "Rapproche exact")
        self.assertIn("Creation conforme; finalisation sur une autre date", row["Observation"])
        self.assertIn("Delai traitement G2 : 2 minute(s)", row["Observation"])
        self.assertTrue(report["anomalies"].empty)

    def test_g2_transaction_time_analysis_counts_days_hours_and_fills_inactive_hours(self) -> None:
        detail = pd.DataFrame(
            [
                {
                    "date": "2026-07-14 09:05:00",
                    "receipt_no": "DAY-001",
                    "currency_code": "CDF",
                    "sens_flux": "Entree",
                    "incluse_synthese": True,
                },
                {
                    "date": "2026-07-14 09:45:00",
                    "receipt_no": "DAY-002",
                    "currency_code": "CDF",
                    "sens_flux": "Sortie",
                    "incluse_synthese": True,
                },
                {
                    "date": "2026-07-14 11:00:00",
                    "receipt_no": "DAY-003",
                    "currency_code": "USD",
                    "sens_flux": "Entree",
                    "incluse_synthese": True,
                },
                {
                    "date": "2026-07-16 09:00:00",
                    "receipt_no": "DAY-004",
                    "currency_code": "CDF",
                    "sens_flux": "Entree",
                    "incluse_synthese": True,
                },
                {
                    "date": "2026-07-16 10:00:00",
                    "receipt_no": "EXCLUDED",
                    "currency_code": "CDF",
                    "sens_flux": "Entree",
                    "incluse_synthese": False,
                },
            ]
        )

        report = build_g2_transaction_time_analysis(detail)
        par_jour = report["par_jour"]
        par_jour_semaine = report["par_jour_semaine"]
        par_heure = report["par_heure"]
        jour_heure = report["jour_heure"]

        self.assertEqual(int(par_jour["nombre_transactions"].sum()), 4)
        self.assertEqual(par_jour["date_transaction"].nunique(), 3)
        july_15 = par_jour["date_transaction"].eq(pd.Timestamp("2026-07-15"))
        self.assertEqual(int(par_jour.loc[july_15, "nombre_transactions"].sum()), 0)

        cdf_entries_at_9 = par_heure.loc[
            par_heure["heure_num"].eq(9)
            & par_heure["currency_code"].eq("CDF")
            & par_heure["sens_flux"].eq("Entree"),
            "nombre_transactions",
        ]
        self.assertEqual(int(cdf_entries_at_9.iloc[0]), 2)
        self.assertEqual(int(par_heure.loc[par_heure["heure_num"].eq(0), "nombre_transactions"].sum()), 0)
        self.assertEqual(set(par_heure["heure_num"]), set(range(24)))
        self.assertEqual(int(jour_heure["nombre_transactions"].sum()), 4)
        self.assertEqual(int(par_jour_semaine["nombre_transactions"].sum()), 4)
        self.assertEqual(set(par_jour_semaine["jour_semaine_num"]), set(range(7)))
        weekday_totals = par_jour_semaine.groupby("jour_semaine", as_index=True)["nombre_transactions"].sum()
        self.assertEqual(int(weekday_totals["Mardi"]), 3)
        self.assertEqual(int(weekday_totals["Mercredi"]), 0)

        export = create_excel_export(
            {
                "transactions_par_jour": par_jour,
                "transactions_par_jour_semaine": par_jour_semaine,
                "transactions_par_heure": par_heure,
                "transactions_jour_heure": jour_heure,
            }
        )
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertEqual(
            workbook.sheet_names,
            [
                "Transactions_Jour",
                "Transactions_Jour_Semaine",
                "Transactions_Heure",
                "Transactions_Jour_Heure",
            ],
        )

    def test_g2_transaction_time_analysis_returns_empty_frames_without_valid_dates(self) -> None:
        report = build_g2_transaction_time_analysis(pd.DataFrame({"date": ["invalide"]}))

        self.assertTrue(report["par_jour"].empty)
        self.assertTrue(report["par_jour_semaine"].empty)
        self.assertTrue(report["par_heure"].empty)
        self.assertTrue(report["jour_heure"].empty)

    def test_mpesa_accounting_analysis_builds_client_balance_from_turbo_only(self) -> None:
        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": "CDF-1",
                        "customer_id": "CLIENT-CDF",
                        "msisdn1": "0811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-CDF",
                        "currency_code": "CDF",
                        "dr": 100,
                        "cr": 0,
                        "bal_before": 0,
                        "bal_after": 100,
                        "ref_no": "REF-CDF",
                        "description": "M-Pesa Depot",
                        "created_at": "2026-07-16 10:00:00",
                    },
                    {
                        "id": "CDF-2",
                        "customer_id": "CLIENT-CDF",
                        "msisdn1": "0811111111",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-CDF",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 100,
                        "bal_before": 200,
                        "bal_after": 300,
                        "ref_no": "REF-CDF",
                        "description": "Epargne depot",
                        "created_at": "2026-07-16 10:00:00",
                    },
                    {
                        "id": "USD-1",
                        "customer_id": "CLIENT-USD",
                        "msisdn1": "0822222222",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "LN-USD",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 20,
                        "bal_before": 0,
                        "bal_after": 20,
                        "ref_no": "REF-USD",
                        "description": "Montant pret",
                        "created_at": "2026-07-16 11:00:00",
                    },
                    {
                        "id": "USD-2",
                        "customer_id": "CLIENT-USD",
                        "msisdn1": "0822222222",
                        "account_type": "PRINCIPLE",
                        "reference_id": "LN-USD",
                        "currency_code": "USD",
                        "dr": 20,
                        "cr": 0,
                        "bal_before": 0,
                        "bal_after": 20,
                        "ref_no": "REF-USD",
                        "description": "Montant principal",
                        "created_at": "2026-07-16 11:00:00",
                    },
                    {
                        "id": "NEXT-DAY",
                        "customer_id": "CLIENT-CDF",
                        "msisdn1": "0811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-CDF",
                        "currency_code": "CDF",
                        "dr": 999,
                        "cr": 0,
                        "bal_before": 100,
                        "bal_after": 1099,
                        "ref_no": "REF-NEXT",
                        "description": "M-Pesa Depot",
                        "created_at": "2026-07-17 08:00:00",
                    },
                ]
            )
        )
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "REF-CDF",
                        "Completion Time": "2026-07-16 10:01:00",
                        "Initiation Time": "2026-07-16 10:00:00",
                        "Opposite Party": "0811111111 - CLIENT CDF",
                        "Currency": "CDF",
                        "Paid In": 900,
                        "Withdrawn": 0,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "REF-USD",
                        "Completion Time": "2026-07-16 11:01:00",
                        "Initiation Time": "2026-07-16 11:00:00",
                        "Opposite Party": "0822222222 - CLIENT USD",
                        "Currency": "USD",
                        "Paid In": 800,
                        "Withdrawn": 0,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                ]
            )
        )
        transactions = enrich_transactions_with_g2_customer_names(transactions, g2)
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=prepare_current_savings(
                pd.DataFrame(
                    [
                        {
                            "customer_id": "CLIENT-CDF",
                            "msisdn": "0811111111",
                            "product_name": "Courant",
                            "account_type": "NORMAL SAVINGS",
                            "balance": 300,
                            "currency_code": "CDF",
                            "created_at": "2026-01-01",
                            "updated_at": "2026-07-17",
                        }
                    ]
                )
            ),
            fixed_savings=pd.DataFrame(),
            loans=prepare_loans(
                pd.DataFrame(
                    [
                        {
                            "loan_id": "LN-USD",
                            "customer_id": "CLIENT-USD",
                            "currency_code": "USD",
                            "loan_balance": 20,
                            "updated_at": "2026-07-17",
                        }
                    ]
                )
            ),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_mpesa_accounting_analysis(
            prepared,
            date_start="2026-07-16",
            date_end="2026-07-16",
        )

        self.assertEqual(set(report["synthese"]["currency_code"]), {"CDF", "USD"})
        self.assertEqual(len(report["journal_ecritures"]), 4)
        self.assertEqual(len(report["journal_operations"]), 2)
        self.assertTrue(report["journal_operations"]["operation_symetrique"].all())
        cdf_summary = report["synthese"].loc[
            report["synthese"]["currency_code"].eq("CDF")
        ].iloc[0]
        self.assertEqual(float(cdf_summary["total_debit"]), 100.0)
        self.assertEqual(float(cdf_summary["total_credit"]), 100.0)
        self.assertEqual(float(cdf_summary["taux_rapprochement_g2_pct"]), 100.0)
        cdf_client = report["balance_clients"].loc[
            report["balance_clients"]["customer_id"].eq("CLIENT-CDF")
        ].iloc[0]
        usd_client = report["balance_clients"].loc[
            report["balance_clients"]["customer_id"].eq("CLIENT-USD")
        ].iloc[0]
        self.assertEqual(float(cdf_client["solde_epargne_courante_observe"]), 300.0)
        self.assertEqual(float(cdf_client["depots_epargne_observes"]), 100.0)
        self.assertEqual(float(cdf_client["retraits_epargne_observes"]), 0.0)
        self.assertEqual(float(usd_client["encours_principal_observe"]), 20.0)
        self.assertEqual(cdf_client["Nom_client"], "CLIENT CDF")
        # Les montants G2 volontairement differents ne remplacent jamais Turbo.
        self.assertNotEqual(float(cdf_summary["total_debit"]), 900.0)

    def test_filtered_turbo_balance_report_recalculates_the_export_scope(self) -> None:
        client_balance = pd.DataFrame(
            [
                {
                    "customer_id": "CLIENT-A",
                    "Nom_client": "CLIENT A",
                    "telephone": "243811111111",
                    "currency_code": "CDF",
                    "nombre_lignes": 2,
                    "nombre_operations": 1,
                    "operations_a_verifier": 0,
                    "total_debit": 100,
                    "total_credit": 100,
                    "depots_epargne_observes": 100,
                    "retraits_epargne_observes": 0,
                },
                {
                    "customer_id": "CLIENT-B",
                    "Nom_client": "CLIENT B",
                    "telephone": "243822222222",
                    "currency_code": "USD",
                    "nombre_lignes": 2,
                    "nombre_operations": 1,
                    "operations_a_verifier": 0,
                    "total_debit": 20,
                    "total_credit": 20,
                    "depots_epargne_observes": 20,
                    "retraits_epargne_observes": 0,
                },
            ]
        )
        journal_entries = pd.DataFrame(
            [
                {
                    "id": "A-1",
                    "customer_id": "CLIENT-A",
                    "currency_code": "CDF",
                    "account_type": "NORMAL SAVINGS",
                    "dr": 0,
                    "cr": 100,
                    "created_at": pd.Timestamp("2026-07-21 08:00:00"),
                    "cle_operation_turbo": "REF-A",
                },
                {
                    "id": "B-1",
                    "customer_id": "CLIENT-B",
                    "currency_code": "USD",
                    "account_type": "NORMAL SAVINGS",
                    "dr": 0,
                    "cr": 20,
                    "created_at": pd.Timestamp("2026-07-21 09:00:00"),
                    "cle_operation_turbo": "REF-B",
                },
            ]
        )
        report = {
            "periode": pd.DataFrame(
                [{"date_debut": pd.Timestamp("2026-07-21"), "date_fin": pd.Timestamp("2026-07-21 23:59:59")}]
            ),
            "balance_clients": client_balance,
            "journal_ecritures": journal_entries,
        }

        filtered = build_filtered_turbo_balance_report(report, client_balance.iloc[[0]])

        self.assertEqual(filtered["balance_clients"]["customer_id"].tolist(), ["CLIENT-A"])
        self.assertEqual(filtered["synthese"]["currency_code"].tolist(), ["CDF"])
        self.assertEqual(float(filtered["synthese"].iloc[0]["depots_epargne_observes"]), 100.0)
        self.assertEqual(int(filtered["balance_comptes"].iloc[0]["nombre_clients"]), 1)
        self.assertEqual(set(filtered["balance_comptes"]["currency_code"]), {"CDF"})

    def test_mpesa_accounting_analysis_degrades_without_g2_and_exports_targeted_sheets(self) -> None:
        prepared = _sample_prepared_data()

        report = build_mpesa_accounting_analysis(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-01",
        )

        self.assertFalse(report["synthese"].empty)
        self.assertFalse(report["balance_clients"].empty)
        self.assertEqual(int(report["controle_g2"]["transactions_g2_chargees"].sum()), 0)
        export = create_excel_export(
            {
                "accounting_summary": report["synthese"],
                "accounting_client_balances": report["balance_clients"],
                "accounting_account_balance": report["balance_comptes"],
                "accounting_operation_journal": report["journal_operations"],
            }
        )
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertEqual(
            workbook.sheet_names,
            [
                "Compta_Synthese_Turbo",
                "Balance_Clients_Turbo",
                "Balance_Comptes_Turbo",
                "Journal_Operations_Turbo",
            ],
        )

    def test_turbo_balance_word_and_pdf_are_direction_ready(self) -> None:
        from docx import Document
        from pypdf import PdfReader

        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 100,
                        "bal_before": 0,
                        "bal_after": 100,
                        "ref_no": "DEP-001",
                        "description": "Epargne depot",
                        "created_at": "2026-07-01 10:00:00",
                    },
                    {
                        "id": 2,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 100,
                        "cr": 0,
                        "bal_before": 0,
                        "bal_after": 0,
                        "ref_no": "DEP-001",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-01 10:00:00",
                    },
                    {
                        "id": 3,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 40,
                        "cr": 0,
                        "bal_before": 100,
                        "bal_after": 60,
                        "ref_no": "RET-001",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-02 11:00:00",
                    },
                    {
                        "id": 4,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 40,
                        "bal_before": 0,
                        "bal_after": 0,
                        "ref_no": "RET-001",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-02 11:00:00",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=pd.DataFrame(),
        )
        report = build_mpesa_accounting_analysis(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-01",
        )

        word = create_turbo_balance_word(
            report,
            period_start="2026-07-01",
            period_end="2026-07-01",
            generated_at=pd.Timestamp("2026-07-02 08:00:00"),
        )
        pdf = create_turbo_balance_pdf(
            report,
            period_start="2026-07-01",
            period_end="2026-07-01",
            generated_at=pd.Timestamp("2026-07-02 08:00:00"),
        )

        self.assertTrue(word.startswith(b"PK"))
        self.assertTrue(pdf.startswith(b"%PDF-"))
        document = Document(BytesIO(word))
        self.assertLess(document.sections[0].page_width, document.sections[0].page_height)
        word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Balance auxiliaire observée", word_text)
        self.assertNotIn("Source des montants", word_text)
        self.assertNotIn("balance générale certifiée", word_text)
        self.assertNotIn("Balance des mouvements par type de compte", word_text)
        criteria_header = document.tables[0]
        self.assertEqual(criteria_header.cell(0, 1).text, "Critères")
        criteria = criteria_header.cell(1, 1).tables[0]
        criteria_labels = [row.cells[0].text for row in criteria.rows]
        self.assertEqual(
            criteria_labels,
            ["Date du :", "Au :", "Clients :", "Périmètre :", "Devise(s) :"],
        )
        client_table = next(
            table
            for table in document.tables
            if table.rows[0].cells[0].text == "Client"
        )
        self.assertEqual(
            [cell.text for cell in client_table.rows[0].cells],
            [
                "Client",
                "Nom du client",
                "Téléphone",
                "Devise",
                "Dépôts épargne",
                "Retraits épargne",
                "Mouvement net épargne",
                "Épargne courante",
                "DAT",
                "Principal crédit",
            ],
        )
        pdf_text = "\n".join(
            page.extract_text() or "" for page in PdfReader(BytesIO(pdf)).pages
        )
        normalized_pdf_text = " ".join(pdf_text.split())
        first_pdf_page = PdfReader(BytesIO(pdf)).pages[0]
        self.assertLess(float(first_pdf_page.mediabox.width), float(first_pdf_page.mediabox.height))
        self.assertIn("Balance auxiliaire observée", normalized_pdf_text)
        self.assertNotIn("Source des montants", normalized_pdf_text)
        self.assertNotIn("balance générale certifiée", normalized_pdf_text)
        self.assertNotIn("Balance des mouvements par type de compte", normalized_pdf_text)
        self.assertNotIn("Type de compte Turbo", normalized_pdf_text)
        self.assertIn("Critères", normalized_pdf_text)
        self.assertIn("Clients :", normalized_pdf_text)
        self.assertIn("Périmètre :", normalized_pdf_text)
        self.assertIn("Devise(s) :", normalized_pdf_text)

    def test_turbo_deposit_withdrawal_pivot_exports_vodacom_style_tracking(self) -> None:
        from docx import Document

        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 100,
                        "bal_before": 0,
                        "bal_after": 100,
                        "ref_no": "DEP-001",
                        "description": "Epargne depot",
                        "created_at": "2026-07-01 10:00:00",
                    },
                    {
                        "id": 2,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 100,
                        "cr": 0,
                        "bal_before": 0,
                        "bal_after": 0,
                        "ref_no": "DEP-001",
                        "description": "M-Pesa Compte",
                        "created_at": "2026-07-01 10:00:00",
                    },
                    {
                        "id": 3,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 40,
                        "cr": 0,
                        "bal_before": 100,
                        "bal_after": 60,
                        "ref_no": "RET-001",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-02 11:00:00",
                    },
                    {
                        "id": 4,
                        "customer_id": "1001",
                        "msisdn1": "243812345678",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SAVE-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 40,
                        "bal_before": 0,
                        "bal_after": 0,
                        "ref_no": "RET-001",
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-02 11:00:00",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=pd.DataFrame(),
        )
        report = build_mpesa_accounting_analysis(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-02",
        )
        pivot_report = build_filtered_turbo_deposit_withdrawal_pivot_report(
            report,
            report["balance_clients"],
        )
        pivot = pivot_report["suivi_depots_retraits_pivot"]

        self.assertEqual(len(pivot), 2)
        self.assertEqual(pivot["operation"].tolist(), ["Depot", "Retrait"])
        self.assertIn("01/07/2026", pivot.columns)
        self.assertIn("02/07/2026", pivot.columns)
        self.assertEqual(pivot["currency_code"].unique().tolist(), ["CDF"])
        summary = pivot_report["synthese"].iloc[0]
        self.assertEqual(int(summary["nombre_jours_periode"]), 2)
        self.assertEqual(int(summary["nombre_clients"]), 1)
        self.assertEqual(float(summary["total_depots"]), 100.0)
        self.assertEqual(float(summary["total_retraits"]), 40.0)
        self.assertEqual(float(summary["solde_net_periode"]), 60.0)
        self.assertEqual(
            pivot.loc[pivot["operation"].eq("Depot"), "score"].iloc[0],
            "1 / 2 jour(s) (50.0%)",
        )

        word = create_turbo_balance_word(
            pivot_report,
            period_start="2026-07-01",
            period_end="2026-07-02",
            generated_at=pd.Timestamp("2026-07-03 08:00:00"),
            balance_by_date=True,
        )
        excel = create_excel_export(
            pivot_report,
            print_orientation="landscape",
        )

        document = Document(BytesIO(word))
        self.assertLess(
            document.sections[0].page_width,
            document.sections[0].page_height,
        )
        word_text = "\n".join(
            paragraph.text for paragraph in document.paragraphs
        )
        self.assertIn("Suivi des dépôts et retraits", word_text)
        criteria = document.tables[0].cell(1, 1).tables[0]
        self.assertEqual(
            [row.cells[0].text for row in criteria.rows],
            [
                "Date du :",
                "Au :",
                "Clients :",
                "Jours couverts :",
                "Périmètre :",
                "Devise(s) :",
            ],
        )
        pivot_table = next(
            table
            for table in document.tables
            if table.rows[0].cells[0].text == "Client"
        )
        self.assertIn("01/07", [cell.text for cell in pivot_table.rows[0].cells])
        self.assertIn("02/07", [cell.text for cell in pivot_table.rows[0].cells])

        workbook = pd.ExcelFile(BytesIO(excel), engine="openpyxl")
        self.assertIn("Synthese", workbook.sheet_names)
        self.assertIn("Suivi_Depots_Retraits", workbook.sheet_names)
        exported_pivot = pd.read_excel(
            BytesIO(excel),
            sheet_name="Suivi_Depots_Retraits",
            engine="openpyxl",
        )
        self.assertIn("01/07/2026", exported_pivot.columns)
        self.assertIn("02/07/2026", exported_pivot.columns)
        self.assertEqual(exported_pivot["operation"].tolist(), ["Depot", "Retrait"])

    def test_mpesa_management_dashboard_builds_actionable_microfinance_views(self) -> None:
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "PILOT-001",
                        "Completion Time": "2026-01-01 09:00:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Paid In": 100,
                        "Withdrawn": 0,
                        "Balance": 1000,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "PILOT-002",
                        "Completion Time": "2026-07-10 10:00:00",
                        "Opposite Party": "0811111111 - CLIENT A",
                        "Currency": "CDF",
                        "Paid In": 200,
                        "Withdrawn": 0,
                        "Balance": 1200,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "PILOT-003",
                        "Completion Time": "2026-07-14 11:00:00",
                        "Opposite Party": "0822222222 - CLIENT B",
                        "Currency": "CDF",
                        "Paid In": 0,
                        "Withdrawn": 50,
                        "Balance": 1150,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouB2C",
                    },
                    {
                        "Receipt No.": "PILOT-004",
                        "Completion Time": "2026-03-01 12:00:00",
                        "Opposite Party": "0833333333 - CLIENT C",
                        "Currency": "USD",
                        "Paid In": 10,
                        "Withdrawn": 0,
                        "Balance": 100,
                        "Transaction Status": "Completed",
                        "Details": "BisouBisouC2B",
                    },
                ]
            )
        )
        loans = prepare_loans(
            pd.DataFrame(
                [
                    {
                        "loan_id": "LN-001",
                        "customer_id": "A",
                        "msisdn1": "0811111111",
                        "currency_code": "CDF",
                        "loan_amount": 1000,
                        "loan_balance": 500,
                        "amount_paid": 500,
                        "due_date": "2026-06-01",
                        "updated_at": "2026-07-15",
                    },
                    {
                        "loan_id": "LN-002",
                        "customer_id": "B",
                        "msisdn1": "0822222222",
                        "currency_code": "CDF",
                        "loan_amount": 500,
                        "loan_balance": 0,
                        "amount_paid": 500,
                        "due_date": "2026-07-01",
                        "updated_at": "2026-07-15",
                    },
                    {
                        "loan_id": "LN-003",
                        "customer_id": "C",
                        "msisdn1": "0833333333",
                        "currency_code": "USD",
                        "loan_amount": 100,
                        "loan_balance": 100,
                        "amount_paid": 0,
                        "due_date": "2026-07-14",
                        "updated_at": "2026-07-15",
                    },
                ]
            )
        )
        fixed = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": "A",
                        "msisdn": "0811111111",
                        "product_name": "1 Month",
                        "account_type": "FIXED SAVINGS",
                        "balance": 200,
                        "currency_code": "CDF",
                        "date_approved": "2026-07-10",
                        "maturity_date": "2026-07-20",
                    },
                    {
                        "customer_id": "B",
                        "msisdn": "0822222222",
                        "product_name": "1 Month",
                        "account_type": "FIXED SAVINGS",
                        "balance": 500,
                        "currency_code": "CDF",
                        "date_approved": "2026-05-30",
                        "maturity_date": "2026-06-30",
                    },
                    {
                        "customer_id": "C",
                        "msisdn": "0833333333",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 50,
                        "currency_code": "USD",
                        "date_approved": "2026-06-01",
                        "maturity_date": "2026-09-01",
                    },
                ]
            )
        )
        perfect = prepare_perfect_clients(
            pd.DataFrame(
                [
                    {"id_client": "PA", "Phone_Prefixe": "243811111111", "nom_complet": "CLIENT A"},
                    {"id_client": "PB", "Phone_Prefixe": "243822222222", "nom_complet": "CLIENT B"},
                    {"id_client": "PD", "Phone_Prefixe": "243844444444", "nom_complet": "CLIENT D"},
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=fixed,
            loans=loans,
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
            perfect_clients=perfect,
        )

        report = build_mpesa_management_dashboard(
            prepared,
            as_of_date="2026-07-15",
            dat_annual_interest_rate_pct=12.0,
        )

        credit_cdf = report["credit_synthese"].loc[
            report["credit_synthese"]["currency_code"].eq("CDF")
        ].iloc[0]
        self.assertEqual(float(credit_cdf["encours_total"]), 500.0)
        self.assertEqual(float(credit_cdf["par_30j_pct"]), 100.0)
        self.assertEqual(int(credit_cdf["credits_retard_30j"]), 1)

        maturity_buckets = set(report["dat_echeances_synthese"]["tranche_echeance"])
        self.assertIn("Echu", maturity_buckets)
        self.assertIn("0 a 7 jours", maturity_buckets)
        self.assertTrue(report["dat_echeances_detail"]["taux_interet_annuel_pct"].eq(12.0).all())
        self.assertTrue(report["dat_echeances_detail"]["interet_estime_echeance"].notna().all())

        # G2 ne doit alimenter aucun montant, meme lorsqu'il est charge.
        self.assertTrue(report["flux_synthese"].empty)
        self.assertTrue(report["activite_clients"].empty)
        self.assertTrue(report["alertes_transactions"].empty)
        self.assertTrue(report["perfect_adoption_synthese"].empty)
        g2_source = report["sources"].loc[
            report["sources"]["source"].eq("Transactions M-PESA_G2")
        ].iloc[0]
        self.assertFalse(bool(g2_source["intervient_dans_les_montants"]))

        export_keys = [
            "credit_synthese",
            "credit_detail",
            "dat_echeances_detail",
        ]
        export = create_excel_export({key: report[key] for key in export_keys if not report[key].empty})
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertEqual(
            workbook.sheet_names,
            [
                "Pilotage_Credit_Turbo",
                "Credits_Risque_Turbo",
                "Echeances_DAT_Turbo",
            ],
        )

    def test_mpesa_statistics_g2_quality_combines_1441_and_15558_without_using_loan_requests_as_b2c_anomalies(self) -> None:
        turbo = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": 100,
                        "msisdn1": "243811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-ENTRY",
                        "currency_code": "USD",
                        "dr": 100,
                        "cr": 0,
                        "bal_before": 100,
                        "bal_after": 0,
                        "ref_no": "G2-ENTRY",
                        "description": "Epargne depot",
                        "created_at": "2026-07-20 09:00:00",
                    },
                    {
                        "id": 2,
                        "customer_id": 100,
                        "msisdn1": "243811111111",
                        "account_type": "NORMAL SAVINGS",
                        "reference_id": "SA-OUTPUT",
                        "currency_code": "USD",
                        "dr": 200,
                        "cr": 0,
                        "bal_before": 300,
                        "bal_after": 100,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-21 10:00:00",
                    },
                    {
                        "id": 3,
                        "customer_id": 100,
                        "msisdn1": "243811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "SA-OUTPUT",
                        "currency_code": "USD",
                        "dr": 0,
                        "cr": 200,
                        "bal_before": 0,
                        "bal_after": 200,
                        "ref_no": pd.NA,
                        "description": "Retrait Vers M-Pesa",
                        "created_at": "2026-07-21 10:00:00",
                    },
                ]
            )
        )
        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "G2-ENTRY",
                        "Initiation Time": "2026-07-20 09:00:00",
                        "Completion Time": "2026-07-20 09:00:00",
                        "Details": "Bisou Bisou C2B",
                        "Reason Type": "BisouBisouC2B",
                        "Transaction Status": "Completed",
                        "Currency": "USD",
                        "Paid In": 100,
                        "Withdrawn": 0,
                        "Opposite Party": "243811111111 - CLIENT TEST",
                        "fichier_source_g2": "ORG_1441__All.xlsx",
                    },
                    {
                        "Receipt No.": "G2-ENTRY-DECLINED",
                        "Initiation Time": "2026-07-20 09:05:00",
                        "Completion Time": "2026-07-20 09:05:00",
                        "Details": "Bisou Bisou C2B",
                        "Reason Type": "BisouBisouC2B",
                        "Transaction Status": "Declined",
                        "Currency": "USD",
                        "Paid In": 50,
                        "Withdrawn": 0,
                        "Opposite Party": "243811111111 - CLIENT TEST",
                        "fichier_source_g2": "ORG_1441__All.xlsx",
                    },
                    {
                        "Receipt No.": "G2-OUTPUT",
                        "Initiation Time": "2026-07-21 10:00:00",
                        "Completion Time": "2026-07-21 10:00:00",
                        "Details": "Bisou Bisou B2C",
                        "Reason Type": "BisouBisouB2C",
                        "Transaction Status": "Completed",
                        "Currency": "USD",
                        "Paid In": 0,
                        "Withdrawn": -200,
                        "Opposite Party": "243811111111 - CLIENT TEST",
                        "fichier_source_g2": "ORG_15558__All.xlsx",
                    },
                    {
                        "Receipt No.": "G2-LOAN",
                        "Initiation Time": "2026-07-22 11:00:00",
                        "Completion Time": "2026-07-22 11:00:00",
                        "Details": "Bisou Bisou Loan payment",
                        "Reason Type": "BisouBisouLoanRequest",
                        "Transaction Status": "Completed",
                        "Currency": "USD",
                        "Paid In": 0,
                        "Withdrawn": -93,
                        "Opposite Party": "243811111111 - CLIENT TEST",
                        "fichier_source_g2": "ORG_15558__All.xlsx",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=turbo,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=g2,
        )

        report = build_mpesa_g2_statistics_quality(
            prepared,
            date_start="2026-07-20",
            date_end="2026-07-22",
        )

        coverage = report["couverture"].iloc[0]
        self.assertEqual(
            coverage["couverture_g2"],
            "Complète - entrées 1441 et sorties 15558",
        )
        quality = report["qualite_rapprochement"].set_index("categorie")
        entry = quality.loc["Entrées et remboursements [1441]"]
        output = quality.loc["Sorties B2C [15558]"]
        loan = quality.loc["Versements de prêts [15558]"]
        self.assertEqual(int(entry["operations_terminees"]), 1)
        self.assertEqual(int(entry["operations_rapprochees"]), 1)
        self.assertEqual(float(entry["taux_rapprochement_pct"]), 100.0)
        self.assertEqual(int(output["operations_terminees"]), 1)
        self.assertEqual(int(output["operations_rapprochees"]), 1)
        self.assertEqual(float(output["taux_rapprochement_pct"]), 100.0)
        self.assertEqual(int(loan["operations_terminees"]), 1)
        self.assertEqual(int(loan["operations_non_rapprochees"]), 0)
        self.assertTrue(pd.isna(loan["taux_rapprochement_pct"]))
        self.assertTrue(report["non_rapprochees"].empty)
        statuses = report["statuts"].set_index("statut_g2")
        self.assertEqual(int(statuses.loc["Completed", "nombre_operations"]), 3)
        self.assertEqual(int(statuses.loc["Declined", "nombre_operations"]), 1)

    def test_mpesa_statistics_report_is_turbo_first_and_exports_word(self) -> None:
        from docx import Document

        base = _sample_customer_transaction_analysis_data()
        customers = prepare_customers(
            pd.DataFrame(
                [
                    {
                        "customer_id": "CLIENT-ANALYSE",
                        "msisdn1": "0812345678",
                        "created_at": "2026-01-15",
                    },
                    {
                        "customer_id": "CLIENT-APRES-PERIODE",
                        "msisdn1": "0899999999",
                        "created_at": "2026-08-05",
                    }
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=base.transactions,
            current_savings=base.current_savings,
            fixed_savings=base.fixed_savings,
            loans=base.loans,
            load_report=base.load_report,
            customers=customers,
            g2_transactions=prepare_g2_transactions(
                pd.DataFrame(
                    [
                        {
                            "Receipt No.": "G2-STATS-IGNORED",
                            "Completion Time": "2026-07-03 08:00:00",
                            "Opposite Party": "0812345678 - CLIENT ANALYSE",
                            "Currency": "CDF",
                            "Paid In": 999999,
                            "Withdrawn": 0,
                            "Balance": 999999,
                            "Transaction Status": "Completed",
                            "Details": "BisouBisouC2B",
                        }
                    ]
                )
            ),
        )

        report = build_mpesa_statistics_report(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-05",
            frequency="Jour",
            comparison_period="Période filtrée",
        )

        source_priority = report["priorite_sources"].set_index("source")
        self.assertEqual(
            source_priority.loc["Transactions [Solution Numérique]", "niveau_importance"],
            "Indispensable",
        )
        self.assertEqual(
            source_priority.loc["Transactions [G2] (facultatif)", "niveau_importance"],
            "Facultatif utile",
        )
        self.assertFalse(
            "G2-STATS-IGNORED" in report["operations_turbo"].get("event_reference", pd.Series(dtype=str)).astype(str).tolist()
        )

        overview = report["vue_ensemble"].loc[
            report["vue_ensemble"]["currency_code"].eq("CDF")
        ].iloc[0]
        self.assertEqual(int(overview["clients_turbo_charges"]), 2)
        self.assertEqual(int(overview["clients_turbo_connus"]), 1)
        self.assertEqual(int(overview["clients_turbo_actifs"]), 1)
        client_indicators = report["clients_indicateurs"].set_index("indicateur")
        self.assertEqual(
            int(client_indicators.loc["Clients du fichier Customers charge", "valeur"]),
            2,
        )
        self.assertGreater(float(overview["volume_total_transactions"]), 0)
        self.assertEqual(float(overview["chiffre_affaires_observe"]), 15.0)

        growth = report["clients_croissance"]
        self.assertEqual(int(growth.iloc[-1]["clients_turbo_cumules"]), 1)
        self.assertIn("epargne_dat_portefeuille", report)
        self.assertFalse(report["epargne_dat_portefeuille"].empty)
        self.assertEqual(report["perimetre_annuel"], "Ensemble des années")

        document = Document(BytesIO(create_mpesa_statistics_word(report)))
        word_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
        )
        self.assertIn("Rapport statistiques - Solution Numérique", word_text)
        self.assertIn("Clients du fichier Customers charge", word_text)
        self.assertIn("1. Clients", word_text)
        self.assertIn("2. Comptes ouverts et comptes bloques", word_text)
        self.assertIn("3. Credits", word_text)
        self.assertIn("4. Transactions", word_text)
        self.assertIn("4.1 Qualité du rapprochement G2", word_text)
        self.assertNotIn("Sources et importance", word_text)
        self.assertIn("Annexe 1. Vue d'ensemble", word_text)
        self.assertNotIn("Annexe 3. Definitions", word_text)
        self.assertIn("Chiffre d'affaires observe", word_text)
        self.assertIn("Solution Numérique uniquement", word_text)
        self.assertIn("Période filtrée", word_text)
        self.assertIn(
            "Comparaison avec la même période de l'année précédente",
            word_text,
        )
        self.assertIn("comparaison_annee_precedente", report)
        self.assertIn("Ensemble des années", document._element.xml)
        self.assertNotIn("Graphiques de synthese", word_text)
        client_growth_tables = [
            table
            for table in document.tables
            if {"Periode", "Nouveaux clients", "Clients cumules"}.issubset(
                {cell.text for cell in table.rows[0].cells}
            )
        ]
        self.assertEqual(len(client_growth_tables), 1)
        self.assertNotIn(
            "Source",
            {cell.text for cell in client_growth_tables[0].rows[0].cells},
        )

    def test_turbo_only_g2_dat_uses_consolidated_loan_repayment_amount(self) -> None:
        transactions = prepare_transactions(
            pd.DataFrame(
                [
                    {
                        "id": 1,
                        "customer_id": "1001",
                        "msisdn1": "243811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "LN-001",
                        "currency_code": "CDF",
                        "dr": 7000,
                        "cr": 0,
                        "bal_before": 20000,
                        "bal_after": 13000,
                        "ref_no": "REP-001",
                        "description": "M-Pesa Remboursement",
                        "created_at": "2026-07-28 17:21:26",
                    },
                    {
                        "id": 2,
                        "customer_id": "1001",
                        "msisdn1": "243811111111",
                        "account_type": "MPESA ACCOUNT",
                        "reference_id": "LN-001",
                        "currency_code": "CDF",
                        "dr": 7000,
                        "cr": 0,
                        "bal_before": 13000,
                        "bal_after": 6000,
                        "ref_no": "REP-001",
                        "description": "M-Pesa Remboursement",
                        "created_at": "2026-07-28 17:21:26",
                    },
                    {
                        "id": 3,
                        "customer_id": "1001",
                        "msisdn1": "243811111111",
                        "account_type": "PRINCIPLE",
                        "reference_id": "LN-001",
                        "currency_code": "CDF",
                        "dr": 0,
                        "cr": 14000,
                        "bal_before": 0,
                        "bal_after": 14000,
                        "ref_no": "REP-001",
                        "description": "Remboursement du principal",
                        "created_at": "2026-07-28 17:21:26",
                    },
                ]
            )
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            customers=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        g2_report = build_g2_daily_savings_report(prepared)
        g2_detail = g2_report["detail"]
        g2_repayment = g2_detail.loc[
            g2_detail["details_rapport"].astype(str).eq("Remboursement prets")
        ]
        finance = build_mpesa_turbo_financial_analysis(
            prepared,
            date_start="2026-07-28",
            date_end="2026-07-28",
        )
        flow = finance["flux_synthese"].iloc[0]

        self.assertEqual(float(g2_repayment["montant"].sum()), 14000.0)
        self.assertEqual(float(flow["remboursements_observes"]), 14000.0)

    def test_year_scope_filters_flows_and_preserves_snapshot_positions(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(
                {
                    "id": ["T24", "T25", "T26", "T27", "T-UNDATED"],
                    "created_at": [
                        "2024-12-31 23:59:59",
                        "2025-01-01 00:00:00",
                        "2026-12-31 23:59:59",
                        "2027-01-01 00:00:00",
                        None,
                    ],
                }
            ),
            g2_transactions=pd.DataFrame(
                {
                    "receipt_no": ["G24", "G25", "G26", "G27"],
                    "completion_time": [
                        "2024-06-01",
                        "2025-06-01",
                        "2026-06-01",
                        "2027-06-01",
                    ],
                }
            ),
            current_savings=pd.DataFrame(
                {
                    "savings_id": ["S24", "S26", "S27", "S-UNDATED"],
                    "date_activated": [
                        "2024-01-01",
                        "2026-05-01",
                        "2027-01-01",
                        None,
                    ],
                }
            ),
            fixed_savings=pd.DataFrame(
                {
                    "savings_id": ["D24", "D27"],
                    "date_approved": ["2024-01-01", "2027-01-01"],
                }
            ),
            fixed_savings_control=pd.DataFrame(
                {
                    "savings_id": ["DC24", "DC27"],
                    "date_approved": ["2024-01-01", "2027-01-01"],
                }
            ),
            loans=pd.DataFrame(
                {
                    "loan_id": ["L24", "L26", "L27"],
                    "created_at": ["2024-01-01", "2026-01-01", "2027-01-01"],
                }
            ),
            customers=pd.DataFrame(
                {
                    "customer_id": ["C24", "C26", "C27"],
                    "created_at": ["2024-01-01", "2026-01-01", "2027-01-01"],
                }
            ),
            perfect_clients=pd.DataFrame({"client_id": ["P1", "P2"]}),
            load_report=pd.DataFrame(),
            cache_fingerprint="uploads-1",
        )

        scoped = scope_mpesa_prepared_data_by_year(
            prepared,
            mode="Plage d'années",
            start_year=2026,
            end_year=2025,
        )

        self.assertEqual(scoped.year_scope_label, "2025-2026")
        self.assertEqual(scoped.year_scope_start, pd.Timestamp("2025-01-01"))
        self.assertEqual(scoped.year_scope_end, pd.Timestamp("2026-12-31"))
        self.assertEqual(scoped.transactions["id"].tolist(), ["T25", "T26"])
        self.assertEqual(scoped.g2_transactions["receipt_no"].tolist(), ["G25", "G26"])
        self.assertEqual(
            scoped.current_savings["savings_id"].tolist(),
            ["S24", "S26", "S-UNDATED"],
        )
        self.assertEqual(scoped.fixed_savings["savings_id"].tolist(), ["D24"])
        self.assertEqual(scoped.fixed_savings_control["savings_id"].tolist(), ["DC24"])
        self.assertEqual(scoped.loans["loan_id"].tolist(), ["L24", "L26"])
        self.assertEqual(scoped.customers["customer_id"].tolist(), ["C24", "C26"])
        self.assertEqual(scoped.perfect_clients["client_id"].tolist(), ["P1", "P2"])
        self.assertIn("|year-scope:2025-2026", scoped.cache_fingerprint)

        full = scope_mpesa_prepared_data_by_year(
            prepared,
            mode="Ensemble des années",
        )
        self.assertEqual(len(full.transactions), len(prepared.transactions))
        self.assertIsNone(full.year_scope_start)
        self.assertIsNone(full.year_scope_end)

    def test_mpesa_comparison_windows_support_microfinance_week_and_filtered_period(self) -> None:
        microfinance = build_mpesa_comparison_windows(
            date_end="2026-07-25",
            date_start="2026-04-20",
            comparison_period="Semaine microfinance (lundi)",
        )
        self.assertEqual(
            microfinance["date_debut_periode_courante"],
            pd.Timestamp("2026-07-20"),
        )
        self.assertEqual(
            microfinance["date_fin_periode_courante"],
            pd.Timestamp("2026-07-25"),
        )
        self.assertEqual(
            microfinance["date_debut_periode_precedente"],
            pd.Timestamp("2026-07-13"),
        )
        self.assertEqual(
            microfinance["date_fin_periode_precedente"],
            pd.Timestamp("2026-07-18"),
        )

        filtered = build_mpesa_comparison_windows(
            date_end="2026-07-25",
            date_start="2026-04-20",
            comparison_period="Période filtrée",
        )
        self.assertEqual(
            filtered["date_debut_periode_courante"],
            pd.Timestamp("2026-04-20"),
        )
        self.assertEqual(
            filtered["date_fin_periode_courante"],
            pd.Timestamp("2026-07-25"),
        )
        self.assertEqual(
            filtered["date_debut_periode_precedente"],
            pd.Timestamp("2026-01-13"),
        )
        self.assertEqual(
            filtered["date_fin_periode_precedente"],
            pd.Timestamp("2026-04-19"),
        )

    def test_mpesa_weekly_comparison_uses_two_consecutive_weeks_and_keeps_currencies_separate(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(
                [
                    {"savings_id": "OPEN-PREV", "customer_id": "C1", "created_at": "2026-07-14"},
                    {"savings_id": "OPEN-CURR", "customer_id": "C2", "created_at": "2026-07-20"},
                ]
            ),
            fixed_savings=pd.DataFrame(
                [
                    {"savings_id": "DAT-PREV", "customer_id": "C1", "date_activated": "2026-07-13"},
                    {"savings_id": "DAT-CURR", "customer_id": "C2", "date_activated": "2026-07-21"},
                ]
            ),
            loans=pd.DataFrame(
                [
                    {
                        "loan_id": "LN-PREV",
                        "customer_id": "C1",
                        "currency_code": "CDF",
                        "loan_amount": 100,
                        "created_at": "2026-07-14",
                    },
                    {
                        "loan_id": "LN-CURR-1",
                        "customer_id": "C2",
                        "currency_code": "CDF",
                        "loan_amount": 200,
                        "created_at": "2026-07-20",
                    },
                    {
                        "loan_id": "LN-CURR-2",
                        "customer_id": "C3",
                        "currency_code": "CDF",
                        "loan_amount": 300,
                        "created_at": "2026-07-22",
                    },
                ]
            ),
            customers=pd.DataFrame(
                [
                    {"msisdn1": "243810000001", "created_at": "2026-07-14"},
                    {"msisdn1": "243810000002", "created_at": "2026-07-20"},
                    {"msisdn1": "243810000003", "created_at": "2026-07-22"},
                ]
            ),
            load_report=pd.DataFrame(),
        )
        events = pd.DataFrame(
            [
                {
                    "event_key": "EV-PREV",
                    "customer_id": "C1",
                    "currency_code": "CDF",
                    "created_at": "2026-07-14 10:00:00",
                    "montant_entree_bisou": 100,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 10,
                    "depot_dat_mpesa": 50,
                },
                {
                    "event_key": "EV-CURR-1",
                    "customer_id": "C2",
                    "currency_code": "CDF",
                    "created_at": "2026-07-20 10:00:00",
                    "montant_entree_bisou": 200,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 20,
                    "depot_dat_mpesa": 100,
                },
                {
                    "event_key": "EV-CURR-2",
                    "customer_id": "C3",
                    "currency_code": "CDF",
                    "created_at": "2026-07-22 10:00:00",
                    "montant_entree_bisou": 0,
                    "montant_sortie_bisou": 100,
                    "remboursement_mpesa": 0,
                    "depot_dat_mpesa": 0,
                },
                {
                    "event_key": "EV-CURR-USD",
                    "customer_id": "C3",
                    "currency_code": "USD",
                    "created_at": "2026-07-22 11:00:00",
                    "montant_entree_bisou": 5,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 0,
                    "depot_dat_mpesa": 0,
                },
            ]
        )
        lines = pd.DataFrame(
            [
                {
                    "account_type": "BISOU COLLECTION",
                    "currency_code": "CDF",
                    "created_at": "2026-07-14",
                    "dr": 10,
                    "cr": 0,
                },
                {
                    "account_type": "BISOU COLLECTION",
                    "currency_code": "CDF",
                    "created_at": "2026-07-20",
                    "dr": 30,
                    "cr": 0,
                },
                {
                    "account_type": "BISOU COLLECTION",
                    "currency_code": "USD",
                    "created_at": "2026-07-22",
                    "dr": 2,
                    "cr": 0,
                },
            ]
        )

        comparison = build_mpesa_weekly_comparison(
            prepared,
            as_of_date="2026-07-22",
            comparison_period="7 jours glissants",
            turbo_events=events,
            turbo_transaction_lines=lines,
        )
        indexed = comparison.set_index(["indicator_key", "currency_code"])

        clients = indexed.loc[("clients_actifs", "")]
        self.assertEqual(float(clients["valeur_semaine_courante"]), 2.0)
        self.assertEqual(float(clients["valeur_semaine_precedente"]), 1.0)
        self.assertEqual(float(clients["evolution_pct"]), 100.0)
        self.assertEqual(
            pd.Timestamp(clients["date_debut_semaine_courante"]),
            pd.Timestamp("2026-07-16"),
        )
        self.assertEqual(
            pd.Timestamp(clients["date_debut_semaine_precedente"]),
            pd.Timestamp("2026-07-09"),
        )

        credits = indexed.loc[("nouveaux_credits", "")]
        self.assertEqual(float(credits["valeur_semaine_courante"]), 2.0)
        self.assertEqual(float(credits["valeur_semaine_precedente"]), 1.0)
        credit_amount = indexed.loc[("montant_nouveaux_credits", "CDF")]
        self.assertEqual(float(credit_amount["valeur_semaine_courante"]), 500.0)
        self.assertEqual(float(credit_amount["valeur_semaine_precedente"]), 100.0)

        volume_cdf = indexed.loc[("volume_transactions", "CDF")]
        volume_usd = indexed.loc[("volume_transactions", "USD")]
        self.assertEqual(float(volume_cdf["valeur_semaine_courante"]), 300.0)
        self.assertEqual(float(volume_cdf["valeur_semaine_precedente"]), 100.0)
        self.assertEqual(float(volume_usd["valeur_semaine_courante"]), 5.0)
        self.assertEqual(float(volume_usd["valeur_semaine_precedente"]), 0.0)
        self.assertTrue(pd.isna(volume_usd["evolution_pct"]))

        turnover_cdf = indexed.loc[("chiffre_affaires_observe", "CDF")]
        turnover_usd = indexed.loc[("chiffre_affaires_observe", "USD")]
        self.assertEqual(float(turnover_cdf["valeur_semaine_courante"]), 30.0)
        self.assertEqual(float(turnover_cdf["valeur_semaine_precedente"]), 10.0)
        self.assertEqual(float(turnover_usd["valeur_semaine_courante"]), 2.0)
        self.assertEqual(float(turnover_usd["valeur_semaine_precedente"]), 0.0)

    def test_mpesa_year_over_year_comparison_uses_same_calendar_dates(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(
                [
                    {
                        "savings_id": "OPEN-2025",
                        "customer_id": "C1",
                        "created_at": "2025-09-01",
                    },
                    {
                        "savings_id": "OPEN-2026",
                        "customer_id": "C2",
                        "created_at": "2026-09-30",
                    },
                ]
            ),
            fixed_savings=pd.DataFrame(
                [
                    {
                        "savings_id": "DAT-2025",
                        "customer_id": "C1",
                        "date_activated": "2025-09-01",
                    },
                    {
                        "savings_id": "DAT-2026",
                        "customer_id": "C2",
                        "date_activated": "2026-09-30",
                    },
                ]
            ),
            loans=pd.DataFrame(
                [
                    {
                        "loan_id": "LN-2025",
                        "customer_id": "C1",
                        "currency_code": "CDF",
                        "loan_amount": 100,
                        "created_at": "2025-09-01",
                    },
                    {
                        "loan_id": "LN-2026",
                        "customer_id": "C2",
                        "currency_code": "CDF",
                        "loan_amount": 150,
                        "created_at": "2026-09-30",
                    },
                ]
            ),
            customers=pd.DataFrame(
                [
                    {
                        "msisdn1": "243810000001",
                        "created_at": "2025-09-01",
                    },
                    {
                        "msisdn1": "243810000002",
                        "created_at": "2026-09-30",
                    },
                ]
            ),
            load_report=pd.DataFrame(),
        )
        events = pd.DataFrame(
            [
                {
                    "event_key": "EV-2025-CDF",
                    "customer_id": "C1",
                    "currency_code": "CDF",
                    "created_at": "2025-09-01 10:00:00",
                    "montant_entree_bisou": 100,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 10,
                    "depot_dat_mpesa": 20,
                },
                {
                    "event_key": "EV-2026-CDF",
                    "customer_id": "C2",
                    "currency_code": "CDF",
                    "created_at": "2026-09-30 10:00:00",
                    "montant_entree_bisou": 160,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 15,
                    "depot_dat_mpesa": 25,
                },
                {
                    "event_key": "EV-2026-USD",
                    "customer_id": "C3",
                    "currency_code": "USD",
                    "created_at": "2026-09-30 11:00:00",
                    "montant_entree_bisou": 10,
                    "montant_sortie_bisou": 0,
                    "remboursement_mpesa": 0,
                    "depot_dat_mpesa": 0,
                },
            ]
        )
        lines = pd.DataFrame(
            [
                {
                    "account_type": "BISOU COLLECTION",
                    "currency_code": "CDF",
                    "created_at": "2025-09-01",
                    "dr": 10,
                    "cr": 0,
                },
                {
                    "account_type": "BISOU COLLECTION",
                    "currency_code": "CDF",
                    "created_at": "2026-09-30",
                    "dr": 15,
                    "cr": 0,
                },
            ]
        )

        comparison = build_mpesa_year_over_year_comparison(
            prepared,
            date_start="2026-09-01",
            date_end="2026-09-30",
            turbo_events=events,
            turbo_transaction_lines=lines,
        )
        indexed = comparison.set_index(["indicator_key", "currency_code"])

        clients = indexed.loc[("clients_actifs", "")]
        self.assertEqual(float(clients["valeur_semaine_courante"]), 2.0)
        self.assertEqual(float(clients["valeur_semaine_precedente"]), 1.0)
        self.assertEqual(float(clients["evolution_pct"]), 100.0)
        self.assertEqual(
            pd.Timestamp(clients["date_debut_semaine_courante"]),
            pd.Timestamp("2026-09-01"),
        )
        self.assertEqual(
            pd.Timestamp(clients["date_debut_semaine_precedente"]),
            pd.Timestamp("2025-09-01"),
        )
        self.assertEqual(
            pd.Timestamp(clients["date_fin_semaine_precedente"]),
            pd.Timestamp("2025-09-30"),
        )
        cdf_volume = indexed.loc[("volume_transactions", "CDF")]
        usd_volume = indexed.loc[("volume_transactions", "USD")]
        self.assertEqual(float(cdf_volume["valeur_semaine_courante"]), 160.0)
        self.assertEqual(float(cdf_volume["valeur_semaine_precedente"]), 100.0)
        self.assertEqual(float(cdf_volume["evolution_pct"]), 60.0)
        self.assertEqual(float(usd_volume["valeur_semaine_courante"]), 10.0)
        self.assertEqual(float(usd_volume["valeur_semaine_precedente"]), 0.0)
        self.assertTrue(pd.isna(usd_volume["evolution_pct"]))
        self.assertTrue(
            comparison["periode_comparaison"]
            .eq("Même période de l'année précédente")
            .all()
        )

    def test_mpesa_statistics_word_keeps_money_by_currency(self) -> None:
        from docx import Document

        report = {
            "date_debut": pd.Timestamp("2026-07-01"),
            "date_fin": pd.Timestamp("2026-07-31"),
            "frequence": "Mois",
            "vue_ensemble": pd.DataFrame(
                [
                    {
                        "currency_code": "CDF",
                        "clients_turbo_connus": 3,
                        "clients_turbo_actifs": 2,
                        "operations": 5,
                        "volume_total_transactions": 100.0,
                        "chiffre_affaires_observe": 4.0,
                    },
                    {
                        "currency_code": "USD",
                        "clients_turbo_connus": 3,
                        "clients_turbo_actifs": 1,
                        "operations": 2,
                        "volume_total_transactions": 10.0,
                        "chiffre_affaires_observe": 2.0,
                    },
                ]
            ),
            "activite_evolution": pd.DataFrame(),
            "clients_croissance": pd.DataFrame(),
            "chiffre_affaires": pd.DataFrame(),
            "epargne_dat_portefeuille": pd.DataFrame(
                [
                    {"currency_code": "CDF", "famille": "Compte ouvert", "nombre_comptes": 2, "solde_total": 100.0},
                    {"currency_code": "USD", "famille": "Compte ouvert", "nombre_comptes": 1, "solde_total": 10.0},
                ]
            ),
            "credit_synthese": pd.DataFrame(
                [
                    {
                        "currency_code": "CDF",
                        "nombre_credits": 1,
                        "nombre_clients": 1,
                        "montant_credits": 100.0,
                        "montant_rembourse": 40.0,
                        "encours_total": 60.0,
                        "encours_retard_30j": 0.0,
                    },
                    {
                        "currency_code": "USD",
                        "nombre_credits": 1,
                        "nombre_clients": 1,
                        "montant_credits": 10.0,
                        "montant_rembourse": 2.0,
                        "encours_total": 8.0,
                        "encours_retard_30j": 1.0,
                    },
                ]
            ),
            "priorite_sources": pd.DataFrame(),
            "definitions": pd.DataFrame(),
        }

        document = Document(BytesIO(create_mpesa_statistics_word(report)))
        word_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
        )

        self.assertIn("aucun montant n'est totalise entre devises", word_text)
        self.assertIn("Volume transactionnel observe | CDF", word_text)
        self.assertIn("Volume transactionnel observe | USD", word_text)
        self.assertIn("Chiffre d'affaires observe | CDF", word_text)
        self.assertIn("Chiffre d'affaires observe | USD", word_text)
        self.assertNotIn("Volume total observe | 110,00", word_text)
        self.assertNotIn("Chiffre d'affaires observe | 6,00", word_text)

    def test_turbo_financial_analysis_uses_one_event_grain_and_never_g2_amounts(self) -> None:
        prepared = _sample_customer_transaction_analysis_data()
        prepared = MpesaPreparedData(
            transactions=prepared.transactions,
            current_savings=prepared.current_savings,
            fixed_savings=prepared.fixed_savings,
            loans=prepared.loans,
            load_report=prepared.load_report,
            g2_transactions=prepare_g2_transactions(
                pd.DataFrame(
                    [
                        {
                            "Receipt No.": "G2-IGNORED",
                            "Completion Time": "2026-07-03 08:00:00",
                            "Opposite Party": "0812345678 - CLIENT TEST",
                            "Currency": "CDF",
                            "Paid In": 9_999_999,
                            "Withdrawn": 0,
                            "Balance": 9_999_999,
                            "Transaction Status": "Completed",
                            "Details": "BisouBisouC2B",
                        }
                    ]
                )
            ),
        )
        journal = build_turbo_operation_events(prepared.transactions)

        report = build_mpesa_turbo_financial_analysis(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-05",
            turbo_events=journal["events"],
            turbo_transaction_lines=journal["lines"],
        )

        self.assertEqual(len(report["operations_turbo"]), 5)
        flow = report["flux_synthese"].iloc[0]
        self.assertEqual(flow["currency_code"], "CDF")
        self.assertEqual(float(flow["montant_entrees"]), 190.0)
        self.assertEqual(float(flow["montant_sorties"]), 100.0)
        self.assertEqual(float(flow["remboursements_observes"]), 40.0)
        self.assertEqual(float(flow["nouveaux_credits_decaissements"]), 100.0)
        new_credit = report["nouveaux_credits_synthese"].iloc[0]
        self.assertEqual(float(new_credit["montant_decaisse_turbo"]), 100.0)
        self.assertEqual(float(new_credit["montant_initial_comptes"]), 100.0)
        self.assertEqual(new_credit["statut_rapprochement"], "Conforme")
        self.assertEqual(float(report["remboursements_synthese"].iloc[0]["montant_rembourse"]), 40.0)
        self.assertTrue(report["dat_sans_credit_actif"].empty)
        self.assertNotIn(9_999_999, report["flux_synthese"].select_dtypes("number").to_numpy())

        export = create_excel_export(
            {
                key: report[key]
                for key in [
                    "flux_synthese",
                    "remboursements_detail",
                    "nouveaux_credits_synthese",
                    "par_tranches_montant",
                    "definitions",
                    "sources",
                ]
            }
        )
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertEqual(
            workbook.sheet_names,
            [
                "Flux_Synthese_Turbo",
                "Remboursements_Pilotage",
                "Nouveaux_Credits_Synthese",
                "PAR_Tranches_Turbo",
                "Definitions_Pilotage",
                "Sources_Pilotage",
            ],
        )

    def test_turbo_financial_analysis_keeps_credit_schema_without_loans(self) -> None:
        prepared_with_loans = _sample_customer_transaction_analysis_data()
        prepared = MpesaPreparedData(
            transactions=prepared_with_loans.transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=prepared_with_loans.load_report,
        )

        report = build_mpesa_turbo_financial_analysis(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-05",
        )

        self.assertFalse(report["flux_synthese"].empty)
        self.assertTrue(report["credit_synthese"].empty)
        self.assertIn("currency_code", report["credit_synthese"].columns)
        self.assertIn("encours_total", report["credit_synthese"].columns)
        self.assertIn("par_30j_pct", report["credit_synthese"].columns)
        self.assertIn("currency_code", report["credit_detail"].columns)

    def test_mpesa_management_dashboard_degrades_cleanly_without_optional_sources(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        report = build_mpesa_management_dashboard(prepared, as_of_date="2026-07-15")

        self.assertTrue(report["credit_synthese"].empty)
        self.assertTrue(report["liquidite_synthese"].empty)
        self.assertTrue(report["activite_clients"].empty)
        self.assertTrue(report["perfect_adoption_detail"].empty)
        self.assertEqual(len(report["sources"]), 5)
        self.assertIn("Customers [Solution Numérique]", report["sources"]["source"].tolist())
        self.assertIn("Transactions M-PESA_G2", report["sources"]["source"].tolist())

    def test_dat_maturity_interest_is_estimated_only_with_a_positive_rate(self) -> None:
        fixed = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": "DAT-1",
                        "msisdn": "0811111111",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 1000,
                        "currency_code": "CDF",
                        "date_approved": "2026-01-01",
                        "maturity_date": "2026-04-01",
                    }
                ]
            )
        )

        with_rate = build_mpesa_dat_maturity_analysis(
            fixed,
            as_of_date="2026-07-15",
            annual_interest_rate_pct=12.0,
        )
        without_rate = build_mpesa_dat_maturity_analysis(
            fixed,
            as_of_date="2026-07-15",
            annual_interest_rate_pct=0.0,
        )

        row = with_rate["detail"].iloc[0]
        self.assertEqual(int(row["duree_contractuelle_jours"]), 90)
        self.assertAlmostEqual(float(row["interet_estime_echeance"]), 1000 * 0.12 * 90 / 365, places=6)
        self.assertAlmostEqual(
            float(row["capital_plus_interet_estime"]),
            1000 + 1000 * 0.12 * 90 / 365,
            places=6,
        )
        self.assertTrue(without_rate["detail"]["interet_estime_echeance"].isna().all())
        self.assertTrue(without_rate["synthese"]["interet_estime_echeance"].isna().all())

    def test_dat_repayment_preparation_defaults_to_eleven_percent(self) -> None:
        fixed = prepare_fixed_savings(
            pd.DataFrame(
                [
                    {
                        "savings_id": "DAT-6M",
                        "customer_id": "CLIENT-6M",
                        "msisdn": "0811111111",
                        "product_name": "6 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 1000,
                        "currency_code": "USD",
                        "date_approved": "2026-01-01",
                        "maturity_date": "2026-07-01",
                    },
                    {
                        "savings_id": "DAT-ECHU",
                        "customer_id": "CLIENT-ECHU",
                        "msisdn": "0822222222",
                        "product_name": "3 Months",
                        "account_type": "FIXED SAVINGS",
                        "balance": 500,
                        "currency_code": "CDF",
                        "date_approved": "2026-03-01",
                        "maturity_date": "2026-06-01",
                    },
                ]
            )
        )

        report = build_mpesa_dat_maturity_analysis(
            fixed,
            as_of_date="2026-06-15",
            preparation_horizon_days=30,
        )
        detail = report["detail"].set_index("savings_id")
        six_months = detail.loc["DAT-6M"]
        expired = detail.loc["DAT-ECHU"]

        self.assertEqual(float(six_months["taux_interet_annuel_pct"]), 11.0)
        self.assertEqual(int(six_months["duree_contractuelle_jours"]), 181)
        self.assertAlmostEqual(
            float(six_months["interet_estime_echeance"]),
            1000 * 0.11 * 181 / 365,
            places=6,
        )
        self.assertAlmostEqual(float(six_months["duree_contractuelle_mois_estimee"]), 6.0, places=1)
        self.assertEqual(int(six_months["jours_avant_echeance"]), 16)
        self.assertTrue(bool(six_months["a_preparer_remboursement"]))
        self.assertEqual(
            six_months["statut_preparation_remboursement"],
            "A preparer sous 30 jours",
        )
        self.assertTrue(bool(expired["a_preparer_remboursement"]))
        self.assertEqual(
            expired["statut_preparation_remboursement"],
            "Echu - remboursement a traiter",
        )

    def test_g2_retention_report_calculates_m1_and_90_days_without_mixing_currencies(self) -> None:
        rows = [
            ("JAN-A", "2026-01-10", "0811111111 - CLIENT A", "CDF", "Completed", "BisouBisouC2B"),
            ("JAN-B", "2026-01-15", "0822222222 - CLIENT B", "CDF", "Completed", "BisouBisouC2B"),
            ("FEB-A", "2026-02-05", "0811111111 - CLIENT A", "CDF", "Completed", "BisouBisouC2B"),
            ("MAR-B", "2026-03-10", "0822222222 - CLIENT B", "CDF", "Completed", "BisouBisouC2BRepayment"),
            ("JAN-U", "2026-01-20", "0833333333 - CLIENT USD", "USD", "Completed", "BisouBisouC2B"),
            ("MAY-C", "2026-05-15", "0844444444 - CLIENT C", "CDF", "Completed", "BisouBisouC2B"),
            ("INTERNAL", "2026-01-12", "0855555555 - INTERNE", "CDF", "Completed", "Super Transaction"),
            ("FAILED", "2026-01-13", "0866666666 - ECHEC", "CDF", "Failed", "BisouBisouC2B"),
        ]
        raw = pd.DataFrame(
            [
                {
                    "Receipt No.": receipt,
                    "Completion Time": date,
                    "Opposite Party": party,
                    "Currency": currency,
                    "Transaction Amount": 1000,
                    "Transaction Status": status,
                    "Details": details,
                }
                for receipt, date, party, currency, status, details in rows
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(raw),
        )

        report = build_g2_retention_report(prepared)
        january = report["mensuelle"].loc[report["mensuelle"]["mois"].eq("2026-01")].set_index("currency_code")

        self.assertEqual(int(january.loc["CDF", "clients_actifs_mois_base"]), 2)
        self.assertEqual(float(january.loc["CDF", "retention_m1_pct"]), 50.0)
        self.assertEqual(float(january.loc["CDF", "retention_90j_pct"]), 100.0)
        self.assertEqual(int(january.loc["USD", "clients_actifs_mois_base"]), 1)
        self.assertEqual(float(january.loc["USD", "retention_m1_pct"]), 0.0)
        self.assertEqual(float(january.loc["USD", "retention_90j_pct"]), 0.0)
        self.assertEqual(len(report["detail_clients"].loc[report["detail_clients"]["mois"].eq("2026-01")]), 3)
        self.assertFalse(report["definitions"].empty)

    def test_g2_retention_report_leaves_incomplete_windows_blank(self) -> None:
        raw = pd.DataFrame(
            [
                {
                    "Receipt No.": "JUL-A",
                    "Completion Time": "2026-07-10",
                    "Opposite Party": "0811111111 - CLIENT A",
                    "Currency": "CDF",
                    "Transaction Amount": 1000,
                    "Transaction Status": "Completed",
                    "Details": "BisouBisouC2B",
                }
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=prepare_g2_transactions(raw),
        )

        row = build_g2_retention_report(prepared)["mensuelle"].iloc[0]

        self.assertFalse(bool(row["eligible_retention_m1"]))
        self.assertFalse(bool(row["eligible_retention_90j"]))
        self.assertTrue(pd.isna(row["retention_m1_pct"]))
        self.assertTrue(pd.isna(row["retention_90j_pct"]))

    def test_g2_retention_report_excludes_non_completed_rows_from_activity(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=pd.DataFrame([{"receipt_no": "SOURCE"}]),
        )
        daily_detail = pd.DataFrame(
            [
                {
                    "date": pd.Timestamp("2026-01-10 09:00:00"),
                    "receipt_no": "COMPLETED-1",
                    "phone_prefixe": "243811111111",
                    "currency_code": "CDF",
                    "details_rapport": "Depot normal",
                    "transaction_status": "Completed",
                    "incluse_synthese": True,
                    "montant_entree": 1000,
                    "montant_sortie": 0,
                },
                {
                    "date": pd.Timestamp("2026-01-10 10:00:00"),
                    "receipt_no": "PENDING-1",
                    "phone_prefixe": "243822222222",
                    "currency_code": "CDF",
                    "details_rapport": "Depot normal",
                    "transaction_status": "Pending",
                    "incluse_synthese": False,
                    "montant_entree": 5000,
                    "montant_sortie": 0,
                },
            ]
        )

        retention = build_g2_retention_report(prepared, daily_detail=daily_detail)

        self.assertEqual(int(retention["mensuelle"].iloc[0]["clients_actifs_mois_base"]), 1)
        self.assertEqual(retention["detail_clients"]["phone_prefixe"].tolist(), ["243811111111"])

    def test_g2_dat_pdf_html_contains_summary_without_client_phone_detail(self) -> None:
        monthly = pd.DataFrame(
            [
                {
                    "periode": pd.Timestamp("2026-01-01"),
                    "mois": "2026-01",
                    "currency_code": "CDF",
                    "clients_actifs_mois_base": 10,
                    "clients_retenus_m1": 6,
                    "retention_m1_pct": 60.0,
                    "clients_retenus_90j": 8,
                    "retention_90j_pct": 80.0,
                    "eligible_retention_m1": True,
                    "eligible_retention_90j": True,
                }
            ]
        )
        report = {
            "rapport_journalier_pivot": pd.DataFrame(
                [{"currency_code": "CDF", "nombre_entrees": 12, "montant_total_entrees": 10000}]
            ),
            "rapport_journalier_synthese": pd.DataFrame(
                [{"currency_code": "CDF", "sens_flux": "Entree", "details_rapport": "DAT", "nombre": 4, "montant": 5000}]
            ),
            "g2_dat": pd.DataFrame([{"customer_id_dat": "1001", "phone_prefixe": "243811111111"}]),
            "retention_mensuelle": monthly,
            "retention_operations": pd.DataFrame(),
            "analysis_date_start": pd.Timestamp("2026-01-01 08:00:00"),
            "analysis_date_end": pd.Timestamp("2026-05-31 17:30:45"),
        }

        html = build_g2_dat_pdf_html(
            report,
            period_text="du 01/01/2026 au 31/05/2026",
            direction_label="Tous <flux>",
            generated_at=pd.Timestamp("2026-07-14 10:00:00"),
        )

        self.assertIn("Synthese executive", html)
        self.assertIn("Retention M+1", html)
        self.assertIn("60.0%", html)
        self.assertIn("Tous &lt;flux&gt;", html)
        self.assertNotIn("243811111111", html)
        self.assertNotIn("Questions ouvertes", html)
        self.assertNotIn("Hypotheses et limites", html)
        self.assertIn("@page { size: A4 portrait;", html)
        self.assertIn('class="report-logo"', html)
        self.assertIn("data:image/png;base64,", html)
        self.assertIn("Critères", html)
        self.assertIn("<th>Date du :</th><td>01/01/2026 08:00:00</td>", html)
        self.assertIn("<th>Au :</th><td>31/05/2026 17:30:45</td>", html)

    def test_g2_dat_excel_uses_portrait_print_layout(self) -> None:
        from openpyxl import load_workbook

        content = create_excel_export(
            {
                "rapport_journalier_synthese": pd.DataFrame(
                    [
                        {
                            "currency_code": "CDF",
                            "sens_flux": "Entree",
                            "details_rapport": "DAT",
                            "nombre": 4,
                            "montant": 5000,
                        }
                    ]
                ),
                "turbo_dat": pd.DataFrame(
                    [{"customer_id_dat": "1001", "statut_rapprochement_dat": "Conforme"}]
                ),
            },
            print_orientation="portrait",
        )
        workbook = load_workbook(BytesIO(content))

        self.assertTrue(workbook.sheetnames)
        for worksheet in workbook.worksheets:
            self.assertEqual(worksheet.page_setup.orientation, "portrait")
            self.assertEqual(worksheet.page_setup.fitToWidth, 1)
            self.assertEqual(worksheet.print_title_rows, "$1:$1")

    def test_g2_dat_word_is_editable_and_uses_the_short_executive_structure(self) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT

        report = {
            "rapport_journalier_pivot": pd.DataFrame(
                [
                    {
                        "currency_code": "CDF",
                        "nombre_entrees": 12,
                        "montant_total_entrees": 10000,
                        "nombre_sorties": 2,
                        "montant_total_sorties": 1000,
                        "solde_net_flux": 9000,
                    }
                ]
            ),
            "rapport_journalier_synthese": pd.DataFrame(
                [{"currency_code": "CDF", "sens_flux": "Entree", "details_rapport": "DAT", "nombre": 4, "montant": 5000}]
            ),
            "g2_dat": pd.DataFrame([{"customer_id_dat": "1001", "statut_rapprochement_dat": "Rapproche"}]),
            "retention_mensuelle": pd.DataFrame(),
            "analysis_date_start": pd.Timestamp("2026-07-13 08:00:00"),
            "analysis_date_end": pd.Timestamp("2026-07-13 18:00:00"),
            "rapport_journalier_detail": pd.DataFrame(
                [
                    {
                        "currency_code": "CDF",
                        "date": pd.Timestamp("2026-07-13 10:15:00"),
                        "receipt_no": "CDF-001",
                        "sens_flux": "Entree",
                        "details_rapport": "DAT",
                        "opposite_party": "243811111111 - CLIENT TEST",
                        "duree": "3 Months",
                        "compte_cree": pd.Timestamp("2026-07-10 08:30:00"),
                        "montant": 12500,
                        "montant_entree": 12500,
                        "montant_sortie": 0,
                        "balance_numeric": 25000,
                        "incluse_synthese": True,
                    },
                    {
                        "currency_code": "USD",
                        "date": pd.Timestamp("2026-07-13 11:00:00"),
                        "receipt_no": "USD-001",
                        "sens_flux": "Sortie",
                        "details_rapport": "Demande de credit",
                        "opposite_party": "243822222222 - CLIENT USD",
                        "duree": "-",
                        "compte_cree": pd.NaT,
                        "montant": 100,
                        "montant_entree": 0,
                        "montant_sortie": 100,
                        "balance_numeric": 900,
                        "incluse_synthese": True,
                    },
                    {
                        "currency_code": "CDF",
                        "date": pd.Timestamp("2026-07-13 12:00:00"),
                        "receipt_no": "DECLINED-001",
                        "sens_flux": "Entree",
                        "details_rapport": "Depot normal",
                        "opposite_party": "243899999999 - CLIENT REFUSE",
                        "montant": 5000,
                        "montant_entree": 5000,
                        "montant_sortie": 0,
                        "transaction_status": "Declined",
                        "statut_transaction_g2": "Declined",
                        "incluse_synthese": False,
                    },
                ]
            ),
        }

        content = create_g2_dat_word(
            report,
            period_text="du 01/01/2026 au 31/01/2026",
            direction_label="Tous",
            generated_at=pd.Timestamp("2026-07-14 10:00:00"),
        )
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        classified_tables = [
            table
            for table in document.tables
            if table.rows and table.rows[0].cells[0].text == "date"
        ]

        self.assertTrue(content.startswith(b"PK"))
        self.assertIn("Rapport - Solution Numérique / M-Pesa", text)
        self.assertIn("Synthese executive", text)
        self.assertIn("Fréquence temporelle", text)
        self.assertIn("Heure la plus fréquente : 10h", text)
        self.assertNotIn("Jour de semaine le plus actif", text)
        self.assertIn("Synthese des flux rapport G2 M-Pesa par devise", text)
        self.assertIn("Point de vigilance", text)
        self.assertIn("CDF : 1 transaction(s) Completed, 1 client(s) distinct(s)", text)
        self.assertIn("USD : 1 transaction(s) Completed, 1 client(s) distinct(s)", text)
        self.assertIn("2 transaction(s) Completed", text)
        self.assertIn("1 transaction(s) d'autres statuts", text)
        self.assertIn("Transactions", text)
        self.assertEqual(len(classified_tables), 1)
        self.assertEqual(
            [cell.text for cell in classified_tables[0].rows[0].cells],
            G2_CLASSIFIED_TRANSACTION_COLUMNS,
        )
        self.assertEqual(classified_tables[0].rows[1].cells[2].text, "CDF")
        self.assertEqual(classified_tables[0].rows[2].cells[2].text, "USD")
        self.assertIn("243811111111", table_text)
        self.assertIn("CLIENT TEST", table_text)
        self.assertIn("13/07/2026 10:15:00", table_text)
        self.assertIn("10/07/2026 08:30:00", table_text)
        self.assertNotIn("DECLINED-001", table_text)
        self.assertNotIn("CLIENT REFUSE", table_text)
        self.assertNotIn("Questions ouvertes", text)
        self.assertGreaterEqual(len(document.tables), 3)
        self.assertGreaterEqual(len(document.sections), 2)
        self.assertGreaterEqual(len(document.inline_shapes), 1)
        self.assertIn("Critères", table_text)
        criteria_table = document.tables[0].cell(1, 1).tables[0]
        criteria_labels = [row.cells[0].text for row in criteria_table.rows]
        criteria_values = [row.cells[1].text for row in criteria_table.rows]
        self.assertEqual(
            criteria_labels,
            ["Date du :", "Au :", "Sens :", "Source :", "Généré le :"],
        )
        self.assertEqual(criteria_values[0], "13/07/2026 08:00:00")
        self.assertEqual(criteria_values[1], "13/07/2026 18:00:00")
        self.assertIn("Sens :", criteria_labels)
        self.assertTrue(
            all(section.orientation == WD_ORIENT.PORTRAIT for section in document.sections)
        )
        self.assertTrue(
            all(section.page_height > section.page_width for section in document.sections)
        )

        multi_day_report = dict(report)
        multi_day_detail = pd.concat(
            [
                report["rapport_journalier_detail"],
                pd.DataFrame(
                    [
                        {
                            "currency_code": "CDF",
                            "date": pd.Timestamp("2026-07-14 09:00:00"),
                            "receipt_no": "CDF-002",
                            "sens_flux": "Entree",
                            "details_rapport": "Depot normal",
                            "opposite_party": "243833333333 - CLIENT MULTI",
                            "montant": 1000,
                            "montant_entree": 1000,
                            "montant_sortie": 0,
                            "incluse_synthese": True,
                        },
                        {
                            "currency_code": "CDF",
                            "date": pd.Timestamp("2026-07-14 09:15:00"),
                            "receipt_no": "CDF-003",
                            "sens_flux": "Entree",
                            "details_rapport": "DAT",
                            "opposite_party": "243844444444 - CLIENT MULTI 2",
                            "montant": 2000,
                            "montant_entree": 2000,
                            "montant_sortie": 0,
                            "incluse_synthese": True,
                        },
                        {
                            "currency_code": "USD",
                            "date": pd.Timestamp("2026-07-14 09:45:00"),
                            "receipt_no": "USD-002",
                            "sens_flux": "Sortie",
                            "details_rapport": "Demande de credit",
                            "opposite_party": "243855555555 - CLIENT MULTI 3",
                            "montant": 50,
                            "montant_entree": 0,
                            "montant_sortie": 50,
                            "incluse_synthese": True,
                        },
                    ]
                ),
            ],
            ignore_index=True,
        )
        time_report = build_g2_transaction_time_analysis(multi_day_detail)
        multi_day_report["rapport_journalier_detail"] = multi_day_detail
        multi_day_report["transactions_par_jour"] = time_report["par_jour"]
        multi_day_report["transactions_par_jour_semaine"] = time_report["par_jour_semaine"]
        multi_day_report["transactions_par_heure"] = time_report["par_heure"]
        multi_day_report["analysis_date_start"] = pd.Timestamp("2026-07-13")
        multi_day_report["analysis_date_end"] = pd.Timestamp("2026-07-14")
        multi_day_content = create_g2_dat_word(
            multi_day_report,
            period_text="du 13/07/2026 au 14/07/2026",
            direction_label="Tous",
            generated_at=pd.Timestamp("2026-07-15 10:00:00"),
        )
        multi_day_document = Document(BytesIO(multi_day_content))
        multi_day_text = "\n".join(paragraph.text for paragraph in multi_day_document.paragraphs)
        self.assertIn("Heure la plus fréquente : 09h, avec 3 transaction(s), soit 60.0%", multi_day_text)
        self.assertIn("Jour de semaine le plus actif : Mardi, avec 3 transaction(s), soit 60.0%", multi_day_text)

        report_without_pivot = dict(report)
        report_without_pivot.pop("rapport_journalier_pivot")
        fallback_content = create_g2_dat_word(
            report_without_pivot,
            period_text="du 01/01/2026 au 31/01/2026",
            direction_label="Tous",
            generated_at=pd.Timestamp("2026-07-14 10:00:00"),
        )
        fallback_document = Document(BytesIO(fallback_content))
        fallback_text = "\n".join(paragraph.text for paragraph in fallback_document.paragraphs)
        flow_tables = [
            table
            for table in fallback_document.tables
            if table.rows and table.rows[0].cells[0].text == "Devise"
        ]
        self.assertIn("Synthese des flux rapport G2 M-Pesa par devise", fallback_text)
        self.assertNotIn("Aucune donnee disponible.", fallback_text)
        self.assertEqual(len(flow_tables), 2)
        self.assertEqual(flow_tables[0].rows[1].cells[0].text, "CDF")
        self.assertEqual(flow_tables[0].rows[2].cells[0].text, "USD")

    def test_g2_dat_word_labels_turbo_only_source_without_simulating_g2_controls(self) -> None:
        from docx import Document

        detail = pd.DataFrame(
            [
                {
                    "currency_code": "CDF",
                    "date": pd.Timestamp("2026-07-15 08:30:00"),
                    "receipt_no": "REF-TURBO",
                    "sens_flux": "Entree",
                    "details_rapport": "Depot normal",
                    "opposite_party": "243811111111",
                    "montant": 1000,
                    "montant_entree": 1000,
                    "montant_sortie": 0,
                    "incluse_synthese": True,
                    "transaction_status": "Comptabilisee Solution Numérique",
                }
            ]
        )
        report = {
            "analysis_source_label": "Solution Numérique",
            "rapport_journalier_pivot": build_entry_pivot(detail),
            "rapport_journalier_synthese": pd.DataFrame(),
            "rapport_journalier_detail": detail,
            "g2_dat": detail.assign(
                statut_rapprochement="Non applicable - Solution Numérique seule",
                est_anomalie=False,
            ),
            "retention_mensuelle": pd.DataFrame(),
        }

        content = create_g2_dat_word(
            report,
            period_text="le 15/07/2026",
            direction_label="Tous",
            generated_at=pd.Timestamp("2026-07-15 10:00:00"),
        )
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn(" Solution Bisou Bisou Digital", text)
        self.assertIn("Synthese des flux Solution Numérique par devise", text)
        self.assertIn("operation(s) comptabilisee(s) dans la Solution Numérique", text)
        self.assertIn("controles croises rapport G2/Solution Numérique sont non applicables", text)
        self.assertNotIn("Rapprochement Receipt No/ref_no", text)

    def test_g2_dat_word_activity_uses_filtered_detail_not_retention_month(self) -> None:
        from docx import Document

        filtered_detail = pd.DataFrame(
            [
                {
                    "currency_code": "CDF",
                    "date": pd.Timestamp("2026-07-15 08:30:00"),
                    "receipt_no": "FILTERED-1",
                    "sens_flux": "Entree",
                    "details_rapport": "Depot normal",
                    "opposite_party": "243811111111 - CLIENT FILTRE",
                    "montant": 1000,
                    "montant_entree": 1000,
                    "montant_sortie": 0,
                    "incluse_synthese": True,
                }
            ]
        )
        report = {
            "rapport_journalier_pivot": build_entry_pivot(filtered_detail),
            "rapport_journalier_synthese": pd.DataFrame(),
            "rapport_journalier_detail": filtered_detail,
            "retention_mensuelle": pd.DataFrame(
                [
                    {
                        "periode": pd.Timestamp("2026-07-01"),
                        "currency_code": "CDF",
                        "clients_actifs_mois_base": 99,
                        "retention_m1_pct": pd.NA,
                        "retention_90j_pct": pd.NA,
                    }
                ]
            ),
        }

        content = create_g2_dat_word(
            report,
            period_text="du 15/07/2026 a 08:00:00 au 15/07/2026 a 09:00:00",
            direction_label="Entrees",
            generated_at=pd.Timestamp("2026-07-15 10:00:00"),
        )
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("CDF : 1 transaction(s) Completed, 1 client(s) distinct(s)", text)
        self.assertNotIn("99 client(s) actif(s)", text)
        self.assertIn("Sens : Entrees", text)

    def test_g2_dat_word_all_sections_use_date_time_direction_and_status_scope(self) -> None:
        from docx import Document

        g2 = prepare_g2_transactions(
            pd.DataFrame(
                [
                    {
                        "Receipt No.": "IN-SCOPE",
                        "Completion Time": "2026-07-15 08:30:00",
                        "Transaction Status": "Completed",
                        "Opposite Party": "0811111111 - CLIENT DANS FILTRE",
                        "Currency": "CDF",
                        "Paid In": 1000,
                        "Withdrawn": 0,
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "OUT-DIRECTION",
                        "Completion Time": "2026-07-15 08:40:00",
                        "Transaction Status": "Completed",
                        "Opposite Party": "0822222222 - CLIENT SORTIE",
                        "Currency": "CDF",
                        "Paid In": 0,
                        "Withdrawn": 200,
                        "Details": "BisouBisouB2C",
                    },
                    {
                        "Receipt No.": "OUT-TIME",
                        "Completion Time": "2026-07-15 12:00:00",
                        "Transaction Status": "Completed",
                        "Opposite Party": "0833333333 - CLIENT HORS HEURE",
                        "Currency": "CDF",
                        "Paid In": 5000,
                        "Withdrawn": 0,
                        "Details": "BisouBisouC2B",
                    },
                    {
                        "Receipt No.": "PENDING-SCOPE",
                        "Completion Time": "2026-07-15 08:50:00",
                        "Transaction Status": "Pending",
                        "Opposite Party": "0844444444 - CLIENT PENDING",
                        "Currency": "CDF",
                        "Paid In": 7000,
                        "Withdrawn": 0,
                        "Details": "BisouBisouC2B",
                    },
                ]
            )
        )
        filtered_g2 = filter_g2_transactions_by_completion_time(
            g2,
            pd.Timestamp("2026-07-15").date(),
            pd.Timestamp("2026-07-15").date(),
            time(8, 0),
            time(9, 0),
        )
        filtered_g2 = filter_g2_transactions_by_direction(filtered_g2, ["Entree"])
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
            g2_transactions=filtered_g2,
        )
        daily = build_g2_daily_savings_report(prepared)
        time_report = build_g2_transaction_time_analysis(daily["detail"])
        retention = build_g2_retention_report(prepared, daily_detail=daily["detail"])
        word_report = {
            "rapport_journalier_pivot": daily["pivot"],
            "rapport_journalier_synthese": daily["synthese"],
            "rapport_journalier_detail": daily["detail"],
            "statuts_g2": daily["statuts"],
            "transactions_par_jour": time_report["par_jour"],
            "transactions_par_jour_semaine": time_report["par_jour_semaine"],
            "transactions_par_heure": time_report["par_heure"],
            "retention_mensuelle": retention["mensuelle"],
            "g2_dat": pd.DataFrame(),
            "analysis_date_start": pd.Timestamp("2026-07-15 08:00:00"),
            "analysis_date_end": pd.Timestamp("2026-07-15 09:00:00"),
        }

        content = create_g2_dat_word(
            word_report,
            period_text="du 15/07/2026 a 08:00:00 au 15/07/2026 a 09:00:00",
            direction_label="Entrees",
            generated_at=pd.Timestamp("2026-07-15 10:00:00"),
        )
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("CDF : 1 transaction(s) Completed, 1 client(s) distinct(s)", text)
        self.assertIn("1 transaction(s) Completed incluse(s)", text)
        self.assertIn("1 transaction(s) d'autres statuts", text)
        self.assertIn("08h, avec 1 transaction(s), soit 100.0% du volume", text)
        self.assertIn("IN-SCOPE", table_text)
        self.assertIn("Depot normal", table_text)
        self.assertNotIn("OUT-DIRECTION", table_text)
        self.assertNotIn("OUT-TIME", table_text)
        self.assertNotIn("PENDING-SCOPE", table_text)

    def test_numeric_column_handles_missing_columns_like_zero_series(self) -> None:
        frame = pd.DataFrame({"customer_id": ["1", "2"]})

        values = numeric_column(frame, "cr")

        self.assertEqual(values.tolist(), [0.0, 0.0])
        self.assertEqual(float(values.sum()), 0.0)

    def test_savings_final_accepts_file_without_balance_column(self) -> None:
        current = prepare_current_savings(
            pd.DataFrame(
                [
                    {
                        "customer_id": 1001,
                        "msisdn": "0812345678",
                        "currency_code": "CDF",
                        "created_at": "2026-07-11",
                    }
                ]
            )
        )

        result = build_savings_final(current, "1001")

        self.assertEqual(result, {})

    def test_diagnostics_accept_transactions_without_amount_columns(self) -> None:
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(
                [
                    {
                        "customer_id": "1",
                        "reference_id": "REF001",
                        "ref_no": "G2REF",
                        "created_at": "2026-07-11",
                        "currency_code": "CDF",
                        "account_type": "MPESA ACCOUNT",
                    }
                ]
            ),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        diagnostics = build_diagnostics(prepared)

        self.assertFalse(diagnostics.empty)
        self.assertIn("Mouvements dr = 0 et cr = 0", diagnostics["controle"].tolist())

    def test_transaction_anomalies_explain_every_flagged_turbo_row(self) -> None:
        transactions = pd.DataFrame(
            [
                {
                    "id": "MISSING-CUSTOMER",
                    "customer_id": "",
                    "reference_id": "REF-1",
                    "currency_code": "CDF",
                    "account_type": "MPESA ACCOUNT",
                    "dr": 10,
                    "cr": 0,
                    "bal_before": 100,
                    "bal_after": 90,
                    "created_at": "2026-07-16 08:00:00",
                },
                {
                    "id": "MISSING-CURRENCY-ZERO",
                    "customer_id": "CLIENT-1",
                    "reference_id": "REF-2",
                    "currency_code": "",
                    "account_type": "NORMAL SAVINGS",
                    "dr": 0,
                    "cr": 0,
                    "bal_before": 100,
                    "bal_after": 100,
                    "created_at": "2026-07-16 09:00:00",
                },
                {
                    "id": "BOTH-SIDES",
                    "customer_id": "CLIENT-2",
                    "reference_id": "REF-3",
                    "currency_code": "USD",
                    "account_type": "MPESA ACCOUNT",
                    "dr": 5,
                    "cr": 7,
                    "bal_before": 100,
                    "bal_after": 102,
                    "created_at": "2026-07-16 10:00:00",
                },
                {
                    "id": "VALID",
                    "customer_id": "CLIENT-3",
                    "reference_id": "REF-4",
                    "currency_code": "CDF",
                    "account_type": "FIXED SAVINGS",
                    "dr": 0,
                    "cr": 25,
                    "bal_before": 100,
                    "bal_after": 125,
                    "created_at": "2026-07-16 11:00:00",
                },
            ]
        )

        anomalies = build_transaction_anomalies(transactions).set_index("id")

        self.assertEqual(
            anomalies.index.tolist(),
            ["MISSING-CUSTOMER", "MISSING-CURRENCY-ZERO", "BOTH-SIDES"],
        )
        self.assertEqual(
            anomalies.loc["MISSING-CUSTOMER", "raison_anomalie"],
            "customer_id manquant",
        )
        self.assertEqual(
            anomalies.loc["MISSING-CURRENCY-ZERO", "raison_anomalie"],
            "currency_code manquant | Mouvement nul : dr = 0 et cr = 0",
        )
        self.assertEqual(
            anomalies.loc["BOTH-SIDES", "raison_anomalie"],
            "Incoherence de sens : dr > 0 et cr > 0",
        )

        zero_movement_only = build_transaction_anomalies(
            transactions,
            selected_controls={"Mouvements dr = 0 et cr = 0"},
        )
        self.assertEqual(
            zero_movement_only["id"].tolist(),
            ["MISSING-CURRENCY-ZERO"],
        )
        self.assertEqual(
            zero_movement_only["raison_anomalie"].tolist(),
            ["Mouvement nul : dr = 0 et cr = 0"],
        )

        no_matching_control = build_transaction_anomalies(
            transactions,
            selected_controls={"Lignes sans reference_id"},
        )
        self.assertTrue(no_matching_control.empty)

    def test_transaction_anomalies_cover_detailed_diagnostic_controls(self) -> None:
        transactions = pd.DataFrame(
            [
                {
                    "id": "MISSING-REFERENCE",
                    "customer_id": "CLIENT-1",
                    "reference_id": "",
                    "currency_code": "CDF",
                    "account_type": "MPESA ACCOUNT",
                    "dr": 10,
                    "cr": 0,
                    "bal_before": 100,
                    "bal_after": 90,
                    "created_at": "2026-07-16 08:00:00",
                },
                {
                    "id": "INVALID-DATE-AND-BALANCE",
                    "customer_id": "CLIENT-2",
                    "reference_id": "REF-2",
                    "currency_code": "USD",
                    "account_type": "UNKNOWN PRODUCT",
                    "dr": 0,
                    "cr": 5,
                    "bal_before": -5,
                    "bal_after": 0,
                    "created_at": "date-invalide",
                },
            ]
        )

        anomalies = build_transaction_anomalies(transactions).set_index("id")

        self.assertEqual(
            anomalies.loc["MISSING-REFERENCE", "raison_anomalie"],
            "reference_id manquant",
        )
        combined_reason = anomalies.loc[
            "INVALID-DATE-AND-BALANCE",
            "raison_anomalie",
        ]
        self.assertIn("Date created_at invalide ou manquante", combined_reason)
        self.assertIn("Solde bal_before ou bal_after negatif", combined_reason)
        self.assertIn("account_type non reference", combined_reason)

        reference_only = build_transaction_anomalies(
            transactions,
            selected_controls={"Lignes sans reference_id"},
        )
        self.assertEqual(reference_only["id"].tolist(), ["MISSING-REFERENCE"])

    def test_repeated_transaction_control_counts_the_detailed_rows(self) -> None:
        transactions = pd.DataFrame(
            [
                {
                    "id": transaction_id,
                    "customer_id": "CLIENT-1",
                    "reference_id": "REF-1",
                    "ref_no": "G2-1",
                    "currency_code": "CDF",
                    "account_type": "MPESA ACCOUNT",
                    "dr": 10,
                    "cr": 0,
                    "bal_before": 100,
                    "bal_after": 90,
                    "created_at": "2026-07-16 08:00:00",
                }
                for transaction_id in ("TX-1", "TX-2")
            ]
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        diagnostics = build_diagnostics(prepared).set_index("controle")
        repeated = diagnostics.loc[
            "Lignes dans des groupes d'ecritures repetees"
        ]
        self.assertEqual(int(repeated["valeur"]), 2)
        self.assertIn("1 groupe(s)", repeated["detail"])

        anomalies = build_transaction_anomalies(
            transactions,
            selected_controls={"Lignes dans des groupes d'ecritures repetees"},
        )
        self.assertEqual(anomalies["id"].tolist(), ["TX-1", "TX-2"])
        self.assertEqual(
            anomalies["raison_anomalie"].unique().tolist(),
            ["Groupe d'ecritures repetees avec le meme type de compte"],
        )

    def test_diagnostics_separates_linked_entries_and_monitors_dat_dates(self) -> None:
        transactions = pd.DataFrame(
            [
                {
                    "id": "TX-1",
                    "customer_id": "1",
                    "reference_id": "REF001",
                    "ref_no": "G2REF",
                    "created_at": "2026-07-15 10:00:00",
                    "currency_code": "CDF",
                    "account_type": "MPESA ACCOUNT",
                    "dr": 0,
                    "cr": 100,
                    "bal_before": 0,
                    "bal_after": 100,
                },
                {
                    "id": "TX-2",
                    "customer_id": "1",
                    "reference_id": "REF001",
                    "ref_no": "G2REF",
                    "created_at": "2026-07-15 10:00:00",
                    "currency_code": "CDF",
                    "account_type": "LOAN ACCOUNT",
                    "dr": 0,
                    "cr": 100,
                    "bal_before": 0,
                    "bal_after": 100,
                },
            ]
        )
        fixed = pd.DataFrame(
            [
                {
                    "customer_id": "1",
                    "balance": 100,
                    "date_approved": "2026-06-07",
                    "maturity_date": "2026-04-22",
                },
                {
                    "customer_id": "2",
                    "balance": 200,
                    "date_approved": "2026-07-01",
                    "maturity_date": "2026-08-01",
                },
            ]
        )
        prepared = MpesaPreparedData(
            transactions=transactions,
            current_savings=pd.DataFrame(),
            fixed_savings=fixed,
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        diagnostics = build_diagnostics(prepared).set_index("controle")

        self.assertEqual(diagnostics.loc["Lignes sans customer_id", "statut"], "OK")
        self.assertEqual(int(diagnostics.loc["Types de comptes a classifier", "valeur"]), 0)
        self.assertEqual(int(diagnostics.loc["Doublons exacts", "valeur"]), 0)
        self.assertEqual(int(diagnostics.loc["Ecritures comptables liees", "valeur"]), 1)
        self.assertEqual(diagnostics.loc["Ecritures comptables liees", "statut"], "Information")
        self.assertEqual(int(diagnostics.loc["Lignes dans des groupes d'ecritures repetees", "valeur"]), 0)
        self.assertEqual(int(diagnostics.loc["DAT - echeance anterieure a l'approbation", "valeur"]), 1)
        self.assertEqual(diagnostics.loc["DAT - echeance anterieure a l'approbation", "statut"], "A surveiller")
        self.assertEqual(int(diagnostics.loc["DAT echus avec solde positif", "valeur"]), 1)
        self.assertEqual(diagnostics.loc["DAT echus avec solde positif", "statut"], "Controle metier")

    def test_mpesa_forecast_keeps_currencies_separate_and_backtests(self) -> None:
        dates = pd.date_range("2026-01-01", periods=120, freq="D")
        events = pd.DataFrame(
            [
                {
                    "created_at": date,
                    "event_key": f"{currency}-{index}",
                    "customer_id": f"C-{index % 10}",
                    "currency_code": currency,
                    "montant_entree_bisou": entry,
                    "montant_sortie_bisou": exit_amount,
                    "remboursement_mpesa": repayment,
                    "remboursement_compte_ouvert": 0.0,
                    "montant_decaisse_client": credit,
                }
                for index, date in enumerate(dates)
                for currency, entry, exit_amount, repayment, credit in [
                    ("CDF", 100.0, 50.0, 20.0, 30.0),
                    ("USD", 1.0, 0.5, 0.2, 0.3),
                ]
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        report = build_mpesa_forecast_report(
            prepared,
            reference_date=dates.max(),
            horizon_days=30,
            confidence_level=80,
            turbo_events=events,
        )

        summary = report["synthese"].set_index(
            ["indicator_key", "currency_code"]
        )
        cdf_volume = summary.loc[("volume_transactions", "CDF")]
        usd_volume = summary.loc[("volume_transactions", "USD")]
        self.assertAlmostEqual(float(cdf_volume["valeur_prevue_horizon"]), 4_500.0)
        self.assertAlmostEqual(float(usd_volume["valeur_prevue_horizon"]), 45.0)
        self.assertLess(float(cdf_volume["wape_pct"]), 0.01)
        self.assertEqual(cdf_volume["qualite_modele"], "Bonne")
        volume_forecast = report["previsions"].loc[
            report["previsions"]["indicator_key"].eq("volume_transactions")
        ]
        self.assertEqual(
            volume_forecast.groupby("currency_code")["date"].nunique().to_dict(),
            {"CDF": 30, "USD": 30},
        )
        self.assertTrue(
            volume_forecast["borne_basse"].le(volume_forecast["prevision"]).all()
        )
        self.assertTrue(
            volume_forecast["borne_haute"].ge(volume_forecast["prevision"]).all()
        )

    def test_mpesa_forecast_dat_schedule_is_deterministic(self) -> None:
        dates = pd.date_range("2026-01-01", periods=120, freq="D")
        events = pd.DataFrame(
            [
                {
                    "created_at": date,
                    "event_key": f"EV-{index}",
                    "customer_id": f"C-{index % 5}",
                    "currency_code": "USD",
                    "montant_entree_bisou": 10.0,
                    "montant_sortie_bisou": 0.0,
                    "remboursement_mpesa": 0.0,
                    "remboursement_compte_ouvert": 0.0,
                    "montant_decaisse_client": 0.0,
                }
                for index, date in enumerate(dates)
            ]
        )
        fixed = pd.DataFrame(
            [
                {
                    "savings_id": "DAT-1",
                    "customer_id": "C-1",
                    "currency_code": "USD",
                    "balance": 100.0,
                    "date_approved": "2026-01-01",
                    "maturity_date": "2026-05-15",
                    "status": "Active",
                },
                {
                    "savings_id": "DAT-2",
                    "customer_id": "C-2",
                    "currency_code": "CDF",
                    "balance": 200.0,
                    "date_approved": "2026-01-01",
                    "maturity_date": "2026-08-15",
                    "status": "Active",
                },
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=fixed,
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        report = build_mpesa_forecast_report(
            prepared,
            reference_date="2026-04-30",
            horizon_days=30,
            confidence_level=95,
            turbo_events=events,
        )

        schedule = report["dat_echeancier"]
        self.assertEqual(schedule["savings_id"].tolist(), ["DAT-1"])
        self.assertEqual(float(schedule.iloc[0]["balance"]), 100.0)
        self.assertGreater(
            float(schedule.iloc[0]["capital_plus_interet_estime"]),
            100.0,
        )

    def test_mpesa_forecast_reports_insufficient_history(self) -> None:
        events = pd.DataFrame(
            [
                {
                    "created_at": date,
                    "event_key": f"EV-{index}",
                    "customer_id": "C-1",
                    "currency_code": "CDF",
                    "montant_entree_bisou": 10.0,
                    "montant_sortie_bisou": 0.0,
                    "remboursement_mpesa": 0.0,
                    "remboursement_compte_ouvert": 0.0,
                    "montant_decaisse_client": 0.0,
                }
                for index, date in enumerate(
                    pd.date_range("2026-07-01", periods=10, freq="D")
                )
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=pd.DataFrame(),
            load_report=build_load_report({}, {}),
        )

        report = build_mpesa_forecast_report(
            prepared,
            reference_date="2026-07-10",
            horizon_days=7,
            turbo_events=events,
        )

        self.assertTrue(report["synthese"].empty)
        self.assertFalse(report["non_calculable"].empty)
        self.assertTrue(
            report["non_calculable"]["motif"].eq("Historique insuffisant").all()
        )


class TestLoanSavingsReconciliation(unittest.TestCase):
    def test_deduced_match_consolidates_savings_once_per_client_currency(self) -> None:
        loans = pd.DataFrame(
            [
                {
                    "loan_id": "LN-1",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "currency_code": "CDF",
                    "loan_amount": 1000,
                    "loan_balance": 400,
                    "amount_paid": 600,
                },
                {
                    "loan_id": "LN-2",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "currency_code": "CDF",
                    "loan_amount": 500,
                    "loan_balance": 100,
                    "amount_paid": 400,
                },
                {
                    "loan_id": "LN-3",
                    "customer_id": "C3",
                    "msisdn1": "0833333333",
                    "currency_code": "USD",
                    "loan_amount": 100,
                    "loan_balance": 100,
                    "amount_paid": 0,
                },
            ]
        )
        current = pd.DataFrame(
            [
                {
                    "id": 10,
                    "savings_id": "SAV-1",
                    "customer_id": "C1",
                    "msisdn": "0811111111",
                    "currency_code": "CDF",
                    "balance": 100,
                    "source_savings_account_complete": True,
                }
            ]
        )
        fixed = pd.DataFrame(
            [
                {
                    "savings_id": "DAT-1",
                    "customer_id": "C1",
                    "currency_code": "CDF",
                    "balance": 200,
                }
            ]
        )

        report = build_loan_savings_reconciliation(loans, current, fixed)
        cdf_summary = report["synthese"].set_index("currency_code").loc["CDF"]
        cdf_client = report["clients"].set_index(["customer_id", "currency_code"]).loc[("C1", "CDF")]

        self.assertEqual(int(cdf_summary["nombre_credits"]), 2)
        self.assertEqual(float(cdf_summary["encours_credit"]), 500.0)
        self.assertEqual(float(cdf_summary["solde_epargne_courante_clients_credit"]), 100.0)
        self.assertEqual(float(cdf_summary["solde_dat_clients_credit"]), 200.0)
        self.assertEqual(float(cdf_summary["taux_rapprochement_pct"]), 100.0)
        self.assertEqual(int(cdf_client["nombre_credits"]), 2)
        self.assertEqual(float(cdf_client["solde_epargne_courante"]), 100.0)
        self.assertEqual(float(cdf_client["solde_dat_positif"]), 200.0)
        self.assertEqual(cdf_client["savings_id_correspondant"], "SAV-1")
        self.assertEqual(set(report["controles"]["loan_id"]), {"LN-3"})
        export = create_excel_export(
            {
                "loan_savings_summary": report["synthese"],
                "loan_savings_clients": report["clients"],
                "loan_savings_detail": report["detail"],
                "loan_savings_controls": report["controles"],
            }
        )
        self.assertEqual(
            pd.ExcelFile(BytesIO(export)).sheet_names,
            [
                "Credit_Epargne_Synthese",
                "Credit_Epargne_Clients",
                "Credit_Epargne_Detail",
                "Controle_Credit_Epargne",
            ],
        )

    def test_direct_identifier_wins_over_ambiguous_customer_currency(self) -> None:
        loans = pd.DataFrame(
            [
                {
                    "loan_id": "LN-DIRECT",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "currency_code": "CDF",
                    "loan_balance": 50,
                    "savings_account_id": 20,
                },
                {
                    "loan_id": "LN-AMBIGUOUS",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "currency_code": "CDF",
                    "loan_balance": 25,
                },
            ]
        )
        current = pd.DataFrame(
            [
                {
                    "id": 20,
                    "savings_id": "SAV-20",
                    "customer_id": "C1",
                    "msisdn": "0811111111",
                    "currency_code": "CDF",
                    "balance": 10,
                },
                {
                    "id": 21,
                    "savings_id": "SAV-21",
                    "customer_id": "C1",
                    "msisdn": "0811111111",
                    "currency_code": "CDF",
                    "balance": 15,
                },
            ]
        )

        detail = build_loan_savings_reconciliation(loans, current)["detail"].set_index("loan_id")

        self.assertEqual(detail.loc["LN-DIRECT", "savings_id_correspondant"], "SAV-20")
        self.assertEqual(
            detail.loc["LN-DIRECT", "statut_controle"],
            "Conforme - correspondance directe",
        )
        self.assertTrue(bool(detail.loc["LN-DIRECT", "liaison_directe_source"]))
        self.assertEqual(detail.loc["LN-AMBIGUOUS", "statut_controle"], "A revoir")
        self.assertIn("Plusieurs comptes courants", detail.loc["LN-AMBIGUOUS", "motif_controle"])

    def test_missing_savings_source_is_not_reported_as_an_operational_anomaly(self) -> None:
        loans = pd.DataFrame(
            [
                {
                    "loan_id": "LN-1",
                    "customer_id": "C1",
                    "currency_code": "CDF",
                    "loan_balance": 50,
                }
            ]
        )

        report = build_loan_savings_reconciliation(loans, pd.DataFrame())

        self.assertEqual(
            report["detail"].iloc[0]["statut_controle"],
            "Non calculable - Savings Account absent",
        )
        self.assertTrue(report["controles"].empty)
        self.assertTrue(pd.isna(report["synthese"].iloc[0]["taux_rapprochement_pct"]))

    def test_credit_cockpit_keeps_snapshot_flux_and_data_gaps_separate(self) -> None:
        loans = pd.DataFrame(
            [
                {
                    "loan_id": "LN-USD-1",
                    "customer_id": "C1",
                    "msisdn1": "243811111111",
                    "currency_code": "USD",
                    "loan_product_id": "P1",
                    "loan_amount": 100,
                    "loan_balance": 60,
                    "amount_paid": 40,
                    "outstanding_principle": 50,
                    "outstanding_interest": 10,
                    "outstanding_penalty_fees": 0,
                    "status_name": "ACTIVE",
                    "defaulted": 0,
                    "is_rollover": 0,
                    "is_grace_period": 0,
                    "due_date": "2026-07-01",
                    "last_repayment_date": "2026-07-10",
                    "created_at": "2026-06-01",
                },
                {
                    "loan_id": "LN-CDF-1",
                    "customer_id": "C2",
                    "msisdn1": "243822222222",
                    "currency_code": "CDF",
                    "loan_product_id": "P2",
                    "loan_amount": 1000,
                    "loan_balance": 1000,
                    "amount_paid": 0,
                    "outstanding_penalty_fees": 25,
                    "status_name": "ACTIVE",
                    "defaulted": 1,
                    "is_rollover": "1.0",
                    "is_grace_period": True,
                    "due_date": "2026-08-10",
                    "created_at": "2026-07-01",
                },
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=pd.DataFrame(),
            fixed_savings=pd.DataFrame(),
            loans=loans,
            load_report=pd.DataFrame(),
            customers=pd.DataFrame(),
        )

        report = build_mpesa_credit_cockpit(
            prepared,
            date_start="2026-07-01",
            date_end="2026-08-05",
        )

        portfolio = report["portefeuille_synthese"].set_index("currency_code")
        self.assertEqual(float(portfolio.loc["USD", "encours_credit"]), 60.0)
        self.assertEqual(float(portfolio.loc["CDF", "encours_credit"]), 1000.0)
        self.assertEqual(int(portfolio.loc["CDF", "prets_defaulted"]), 1)
        self.assertEqual(int(portfolio.loc["CDF", "prets_rollover"]), 1)
        self.assertEqual(int(portfolio.loc["CDF", "prets_grace_period"]), 1)

        maturity = report["echeances_synthese"].set_index(["currency_code", "tranche_echeance"])
        self.assertEqual(float(maturity.loc[("USD", "Echu"), "encours_credit"]), 60.0)
        self.assertEqual(float(maturity.loc[("CDF", "0-7 jours"), "encours_credit"]), 1000.0)

        catalogue = report["catalogue_kpi"].set_index("kpi")
        self.assertEqual(catalogue.loc["par_reglementaire_detaille", "statut"], "data_gap")
        self.assertIn("prets_par_simplifie_30j", report["listes_action"])
        self.assertTrue(
            report["statuts_portefeuille"]["valeur_statut"].map(type).eq(str).all()
        )

    def test_credit_cockpit_excel_export_contains_credit_sheets(self) -> None:
        report = {
            "credit_vue_ensemble": pd.DataFrame(
                [{"indicateur": "encours_credit", "currency_code": "USD", "valeur": 60}]
            ),
            "credit_catalogue_kpi": pd.DataFrame(
                [{"kpi": "par_reglementaire_detaille", "statut": "data_gap"}]
            ),
            "credit_liste_prets_defaulted": pd.DataFrame(
                [{"loan_id": "LN-1", "currency_code": "USD"}]
            ),
        }

        export = create_excel_export(report, rename_user_columns=True)
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")

        self.assertIn("Credit_Vue_Ensemble", workbook.sheet_names)
        self.assertIn("Credit_Catalogue_KPI", workbook.sheet_names)
        self.assertIn("Liste_Prets_Defaulted", workbook.sheet_names)
        exported = pd.read_excel(workbook, sheet_name="Liste_Prets_Defaulted")
        self.assertIn("numero_pret", exported.columns)

    def test_savings_cockpit_separates_snapshot_flux_and_currencies(self) -> None:
        savings = pd.DataFrame(
            [
                {
                    "savings_id": "SAV-USD-1",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "currency_code": "USD",
                    "balance": 25,
                    "status": "active",
                    "created_at": "2026-07-01",
                    "date_activated": "2026-07-01",
                },
                {
                    "savings_id": "SAV-USD-0",
                    "customer_id": "C2",
                    "msisdn1": "0822222222",
                    "product_name": "Open Savings",
                    "product_description": "Current account",
                    "currency_code": "USD",
                    "balance": 0,
                    "status": "active",
                    "created_at": "2026-07-02",
                },
                {
                    "savings_id": "DAT-USD-1",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "currency_code": "USD",
                    "balance": 100,
                    "status": "active",
                    "date_approved": "2026-07-01",
                    "date_activated": "2026-07-01",
                    "maturity_date": "2026-07-31",
                    "interest_earned": 1.25,
                    "created_at": "2026-07-01",
                },
                {
                    "savings_id": "DAT-CDF-0",
                    "customer_id": "C3",
                    "msisdn1": "0833333333",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "currency_code": "CDF",
                    "balance": 0,
                    "status": "withdrawal",
                    "date_approved": "2026-06-01",
                    "maturity_date": "2026-07-01",
                    "created_at": "2026-06-01",
                },
            ]
        )
        transactions = pd.DataFrame(
            [
                {
                    "id": 1,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "SAV-USD-1",
                    "currency_code": "USD",
                    "dr": 10,
                    "cr": 0,
                    "bal_before": 100,
                    "bal_after": 90,
                    "ref_no": "DEP-1",
                    "description": "M-Pesa Depot",
                    "created_at": "2026-07-10 10:00:00",
                },
                {
                    "id": 2,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "NORMAL SAVINGS",
                    "reference_id": "SAV-USD-1",
                    "currency_code": "USD",
                    "dr": 0,
                    "cr": 10,
                    "bal_before": 15,
                    "bal_after": 25,
                    "ref_no": "DEP-1",
                    "description": "Epargne depot",
                    "created_at": "2026-07-10 10:00:00",
                },
                {
                    "id": 3,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "NORMAL SAVINGS",
                    "reference_id": "SAV-USD-1",
                    "currency_code": "USD",
                    "dr": 5,
                    "cr": 0,
                    "bal_before": 25,
                    "bal_after": 20,
                    "ref_no": "",
                    "description": "Retrait Vers M-Pesa",
                    "created_at": "2026-07-11 11:00:00",
                },
                {
                    "id": 4,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "SAV-USD-1",
                    "currency_code": "USD",
                    "dr": 0,
                    "cr": 5,
                    "bal_before": 90,
                    "bal_after": 95,
                    "ref_no": "",
                    "description": "Retrait Vers M-Pesa",
                    "created_at": "2026-07-11 11:00:00",
                },
                {
                    "id": 5,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "MPESA ACCOUNT",
                    "reference_id": "DAT-USD-1",
                    "currency_code": "USD",
                    "dr": 100,
                    "cr": 0,
                    "bal_before": 95,
                    "bal_after": -5,
                    "ref_no": "DAT-1",
                    "description": "M-Pesa Compte",
                    "created_at": "2026-07-12 12:00:00",
                },
                {
                    "id": 6,
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "account_type": "FIXED SAVINGS",
                    "reference_id": "DAT-USD-1",
                    "currency_code": "USD",
                    "dr": 0,
                    "cr": 100,
                    "bal_before": 0,
                    "bal_after": 100,
                    "ref_no": "DAT-1",
                    "description": "Depot Bloque",
                    "created_at": "2026-07-12 12:00:00",
                },
            ]
        )
        prepared = MpesaPreparedData(
            transactions=prepare_transactions(transactions),
            current_savings=prepare_current_savings(savings),
            fixed_savings=prepare_fixed_savings_from_accounts(savings),
            loans=pd.DataFrame(),
            g2_transactions=pd.DataFrame(),
            customers=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )

        report = build_mpesa_savings_cockpit(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-31",
            annual_interest_rate_pct=11,
            maturity_horizon_days=30,
            large_savings_thresholds={"USD": 50, "CDF": 50000},
        )

        portfolio = report["portefeuille_synthese"].set_index(["currency_code", "famille_epargne"])
        self.assertEqual(float(portfolio.loc[("USD", "Compte ouvert"), "encours_actuel"]), 25.0)
        self.assertEqual(float(portfolio.loc[("USD", "Compte ouvert"), "comptes_solde_nul"]), 1.0)
        self.assertEqual(float(portfolio.loc[("USD", "DAT"), "encours_actuel"]), 100.0)
        self.assertEqual(float(portfolio.loc[("CDF", "DAT"), "encours_actuel"]), 0.0)

        flows = report["flux_synthese"].set_index("currency_code")
        self.assertEqual(float(flows.loc["USD", "montant_depots_compte_ouvert"]), 10.0)
        self.assertEqual(float(flows.loc["USD", "montant_depots_dat"]), 100.0)
        self.assertEqual(float(flows.loc["USD", "montant_retraits"]), 5.0)
        self.assertEqual(float(flows.loc["USD", "flux_net_epargne"]), 105.0)
        self.assertNotIn("CDF", flows.index)

        dat = report["dat_detail"].set_index("savings_id")
        self.assertAlmostEqual(
            float(dat.loc["DAT-USD-1", "interet_estime"]),
            100 * 0.11 * 30 / 365,
            places=6,
        )
        self.assertEqual(
            report["catalogue_kpi"].set_index("kpi").loc["renouvellement_dat", "statut"],
            "data_gap",
        )
        self.assertIn("forte_epargne_sans_credit", report["listes_action"])
        self.assertFalse(report["opportunites"].empty)

    def test_savings_cockpit_allows_zero_interest_and_exports(self) -> None:
        savings = pd.DataFrame(
            [
                {
                    "savings_id": "DAT-USD-1",
                    "customer_id": "C1",
                    "msisdn1": "0811111111",
                    "product_name": "1 Month",
                    "product_description": "1 Month Fixed Account",
                    "currency_code": "USD",
                    "balance": 100,
                    "status": "active",
                    "date_approved": "2026-07-01",
                    "maturity_date": "2026-07-31",
                    "created_at": "2026-07-01",
                }
            ]
        )
        prepared = MpesaPreparedData(
            transactions=pd.DataFrame(),
            current_savings=prepare_current_savings(savings),
            fixed_savings=prepare_fixed_savings_from_accounts(savings),
            loans=pd.DataFrame(),
            load_report=pd.DataFrame(),
        )
        report = build_mpesa_savings_cockpit(
            prepared,
            date_start="2026-07-01",
            date_end="2026-07-31",
            annual_interest_rate_pct=0,
        )

        dat = report["dat_detail"].iloc[0]
        self.assertTrue(pd.isna(dat["interet_estime"]))
        quality = report["qualite_donnees"].set_index("controle")
        self.assertEqual(quality.loc["Transactions potentiellement plafonnees", "statut"], "OK")

        export = create_excel_export(
            {
                "epargne_vue_ensemble": report["vue_ensemble"],
                "epargne_portefeuille_detail": report["portefeuille_detail"],
                "epargne_catalogue_kpi": report["catalogue_kpi"],
            },
            rename_user_columns=True,
        )
        workbook = pd.ExcelFile(BytesIO(export), engine="openpyxl")
        self.assertIn("Epargne_Vue_Ensemble", workbook.sheet_names)
        self.assertIn("Epargne_Detail", workbook.sheet_names)
        exported = pd.read_excel(workbook, sheet_name="Epargne_Detail")
        self.assertIn("numero_compte", exported.columns)


if __name__ == "__main__":
    unittest.main()
